# app/routers/plots.py

from fastapi import APIRouter, Depends, Body, HTTPException, BackgroundTasks, Query, Request
from pydantic import BaseModel
from sqlalchemy.orm import Session
from geoalchemy2.shape import from_shape
from shapely.geometry import Polygon, shape, Point, box, mapping
from shapely import wkt as shapely_wkt
from shapely.affinity import rotate
from shapely.ops import unary_union, snap
from sqlalchemy import text
from fastapi.responses import FileResponse
from app.schemas.plot_create import PlotCreateRequest
from typing import Optional, Union, List, Any

import os
import tempfile
import glob
import re
import shutil
import zipfile
import csv
import time
import json
import hashlib
import math
import textwrap
import uuid
import threading
import mimetypes
import logging
from threading import Lock

logger = logging.getLogger(__name__)
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches

from app.db import SessionLocal
from app.models.plot import Plot
from app.models.plot_buffer import PlotBuffer
from app.utils.pdf import generate_plot_report_pdf
from app.utils.survey_auth_security import require_survey_session, resolve_survey_session
from app.utils.survey_activity import ensure_survey_activity_table, log_survey_activity
from app.utils.map_renderer_layout import (
    render_plot_map_layout,
    get_paper_config,
    parse_scale_ratio,
    is_auto_scale_text,
    compute_fit_scale_ratio,
    apply_true_scale,
    annotate_vertices,
    draw_building_hatch,
    draw_fences,
    build_fence_avoid_geom,
    add_north_arrow,
    _collect_connected_road_edge_lines,
    format_area_display,
)
from app.utils.back_computation import compute_back_computation
from app.utils.back_computation_pdf import render_back_computation_pdf
from app.utils.coordinate_converter import resolve_coordinate_system_key, validate_nigeria_bounds
from shapely import wkb
from shapely.errors import GEOSException
import geopandas as gpd
from app.utils.dwg_exporter import export_survey_plan_to_dxf
from app.utils.r2_exports import upload_export_file_best_effort


from app.utils.orthophoto_renderer import (
    render_orthophoto_png,
    render_orthophoto_pdf_from_png
)

router = APIRouter(prefix="/plots", tags=["plots"])

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
PREVIEW_CACHE_DIR = os.path.join(REPORTS_DIR, "previews_cache")
PREVIEW_CACHE_TTL_SECONDS = max(30, int(os.getenv("PLOT_PREVIEW_CACHE_TTL_SECONDS", "180")))
PREVIEW_CACHE_MAX_FILES_PER_PLOT = max(5, int(os.getenv("PLOT_PREVIEW_CACHE_MAX_FILES_PER_PLOT", "24")))
PREVIEW_LAYOUT_VERSION = "survey_layout_2026_08_13_autofit_preview_v86"
SURVEY_REPORT_RENDER_VERSION = "survey_report_2026_03_23_clockwise_v1"
CLEAN_COPY_RENDER_VERSION = "clean_copy_2026_03_20_layout_v14"
PLOT_EXPORT_JOB_STATUS_VALUES = {"queued", "running", "completed", "failed"}
STANDARD_SURVEY_SCALE_DENOMINATORS = (250, 500, 1000, 1250, 2000, 2500, 5000, 10000, 12500, 20000, 25000, 50000)


class PlotGeometryUpdateRequest(BaseModel):
    coordinates: List[List[float]]

# Coordinate system EPSG codes mapping
COORDINATE_SYSTEMS = {
    "wgs84": 4326,
    "wgs84_nigeria_meters": 32632,
    "utm_31n": 32631,
    "utm_32n": 32632,
    "utm_33n": 32633,
    "minna_31": 26331,
    "minna_32": 26332,
    "minna_33": 26333,
    # Ghana - most of the country falls in UTM zone 30N; Leigon / Ghana Metre Grid (EPSG:25000) is
    # the national grid in use since 1978, superseding the older Accra / Ghana National Grid.
    "ghana_utm_30n": 32630,
    "ghana_leigon_grid": 25000,
    # Uganda spans UTM zones 35N (west of 30E) and 36N (east of 30E); Arc 1960 is the pre-GPS local
    # datum still used on older cadastral records, WGS84 UTM is the modern GPS-compatible option.
    "uganda_utm_35n": 32635,
    "uganda_utm_36n": 32636,
    "uganda_arc1960_35n": 21095,
    "uganda_arc1960_36n": 21096,
    # Uganda dips south of the equator too (Lake Victoria/Masaka/Kabale area) - EPSG explicitly
    # documents 21035/21036 as covering "Uganda - south of equator" for west/east of 30E
    # respectively (verified against epsg.io), so these are additions alongside the *N zones
    # above, not replacements. Zone 37S (EPSG:21037) is Kenya/Tanzania only, not Uganda's
    # territory, so it's deliberately not included here.
    "uganda_utm_35s": 32735,
    "uganda_utm_36s": 32736,
    "uganda_arc1960_35s": 21035,
    "uganda_arc1960_36s": 21036,
}

COORDINATE_SYSTEM_NAMES = {
    "wgs84": "WGS84 (Lat/Lon)",
    "wgs84_nigeria_meters": "WGS84 Nigeria Metres",
    "utm_31n": "UTM Zone 31N",
    "utm_32n": "UTM Zone 32N",
    "utm_33n": "UTM Zone 33N",
    "minna_31": "Minna Datum Zone 31",
    "minna_32": "Minna Datum Zone 32",
    "minna_33": "Minna Datum Zone 33",
    "ghana_utm_30n": "Ghana UTM Zone 30N",
    "ghana_leigon_grid": "Ghana Leigon National Grid",
    "uganda_utm_35n": "Uganda UTM Zone 35N",
    "uganda_utm_36n": "Uganda UTM Zone 36N",
    "uganda_arc1960_35n": "Uganda Arc 1960 Zone 35N",
    "uganda_arc1960_36n": "Uganda Arc 1960 Zone 36N",
    "uganda_utm_35s": "Uganda UTM Zone 35S",
    "uganda_utm_36s": "Uganda UTM Zone 36S",
    "uganda_arc1960_35s": "Uganda Arc 1960 Zone 35S",
    "uganda_arc1960_36s": "Uganda Arc 1960 Zone 36S",
}

# Which country a coordinate system key belongs to - drives the country-grouped picker on the
# frontend and (via plot centroid, not this map) the Overpass-vs-local-table feature detection
# dispatch in _run_plot_feature_detection.
COORDINATE_SYSTEM_COUNTRY = {
    "wgs84": "Global",
    "wgs84_nigeria_meters": "Nigeria",
    "utm_31n": "Nigeria",
    "utm_32n": "Nigeria",
    "utm_33n": "Nigeria",
    "minna_31": "Nigeria",
    "minna_32": "Nigeria",
    "minna_33": "Nigeria",
    "ghana_utm_30n": "Ghana",
    "ghana_leigon_grid": "Ghana",
    "uganda_utm_35n": "Uganda",
    "uganda_utm_36n": "Uganda",
    "uganda_arc1960_35n": "Uganda",
    "uganda_arc1960_36n": "Uganda",
    "uganda_utm_35s": "Uganda",
    "uganda_utm_36s": "Uganda",
    "uganda_arc1960_35s": "Uganda",
    "uganda_arc1960_36s": "Uganda",
}

DEFAULT_CERTIFICATION_STATEMENT = (
    "I hereby certify that this survey plan is a true representation of the survey "
    "executed by me and conforms with the regulations of surveying profession."
)
DEFAULT_TEMPLATE_NAME = "general"
DEFAULT_ADAMAWA_AUTHORITY_TITLE = "SURVEYOR GENERAL"
DEFAULT_ADAMAWA_AUTHORITY_DATE = "November, 2024"
DEFAULT_ADAMAWA_ORIGIN_TEXT = "ORIGIN:- WGS 84 UTM ZONE 33N"
DEFAULT_ADAMAWA_TOPO_SHEET_TEXT = "BASED ON GIREI TOPO SHEET 197 NE"
DEFAULT_ADAMAWA_DISCLAIMER_TEXT = (
    "Detail shewn not the result of accurate survey. All bearing and distances shewn on this plan "
    "have been computed from registered Co-ordinates."
)
DEFAULT_TECHNICAL_REPORT_COMPUTATION_SOFTWARE = "AutoCAD software"
DEFAULT_TECHNICAL_REPORT_PLOTTING_SOFTWARE = "AutoCAD software"
DEFAULT_TECHNICAL_REPORT_GENERAL_OBSERVATION = "The work was hitch-free."

_PLOTS_SCHEMA_READY = False
_PLOTS_SCHEMA_LOCK = Lock()
_PLOT_META_TABLE_READY = False
_PLOT_META_TABLE_LOCK = Lock()


def ensure_plots_schema_once(db: Session):
    global _PLOTS_SCHEMA_READY
    if _PLOTS_SCHEMA_READY:
        return
    with _PLOTS_SCHEMA_LOCK:
        if _PLOTS_SCHEMA_READY:
            return
        ensure_plot_meta_table(db)
        ensure_plot_feature_overrides_table(db)
        ensure_plot_subdivision_tables(db)
        ensure_plot_export_jobs_table(db)
        ensure_plot_query_indexes(db)
        ensure_plot_idempotency_columns(db)
        ensure_plot_ownership_column(db)
        _PLOTS_SCHEMA_READY = True


def get_db():
    db = SessionLocal()
    try:
        ensure_plots_schema_once(db)
        yield db
    finally:
        db.close()


def _table_exists(db: Session, table_name: str) -> bool:
    try:
        return db.execute(text("SELECT to_regclass(:table_name)"), {"table_name": table_name}).scalar() is not None
    except Exception:
        return False


def _safe_run_ddl(db: Session, ddl_sql: str):
    try:
        db.execute(text(ddl_sql))
        db.commit()
    except Exception:
        db.rollback()


def ensure_plot_query_indexes(db: Session):
    index_plan = [
        ("public.plots", "CREATE INDEX IF NOT EXISTS idx_plots_geom ON plots USING GIST (geom)"),
        ("public.plot_buffers", "CREATE INDEX IF NOT EXISTS idx_plot_buffers_plot_id ON plot_buffers(plot_id)"),
        ("public.plot_buffers", "CREATE INDEX IF NOT EXISTS idx_plot_buffers_geom ON plot_buffers USING GIST (geom)"),
        ("public.detected_features", "CREATE INDEX IF NOT EXISTS idx_detected_features_plot_id ON detected_features(plot_id)"),
        ("public.detected_features", "CREATE INDEX IF NOT EXISTS idx_detected_features_plot_type ON detected_features(plot_id, feature_type)"),
        ("public.detected_features", "CREATE INDEX IF NOT EXISTS idx_detected_features_geom ON detected_features USING GIST (geom)"),
        ("public.lines", "CREATE INDEX IF NOT EXISTS idx_lines_geom ON lines USING GIST (geom)"),
        ("public.lines", "CREATE INDEX IF NOT EXISTS idx_lines_highway_not_null ON lines(highway) WHERE highway IS NOT NULL"),
        ("public.lines", "CREATE INDEX IF NOT EXISTS idx_lines_waterway_not_null ON lines(waterway) WHERE waterway IS NOT NULL"),
        ("public.multipolygons", "CREATE INDEX IF NOT EXISTS idx_multipolygons_geom ON multipolygons USING GIST (geom)"),
        ("public.multipolygons", "CREATE INDEX IF NOT EXISTS idx_multipolygons_building_not_null ON multipolygons(building) WHERE building IS NOT NULL"),
        ("public.multilinestrings", "CREATE INDEX IF NOT EXISTS idx_multilinestrings_geom ON multilinestrings USING GIST (geom)"),
        ("public.multilinestrings", "CREATE INDEX IF NOT EXISTS idx_multilinestrings_type ON multilinestrings(type)"),
    ]
    for table_name, ddl_sql in index_plan:
        if _table_exists(db, table_name):
            _safe_run_ddl(db, ddl_sql)

    # `name`/`subtype` on detected_features - nullable, so Nigeria rows (which never set them,
    # since Nigeria's live "lines"/"multipolygons" tables already carry name/highway themselves)
    # are unaffected. Populated only by the Overpass-backed regional cache for non-Nigeria
    # countries (see app/utils/osm_overpass.py), so a road detected there can still show its real
    # OSM name on the general template's final export the same way a Nigeria road already can.
    if _table_exists(db, "public.detected_features"):
        _safe_run_ddl(db, "ALTER TABLE detected_features ADD COLUMN IF NOT EXISTS name TEXT")
        _safe_run_ddl(db, "ALTER TABLE detected_features ADD COLUMN IF NOT EXISTS subtype TEXT")


def ensure_plot_meta_table(db: Session):
    # This is called from several hot paths (create_plot, upsert_plot_meta - itself called on
    # every metadata save and every preview/orthophoto render), not just once at startup like
    # the rest of the schema bootstrap. The CREATE + ~38 ALTER statements below are a no-op after
    # the first successful run, so guard them the same way ensure_plots_schema_once already does
    # to avoid paying that DDL round-trip cost on every single save/preview request.
    global _PLOT_META_TABLE_READY
    if _PLOT_META_TABLE_READY:
        return
    with _PLOT_META_TABLE_LOCK:
        if _PLOT_META_TABLE_READY:
            return
        _ensure_plot_meta_table_impl(db)
        _PLOT_META_TABLE_READY = True


def _ensure_plot_meta_table_impl(db: Session):
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS plot_meta (
            plot_id INTEGER PRIMARY KEY REFERENCES plots(id) ON DELETE CASCADE,
            title_text VARCHAR(255),
            location_text TEXT,
            lga_text TEXT,
            state_text TEXT,
            surveyor_name TEXT,
            surveyor_rank TEXT,
            certification_statement TEXT,
            scale_text VARCHAR(50),
            paper_size VARCHAR(10),
            coordinate_system VARCHAR(20),
            template_name VARCHAR(40) DEFAULT 'general',
            adamawa_rof_no TEXT,
            adamawa_owner_name TEXT,
            adamawa_authority_title TEXT,
            adamawa_authority_date_text TEXT,
            adamawa_control_point_name TEXT,
            adamawa_northing TEXT,
            adamawa_easting TEXT,
            adamawa_elevation TEXT,
            adamawa_origin_text TEXT,
            adamawa_topo_sheet_text TEXT,
            adamawa_computation_no TEXT,
            adamawa_cadastral_sheet_no TEXT,
            adamawa_plan_no TEXT,
            adamawa_surveyed_by_text TEXT,
            adamawa_disclaimer_text TEXT,
            cadastral_plan_no TEXT,
            cadastral_area_name TEXT,
            cadastral_datum_text TEXT,
            cadastral_firm_block_text TEXT,
            fct_file_no TEXT,
            fct_district TEXT,
            fct_cadastral_zone TEXT,
            fct_origin_beacon_text TEXT,
            fct_cadastral_map_ref TEXT,
            fct_title_prefix TEXT,
            technical_report_instruments JSONB DEFAULT '[]',
            technical_report_dgps_type TEXT,
            technical_report_num_surveyors INTEGER,
            technical_report_num_technical_officers INTEGER,
            technical_report_num_labourers INTEGER,
            technical_report_recce_text TEXT,
            technical_report_demarcation_text TEXT,
            technical_report_computation_software_text TEXT,
            technical_report_plotting_software_text TEXT,
            technical_report_general_observation_text TEXT,
            elevation_points JSONB DEFAULT '[]',
            survey_input_coordinates JSONB DEFAULT '[]',
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    # Ensure columns exist if the table was created previously with missing fields
    columns_to_add = [
        ("title_text", "VARCHAR(255)"),
        ("location_text", "TEXT"),
        ("lga_text", "TEXT"),
        ("state_text", "TEXT"),
        ("surveyor_name", "TEXT"),
        ("surveyor_rank", "TEXT"),
        ("certification_statement", "TEXT"),
        ("scale_text", "VARCHAR(50)"),
        ("paper_size", "VARCHAR(10)"),
        ("coordinate_system", "VARCHAR(20)"),
        ("template_name", "VARCHAR(40) DEFAULT 'general'"),
        ("adamawa_rof_no", "TEXT"),
        ("adamawa_owner_name", "TEXT"),
        ("adamawa_authority_title", "TEXT"),
        ("adamawa_authority_date_text", "TEXT"),
        ("adamawa_control_point_name", "TEXT"),
        ("adamawa_northing", "TEXT"),
        ("adamawa_easting", "TEXT"),
        ("adamawa_elevation", "TEXT"),
        ("adamawa_origin_text", "TEXT"),
        ("adamawa_topo_sheet_text", "TEXT"),
        ("adamawa_computation_no", "TEXT"),
        ("adamawa_cadastral_sheet_no", "TEXT"),
        ("adamawa_plan_no", "TEXT"),
        ("adamawa_surveyed_by_text", "TEXT"),
        ("adamawa_disclaimer_text", "TEXT"),
        ("cadastral_plan_no", "TEXT"),
        ("cadastral_area_name", "TEXT"),
        ("cadastral_datum_text", "TEXT"),
        ("cadastral_firm_block_text", "TEXT"),
        ("fct_file_no", "TEXT"),
        ("fct_district", "TEXT"),
        ("fct_cadastral_zone", "TEXT"),
        ("fct_origin_beacon_text", "TEXT"),
        ("fct_cadastral_map_ref", "TEXT"),
        ("fct_title_prefix", "TEXT"),
        ("technical_report_instruments", "JSONB DEFAULT '[]'"),
        ("technical_report_dgps_type", "TEXT"),
        ("technical_report_num_surveyors", "INTEGER"),
        ("technical_report_num_technical_officers", "INTEGER"),
        ("technical_report_num_labourers", "INTEGER"),
        ("technical_report_recce_text", "TEXT"),
        ("technical_report_demarcation_text", "TEXT"),
        ("technical_report_computation_software_text", "TEXT"),
        ("technical_report_plotting_software_text", "TEXT"),
        ("technical_report_general_observation_text", "TEXT"),
        ("elevation_points", "JSONB DEFAULT '[]'"),
        ("survey_input_coordinates", "JSONB DEFAULT '[]'"),
        ("created_at", "TIMESTAMP DEFAULT NOW()"),
        ("updated_at", "TIMESTAMP DEFAULT NOW()"),
    ]
    for col_name, col_type in columns_to_add:
        try:
            db.execute(text(f"ALTER TABLE plot_meta ADD COLUMN IF NOT EXISTS {col_name} {col_type}"))
        except Exception:
            # Ignore for databases that don't support IF NOT EXISTS or already have column
            pass
    db.commit()


def ensure_plots_created_at(db: Session):
    try:
        db.execute(text("ALTER TABLE plots ADD COLUMN IF NOT EXISTS created_at TIMESTAMP DEFAULT NOW()"))
        db.commit()
    except Exception:
        db.rollback()


def ensure_plot_idempotency_columns(db: Session):
    # Lets a client resend the same "create plot" / "save feature edit" request (e.g. after a
    # dropped response on a flaky connection) without risking a duplicate row - the caller sends
    # a per-attempt id and re-sends the SAME id on retry; a genuinely new action gets a new id.
    # A partial unique index (rather than a plain UNIQUE column) means requests that don't send an
    # id at all (NULL) never collide with each other, so older/other callers are unaffected.
    try:
        db.execute(text("ALTER TABLE plots ADD COLUMN IF NOT EXISTS client_request_id TEXT"))
        db.execute(text(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_plots_client_request_id "
            "ON plots (client_request_id) WHERE client_request_id IS NOT NULL"
        ))
        db.execute(text("ALTER TABLE plot_feature_overrides ADD COLUMN IF NOT EXISTS client_request_id TEXT"))
        db.execute(text(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_plot_feature_overrides_client_request_id "
            "ON plot_feature_overrides (client_request_id) WHERE client_request_id IS NOT NULL"
        ))
        db.commit()
    except Exception:
        db.rollback()


def ensure_plot_ownership_column(db: Session):
    # Nullable so anonymous plot creation (the "value first" gate-free flow) is unaffected -
    # only plots created by (or later claimed by) a signed-in Survey user get this set.
    try:
        db.execute(text("ALTER TABLE plots ADD COLUMN IF NOT EXISTS owner_user_id BIGINT"))
        db.execute(text("CREATE INDEX IF NOT EXISTS idx_plots_owner_user_id ON plots (owner_user_id)"))
        db.commit()
    except Exception:
        db.rollback()


def ensure_plot_feature_overrides_table(db: Session):
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS plot_feature_overrides (
            id SERIAL PRIMARY KEY,
            plot_id INTEGER NOT NULL REFERENCES plots(id) ON DELETE CASCADE,
            feature_type TEXT NOT NULL CHECK (feature_type IN ('road', 'building', 'river', 'fence')),
            action TEXT NOT NULL CHECK (action IN ('add', 'delete', 'update')),
            name TEXT,
            width_m DOUBLE PRECISION,
            geom GEOMETRY,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    try:
        db.execute(text("ALTER TABLE plot_feature_overrides ADD COLUMN IF NOT EXISTS width_m DOUBLE PRECISION"))
        # Upgrade legacy CHECK constraints that do not include fence
        check_rows = db.execute(text("""
            SELECT c.conname AS name, pg_get_constraintdef(c.oid) AS definition
            FROM pg_constraint c
            JOIN pg_class t ON t.oid = c.conrelid
            JOIN pg_namespace n ON n.oid = t.relnamespace
            WHERE t.relname = 'plot_feature_overrides'
              AND c.contype = 'c'
        """)).mappings().all()
        has_fence_check = False
        for row in check_rows:
            definition = (row.get("definition") or "").lower()
            if "feature_type" not in definition:
                continue
            if "fence" in definition:
                has_fence_check = True
                continue
            constraint_name = (row.get("name") or "").replace('"', '""')
            if constraint_name:
                db.execute(text(f'ALTER TABLE plot_feature_overrides DROP CONSTRAINT IF EXISTS "{constraint_name}"'))
        if not has_fence_check:
            db.execute(text("""
                ALTER TABLE plot_feature_overrides
                ADD CONSTRAINT plot_feature_overrides_feature_type_check
                CHECK (feature_type IN ('road', 'building', 'river', 'fence'))
            """))
        db.execute(text("CREATE INDEX IF NOT EXISTS idx_plot_feature_overrides_plot_id ON plot_feature_overrides(plot_id)"))
        db.execute(text("CREATE INDEX IF NOT EXISTS idx_plot_feature_overrides_geom ON plot_feature_overrides USING GIST (geom)"))
        db.commit()
    except Exception:
        db.rollback()


def ensure_plot_subdivision_tables(db: Session):
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS plot_subdivision_batches (
                id SERIAL PRIMARY KEY,
                parent_plot_id INTEGER NOT NULL REFERENCES plots(id) ON DELETE CASCADE,
                estate_name TEXT,
                method TEXT NOT NULL,
                requested_count INTEGER,
                target_area_m2 DOUBLE PRECISION,
                orientation_deg DOUBLE PRECISION DEFAULT 0,
                generated_count INTEGER DEFAULT 0,
                total_area_m2 DOUBLE PRECISION DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'completed',
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
            """
        )
    )
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS plot_subdivision_items (
                id SERIAL PRIMARY KEY,
                batch_id INTEGER NOT NULL REFERENCES plot_subdivision_batches(id) ON DELETE CASCADE,
                child_plot_id INTEGER NOT NULL REFERENCES plots(id) ON DELETE CASCADE,
                lot_no TEXT,
                area_m2 DOUBLE PRECISION,
                created_at TIMESTAMP DEFAULT NOW()
            )
            """
        )
    )
    for col_name, col_type in [
        ("parent_plot_id", "INTEGER"),
        ("subdivision_batch_id", "INTEGER"),
        ("subdivision_lot_no", "TEXT"),
        ("estate_name", "TEXT"),
    ]:
        try:
            db.execute(text(f"ALTER TABLE plot_meta ADD COLUMN IF NOT EXISTS {col_name} {col_type}"))
        except Exception:
            pass

    # Added for road-exclusion and fixed-dimension-lot subdivision - nullable/defaulted so
    # existing batch rows and readers are unaffected.
    for col_name, col_type in [
        ("exclude_road", "BOOLEAN DEFAULT FALSE"),
        ("road_width_m", "DOUBLE PRECISION"),
        ("lot_width_m", "DOUBLE PRECISION"),
        ("lot_height_m", "DOUBLE PRECISION"),
        ("dimension_unit", "TEXT"),
        ("leftover_area_m2", "DOUBLE PRECISION"),
        ("excluded_area_m2", "DOUBLE PRECISION"),
        ("leftover_geojson", "JSONB"),
        ("excluded_geojson", "JSONB"),
    ]:
        try:
            db.execute(text(f"ALTER TABLE plot_subdivision_batches ADD COLUMN IF NOT EXISTS {col_name} {col_type}"))
        except Exception:
            pass

    index_sql = [
        "CREATE INDEX IF NOT EXISTS idx_plot_sub_batches_parent ON plot_subdivision_batches(parent_plot_id, created_at DESC)",
        "CREATE INDEX IF NOT EXISTS idx_plot_sub_items_batch ON plot_subdivision_items(batch_id)",
        "CREATE INDEX IF NOT EXISTS idx_plot_sub_items_child ON plot_subdivision_items(child_plot_id)",
        "CREATE INDEX IF NOT EXISTS idx_plot_meta_parent_plot_id ON plot_meta(parent_plot_id)",
        "CREATE INDEX IF NOT EXISTS idx_plot_meta_subdivision_batch ON plot_meta(subdivision_batch_id)",
    ]
    for ddl in index_sql:
        try:
            db.execute(text(ddl))
        except Exception:
            pass
    db.commit()


def ensure_plot_export_jobs_table(db: Session):
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS plot_export_jobs (
                id TEXT PRIMARY KEY,
                export_type TEXT NOT NULL,
                plot_id INTEGER REFERENCES plots(id) ON DELETE CASCADE,
                subdivision_batch_id INTEGER REFERENCES plot_subdivision_batches(id) ON DELETE CASCADE,
                cache_key TEXT,
                status TEXT NOT NULL DEFAULT 'queued',
                file_name TEXT,
                local_path TEXT,
                request_payload JSONB,
                error_text TEXT,
                started_at TIMESTAMP,
                completed_at TIMESTAMP,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
            """
        )
    )
    try:
        db.execute(text("ALTER TABLE plot_export_jobs ADD COLUMN IF NOT EXISTS cache_key TEXT"))
        db.execute(text("ALTER TABLE plot_export_jobs ADD COLUMN IF NOT EXISTS request_payload JSONB"))
    except Exception:
        pass
    for ddl in [
        "CREATE INDEX IF NOT EXISTS idx_plot_export_jobs_status_created ON plot_export_jobs(status, created_at DESC)",
        "CREATE INDEX IF NOT EXISTS idx_plot_export_jobs_batch_created ON plot_export_jobs(subdivision_batch_id, created_at DESC)",
        "CREATE INDEX IF NOT EXISTS idx_plot_export_jobs_type_batch_created ON plot_export_jobs(export_type, subdivision_batch_id, created_at DESC)",
        "CREATE INDEX IF NOT EXISTS idx_plot_export_jobs_plot_cache_created ON plot_export_jobs(plot_id, export_type, cache_key, created_at DESC)",
    ]:
        try:
            db.execute(text(ddl))
        except Exception:
            pass
    db.commit()


def upsert_plot_meta(
    db: Session,
    plot_id: int,
    title_text: Optional[str] = None,
    location_text: Optional[str] = None,
    lga_text: Optional[str] = None,
    state_text: Optional[str] = None,
    surveyor_name: Optional[str] = None,
    surveyor_rank: Optional[str] = None,
    certification_statement: Optional[str] = None,
    scale_text: Optional[str] = None,
    paper_size: Optional[str] = None,
    coordinate_system: Optional[str] = None,
    template_name: Optional[str] = None,
    adamawa_rof_no: Optional[str] = None,
    adamawa_owner_name: Optional[str] = None,
    adamawa_authority_title: Optional[str] = None,
    adamawa_authority_date_text: Optional[str] = None,
    adamawa_control_point_name: Optional[str] = None,
    adamawa_northing: Optional[str] = None,
    adamawa_easting: Optional[str] = None,
    adamawa_elevation: Optional[str] = None,
    adamawa_origin_text: Optional[str] = None,
    adamawa_topo_sheet_text: Optional[str] = None,
    adamawa_computation_no: Optional[str] = None,
    adamawa_cadastral_sheet_no: Optional[str] = None,
    adamawa_plan_no: Optional[str] = None,
    adamawa_surveyed_by_text: Optional[str] = None,
    adamawa_disclaimer_text: Optional[str] = None,
    cadastral_plan_no: Optional[str] = None,
    cadastral_area_name: Optional[str] = None,
    cadastral_datum_text: Optional[str] = None,
    cadastral_firm_block_text: Optional[str] = None,
    fct_file_no: Optional[str] = None,
    fct_district: Optional[str] = None,
    fct_cadastral_zone: Optional[str] = None,
    fct_origin_beacon_text: Optional[str] = None,
    fct_cadastral_map_ref: Optional[str] = None,
    fct_title_prefix: Optional[str] = None,
    technical_report_instruments: Optional[list] = None,
    technical_report_dgps_type: Optional[str] = None,
    technical_report_num_surveyors: Optional[int] = None,
    technical_report_num_technical_officers: Optional[int] = None,
    technical_report_num_labourers: Optional[int] = None,
    technical_report_recce_text: Optional[str] = None,
    technical_report_demarcation_text: Optional[str] = None,
    technical_report_computation_software_text: Optional[str] = None,
    technical_report_plotting_software_text: Optional[str] = None,
    technical_report_general_observation_text: Optional[str] = None,
    elevation_points: Optional[list] = None,
    survey_input_coordinates: Optional[list] = None,
    commit: bool = True,
):
    ensure_plot_meta_table(db)

    technical_report_instruments_json = (
        json.dumps(technical_report_instruments) if technical_report_instruments is not None else None
    )
    elevation_points_json = json.dumps(elevation_points) if elevation_points is not None else None
    survey_input_coordinates_json = (
        json.dumps(survey_input_coordinates) if survey_input_coordinates is not None else None
    )

    db.execute(text("""
        INSERT INTO plot_meta (
            plot_id, title_text, location_text, lga_text, state_text,
            surveyor_name, surveyor_rank, certification_statement, scale_text, paper_size, coordinate_system,
            template_name, adamawa_rof_no, adamawa_owner_name, adamawa_authority_title, adamawa_authority_date_text,
            adamawa_control_point_name, adamawa_northing, adamawa_easting, adamawa_elevation, adamawa_origin_text,
            adamawa_topo_sheet_text, adamawa_computation_no, adamawa_cadastral_sheet_no, adamawa_plan_no,
            adamawa_surveyed_by_text, adamawa_disclaimer_text,
            cadastral_plan_no, cadastral_area_name, cadastral_datum_text, cadastral_firm_block_text,
            fct_file_no, fct_district, fct_cadastral_zone, fct_origin_beacon_text, fct_cadastral_map_ref,
            fct_title_prefix,
            technical_report_instruments, technical_report_dgps_type,
            technical_report_num_surveyors, technical_report_num_technical_officers,
            technical_report_num_labourers, technical_report_recce_text,
            technical_report_demarcation_text, technical_report_computation_software_text,
            technical_report_plotting_software_text, technical_report_general_observation_text,
            elevation_points, survey_input_coordinates
        )
        VALUES (
            :plot_id, :title_text, :location_text, :lga_text, :state_text,
            :surveyor_name, :surveyor_rank, :certification_statement, :scale_text, :paper_size, :coordinate_system,
            :template_name, :adamawa_rof_no, :adamawa_owner_name, :adamawa_authority_title, :adamawa_authority_date_text,
            :adamawa_control_point_name, :adamawa_northing, :adamawa_easting, :adamawa_elevation, :adamawa_origin_text,
            :adamawa_topo_sheet_text, :adamawa_computation_no, :adamawa_cadastral_sheet_no, :adamawa_plan_no,
            :adamawa_surveyed_by_text, :adamawa_disclaimer_text,
            :cadastral_plan_no, :cadastral_area_name, :cadastral_datum_text, :cadastral_firm_block_text,
            :fct_file_no, :fct_district, :fct_cadastral_zone, :fct_origin_beacon_text, :fct_cadastral_map_ref,
            :fct_title_prefix,
            CAST(COALESCE(:technical_report_instruments, '[]') AS JSONB), :technical_report_dgps_type,
            :technical_report_num_surveyors, :technical_report_num_technical_officers,
            :technical_report_num_labourers, :technical_report_recce_text,
            :technical_report_demarcation_text, :technical_report_computation_software_text,
            :technical_report_plotting_software_text, :technical_report_general_observation_text,
            CAST(COALESCE(:elevation_points, '[]') AS JSONB),
            CAST(COALESCE(:survey_input_coordinates, '[]') AS JSONB)
        )
        ON CONFLICT (plot_id) DO UPDATE SET
            title_text = COALESCE(NULLIF(EXCLUDED.title_text, ''), plot_meta.title_text),
            location_text = COALESCE(NULLIF(EXCLUDED.location_text, ''), plot_meta.location_text),
            lga_text = COALESCE(NULLIF(EXCLUDED.lga_text, ''), plot_meta.lga_text),
            state_text = COALESCE(NULLIF(EXCLUDED.state_text, ''), plot_meta.state_text),
            surveyor_name = COALESCE(NULLIF(EXCLUDED.surveyor_name, ''), plot_meta.surveyor_name),
            surveyor_rank = COALESCE(NULLIF(EXCLUDED.surveyor_rank, ''), plot_meta.surveyor_rank),
            certification_statement = COALESCE(NULLIF(EXCLUDED.certification_statement, ''), plot_meta.certification_statement),
            scale_text = COALESCE(NULLIF(EXCLUDED.scale_text, ''), plot_meta.scale_text),
            paper_size = COALESCE(NULLIF(EXCLUDED.paper_size, ''), plot_meta.paper_size),
            coordinate_system = COALESCE(NULLIF(EXCLUDED.coordinate_system, ''), plot_meta.coordinate_system),
            template_name = COALESCE(NULLIF(EXCLUDED.template_name, ''), plot_meta.template_name),
            adamawa_rof_no = COALESCE(NULLIF(EXCLUDED.adamawa_rof_no, ''), plot_meta.adamawa_rof_no),
            adamawa_owner_name = COALESCE(NULLIF(EXCLUDED.adamawa_owner_name, ''), plot_meta.adamawa_owner_name),
            adamawa_authority_title = COALESCE(NULLIF(EXCLUDED.adamawa_authority_title, ''), plot_meta.adamawa_authority_title),
            adamawa_authority_date_text = COALESCE(NULLIF(EXCLUDED.adamawa_authority_date_text, ''), plot_meta.adamawa_authority_date_text),
            adamawa_control_point_name = COALESCE(NULLIF(EXCLUDED.adamawa_control_point_name, ''), plot_meta.adamawa_control_point_name),
            adamawa_northing = COALESCE(NULLIF(EXCLUDED.adamawa_northing, ''), plot_meta.adamawa_northing),
            adamawa_easting = COALESCE(NULLIF(EXCLUDED.adamawa_easting, ''), plot_meta.adamawa_easting),
            adamawa_elevation = COALESCE(NULLIF(EXCLUDED.adamawa_elevation, ''), plot_meta.adamawa_elevation),
            adamawa_origin_text = COALESCE(NULLIF(EXCLUDED.adamawa_origin_text, ''), plot_meta.adamawa_origin_text),
            adamawa_topo_sheet_text = COALESCE(NULLIF(EXCLUDED.adamawa_topo_sheet_text, ''), plot_meta.adamawa_topo_sheet_text),
            adamawa_computation_no = COALESCE(NULLIF(EXCLUDED.adamawa_computation_no, ''), plot_meta.adamawa_computation_no),
            adamawa_cadastral_sheet_no = COALESCE(NULLIF(EXCLUDED.adamawa_cadastral_sheet_no, ''), plot_meta.adamawa_cadastral_sheet_no),
            adamawa_plan_no = COALESCE(NULLIF(EXCLUDED.adamawa_plan_no, ''), plot_meta.adamawa_plan_no),
            adamawa_surveyed_by_text = COALESCE(NULLIF(EXCLUDED.adamawa_surveyed_by_text, ''), plot_meta.adamawa_surveyed_by_text),
            adamawa_disclaimer_text = COALESCE(NULLIF(EXCLUDED.adamawa_disclaimer_text, ''), plot_meta.adamawa_disclaimer_text),
            cadastral_plan_no = COALESCE(NULLIF(EXCLUDED.cadastral_plan_no, ''), plot_meta.cadastral_plan_no),
            cadastral_area_name = COALESCE(NULLIF(EXCLUDED.cadastral_area_name, ''), plot_meta.cadastral_area_name),
            cadastral_datum_text = COALESCE(NULLIF(EXCLUDED.cadastral_datum_text, ''), plot_meta.cadastral_datum_text),
            cadastral_firm_block_text = COALESCE(NULLIF(EXCLUDED.cadastral_firm_block_text, ''), plot_meta.cadastral_firm_block_text),
            fct_file_no = COALESCE(NULLIF(EXCLUDED.fct_file_no, ''), plot_meta.fct_file_no),
            fct_district = COALESCE(NULLIF(EXCLUDED.fct_district, ''), plot_meta.fct_district),
            fct_cadastral_zone = COALESCE(NULLIF(EXCLUDED.fct_cadastral_zone, ''), plot_meta.fct_cadastral_zone),
            fct_origin_beacon_text = COALESCE(NULLIF(EXCLUDED.fct_origin_beacon_text, ''), plot_meta.fct_origin_beacon_text),
            fct_cadastral_map_ref = COALESCE(NULLIF(EXCLUDED.fct_cadastral_map_ref, ''), plot_meta.fct_cadastral_map_ref),
            fct_title_prefix = COALESCE(NULLIF(EXCLUDED.fct_title_prefix, ''), plot_meta.fct_title_prefix),
            technical_report_instruments = COALESCE(CAST(:technical_report_instruments AS JSONB), plot_meta.technical_report_instruments),
            technical_report_dgps_type = COALESCE(NULLIF(EXCLUDED.technical_report_dgps_type, ''), plot_meta.technical_report_dgps_type),
            technical_report_num_surveyors = COALESCE(EXCLUDED.technical_report_num_surveyors, plot_meta.technical_report_num_surveyors),
            technical_report_num_technical_officers = COALESCE(EXCLUDED.technical_report_num_technical_officers, plot_meta.technical_report_num_technical_officers),
            technical_report_num_labourers = COALESCE(EXCLUDED.technical_report_num_labourers, plot_meta.technical_report_num_labourers),
            technical_report_recce_text = COALESCE(NULLIF(EXCLUDED.technical_report_recce_text, ''), plot_meta.technical_report_recce_text),
            technical_report_demarcation_text = COALESCE(NULLIF(EXCLUDED.technical_report_demarcation_text, ''), plot_meta.technical_report_demarcation_text),
            technical_report_computation_software_text = COALESCE(NULLIF(EXCLUDED.technical_report_computation_software_text, ''), plot_meta.technical_report_computation_software_text),
            technical_report_plotting_software_text = COALESCE(NULLIF(EXCLUDED.technical_report_plotting_software_text, ''), plot_meta.technical_report_plotting_software_text),
            technical_report_general_observation_text = COALESCE(NULLIF(EXCLUDED.technical_report_general_observation_text, ''), plot_meta.technical_report_general_observation_text),
            elevation_points = COALESCE(CAST(:elevation_points AS JSONB), plot_meta.elevation_points),
            survey_input_coordinates = COALESCE(CAST(:survey_input_coordinates AS JSONB), plot_meta.survey_input_coordinates),
            updated_at = NOW()
    """), {
        "plot_id": plot_id,
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "certification_statement": certification_statement,
        "scale_text": scale_text,
        "paper_size": paper_size,
        "coordinate_system": coordinate_system,
        "template_name": template_name,
        "adamawa_rof_no": adamawa_rof_no,
        "adamawa_owner_name": adamawa_owner_name,
        "adamawa_authority_title": adamawa_authority_title,
        "adamawa_authority_date_text": adamawa_authority_date_text,
        "adamawa_control_point_name": adamawa_control_point_name,
        "adamawa_northing": adamawa_northing,
        "adamawa_easting": adamawa_easting,
        "adamawa_elevation": adamawa_elevation,
        "adamawa_origin_text": adamawa_origin_text,
        "adamawa_topo_sheet_text": adamawa_topo_sheet_text,
        "adamawa_computation_no": adamawa_computation_no,
        "adamawa_cadastral_sheet_no": adamawa_cadastral_sheet_no,
        "adamawa_plan_no": adamawa_plan_no,
        "adamawa_surveyed_by_text": adamawa_surveyed_by_text,
        "adamawa_disclaimer_text": adamawa_disclaimer_text,
        "cadastral_plan_no": cadastral_plan_no,
        "cadastral_area_name": cadastral_area_name,
        "cadastral_datum_text": cadastral_datum_text,
        "cadastral_firm_block_text": cadastral_firm_block_text,
        "fct_file_no": fct_file_no,
        "fct_district": fct_district,
        "fct_cadastral_zone": fct_cadastral_zone,
        "fct_origin_beacon_text": fct_origin_beacon_text,
        "fct_cadastral_map_ref": fct_cadastral_map_ref,
        "fct_title_prefix": fct_title_prefix,
        "fct_title_prefix": fct_title_prefix,
        "technical_report_instruments": technical_report_instruments_json,
        "technical_report_dgps_type": technical_report_dgps_type,
        "technical_report_num_surveyors": technical_report_num_surveyors,
        "technical_report_num_technical_officers": technical_report_num_technical_officers,
        "technical_report_num_labourers": technical_report_num_labourers,
        "technical_report_recce_text": technical_report_recce_text,
        "technical_report_demarcation_text": technical_report_demarcation_text,
        "technical_report_computation_software_text": technical_report_computation_software_text,
        "technical_report_plotting_software_text": technical_report_plotting_software_text,
        "technical_report_general_observation_text": technical_report_general_observation_text,
        "elevation_points": elevation_points_json,
        "survey_input_coordinates": survey_input_coordinates_json,
    })
    if commit:
        db.commit()


def get_plot_meta(db: Session, plot_id: int) -> dict:
    # Guards against a freshly-deployed process whose first plot_meta-touching request is a
    # read (e.g. report/preview) rather than a save - without this, a column added in code but
    # not yet migrated onto the live table (ensure_plot_meta_table is otherwise only called from
    # upsert_plot_meta/create_plot) crashes this SELECT with "column does not exist", which surfaces
    # to the browser as an opaque CORS/network error since the response never gets CORS headers.
    # Cheap after the first call in a process's lifetime (see ensure_plot_meta_table's cached flag).
    ensure_plot_meta_table(db)
    row = db.execute(text("""
        SELECT title_text, location_text, lga_text, state_text,
               surveyor_name, surveyor_rank, certification_statement, scale_text, paper_size, coordinate_system,
               template_name, adamawa_rof_no, adamawa_owner_name, adamawa_authority_title, adamawa_authority_date_text,
               adamawa_control_point_name, adamawa_northing, adamawa_easting, adamawa_elevation, adamawa_origin_text,
               adamawa_topo_sheet_text, adamawa_computation_no, adamawa_cadastral_sheet_no, adamawa_plan_no,
               adamawa_surveyed_by_text, adamawa_disclaimer_text,
               cadastral_plan_no, cadastral_area_name, cadastral_datum_text, cadastral_firm_block_text,
               fct_file_no, fct_district, fct_cadastral_zone, fct_origin_beacon_text, fct_cadastral_map_ref,
               fct_title_prefix,
               parent_plot_id, subdivision_batch_id, subdivision_lot_no, estate_name,
               technical_report_instruments, technical_report_dgps_type,
               technical_report_num_surveyors, technical_report_num_technical_officers,
               technical_report_num_labourers, technical_report_recce_text,
               technical_report_demarcation_text, technical_report_computation_software_text,
               technical_report_plotting_software_text, technical_report_general_observation_text,
               elevation_points, survey_input_coordinates
        FROM plot_meta
        WHERE plot_id = :plot_id
    """), {"plot_id": plot_id}).mappings().first()
    if not row:
        return {
            "title_text": "SURVEY PLAN",
            "location_text": "",
            "lga_text": "",
            "state_text": "",
            "surveyor_name": "",
            "surveyor_rank": "",
            "certification_statement": DEFAULT_CERTIFICATION_STATEMENT,
            "scale_text": "1 : 1000",
            "paper_size": "A4",
            "coordinate_system": "wgs84",
            "template_name": DEFAULT_TEMPLATE_NAME,
            "adamawa_rof_no": "",
            "adamawa_owner_name": "",
            "adamawa_authority_title": DEFAULT_ADAMAWA_AUTHORITY_TITLE,
            "adamawa_authority_date_text": DEFAULT_ADAMAWA_AUTHORITY_DATE,
            "adamawa_control_point_name": "",
            "adamawa_northing": "",
            "adamawa_easting": "",
            "adamawa_elevation": "",
            "adamawa_origin_text": DEFAULT_ADAMAWA_ORIGIN_TEXT,
            "adamawa_topo_sheet_text": DEFAULT_ADAMAWA_TOPO_SHEET_TEXT,
            "adamawa_computation_no": "",
            "adamawa_cadastral_sheet_no": "",
            "adamawa_plan_no": "",
            "adamawa_surveyed_by_text": "",
            "adamawa_disclaimer_text": DEFAULT_ADAMAWA_DISCLAIMER_TEXT,
            "cadastral_plan_no": "",
            "cadastral_area_name": "",
            "cadastral_datum_text": "",
            "cadastral_firm_block_text": "",
            "fct_file_no": "",
            "fct_district": "",
            "fct_cadastral_zone": "",
            "fct_origin_beacon_text": "",
            "fct_cadastral_map_ref": "",
            "fct_title_prefix": "",
            "parent_plot_id": None,
            "subdivision_batch_id": None,
            "subdivision_lot_no": "",
            "estate_name": "",
            "technical_report_instruments": [],
            "technical_report_dgps_type": "",
            "technical_report_num_surveyors": None,
            "technical_report_num_technical_officers": None,
            "technical_report_num_labourers": None,
            "technical_report_recce_text": "",
            "technical_report_demarcation_text": "",
            "technical_report_computation_software_text": DEFAULT_TECHNICAL_REPORT_COMPUTATION_SOFTWARE,
            "technical_report_plotting_software_text": DEFAULT_TECHNICAL_REPORT_PLOTTING_SOFTWARE,
            "technical_report_general_observation_text": DEFAULT_TECHNICAL_REPORT_GENERAL_OBSERVATION,
            "elevation_points": [],
            "survey_input_coordinates": [],
        }
    raw_instruments = row.get("technical_report_instruments")
    if isinstance(raw_instruments, str):
        try:
            raw_instruments = json.loads(raw_instruments)
        except Exception:
            raw_instruments = []
    if not isinstance(raw_instruments, list):
        raw_instruments = []
    raw_elevation_points = row.get("elevation_points")
    if isinstance(raw_elevation_points, str):
        try:
            raw_elevation_points = json.loads(raw_elevation_points)
        except Exception:
            raw_elevation_points = []
    if not isinstance(raw_elevation_points, list):
        raw_elevation_points = []
    raw_survey_input_coordinates = row.get("survey_input_coordinates")
    if isinstance(raw_survey_input_coordinates, str):
        try:
            raw_survey_input_coordinates = json.loads(raw_survey_input_coordinates)
        except Exception:
            raw_survey_input_coordinates = []
    if not isinstance(raw_survey_input_coordinates, list):
        raw_survey_input_coordinates = []
    return {
        "title_text": row.get("title_text") or "SURVEY PLAN",
        "location_text": row.get("location_text") or "",
        "lga_text": row.get("lga_text") or "",
        "state_text": row.get("state_text") or "",
        "surveyor_name": row.get("surveyor_name") or "",
        "surveyor_rank": row.get("surveyor_rank") or "",
        "certification_statement": row.get("certification_statement") or DEFAULT_CERTIFICATION_STATEMENT,
        "scale_text": row.get("scale_text") or "1 : 1000",
        "paper_size": row.get("paper_size") or "A4",
        "coordinate_system": row.get("coordinate_system") or "wgs84",
        "template_name": row.get("template_name") or DEFAULT_TEMPLATE_NAME,
        "adamawa_rof_no": row.get("adamawa_rof_no") or "",
        "adamawa_owner_name": row.get("adamawa_owner_name") or "",
        "adamawa_authority_title": row.get("adamawa_authority_title") or DEFAULT_ADAMAWA_AUTHORITY_TITLE,
        "adamawa_authority_date_text": row.get("adamawa_authority_date_text") or DEFAULT_ADAMAWA_AUTHORITY_DATE,
        "adamawa_control_point_name": row.get("adamawa_control_point_name") or "",
        "adamawa_northing": row.get("adamawa_northing") or "",
        "adamawa_easting": row.get("adamawa_easting") or "",
        "adamawa_elevation": row.get("adamawa_elevation") or "",
        "adamawa_origin_text": row.get("adamawa_origin_text") or DEFAULT_ADAMAWA_ORIGIN_TEXT,
        "adamawa_topo_sheet_text": row.get("adamawa_topo_sheet_text") or DEFAULT_ADAMAWA_TOPO_SHEET_TEXT,
        "adamawa_computation_no": row.get("adamawa_computation_no") or "",
        "adamawa_cadastral_sheet_no": row.get("adamawa_cadastral_sheet_no") or "",
        "adamawa_plan_no": row.get("adamawa_plan_no") or "",
        "adamawa_surveyed_by_text": row.get("adamawa_surveyed_by_text") or "",
        "adamawa_disclaimer_text": row.get("adamawa_disclaimer_text") or DEFAULT_ADAMAWA_DISCLAIMER_TEXT,
        "cadastral_plan_no": row.get("cadastral_plan_no") or "",
        "cadastral_area_name": row.get("cadastral_area_name") or "",
        "cadastral_datum_text": row.get("cadastral_datum_text") or "",
        "cadastral_firm_block_text": row.get("cadastral_firm_block_text") or "",
        "fct_file_no": row.get("fct_file_no") or "",
        "fct_district": row.get("fct_district") or "",
        "fct_cadastral_zone": row.get("fct_cadastral_zone") or "",
        "fct_origin_beacon_text": row.get("fct_origin_beacon_text") or "",
        "fct_cadastral_map_ref": row.get("fct_cadastral_map_ref") or "",
        "fct_title_prefix": row.get("fct_title_prefix") or "",
        "parent_plot_id": row.get("parent_plot_id"),
        "subdivision_batch_id": row.get("subdivision_batch_id"),
        "subdivision_lot_no": row.get("subdivision_lot_no") or "",
        "estate_name": row.get("estate_name") or "",
        "technical_report_instruments": raw_instruments,
        "technical_report_dgps_type": row.get("technical_report_dgps_type") or "",
        "technical_report_num_surveyors": row.get("technical_report_num_surveyors"),
        "technical_report_num_technical_officers": row.get("technical_report_num_technical_officers"),
        "technical_report_num_labourers": row.get("technical_report_num_labourers"),
        "technical_report_recce_text": row.get("technical_report_recce_text") or "",
        "technical_report_demarcation_text": row.get("technical_report_demarcation_text") or "",
        "technical_report_computation_software_text": row.get("technical_report_computation_software_text") or DEFAULT_TECHNICAL_REPORT_COMPUTATION_SOFTWARE,
        "technical_report_plotting_software_text": row.get("technical_report_plotting_software_text") or DEFAULT_TECHNICAL_REPORT_PLOTTING_SOFTWARE,
        "technical_report_general_observation_text": row.get("technical_report_general_observation_text") or DEFAULT_TECHNICAL_REPORT_GENERAL_OBSERVATION,
        "elevation_points": raw_elevation_points,
        "survey_input_coordinates": raw_survey_input_coordinates,
    }


def safe_remove(path: str):
    try:
        os.remove(path)
    except FileNotFoundError:
        pass
    except Exception:
        pass


def _pdf_response_with_r2(
    local_pdf_path: str,
    filename: str,
    *,
    category: str,
    project_id: int | None = None,
):
    upload_meta = upload_export_file_best_effort(
        local_pdf_path,
        filename,
        category=category,
        project_id=project_id,
        content_type="application/pdf",
    )
    response = FileResponse(local_pdf_path, media_type="application/pdf", filename=filename)
    if upload_meta:
        object_key = upload_meta.get("object_key")
        public_url = upload_meta.get("public_url")
        if object_key:
            response.headers["X-LandCheck-R2-Key"] = str(object_key)
        if public_url:
            response.headers["X-LandCheck-R2-Url"] = str(public_url)
    return response


def _docx_response_with_r2(
    local_docx_path: str,
    filename: str,
    *,
    category: str,
    project_id: int | None = None,
):
    docx_media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    upload_meta = upload_export_file_best_effort(
        local_docx_path,
        filename,
        category=category,
        project_id=project_id,
        content_type=docx_media_type,
    )
    response = FileResponse(local_docx_path, media_type=docx_media_type, filename=filename)
    if upload_meta:
        object_key = upload_meta.get("object_key")
        public_url = upload_meta.get("public_url")
        if object_key:
            response.headers["X-LandCheck-R2-Key"] = str(object_key)
        if public_url:
            response.headers["X-LandCheck-R2-Url"] = str(public_url)
    return response


_LOWER_ROMAN_NUMERALS = [
    "i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x",
    "xi", "xii", "xiii", "xiv", "xv", "xvi", "xvii", "xviii", "xix", "xx",
]


def _to_lower_roman(n: int) -> str:
    if 1 <= n <= len(_LOWER_ROMAN_NUMERALS):
        return _LOWER_ROMAN_NUMERALS[n - 1]
    return str(n)


def _render_technical_report_docx(meta: dict, area_m2: float, output_path: str) -> None:
    """Builds the Adamawa OSG "Survey Technical Report" narrative document, matching the
    structure/style of the sample report the template is based on (title page, ASSIGNMENT page,
    INSTRUMENT/PERSONNEL/RECCE/DEMARCATION/METHOD page, COMPUTATIONS/PLOTTING/OBSERVATION page).
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:
        raise HTTPException(status_code=501, detail="Technical report export requires python-docx.") from exc

    owner_name = str(meta.get("adamawa_owner_name") or "").strip() or "THE APPLICANT"
    location_text = str(meta.get("location_text") or "").strip()
    lga_text = str(meta.get("lga_text") or "").strip()
    state_text = str(meta.get("state_text") or "ADAMAWA STATE").strip()
    surveyor_name = str(meta.get("surveyor_name") or "").strip()
    surveyor_rank = str(meta.get("surveyor_rank") or "").strip()
    topo_sheet_text = str(meta.get("adamawa_topo_sheet_text") or "").strip()
    authority_title = str(meta.get("adamawa_authority_title") or DEFAULT_ADAMAWA_AUTHORITY_TITLE).strip()
    control_point_name = str(meta.get("adamawa_control_point_name") or "").strip() or "the reference control point"

    instruments = meta.get("technical_report_instruments") or []
    dgps_type = str(meta.get("technical_report_dgps_type") or "").strip()
    num_surveyors = int(meta.get("technical_report_num_surveyors") or 0)
    num_technical_officers = int(meta.get("technical_report_num_technical_officers") or 0)
    num_labourers = int(meta.get("technical_report_num_labourers") or 0)
    recce_text = str(meta.get("technical_report_recce_text") or "").strip()
    demarcation_text = str(meta.get("technical_report_demarcation_text") or "").strip()
    computation_software_text = str(
        meta.get("technical_report_computation_software_text") or DEFAULT_TECHNICAL_REPORT_COMPUTATION_SOFTWARE
    ).strip()
    plotting_software_text = str(
        meta.get("technical_report_plotting_software_text") or DEFAULT_TECHNICAL_REPORT_PLOTTING_SOFTWARE
    ).strip()
    general_observation_text = str(
        meta.get("technical_report_general_observation_text") or DEFAULT_TECHNICAL_REPORT_GENERAL_OBSERVATION
    ).strip()

    surveyor_full = f"SURV. {surveyor_name}" if surveyor_name else "SURV."
    surveyor_with_rank = f"{surveyor_full} ({surveyor_rank})" if surveyor_rank else surveyor_full

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    def add_line(text_value: str, *, bold: bool = False, center: bool = True, size: int = 13, space_after: int = 6):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(space_after)
        run = paragraph.add_run(text_value)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        return paragraph

    def add_heading(text_value: str):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text_value)
        run.bold = True
        run.underline = True
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        return paragraph

    def add_body(text_value: str):
        paragraph = document.add_paragraph(text_value)
        paragraph.paragraph_format.space_after = Pt(10)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        return paragraph

    # Text after the tab must land on the same column for every item regardless of label width
    # ("i." vs "vii."/"viii.") - a bare "\t" alone snaps to Word's default tab grid, which puts
    # each label's very next tick at a different offset, producing a ragged column. A hanging
    # indent (first line pulled back to the label position) plus an explicit tab stop at the text
    # column fixes that: every label sits flush left, every item's text starts at the same x.
    LIST_LABEL_INDENT = Pt(24)
    LIST_TEXT_INDENT = Pt(54)

    def add_list_item(label: str, text_value: str):
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = LIST_TEXT_INDENT
        paragraph_format.first_line_indent = LIST_LABEL_INDENT - LIST_TEXT_INDENT
        paragraph_format.space_after = Pt(4)
        paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_INDENT)
        run = paragraph.add_run(f"{label}.\t{text_value}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return paragraph

    def add_bullet(text_value: str):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"• {text_value}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return paragraph

    # ---- Page 1: title page ----
    add_line(f"SURVEY TECHNICAL REPORT IN RESPECT OF {owner_name.upper()}", bold=True, size=14)
    add_line(f"LOCATED AT {location_text.upper()}", bold=True, size=13)
    add_line(f"{lga_text.upper()} LOCAL GOVERNMENT AREA, {state_text.upper()}", bold=True, size=13)
    if topo_sheet_text:
        add_line(topo_sheet_text.upper(), bold=True, size=13, space_after=24)
    add_line("SUBMITTED", bold=True, space_after=24)
    add_line("TO", bold=True, space_after=24)
    for line in authority_title.splitlines():
        if line.strip():
            add_line(line.strip().upper(), bold=True, space_after=4)
    add_line("BY", bold=True, space_after=24)
    add_line(surveyor_with_rank.upper(), bold=True, space_after=4)
    add_line("OFFICE OF THE SURVEYOR GENERAL", bold=True, space_after=4)
    add_line(state_text.upper(), bold=True)
    document.add_page_break()

    # ---- Page 2: assignment ----
    add_line(f"A SURVEY REPORT ON DEMARCATION AND SURVEY IN RESPECT OF {owner_name.upper()}", bold=True, size=13, space_after=4)
    add_line(f"LOCATED AT {location_text.upper()}, {lga_text.upper()} LOCAL GOVERNMENT AREA, {state_text.upper()}", bold=True, size=12, space_after=18)
    add_heading("ASSIGNMENT")
    add_body(
        f"The instruction is to carry out the demarcation and re-survey of item in respect of "
        f"{owner_name} Located At {location_text}, {lga_text} Local Government Area, {state_text}, "
        f"Instruction to Survey (I to S) is here by attached."
    )
    add_body("The above job is fully completed and the following are hereby attached for submissions;")
    for idx, attachment in enumerate([
        "Instruction to Survey",
        "Brief report on the job",
        "Work diagram",
        "Observation sheet",
        "Area/back computation sheet",
        "Orthophoto of the survey area",
        "Survey plan",
    ], start=1):
        add_list_item(_to_lower_roman(idx), attachment)
    add_body("Thanks")
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    closing_run = closing.add_run("Yours faithfully")
    closing_run.font.name = "Times New Roman"
    closing_run.font.size = Pt(12)
    for line in ["", surveyor_with_rank]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    document.add_page_break()

    # ---- Page 3+: instruments, personnel, recce, demarcation, method ----
    add_heading("INSTRUMENT / EQUIPMENT USED")
    equipment_lines = list(instruments)
    if dgps_type:
        equipment_lines = [f"{dgps_type} Differential Global Positioning System (DGPS), with its accessories was used to carry out the survey"] + equipment_lines
    if not equipment_lines:
        equipment_lines = ["Standard survey equipment appropriate to the job"]
    for idx, item in enumerate(equipment_lines, start=1):
        add_list_item(_to_lower_roman(idx), item)

    add_heading("SURVEY PARTY/PERSONNEL")
    add_body(
        f"{num_surveyors or 'One'} surveyor{'s' if num_surveyors > 1 else ''}, "
        f"{num_technical_officers or 'one'} technical officer{'s' if num_technical_officers > 1 else ''} "
        f"and {num_labourers or 'two'} labourer{'s' if num_labourers != 1 else ''}"
    )

    add_heading("RECONNAISSANCE (RECCE)")
    add_body(f"Reconnaissance survey was carried out by the Surveyor as shown in the work diagram.")
    if recce_text:
        add_body(
            f"{control_point_name} {recce_text} and was used as connection point for the survey as base station."
        )

    add_heading("DEMARCATION")
    if demarcation_text:
        add_body(
            f"The entire corner point to be surveyed was demarcated with property beacon from "
            f"{demarcation_text}, was capped and numbered appropriately as could be seen on the work diagram."
        )
    else:
        add_body(
            "The entire corner points to be surveyed were demarcated with property beacons, capped and "
            "numbered appropriately as could be seen on the work diagram."
        )

    add_heading("SURVEY METHOD")
    dgps_label = dgps_type or "Differential Global Positioning System (DGPS)"
    add_body(
        f"{dgps_label} was set up at a second order control point with the prefix number {control_point_name} "
        "and levelled as its base station, the master was allowed for some time to receive satellite signal, "
        "while the Rover was also set up on its pole at 2 meter height."
    )

    add_heading("MODE OF OBSERVATION")
    add_body("RTK mode was used as the operational method of observation for in capturing field data.")
    add_body(f"In setting up the {dgps_label} the following procedure was followed:")
    for step in [
        "Power on the master",
        "Navigate to App to create your project",
        "Set the project name",
        "Select configuration",
        "Connect the base with the Bluetooth name",
        "Go to base and select",
        "Then input the coordinate of the control point (coordinate of the tie) in the master. i.e NEH",
        "Input the height of the station",
        "Enter ok",
        "Select start",
        "Then disconnect the base",
        "Connect to rover",
        "Press start",
        "At the extreme top it will display fixed (that means it is ready for work but when it shows "
        "single or float then you allow it to completely receive the satellite until it displays fixed)",
        "Then select survey",
        "Press survey then it displays a count down from 5 seconds to zero on the screen",
        "Edit the point name that will display on the screen",
        "Then go to the next point and select survey again and wait for it to count down from 5 to zero",
        "Then go round all the beacons along the boundary of the land to be surveyed and pick",
        "After going round all your interested points to be surveyed",
        "Select export to export the data to either computer or phone",
        "Highlight all the display box on the screen",
        "Edit the file name",
        "Then press ok",
    ]:
        add_bullet(step)

    add_heading("SURVEY COMPUTATIONS")
    add_body(
        f"{computation_software_text} was used for back & area computation to serve as a check to the values "
        f"obtained by the use of the software. The computed area is A={area_m2:.3f} Sqm."
    )

    add_heading("PLOTTING")
    add_body(f"{plotting_software_text} was used for survey plan production.")

    add_heading("GENERAL OBSERVATION/ SUGGESTION")
    add_body(general_observation_text)

    add_line(surveyor_with_rank, center=False, bold=False, size=12, space_after=2)
    add_line("Sign………………………...", center=False, bold=False, size=12, space_after=2)
    add_line("Date ……………………….", center=False, bold=False, size=12, space_after=2)

    document.save(output_path)


def safe_rmtree(path: str):
    try:
        shutil.rmtree(path, ignore_errors=True)
    except Exception:
        pass


def resolve_existing_path(candidates: list[str]) -> str | None:
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def cleanup_preview_files(plot_id: int):
    patterns = [
        # .jpg is the current orthophoto/topo output format; .png patterns are kept so any
        # leftover files from before that switch still get cleaned up.
        os.path.join(REPORTS_DIR, "orthophoto", f"plot_{plot_id}_orthophoto_*_preview*.jpg"),
        os.path.join(REPORTS_DIR, "orthophoto", f"plot_{plot_id}_orthophoto_preview*.jpg"),
        os.path.join(REPORTS_DIR, "orthophoto", f"plot_{plot_id}_orthophoto_*_preview*.png"),
        os.path.join(REPORTS_DIR, "orthophoto", f"plot_{plot_id}_orthophoto_preview*.png"),
        os.path.join(REPORTS_DIR, "previews", f"plot_{plot_id}_preview*.png"),
        f"app/reports/orthophoto/plot_{plot_id}_orthophoto_*_preview*.jpg",
        f"app/reports/orthophoto/plot_{plot_id}_orthophoto_preview*.jpg",
        f"app/reports/orthophoto/plot_{plot_id}_orthophoto_*_preview*.png",
        f"app/reports/orthophoto/plot_{plot_id}_orthophoto_preview*.png",
        f"app/reports/previews/plot_{plot_id}_preview*.png",
    ]
    for pattern in patterns:
        for path in glob.glob(pattern):
            safe_remove(path)


def build_preview_revision_token(db: Session, plot_id: int) -> dict:
    feature_row = db.execute(text("""
        SELECT COALESCE(MAX(id), 0) AS max_id, COUNT(*) AS count
        FROM detected_features
        WHERE plot_id = :plot_id
    """), {"plot_id": plot_id}).mappings().first() if _table_exists(db, "public.detected_features") else {"max_id": 0, "count": 0}

    override_row = db.execute(text("""
        SELECT
            COALESCE(MAX(EXTRACT(EPOCH FROM updated_at)), 0)::BIGINT AS updated_epoch,
            COUNT(*) AS count
        FROM plot_feature_overrides
        WHERE plot_id = :plot_id
    """), {"plot_id": plot_id}).mappings().first() if _table_exists(db, "public.plot_feature_overrides") else {"updated_epoch": 0, "count": 0}

    geom_hex = db.execute(
        text("SELECT encode(ST_AsEWKB(geom), 'hex') FROM plots WHERE id = :plot_id"),
        {"plot_id": plot_id},
    ).scalar() if _table_exists(db, "public.plots") else ""

    return {
        "features_max_id": int((feature_row or {}).get("max_id") or 0),
        "features_count": int((feature_row or {}).get("count") or 0),
        "overrides_updated_epoch": int((override_row or {}).get("updated_epoch") or 0),
        "overrides_count": int((override_row or {}).get("count") or 0),
        "plot_geom_hash": hashlib.sha256((geom_hex or "").encode("utf-8")).hexdigest() if geom_hex else "",
    }


def build_preview_cache_key(plot_id: int, payload: dict, revision_token: dict) -> str:
    packed = json.dumps(
        {
            "plot_id": plot_id,
            "payload": payload,
            "revision": revision_token,
        },
        sort_keys=True,
        separators=(",", ":"),
        default=str,
    )
    return hashlib.sha256(packed.encode("utf-8")).hexdigest()


def build_plot_geom_revision_token(db: Session, plot_id: int) -> dict:
    geom_hex = db.execute(
        text("SELECT encode(ST_AsEWKB(geom), 'hex') FROM plots WHERE id = :plot_id"),
        {"plot_id": plot_id},
    ).scalar() if _table_exists(db, "public.plots") else ""
    return {
        "plot_geom_hash": hashlib.sha256((geom_hex or "").encode("utf-8")).hexdigest() if geom_hex else "",
    }


def preview_cache_path(plot_id: int, cache_key: str, variant: str = "preview", extension: str = "png") -> str:
    os.makedirs(PREVIEW_CACHE_DIR, exist_ok=True)
    safe_variant = re.sub(r"[^a-zA-Z0-9_-]", "_", str(variant or "preview"))
    safe_extension = re.sub(r"[^a-zA-Z0-9]", "", str(extension or "png")) or "png"
    return os.path.join(PREVIEW_CACHE_DIR, f"plot_{plot_id}_{safe_variant}_{cache_key}.{safe_extension}")


def get_cached_preview_path(plot_id: int, cache_key: str, variant: str = "preview", extension: str = "png") -> str | None:
    cache_path = preview_cache_path(plot_id, cache_key, variant=variant, extension=extension)
    if not os.path.exists(cache_path):
        return None
    try:
        age_s = max(0.0, time.time() - os.path.getmtime(cache_path))
        if age_s > PREVIEW_CACHE_TTL_SECONDS:
            safe_remove(cache_path)
            return None
        return cache_path
    except Exception:
        return None


def prune_preview_cache(plot_id: int, variant: str | None = None, extension: str = "*"):
    os.makedirs(PREVIEW_CACHE_DIR, exist_ok=True)
    safe_extension = re.sub(r"[^a-zA-Z0-9*]", "", str(extension or "*")) or "*"
    if variant:
        safe_variant = re.sub(r"[^a-zA-Z0-9_-]", "_", str(variant))
        pattern = os.path.join(PREVIEW_CACHE_DIR, f"plot_{plot_id}_{safe_variant}_*.{safe_extension}")
    else:
        pattern = os.path.join(PREVIEW_CACHE_DIR, f"plot_{plot_id}_*.{safe_extension}")
    files = []
    for path in glob.glob(pattern):
        try:
            files.append((path, os.path.getmtime(path)))
        except Exception:
            safe_remove(path)
    files.sort(key=lambda item: item[1], reverse=True)
    now = time.time()
    for idx, (path, modified_at) in enumerate(files):
        too_old = (now - modified_at) > (PREVIEW_CACHE_TTL_SECONDS * 6)
        overflow = idx >= PREVIEW_CACHE_MAX_FILES_PER_PLOT
        if too_old or overflow:
            safe_remove(path)


def _coerce_positive_int(value: Any) -> int | None:
    try:
        if value is None or value == "":
            return None
        parsed = int(value)
        if parsed <= 0:
            return None
        return parsed
    except Exception:
        return None


def _coerce_float(value: Any, default: float = 0.0) -> float:
    try:
        if value is None or value == "":
            return float(default)
        parsed = float(value)
        if math.isnan(parsed) or math.isinf(parsed):
            return float(default)
        return float(parsed)
    except Exception:
        return float(default)


def _station_name(index: int) -> str:
    name = ""
    num = int(index)
    while True:
        name = chr(65 + (num % 26)) + name
        num = (num // 26) - 1
        if num < 0:
            break
    return name


def _polygon_parts(geom: Any) -> list[Polygon]:
    if geom is None:
        return []
    if isinstance(geom, Polygon):
        return [geom]
    if geom.geom_type == "MultiPolygon":
        return [g for g in geom.geoms if isinstance(g, Polygon)]
    if geom.geom_type == "GeometryCollection":
        out: list[Polygon] = []
        for g in geom.geoms:
            if isinstance(g, Polygon):
                out.append(g)
            elif g.geom_type == "MultiPolygon":
                out.extend([p for p in g.geoms if isinstance(p, Polygon)])
        return out
    return []


def _largest_polygon(geom: Any) -> Polygon | None:
    parts = [p for p in _polygon_parts(geom) if p.area > 1e-9]
    if not parts:
        return None
    return max(parts, key=lambda p: p.area)


def _clean_single_polygon(geom: Any) -> Polygon | None:
    if geom is None:
        return None
    cleaned = geom.buffer(0)
    return _largest_polygon(cleaned)


def _metric_epsg_for_wgs84_polygon(poly_wgs84: Polygon) -> int:
    centroid = poly_wgs84.centroid
    zone = int((centroid.x + 180) / 6) + 1
    zone = max(1, min(zone, 60))
    return (32600 + zone) if centroid.y >= 0 else (32700 + zone)


def _epsg_display_name(epsg_code: int, fallback_key: str = "wgs84") -> str:
    try:
        epsg = int(epsg_code or 4326)
    except Exception:
        epsg = 4326
    if 32600 < epsg < 32700:
        return f"UTM Zone {epsg - 32600}N"
    if 32700 < epsg < 32800:
        return f"UTM Zone {epsg - 32700}S"
    return COORDINATE_SYSTEM_NAMES.get(str(fallback_key or "wgs84").strip().lower(), f"EPSG:{epsg}")


def _resolve_survey_render_crs(
    coordinate_system: str | None,
    plot_geom_wgs84: Polygon | None,
) -> tuple[str, int, str]:
    selected_key = str(coordinate_system or "wgs84").strip().lower() or "wgs84"
    resolved_key = selected_key
    epsg_code: int | None = None

    if plot_geom_wgs84 is not None and not plot_geom_wgs84.is_empty:
        centroid = plot_geom_wgs84.centroid
        if selected_key == "wgs84":
            metric_epsg = int(_metric_epsg_for_wgs84_polygon(plot_geom_wgs84))
            named_key = {
                32631: "utm_31n",
                32632: "utm_32n",
                32633: "utm_33n",
            }.get(metric_epsg)
            if named_key is not None:
                resolved_key = named_key
            else:
                # No named Nigeria zone matches - the plot is outside Nigeria. Use the
                # auto-picked global UTM EPSG directly (any longitude/hemisphere) instead of
                # silently falling back to WGS84 (EPSG:4326) itself below, which would make
                # every "metric" area/render calculation downstream nonsensical (square degrees,
                # not square metres) for any non-Nigeria plot left on plain "WGS84" rather than
                # an explicit projected system.
                epsg_code = metric_epsg
        else:
            resolved_key = resolve_coordinate_system_key(
                selected_key,
                float(centroid.x),
                float(centroid.y),
            )

    if epsg_code is None:
        epsg_code = COORDINATE_SYSTEMS.get(resolved_key, COORDINATE_SYSTEMS.get(selected_key, 4326))

    if selected_key == "wgs84_nigeria_meters":
        crs_name = f"WGS84 Nigeria Metres ({_epsg_display_name(epsg_code, resolved_key)})"
        render_key = resolved_key
    elif selected_key == "wgs84":
        crs_name = _epsg_display_name(epsg_code, resolved_key)
        render_key = selected_key
    else:
        crs_name = COORDINATE_SYSTEM_NAMES.get(resolved_key, _epsg_display_name(epsg_code, resolved_key))
        render_key = selected_key

    return render_key, int(epsg_code), crs_name


def _survey_template_map_frame(template_name: str | None) -> tuple[float, float]:
    template = str(template_name or DEFAULT_TEMPLATE_NAME).strip().lower()
    if template == "adamawa_osg":
        return 0.84, 0.555
    if template in {"akwa_ibom_osg", "cross_river_osg", "rivers_osg"}:
        return 0.84, 0.455
    if template == "fct_abuja_osg":
        return 0.66, 0.48
    return 0.80, 0.45


def _select_standard_survey_scale(fitted_ratio: int) -> int:
    """Choose the first recognised scale that still contains the fitted map extent."""
    required = max(1, int(fitted_ratio or 1))
    for scale in STANDARD_SURVEY_SCALE_DENOMINATORS:
        if scale >= required:
            return scale
    return STANDARD_SURVEY_SCALE_DENOMINATORS[-1]


def _split_polygon_once_by_area(poly_metric: Polygon, target_area: float) -> tuple[Polygon, Polygon]:
    poly = _clean_single_polygon(poly_metric)
    if poly is None:
        raise HTTPException(status_code=400, detail="Subdivision failed: invalid mother parcel geometry.")

    minx, miny, maxx, maxy = poly.bounds
    if maxx - minx <= 1e-8:
        raise HTTPException(status_code=400, detail="Subdivision failed: mother parcel width too small for automatic split.")

    pad_x = max(1.0, (maxx - minx) * 0.2)
    pad_y = max(1.0, (maxy - miny) * 0.2)
    lo = minx + 1e-9
    hi = maxx - 1e-9
    best_x = (lo + hi) / 2.0
    best_diff = float("inf")
    best_left: Polygon | None = None

    for _ in range(56):
        mid = (lo + hi) / 2.0
        left_mask = box(minx - pad_x, miny - pad_y, mid, maxy + pad_y)
        left_geom = _clean_single_polygon(poly.intersection(left_mask))
        left_area = float(left_geom.area if left_geom is not None else 0.0)
        diff = abs(left_area - target_area)
        if diff < best_diff and left_geom is not None:
            best_diff = diff
            best_x = mid
            best_left = left_geom
        if left_area < target_area:
            lo = mid
        else:
            hi = mid

    if best_left is None or best_left.area <= 1e-6:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to compute left split parcel.")

    right_mask = box(best_x, miny - pad_y, maxx + pad_x, maxy + pad_y)
    right_geom = _clean_single_polygon(poly.intersection(right_mask))
    if right_geom is None or right_geom.area <= 1e-6:
        # fallback using difference if masking loses a tiny seam
        right_geom = _clean_single_polygon(poly.difference(best_left))
    if right_geom is None or right_geom.area <= 1e-6:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to compute right split parcel.")

    return best_left, right_geom


def _split_polygon_by_distance_from(poly_metric: Polygon, reference_geom: Any, target_area: float) -> tuple[Polygon, Polygon]:
    """Splits poly into (near, far) relative to reference_geom - `near` is the portion within some
    distance of reference_geom, sized (via binary search on that distance) to have area close to
    target_area. Unlike _split_polygon_once_by_area's fixed left-to-right sweep, this always pulls
    from whichever edge of the polygon actually touches reference_geom, wherever that is.
    """
    poly = _clean_single_polygon(poly_metric)
    if poly is None:
        raise HTTPException(status_code=400, detail="Subdivision failed: invalid geometry for road-adjacent split.")

    minx, miny, maxx, maxy = poly.bounds
    lo, hi = 0.0, math.hypot(maxx - minx, maxy - miny) + 1.0
    best_near: Polygon | None = None
    best_diff = float("inf")

    for _ in range(50):
        mid = (lo + hi) / 2.0
        try:
            near = _clean_single_polygon(poly.intersection(reference_geom.buffer(mid)))
        except Exception:
            near = None
        near_area = float(near.area) if near is not None else 0.0
        diff = abs(near_area - target_area)
        if diff < best_diff and near is not None:
            best_diff = diff
            best_near = near
        if near_area < target_area:
            lo = mid
        else:
            hi = mid

    if best_near is None or best_near.area <= 1e-6:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to compute road-adjacent split.")

    far = _clean_single_polygon(poly.difference(best_near))
    if far is None or far.area <= 1e-6:
        far = _clean_single_polygon(poly.difference(best_near.buffer(1e-6)))
    if far is None or far.area <= 1e-6:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to compute the remaining parcel after the road-adjacent split.")

    return best_near, far


def _subdivide_polygon_equal_count(poly_metric: Polygon, split_count: int, orientation_deg: float) -> list[Polygon]:
    if split_count < 2:
        raise HTTPException(status_code=400, detail="Subdivision requires at least 2 derived plots.")
    base_poly = _clean_single_polygon(poly_metric)
    if base_poly is None:
        raise HTTPException(status_code=400, detail="Mother parcel geometry is invalid.")

    rotated = _clean_single_polygon(rotate(base_poly, -orientation_deg, origin="centroid", use_radians=False))
    if rotated is None:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to rotate mother parcel.")

    pieces: list[Polygon] = []
    remaining = rotated
    for idx in range(split_count - 1):
        pending_slots = split_count - idx
        target_area = float(remaining.area / pending_slots)
        left_piece, right_piece = _split_polygon_once_by_area(remaining, target_area)
        pieces.append(left_piece)
        remaining = right_piece
    pieces.append(remaining)

    out: list[Polygon] = []
    for piece in pieces:
        restored = _clean_single_polygon(rotate(piece, orientation_deg, origin=base_poly.centroid, use_radians=False))
        if restored is None or restored.area <= 1e-8:
            raise HTTPException(status_code=400, detail="Subdivision failed: generated a degenerate parcel.")
        out.append(restored)
    return out


def _normalize_fraction_weights(values: Any) -> list[float]:
    if not isinstance(values, (list, tuple)):
        return []
    out: list[float] = []
    for raw in values:
        parsed = _coerce_float(raw, 0.0)
        if parsed > 0:
            out.append(float(parsed))
    return out


def _normalize_fraction_breaks(values: Any) -> list[float]:
    if not isinstance(values, (list, tuple)):
        return []
    raw_breaks: list[float] = []
    for raw in values:
        parsed = _coerce_float(raw, float("nan"))
        if not math.isfinite(parsed):
            continue
        # Allow percentages (0-100) and normalized ratios (0-1).
        if parsed > 1.0 and parsed <= 100.0:
            parsed = parsed / 100.0
        if 0.0 < parsed < 1.0:
            raw_breaks.append(float(parsed))
    if not raw_breaks:
        return []
    raw_breaks.sort()
    deduped: list[float] = []
    for value in raw_breaks:
        if not deduped or abs(value - deduped[-1]) > 1e-6:
            deduped.append(value)
    return deduped


def _breaks_to_weights(breaks: list[float]) -> list[float]:
    if not breaks:
        return []
    weights: list[float] = []
    prev = 0.0
    for value in breaks:
        weights.append(max(value - prev, 0.0))
        prev = value
    weights.append(max(1.0 - prev, 0.0))
    return [w for w in weights if w > 0]


def _weights_to_breaks(weights: list[float]) -> list[float]:
    if not weights or len(weights) < 2:
        return []
    total = sum(weights)
    if total <= 0:
        return []
    acc = 0.0
    out: list[float] = []
    for weight in weights[:-1]:
        acc += float(weight) / total
        out.append(max(0.0, min(1.0, acc)))
    return out


def _subdivide_polygon_weighted(poly_metric: Polygon, weights: list[float], orientation_deg: float) -> list[Polygon]:
    normalized_weights = [float(w) for w in weights if float(w) > 0]
    if len(normalized_weights) < 2:
        raise HTTPException(status_code=400, detail="Fraction subdivision requires at least 2 positive fraction values.")

    base_poly = _clean_single_polygon(poly_metric)
    if base_poly is None:
        raise HTTPException(status_code=400, detail="Mother parcel geometry is invalid.")

    rotated = _clean_single_polygon(rotate(base_poly, -orientation_deg, origin="centroid", use_radians=False))
    if rotated is None:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to rotate mother parcel.")

    pieces: list[Polygon] = []
    remaining = rotated
    remaining_weight = float(sum(normalized_weights))
    if remaining_weight <= 0:
        raise HTTPException(status_code=400, detail="Fraction subdivision weights are invalid.")

    for weight in normalized_weights[:-1]:
        weight = float(weight)
        if remaining_weight <= 1e-9:
            break
        share = max(0.0, min(1.0, weight / remaining_weight))
        target_area = max(1e-9, float(remaining.area) * share)
        left_piece, right_piece = _split_polygon_once_by_area(remaining, target_area)
        pieces.append(left_piece)
        remaining = right_piece
        remaining_weight -= weight

    pieces.append(remaining)

    out: list[Polygon] = []
    for piece in pieces:
        restored = _clean_single_polygon(rotate(piece, orientation_deg, origin=base_poly.centroid, use_radians=False))
        if restored is None or restored.area <= 1e-8:
            raise HTTPException(status_code=400, detail="Subdivision failed: generated a degenerate parcel.")
        out.append(restored)
    return out


def _load_plot_road_segments_wgs84(db: Session, plot_id: int) -> list[dict]:
    """Road centerlines for this plot in WGS84 - OSM-detected roads clipped to the plot buffer,
    plus/minus whatever the Feature Editor's road overrides (add/delete/update) say. Same road
    definition get_plot_features_geojson uses for the CAD editor, kept as an independent, minimal
    copy here rather than a shared extraction - that endpoint has since grown scale-aware label
    decoration (position hints, segment keys) this doesn't need.

    Each entry carries enough identity for a caller to act on "just this one road" - a detected/
    OSM segment has no row of its own to delete, so it's identified by a stable hash of its own
    geometry (the same way one gets suppressed today: a matching "delete" override), while an
    "add" override carries its real plot_feature_overrides.id, deletable directly.
    Returns [{"id": str, "source": "detected"|"override", "override_id": int|None, "geom": LineString}, ...]
    """
    road_rows = db.execute(text("""
        WITH roads AS (
            SELECT
                CASE
                    WHEN ST_SRID(r.geom) = 4326 THEN r.geom
                    WHEN ST_SRID(r.geom) = 0 THEN ST_SetSRID(r.geom, 4326)
                    ELSE ST_Transform(r.geom, 4326)
                END AS geom
            FROM lines r
            WHERE r.highway IS NOT NULL
        ),
        clipped AS (
            SELECT (ST_Dump(ST_Intersection(roads.geom, b.geom))).geom AS geom
            FROM roads
            JOIN plot_buffers b ON b.plot_id = :plot_id
            WHERE ST_Intersects(roads.geom, b.geom)
        )
        SELECT ST_AsGeoJSON(geom) AS geojson
        FROM clipped
        WHERE ST_GeometryType(geom) = 'ST_LineString'
    """), {"plot_id": plot_id}).fetchall()

    def _detected_segment(geom) -> dict:
        try:
            seg_id = "detected:" + hashlib.sha1(bytes(geom.wkb)).hexdigest()[:16]
        except Exception:
            seg_id = "detected:" + hashlib.sha1(str(geom).encode("utf-8")).hexdigest()[:16]
        return {"id": seg_id, "source": "detected", "override_id": None, "geom": geom}

    segments: list[dict] = []
    for r in road_rows:
        if not r.geojson:
            continue
        try:
            segments.append(_detected_segment(shape(json.loads(r.geojson))))
        except Exception:
            continue

    if not segments:
        # Fallback for plots without a plot_buffers row yet (mirrors get_plot_features_geojson).
        feature_rows = db.execute(text("""
            SELECT ST_AsGeoJSON(geom) AS geojson
            FROM detected_features
            WHERE plot_id = :plot_id AND feature_type = 'road'
        """), {"plot_id": plot_id}).fetchall()
        for r in feature_rows:
            if not r.geojson:
                continue
            try:
                segments.append(_detected_segment(shape(json.loads(r.geojson))))
            except Exception:
                continue

    override_rows = db.execute(text("""
        SELECT id, action, ST_AsGeoJSON(geom) AS geojson
        FROM plot_feature_overrides
        WHERE plot_id = :plot_id AND feature_type = 'road'
    """), {"plot_id": plot_id}).fetchall()

    def _line_replaced_by(geom, override_geom, tol_deg: float = 0.00001) -> bool:
        try:
            total_len = max(getattr(geom, "length", 0.0), 1e-9)
            uncovered = geom.difference(override_geom.buffer(tol_deg))
            uncovered_len = getattr(uncovered, "length", 0.0)
            return uncovered_len < total_len * 0.1
        except Exception:
            return geom.intersects(override_geom)

    for r in override_rows:
        if not r.geojson:
            continue
        try:
            override_geom = shape(json.loads(r.geojson))
        except Exception:
            continue
        action = str(r.action or "").strip().lower()
        if action in ("delete", "update"):
            segments = [s for s in segments if not _line_replaced_by(s["geom"], override_geom)]
        if action in ("add", "update"):
            segments.append({
                "id": f"override:{r.id}",
                "source": "override",
                "override_id": int(r.id),
                "geom": override_geom,
            })

    return segments


def _build_road_exclusion_geom(
    db: Session, plot_id: int, parent_metric: Polygon, metric_epsg: int, road_width_m: float,
) -> Any:
    """The road-reserve corridor to carve out of the mother parcel before subdividing: every road
    line on this plot (see _load_plot_road_segments_wgs84), buffered by half the requested width
    on each side, unioned, and clipped to the parcel itself. None if there's no road on this plot.
    """
    lines_wgs84 = [seg["geom"] for seg in _load_plot_road_segments_wgs84(db, plot_id)]
    if not lines_wgs84:
        return None
    try:
        gdf_lines = gpd.GeoDataFrame(geometry=lines_wgs84, crs="EPSG:4326").to_crs(epsg=metric_epsg)
    except Exception:
        return None
    half_width = max(0.1, float(road_width_m) / 2.0)
    buffered = [
        geom.buffer(half_width) for geom in gdf_lines.geometry.tolist() if geom is not None and not geom.is_empty
    ]
    if not buffered:
        return None
    try:
        corridor = unary_union(buffered).intersection(parent_metric)
    except Exception:
        return None
    if corridor is None or corridor.is_empty or corridor.area <= 1e-6:
        return None
    return corridor


def _subdivide_polygon_by_dimension(
    poly_metric: Polygon, lot_width_m: float, lot_height_m: float, orientation_deg: float,
) -> tuple[list[Polygon], Any]:
    """Tiles poly_metric with lot_width_m x lot_height_m rectangles (grid-aligned to
    orientation_deg), keeping only cells that land essentially whole inside the parcel. Returns
    (full_lots, leftover) where leftover = poly - union(full_lots) is computed as an exact
    difference (not accumulated cell-by-cell), so lots + leftover always exactly reconstruct the
    input polygon with no gap or double-count. leftover may be a MultiPolygon (several disjoint
    unusable slivers), so it's deliberately never passed through _clean_single_polygon, which
    would silently keep only the single largest piece.
    """
    if lot_width_m <= 0 or lot_height_m <= 0:
        raise HTTPException(status_code=400, detail="Lot width and height must be positive.")

    base_poly = _clean_single_polygon(poly_metric)
    if base_poly is None:
        raise HTTPException(status_code=400, detail="Mother parcel geometry is invalid.")

    rotated = _clean_single_polygon(rotate(base_poly, -orientation_deg, origin="centroid", use_radians=False))
    if rotated is None:
        raise HTTPException(status_code=400, detail="Subdivision failed: unable to rotate mother parcel.")

    minx, miny, maxx, maxy = rotated.bounds
    cell_area = lot_width_m * lot_height_m
    bounds_area = max(0.0, (maxx - minx) * (maxy - miny))
    if cell_area > 0 and bounds_area / cell_area > 5000:
        raise HTTPException(
            status_code=400,
            detail="This lot size would generate far too many plots for this parcel. Increase lot size or pick a smaller area.",
        )

    full_lots: list[Polygon] = []
    full_area_threshold = cell_area * 0.985
    y = miny
    while y < maxy:
        x = minx
        while x < maxx:
            cell = box(x, y, x + lot_width_m, y + lot_height_m)
            try:
                clipped = _clean_single_polygon(rotated.intersection(cell))
            except Exception:
                clipped = None
            if clipped is not None and clipped.area >= full_area_threshold:
                full_lots.append(clipped)
                if len(full_lots) > 500:
                    raise HTTPException(
                        status_code=400,
                        detail="This lot size would generate over 500 plots. Increase lot size or split the parcel first.",
                    )
            x += lot_width_m
        y += lot_height_m

    if not full_lots:
        raise HTTPException(
            status_code=400,
            detail="No full-size lot fits inside this parcel at the given dimensions. Reduce the lot size.",
        )

    # Row-major reading order (bottom-to-top rows, left-to-right within a row) in the rotated
    # frame, so lot numbering reads naturally instead of following grid-scan/insertion order.
    full_lots.sort(key=lambda p: (round(p.centroid.y / max(lot_height_m, 1e-6)), p.centroid.x))

    leftover_rotated = None
    try:
        diff = rotated.difference(unary_union(full_lots))
        if not diff.is_valid:
            diff = diff.buffer(0)
        if diff is not None and not diff.is_empty and diff.area > 1e-6:
            leftover_rotated = diff
    except Exception:
        leftover_rotated = None

    out_lots: list[Polygon] = []
    for piece in full_lots:
        restored = _clean_single_polygon(rotate(piece, orientation_deg, origin=base_poly.centroid, use_radians=False))
        if restored is None or restored.area <= 1e-8:
            continue
        out_lots.append(restored)
    if not out_lots:
        raise HTTPException(status_code=400, detail="Subdivision failed: generated only degenerate lots.")

    out_leftover = None
    if leftover_rotated is not None:
        restored_leftover = rotate(leftover_rotated, orientation_deg, origin=base_poly.centroid, use_radians=False)
        out_leftover = restored_leftover if not restored_leftover.is_empty else None

    return out_lots, out_leftover


def _distribute_count_across_parts(parts: list[Polygon], total_count: int) -> list[int]:
    """How many *whole* target-sized lots fit in each part (floor, never rounded up) - a part
    smaller than one target lot gets 0 here on purpose; its entire area instead becomes part of
    the merged "combined" lot built by _subdivide_polygon_multi_part_equal_count, rather than
    either being force-fit into one oversized/undersized lot of its own or left unused.
    """
    total_area = sum(max(p.area, 1e-9) for p in parts)
    target = total_area / max(total_count, 1)
    return [int(math.floor(p.area / target)) if target > 0 else 0 for p in parts]


def _subdivide_polygon_multi_part_equal_count(
    parts: list[Polygon], total_count: int, orientation_deg: float, road_reference_geom: Any = None
) -> list[Polygon]:
    """Splits total_count *equal-sized* lots out of the combined area on both sides of an
    excluded road, exactly as if the road weren't there. Each part is cut into as many
    target-sized whole lots as fully fit; whatever's left over in each part (a part too small for
    even one whole lot ends up entirely "left over") is pooled together across every part and
    merged into one combined lot - not shown as unused "leftover" land, and not force-fit onto
    whichever side happens to have it, but folded into the split so every one of the total_count
    lots comes out the same size. That combined lot may end up as a MultiPolygon (spanning both
    sides of the road) - callers need to handle that when serializing lot geometry.

    When road_reference_geom (the excluded road corridor) is given, each part's remainder is cut
    from whichever portion of it actually sits nearest the road, not an arbitrary corner - so the
    piece that goes on to merge with the small side's own remainder sits right next to it on the
    map, reading as one lot straddling the road, instead of two unrelated-looking fragments.
    """
    total_area = sum(max(p.area, 1e-9) for p in parts)
    target = total_area / max(total_count, 1)
    whole_counts = _distribute_count_across_parts(parts, total_count)

    out: list[Any] = []
    remainder_pieces: list[Any] = []
    whole_lots_used = 0
    for part, whole_count in zip(parts, whole_counts):
        if whole_count <= 0:
            remainder_pieces.append(part)
            continue
        whole_lots_used += whole_count
        remainder_target_area = part.area - whole_count * target
        clean_part = part
        if remainder_target_area > target * 0.01:
            if road_reference_geom is not None:
                try:
                    remainder_piece, clean_part = _split_polygon_by_distance_from(
                        part, road_reference_geom, remainder_target_area
                    )
                    remainder_pieces.append(remainder_piece)
                except HTTPException:
                    clean_part = part
            else:
                base_poly = _clean_single_polygon(part)
                rotated = _clean_single_polygon(rotate(base_poly, -orientation_deg, origin="centroid", use_radians=False)) if base_poly is not None else None
                if rotated is not None:
                    try:
                        remainder_rot, clean_rot = _split_polygon_once_by_area(rotated, remainder_target_area)
                        remainder_pieces.append(
                            rotate(remainder_rot, orientation_deg, origin=base_poly.centroid, use_radians=False)
                        )
                        clean_part = rotate(clean_rot, orientation_deg, origin=base_poly.centroid, use_radians=False)
                    except HTTPException:
                        pass
        if whole_count == 1:
            cleaned = _clean_single_polygon(clean_part)
            if cleaned is not None and cleaned.area > 1e-8:
                out.append(cleaned)
        else:
            out.extend(_subdivide_polygon_equal_count(clean_part, whole_count, orientation_deg))

    deficit = total_count - whole_lots_used
    if remainder_pieces:
        if deficit == 1:
            combined = unary_union(remainder_pieces)
            if combined is not None and not combined.is_empty and combined.area > 1e-8:
                out.append(combined)
        else:
            # More than one extra lot's worth of pooled remainder (rare - needs 3+ sides with
            # modest fractional shares each) - the area-bisection helper this function otherwise
            # relies on assumes one contiguous polygon, so a scattered multi-part pool can't be
            # safely bisected further here. Hand each remainder back as its own lot instead of
            # dropping it, rather than risk mis-splitting a MultiPolygon.
            for piece in remainder_pieces:
                cleaned = piece if piece.geom_type in ("Polygon", "MultiPolygon") else _clean_single_polygon(piece)
                if cleaned is not None and not cleaned.is_empty and cleaned.area > 1e-8:
                    out.append(cleaned)
    return out


def _compute_subdivision_payload(
    parent_plot_id: int,
    parent_geom_wgs84: Polygon,
    db: Session,
    *,
    method: str,
    split_count: int | None,
    target_area_m2: float | None,
    orientation_deg: float,
    lot_prefix: str,
    fraction_weights: list[float] | None = None,
    fraction_breaks: list[float] | None = None,
    custom_areas_m2: list[float] | None = None,
    lot_names: list[str] | None = None,
    exclude_road: bool = False,
    road_width_m: float = 10.0,
    lot_width: float | None = None,
    lot_height: float | None = None,
    dimension_unit: str = "m",
) -> dict:
    method_key = (method or "by_count").strip().lower()
    valid_methods = {"by_count", "by_area", "by_fraction", "by_custom_area", "by_dimension"}
    if method_key not in valid_methods:
        raise HTTPException(
            status_code=400,
            detail="Invalid subdivision method. Use 'by_count', 'by_area', 'by_fraction', 'by_custom_area', or 'by_dimension'.",
        )

    metric_epsg = _metric_epsg_for_wgs84_polygon(parent_geom_wgs84)
    gdf_metric = gpd.GeoDataFrame(geometry=[parent_geom_wgs84], crs="EPSG:4326").to_crs(epsg=metric_epsg)
    parent_metric = _clean_single_polygon(gdf_metric.geometry.iloc[0])
    if parent_metric is None or parent_metric.area <= 1e-6:
        raise HTTPException(status_code=400, detail="Mother parcel area is too small for subdivision.")

    total_area_m2 = float(parent_metric.area)

    # Optionally carve the access road out of the parcel before running whichever split method
    # was requested, so a road never ends up sliced across multiple lots. The excluded corridor
    # is reported separately (excluded_geojson/excluded_area_m2), never as a numbered lot. An
    # internal estate road commonly splits the parcel into two or more disconnected pieces (not
    # just a frontage strip along one edge) - each piece is then subdivided independently rather
    # than rejected outright, for the methods where "which piece gets which lot" doesn't need an
    # explicit answer (by_count/by_area/by_dimension). by_fraction/by_custom_area assign a
    # specific share to a specific position, which doesn't have an unambiguous meaning once the
    # parcel is in separate pieces, so those two still require a single contiguous area.
    working_parts: list[Polygon] = [parent_metric]
    excluded_geom_metric = None
    road_segments_out: list[dict] = []
    unit_key = str(dimension_unit or "m").strip().lower()
    ft_to_m = 0.3048
    if exclude_road:
        # Reported back so the subdivision preview can show/act on individual roads (delete just
        # one) rather than only the final merged corridor - see _load_plot_road_segments_wgs84.
        try:
            for seg in _load_plot_road_segments_wgs84(db, parent_plot_id):
                road_segments_out.append({
                    "id": seg["id"],
                    "source": seg["source"],
                    "override_id": seg["override_id"],
                    "geojson": {"type": "Feature", "properties": {}, "geometry": mapping(seg["geom"])},
                })
        except Exception:
            road_segments_out = []
        try:
            corridor = _build_road_exclusion_geom(db, parent_plot_id, parent_metric, metric_epsg, road_width_m)
        except Exception:
            corridor = None
        if corridor is not None and not corridor.is_valid:
            try:
                corridor = corridor.buffer(0)
            except Exception:
                corridor = None
        if corridor is not None:
            try:
                remainder = parent_metric.difference(corridor)
            except GEOSException:
                # Real OSM road geometry can be messy enough that a direct difference hits a GEOS
                # topology exception - nudging both operands through buffer(0) first fixes almost
                # all of these without changing the result in any way that matters here.
                remainder = _clean_single_polygon(parent_metric.buffer(0)).difference(corridor.buffer(0))
            significant_parts = [p for p in _polygon_parts(remainder) if p.area > total_area_m2 * 0.01]
            if len(significant_parts) > 1 and method_key in ("by_fraction", "by_custom_area"):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Excluding this road splits the parcel into {len(significant_parts)} separate areas, "
                        "which fraction/custom-area splitting can't divide across. Use 'Split by number of "
                        "plots', 'Split by target plot area', or the fixed-size method instead, or subdivide "
                        "each side of the road separately."
                    ),
                )
            if significant_parts:
                # Left-to-right by centroid so lot numbering reads in a sensible order across sides.
                significant_parts.sort(key=lambda p: p.centroid.x)
                working_parts = significant_parts
                excluded_geom_metric = corridor

    working_metric = working_parts[0]
    working_area_m2 = float(sum(p.area for p in working_parts))
    resolved_count = _coerce_positive_int(split_count)
    target_area = _coerce_float(target_area_m2, 0.0)
    effective_fraction_weights: list[float] = []
    effective_fraction_breaks: list[float] = []
    effective_custom_areas_m2: list[float] = []
    leftover_metric = None
    lot_width_m = None
    lot_height_m = None
    lot_count_balanced = False
    if method_key == "by_count":
        if resolved_count is None or resolved_count < 2:
            raise HTTPException(status_code=400, detail="For 'by_count', set derived plot count to 2 or more.")
        if len(working_parts) > 1:
            pieces_wgs84 = _subdivide_polygon_multi_part_equal_count(
                working_parts, int(resolved_count), orientation_deg, excluded_geom_metric
            )
            lot_count_balanced = len(pieces_wgs84) != int(resolved_count)
        else:
            pieces_wgs84 = _subdivide_polygon_equal_count(working_metric, int(resolved_count), orientation_deg)
    elif method_key == "by_dimension":
        raw_width = _coerce_float(lot_width, 0.0)
        raw_height = _coerce_float(lot_height, 0.0)
        if raw_width <= 0 or raw_height <= 0:
            raise HTTPException(status_code=400, detail="For 'by_dimension', provide a positive lot width and height.")
        unit_mult = ft_to_m if unit_key == "ft" else 1.0
        lot_width_m = raw_width * unit_mult
        lot_height_m = raw_height * unit_mult
        pieces_wgs84 = []
        leftover_parts = []
        for part in working_parts:
            try:
                part_lots, part_leftover = _subdivide_polygon_by_dimension(
                    part, lot_width_m, lot_height_m, orientation_deg
                )
                pieces_wgs84.extend(part_lots)
                if part_leftover is not None:
                    leftover_parts.append(part_leftover)
            except HTTPException:
                # This piece is too small for even one full-size lot - treat the whole thing as
                # leftover rather than failing the entire subdivision over one small side.
                leftover_parts.append(part)
        if not pieces_wgs84:
            raise HTTPException(
                status_code=400,
                detail="No full-size lot fits inside this parcel at the given dimensions. Reduce the lot size.",
            )
        leftover_metric = unary_union(leftover_parts) if leftover_parts else None
    else:
        if method_key == "by_area":
            if target_area <= 0:
                raise HTTPException(status_code=400, detail="For 'by_area', provide a positive target plot size (sqm).")
            if len(working_parts) > 1:
                # The entire combined area (both sides of the excluded road) is divided up with
                # nothing left over - every part gets at least one lot, even a part smaller than
                # target_area, same "use it all" principle as by_count above.
                pieces_wgs84 = []
                for part in working_parts:
                    part_count = max(1, min(500, int(round(part.area / target_area))))
                    if part_count == 1:
                        cleaned = _clean_single_polygon(part)
                        if cleaned is not None and cleaned.area > 1e-8:
                            pieces_wgs84.append(cleaned)
                        continue
                    pieces_wgs84.extend(_subdivide_polygon_equal_count(part, part_count, orientation_deg))
                resolved_count = min(max(2, len(pieces_wgs84)), 500)
            else:
                approx_count = int(round(working_area_m2 / target_area))
                resolved_count = max(2, approx_count)
                resolved_count = min(int(resolved_count or 2), 500)
                pieces_wgs84 = _subdivide_polygon_equal_count(working_metric, resolved_count, orientation_deg)
        elif method_key == "by_fraction":
            normalized_breaks = _normalize_fraction_breaks(fraction_breaks)
            normalized_weights = _normalize_fraction_weights(fraction_weights)
            if normalized_breaks:
                effective_fraction_breaks = normalized_breaks
                effective_fraction_weights = _breaks_to_weights(normalized_breaks)
            elif len(normalized_weights) >= 2:
                effective_fraction_weights = normalized_weights
                effective_fraction_breaks = _weights_to_breaks(normalized_weights)
            else:
                raise HTTPException(
                    status_code=400,
                    detail="For 'by_fraction', provide fraction values (for example: [2,3,5]) or break positions (for example: [0.3,0.65]).",
                )
            if len(effective_fraction_weights) < 2:
                raise HTTPException(status_code=400, detail="Fraction subdivision requires at least two resulting plots.")
            resolved_count = min(len(effective_fraction_weights), 500)
            effective_fraction_weights = effective_fraction_weights[:resolved_count]
            effective_fraction_breaks = _weights_to_breaks(effective_fraction_weights)
            pieces_wgs84 = _subdivide_polygon_weighted(working_metric, effective_fraction_weights, orientation_deg)
        else:
            normalized_custom_areas = _normalize_fraction_weights(custom_areas_m2)
            if resolved_count is not None and resolved_count >= 2 and len(normalized_custom_areas) != resolved_count:
                raise HTTPException(
                    status_code=400,
                    detail=f"For 'by_custom_area', provide exactly {resolved_count} custom lot areas.",
                )
            if len(normalized_custom_areas) < 2:
                raise HTTPException(
                    status_code=400,
                    detail="For 'by_custom_area', provide at least two positive custom lot areas (sqm).",
                )

            allocated_sum = float(sum(normalized_custom_areas))
            tolerance_m2 = 0.01
            if allocated_sum > working_area_m2 + tolerance_m2:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Custom areas exceed the available parcel area by "
                        f"{allocated_sum - working_area_m2:.2f} sqm. Reduce allocations."
                    ),
                )
            if allocated_sum < working_area_m2 - tolerance_m2:
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Custom areas do not fully allocate the available parcel area. Remaining "
                        f"{working_area_m2 - allocated_sum:.2f} sqm. Adjust allocations to match total area."
                    ),
                )

            resolved_count = min(len(normalized_custom_areas), 500)
            effective_custom_areas_m2 = normalized_custom_areas[:resolved_count]
            pieces_wgs84 = _subdivide_polygon_weighted(working_metric, effective_custom_areas_m2, orientation_deg)

    resolved_count = min(int(resolved_count or 2), 500)
    gdf_out = gpd.GeoDataFrame(geometry=pieces_wgs84, crs=f"EPSG:{metric_epsg}").to_crs(epsg=4326)

    safe_prefix = re.sub(r"[^A-Za-z0-9]+", "", str(lot_prefix or "LOT").upper()) or "LOT"
    custom_lot_names = list(lot_names or [])
    plots: list[dict] = []
    derived_total = 0.0
    output_idx = 0
    for source_idx, (poly_wgs, poly_metric) in enumerate(zip(gdf_out.geometry.tolist(), pieces_wgs84)):
        custom_name = str(custom_lot_names[source_idx] or "").strip() if source_idx < len(custom_lot_names) else ""
        base_plot_no = custom_name or f"{safe_prefix}-{source_idx + 1:03d}"

        # A lot combining both sides of an excluded road (see
        # _subdivide_polygon_multi_part_equal_count) is a MultiPolygon here - every *saved* plot
        # is a single contiguous parcel at the database level (the CAD editor, survey plan
        # renderer, etc. all assume that too), so each part becomes its own ordinary lot entry
        # sharing one base lot number (LOT-004A / LOT-004B via combined_group) instead of one
        # entry carrying a geometry most of the app can't handle.
        if poly_wgs.geom_type == "MultiPolygon":
            wgs_parts = list(poly_wgs.geoms)
            metric_parts = list(poly_metric.geoms) if poly_metric.geom_type == "MultiPolygon" else [poly_metric]
        else:
            wgs_parts = [poly_wgs]
            metric_parts = [poly_metric]

        is_combined = len(wgs_parts) > 1
        for part_idx, (part_wgs, part_metric) in enumerate(zip(wgs_parts, metric_parts)):
            output_idx += 1
            area_m2 = float(max(part_metric.area, 0.0))
            derived_total += area_m2
            ring = [[float(x), float(y)] for x, y in list(part_wgs.exterior.coords)]
            plot_no = f"{base_plot_no}{chr(65 + part_idx)}" if is_combined else base_plot_no
            plots.append(
                {
                    "index": output_idx,
                    "lot_no": plot_no,
                    "area_m2": round(area_m2, 2),
                    "area_hectares": round(area_m2 / 10000.0, 4),
                    "geometry": mapping(part_wgs),
                    "station_names": [_station_name(i) for i in range(max(0, len(ring) - 1))],
                    "combined_group": base_plot_no if is_combined else None,
                }
            )

    feature_collection = {
        "type": "FeatureCollection",
        "features": [
            {
                "type": "Feature",
                "properties": {
                    "lot_no": p["lot_no"],
                    "index": p["index"],
                    "area_m2": p["area_m2"],
                    "area_hectares": p["area_hectares"],
                    "combined_group": p.get("combined_group"),
                },
                "geometry": p["geometry"],
            }
            for p in plots
        ],
    }

    def _to_wgs84_geojson(geom_metric):
        # geom_metric may be a Polygon or MultiPolygon (a road split into two reserve strips,
        # several disjoint dimension-tiling leftover slivers) - mapping() handles either directly
        # rather than forcing it through the single-polygon-only helpers used for actual lots.
        if geom_metric is None or geom_metric.is_empty:
            return None, None
        try:
            geom_wgs84 = gpd.GeoSeries([geom_metric], crs=f"EPSG:{metric_epsg}").to_crs(epsg=4326).iloc[0]
        except Exception:
            return None, None
        return {"type": "Feature", "properties": {}, "geometry": mapping(geom_wgs84)}, round(float(geom_metric.area), 2)

    excluded_geojson, excluded_area_m2 = _to_wgs84_geojson(excluded_geom_metric)
    leftover_geojson, leftover_area_m2 = _to_wgs84_geojson(leftover_metric)

    return {
        "parent_plot_id": int(parent_plot_id),
        "method": method_key,
        "metric_epsg": int(metric_epsg),
        "orientation_deg": round(float(orientation_deg), 3),
        "requested_count": int(split_count or 0) if split_count is not None else None,
        "target_area_m2": round(float(target_area), 2) if target_area > 0 else None,
        "fraction_weights": [round(float(w), 6) for w in effective_fraction_weights] if effective_fraction_weights else None,
        "fraction_breaks": [round(float(v), 6) for v in effective_fraction_breaks] if effective_fraction_breaks else None,
        "custom_areas_m2": [round(float(v), 4) for v in effective_custom_areas_m2] if effective_custom_areas_m2 else None,
        "lot_width_m": round(float(lot_width_m), 4) if lot_width_m else None,
        "lot_height_m": round(float(lot_height_m), 4) if lot_height_m else None,
        "dimension_unit": unit_key if lot_width_m else None,
        "exclude_road": bool(exclude_road),
        "road_width_m": round(float(road_width_m), 2) if exclude_road else None,
        "excluded_geojson": excluded_geojson,
        "excluded_area_m2": excluded_area_m2,
        "road_segments": road_segments_out,
        "leftover_geojson": leftover_geojson,
        "leftover_area_m2": leftover_area_m2,
        "lot_count_balanced": lot_count_balanced,
        "resolved_count": len(plots),
        "total_area_m2": round(total_area_m2, 2),
        "derived_total_area_m2": round(derived_total, 2),
        "area_imbalance_m2": round(total_area_m2 - derived_total, 4),
        "plots": plots,
        "preview_geojson": feature_collection,
    }


def _ensure_plot_meta_row(db: Session, plot_id: int):
    db.execute(
        text(
            """
            INSERT INTO plot_meta (plot_id)
            VALUES (:plot_id)
            ON CONFLICT (plot_id) DO NOTHING
            """
        ),
        {"plot_id": int(plot_id)},
    )


def _set_plot_subdivision_meta(
    db: Session,
    *,
    child_plot_id: int,
    parent_plot_id: int,
    batch_id: int,
    lot_no: str,
    estate_name: str | None,
):
    _ensure_plot_meta_row(db, child_plot_id)
    db.execute(
        text(
            """
            UPDATE plot_meta
            SET parent_plot_id = :parent_plot_id,
                subdivision_batch_id = :batch_id,
                subdivision_lot_no = :lot_no,
                estate_name = :estate_name,
                updated_at = NOW()
            WHERE plot_id = :plot_id
            """
        ),
        {
            "plot_id": int(child_plot_id),
            "parent_plot_id": int(parent_plot_id),
            "batch_id": int(batch_id),
            "lot_no": str(lot_no or ""),
            "estate_name": str(estate_name or ""),
        },
    )


def _apply_child_plot_meta(
    db: Session,
    *,
    plot_id: int,
    title_text: str,
    parent_meta: dict,
    adamawa_owner_name: str | None = None,
):
    _ensure_plot_meta_row(db, plot_id)
    db.execute(
        text(
            """
            UPDATE plot_meta
            SET title_text = :title_text,
                location_text = :location_text,
                lga_text = :lga_text,
                state_text = :state_text,
                surveyor_name = :surveyor_name,
                surveyor_rank = :surveyor_rank,
                certification_statement = :certification_statement,
                scale_text = :scale_text,
                paper_size = :paper_size,
                coordinate_system = :coordinate_system,
                template_name = :template_name,
                adamawa_rof_no = :adamawa_rof_no,
                adamawa_owner_name = :adamawa_owner_name,
                adamawa_authority_title = :adamawa_authority_title,
                adamawa_authority_date_text = :adamawa_authority_date_text,
                adamawa_control_point_name = :adamawa_control_point_name,
                adamawa_northing = :adamawa_northing,
                adamawa_easting = :adamawa_easting,
                adamawa_elevation = :adamawa_elevation,
                adamawa_origin_text = :adamawa_origin_text,
                adamawa_topo_sheet_text = :adamawa_topo_sheet_text,
                adamawa_computation_no = :adamawa_computation_no,
                adamawa_cadastral_sheet_no = :adamawa_cadastral_sheet_no,
                adamawa_plan_no = :adamawa_plan_no,
                adamawa_surveyed_by_text = :adamawa_surveyed_by_text,
                adamawa_disclaimer_text = :adamawa_disclaimer_text,
                cadastral_plan_no = :cadastral_plan_no,
                cadastral_area_name = :cadastral_area_name,
                cadastral_datum_text = :cadastral_datum_text,
                cadastral_firm_block_text = :cadastral_firm_block_text,
                fct_file_no = :fct_file_no,
                fct_district = :fct_district,
                fct_cadastral_zone = :fct_cadastral_zone,
                fct_origin_beacon_text = :fct_origin_beacon_text,
                fct_cadastral_map_ref = :fct_cadastral_map_ref,
                fct_title_prefix = :fct_title_prefix,
                updated_at = NOW()
            WHERE plot_id = :plot_id
            """
        ),
        {
            "plot_id": int(plot_id),
            "title_text": title_text,
            "location_text": parent_meta.get("location_text") or None,
            "lga_text": parent_meta.get("lga_text") or None,
            "state_text": parent_meta.get("state_text") or None,
            "surveyor_name": parent_meta.get("surveyor_name") or None,
            "surveyor_rank": parent_meta.get("surveyor_rank") or None,
            "certification_statement": parent_meta.get("certification_statement") or DEFAULT_CERTIFICATION_STATEMENT,
            "scale_text": parent_meta.get("scale_text") or "1 : 1000",
            "paper_size": parent_meta.get("paper_size") or "A4",
            "coordinate_system": parent_meta.get("coordinate_system") or "wgs84",
            "template_name": parent_meta.get("template_name") or DEFAULT_TEMPLATE_NAME,
            "adamawa_rof_no": parent_meta.get("adamawa_rof_no") or "",
            "adamawa_owner_name": (
                str(adamawa_owner_name or "").strip()
                or parent_meta.get("adamawa_owner_name")
                or ""
            ),
            "adamawa_authority_title": parent_meta.get("adamawa_authority_title") or DEFAULT_ADAMAWA_AUTHORITY_TITLE,
            "adamawa_authority_date_text": parent_meta.get("adamawa_authority_date_text") or DEFAULT_ADAMAWA_AUTHORITY_DATE,
            "adamawa_control_point_name": "",
            "adamawa_northing": "",
            "adamawa_easting": "",
            "adamawa_elevation": "",
            "adamawa_origin_text": parent_meta.get("adamawa_origin_text") or DEFAULT_ADAMAWA_ORIGIN_TEXT,
            "adamawa_topo_sheet_text": parent_meta.get("adamawa_topo_sheet_text") or DEFAULT_ADAMAWA_TOPO_SHEET_TEXT,
            "adamawa_computation_no": parent_meta.get("adamawa_computation_no") or "",
            "adamawa_cadastral_sheet_no": parent_meta.get("adamawa_cadastral_sheet_no") or "",
            "adamawa_plan_no": parent_meta.get("adamawa_plan_no") or "",
            "adamawa_surveyed_by_text": parent_meta.get("adamawa_surveyed_by_text") or "",
            "adamawa_disclaimer_text": parent_meta.get("adamawa_disclaimer_text") or DEFAULT_ADAMAWA_DISCLAIMER_TEXT,
            # Plan number is per-lot (left blank for the surveyor to fill in individually);
            # locality, datum, and firm details are shared across the whole subdivision.
            "cadastral_plan_no": "",
            "cadastral_area_name": parent_meta.get("cadastral_area_name") or "",
            "cadastral_datum_text": parent_meta.get("cadastral_datum_text") or "",
            "cadastral_firm_block_text": parent_meta.get("cadastral_firm_block_text") or "",
            # File No is per-lot (left blank for the surveyor to fill in individually); district,
            # cadastral zone, origin beacon, and map reference are shared across the subdivision.
            "fct_file_no": "",
            "fct_district": parent_meta.get("fct_district") or "",
            "fct_cadastral_zone": parent_meta.get("fct_cadastral_zone") or "",
            "fct_origin_beacon_text": parent_meta.get("fct_origin_beacon_text") or "",
            "fct_cadastral_map_ref": parent_meta.get("fct_cadastral_map_ref") or "",
            "fct_title_prefix": parent_meta.get("fct_title_prefix") or "",
        },
    )


def _run_plot_feature_detection(db: Session, plot_id: int):
    db.execute(
        text(
            """
            INSERT INTO plot_buffers (plot_id, geom)
            SELECT :plot_id,
                   ST_Buffer(geom::geography, 50)::geometry
            FROM plots
            WHERE id = :plot_id
            ON CONFLICT DO NOTHING
            """
        ),
        {"plot_id": int(plot_id)},
    )

    db.execute(
        text(
            """
            DELETE FROM detected_features
            WHERE plot_id = :plot_id
            """
        ),
        {"plot_id": int(plot_id)},
    )

    # `multipolygons`/`lines` below are a one-time bulk import of a Nigeria OSM extract - real,
    # fast, local data, but only for Nigeria. A plot centered outside Nigeria's bounds gets its
    # buildings/roads/rivers from the Overpass-backed regional cache instead (see
    # app/utils/osm_overpass.py for why: per-country bulk imports for every country a surveyor
    # might ever use aren't worth the storage until a country shows real, sustained usage).
    centroid_row = db.execute(
        text("SELECT ST_X(ST_Centroid(geom)) AS lon, ST_Y(ST_Centroid(geom)) AS lat FROM plots WHERE id = :plot_id"),
        {"plot_id": int(plot_id)},
    ).mappings().first()
    centroid_lon = float(centroid_row["lon"]) if centroid_row and centroid_row["lon"] is not None else None
    centroid_lat = float(centroid_row["lat"]) if centroid_row and centroid_row["lat"] is not None else None

    if centroid_lon is not None and centroid_lat is not None and not validate_nigeria_bounds(centroid_lon, centroid_lat):
        # Deliberately synchronous, not deferred to BackgroundTasks. This used to run in the
        # background so plot creation could return instantly, but that made feature detection a
        # race against the user's own next request: the frontend's first preview/render (which
        # reads detected_features directly) almost always got there before the background job
        # finished, showing zero features - only a *second* preview, moments later once the
        # background job had caught up, showed the real ones. "Detected nothing" and "haven't
        # checked yet" must never look the same to the user. osm_overpass.py's own bucket cache
        # (90-day TTL) already makes this fast for every plot after the first one in a given
        # ~2.2km area - only a genuinely cold area pays the full Overpass round-trip, which
        # OVERPASS_HARD_TIMEOUT_S still bounds. Paying that latency once, inline, so the response
        # the user gets back is already correct, is the whole point.
        try:
            from app.utils.osm_overpass import run_overpass_feature_detection
            run_overpass_feature_detection(db, int(plot_id), centroid_lat, centroid_lon)
        except Exception:
            logger.warning("Overpass feature detection failed for plot %s", plot_id, exc_info=True)
        return

    # Buildings
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'building', 'inside', m.geom
            FROM multipolygons m
            JOIN plots p ON p.id = :plot_id
            WHERE m.building IS NOT NULL
              AND ST_Intersects(m.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'building', 'buffer', m.geom
            FROM multipolygons m
            JOIN plot_buffers b ON b.plot_id = :plot_id
            JOIN plots p ON p.id = :plot_id
            WHERE m.building IS NOT NULL
              AND ST_Intersects(m.geom, b.geom)
              AND NOT ST_Intersects(m.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )

    # Roads
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'road', 'inside', r.geom
            FROM (
                SELECT geom FROM lines WHERE highway IS NOT NULL
                UNION ALL
                SELECT geom FROM multilinestrings
                WHERE type = 'highway' OR other_tags LIKE '%highway%'
            ) r
            JOIN plots p ON p.id = :plot_id
            WHERE ST_Intersects(r.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'road', 'buffer', r.geom
            FROM (
                SELECT geom FROM lines WHERE highway IS NOT NULL
                UNION ALL
                SELECT geom FROM multilinestrings
                WHERE type = 'highway' OR other_tags LIKE '%highway%'
            ) r
            JOIN plot_buffers b ON b.plot_id = :plot_id
            JOIN plots p ON p.id = :plot_id
            WHERE ST_Intersects(r.geom, b.geom)
              AND NOT ST_Intersects(r.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )

    # Rivers
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'river', 'inside', r.geom
            FROM (
                SELECT geom FROM lines WHERE waterway IS NOT NULL
                UNION ALL
                SELECT geom FROM multilinestrings
                WHERE type = 'waterway' OR other_tags LIKE '%waterway%'
            ) r
            JOIN plots p ON p.id = :plot_id
            WHERE ST_Intersects(r.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )
    db.execute(
        text(
            """
            INSERT INTO detected_features (plot_id, feature_type, location, geom)
            SELECT :plot_id, 'river', 'buffer', r.geom
            FROM (
                SELECT geom FROM lines WHERE waterway IS NOT NULL
                UNION ALL
                SELECT geom FROM multilinestrings
                WHERE type = 'waterway' OR other_tags LIKE '%waterway%'
            ) r
            JOIN plot_buffers b ON b.plot_id = :plot_id
            JOIN plots p ON p.id = :plot_id
            WHERE ST_Intersects(r.geom, b.geom)
              AND NOT ST_Intersects(r.geom, p.geom)
            """
        ),
        {"plot_id": int(plot_id)},
    )


def _load_plot_polygon_wgs84(db: Session, plot_id: int) -> Polygon:
    plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id = :id"), {"id": int(plot_id)}).scalar()
    if plot_wkb is None:
        raise HTTPException(status_code=404, detail="Mother parcel not found.")
    try:
        geom = wkb.loads(plot_wkb)
    except Exception:
        raise HTTPException(status_code=400, detail="Mother parcel geometry is invalid.")
    poly = _clean_single_polygon(geom)
    if poly is None:
        raise HTTPException(status_code=400, detail="Mother parcel geometry is invalid.")
    return poly


def _normalize_survey_input_coordinates(raw: Any) -> list[dict]:
    if raw is None:
        return []
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except Exception:
            return []
    if not isinstance(raw, list):
        return []

    normalized: list[dict] = []
    for index, item in enumerate(raw):
        if not isinstance(item, dict):
            continue
        try:
            x = float(item.get("x"))
            y = float(item.get("y"))
        except Exception:
            continue
        station = str(item.get("station") or item.get("name") or f"P{index + 1}").strip() or f"P{index + 1}"
        normalized.append(
            {
                "station": station,
                "x": x,
                "y": y,
                "height": item.get("height"),
                # Missing/omitted (older saved drafts, or any caller that never set it) defaults
                # to a boundary point, matching pre-existing behavior where every entered point
                # was always a boundary vertex.
                "is_boundary": item.get("is_boundary") is not False,
            }
        )
    return normalized


def _build_exact_measurement_polygon(
    survey_input_coordinates: Any,
    coordinate_system: str | None,
    target_epsg: int,
) -> tuple[Polygon | None, float | None, list[dict], str | None]:
    normalized = _normalize_survey_input_coordinates(survey_input_coordinates)
    # survey_input_coordinates can hold spot-height-only samples alongside the boundary corners
    # (see the "Spot Heights" CSV import step) - only the boundary-flagged points form the ring;
    # everything else exists purely as elevation reference data and must never distort the shape.
    boundary_points = [item for item in normalized if item.get("is_boundary") is not False]
    if len(boundary_points) < 3:
        return None, None, [], None

    selected_key = str(coordinate_system or "wgs84").strip().lower() or "wgs84"
    coords = [(float(item["x"]), float(item["y"])) for item in boundary_points]
    if coords[0] != coords[-1]:
        coords.append(coords[0])

    if selected_key == "wgs84":
        src_epsg = 4326
        resolved_key = "wgs84"
    else:
        sample_x, sample_y = coords[0]
        resolved_key = resolve_coordinate_system_key(selected_key, sample_x, sample_y)
        src_epsg = int(COORDINATE_SYSTEMS.get(resolved_key, COORDINATE_SYSTEMS.get(selected_key, 4326)))

    poly = Polygon(coords)
    poly = _clean_single_polygon(poly)
    if poly is None:
        return None, None, boundary_points, resolved_key

    gdf_poly = gpd.GeoDataFrame(geometry=[poly], crs=f"EPSG:{src_epsg}")
    if int(src_epsg) != int(target_epsg):
        gdf_poly = gdf_poly.to_crs(epsg=int(target_epsg))
    projected_poly = gdf_poly.geometry.iloc[0]
    projected_poly = _clean_single_polygon(projected_poly)
    if projected_poly is None:
        return None, None, boundary_points, resolved_key

    return projected_poly, float(projected_poly.area), boundary_points, resolved_key


def _resolve_measurement_polygon_context(
    plot_geom_wgs84: Polygon,
    coordinate_system: str | None,
    survey_input_coordinates: Any = None,
) -> dict:
    render_coordinate_system, epsg_code, crs_name = _resolve_survey_render_crs(
        coordinate_system,
        plot_geom_wgs84,
    )
    exact_poly, exact_area_m2, normalized_points, resolved_exact_key = _build_exact_measurement_polygon(
        survey_input_coordinates,
        coordinate_system,
        epsg_code,
    )
    if exact_poly is not None and exact_area_m2 is not None:
        return {
            "render_coordinate_system": render_coordinate_system,
            "epsg_code": int(epsg_code),
            "crs_name": crs_name,
            "measurement_polygon": exact_poly,
            "measurement_area_m2": float(exact_area_m2),
            "survey_input_coordinates": normalized_points,
            "used_exact_input": True,
            "resolved_exact_key": resolved_exact_key,
        }

    fallback_poly = gpd.GeoSeries([plot_geom_wgs84], crs="EPSG:4326").to_crs(epsg=epsg_code).iloc[0]
    fallback_poly = _clean_single_polygon(fallback_poly)
    if fallback_poly is None:
        raise HTTPException(status_code=400, detail="Plot geometry could not be projected for measurement.")
    return {
        "render_coordinate_system": render_coordinate_system,
        "epsg_code": int(epsg_code),
        "crs_name": crs_name,
        "measurement_polygon": fallback_poly,
        "measurement_area_m2": float(fallback_poly.area),
        "survey_input_coordinates": normalized_points,
        "used_exact_input": False,
        "resolved_exact_key": None,
    }


def _safe_filename_fragment(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", str(value or "")).strip("._-")
    return cleaned or fallback


def _build_plot_export_filename(plot_id, meta: dict | None, doc_label: str, ext: str) -> str:
    """Builds a human-readable export filename from whichever identity fields are actually set
    (applicant/title, location) instead of the opaque `plot_<id>_...` names these endpoints used
    to return - mirrors the frontend's buildExportFilename/surveyPlanIdentitySegments convention
    (SurveyPlan.tsx) so the same plot produces a same-shaped name whichever path it's downloaded
    through (this backend name is also what becomes the R2 storage object key). Falls back to
    Plot_<id> only when nothing identifying has been filled in yet.
    """
    meta = meta or {}
    title = str(meta.get("title_text") or "").strip()
    if title.upper() == "SURVEY PLAN":
        # The literal placeholder default on the general template, not a real identifying name.
        title = ""
    location = str(meta.get("location_text") or "").strip()
    parts = [frag for frag in (_safe_filename_fragment(title, ""), _safe_filename_fragment(location, "")) if frag]
    identity = "_".join(parts) if parts else f"Plot_{plot_id}"
    return f"{identity}_{doc_label}.{ext}"


def _format_coordinate_number(value: float, decimals: int) -> str:
    try:
        numeric = float(value)
    except Exception:
        numeric = 0.0
    # Plain "." decimal, no thousands separator and no Excel apostrophe/formula wrapper - this
    # file is named for DGPS/GIS ingestion, and every consumer of that kind (QGIS, Civil 3D, DGPS
    # receivers) expects a bare numeric token, not an Excel-only display trick.
    return f"{numeric:.{decimals}f}"


def _normalize_lot_key(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).lower()


def _extract_clean_copy_area_overrides(raw_items: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    if not isinstance(raw_items, (list, tuple)):
        return out

    for raw in raw_items:
        if not isinstance(raw, dict):
            continue

        label = str(raw.get("label") or "").strip()
        if not label:
            continue

        lot_no_raw = raw.get("lot_no")
        child_plot_id_raw = raw.get("child_plot_id")
        lot_key = _normalize_lot_key(lot_no_raw)

        if lot_key:
            out[f"lot:{lot_key}"] = label
            continue

        try:
            child_plot_id = int(child_plot_id_raw)
        except Exception:
            child_plot_id = 0
        if child_plot_id > 0:
            out[f"child:{child_plot_id}"] = label

    return out


def _resolve_clean_copy_area_label(
    lot_no: str,
    child_plot_id: int,
    area_m2: float,
    overrides: dict[str, str] | None,
) -> str:
    override_map = overrides or {}
    lot_key = _normalize_lot_key(lot_no)
    if lot_key:
        label = override_map.get(f"lot:{lot_key}")
        if label:
            return label
    label = override_map.get(f"child:{int(child_plot_id)}")
    if label:
        return label
    return format_area_display(float(area_m2))


def _iter_line_geometries_for_clean_copy(geom: Any):
    if geom is None or getattr(geom, "is_empty", False):
        return
    gtype = getattr(geom, "geom_type", "")
    if gtype in ("LineString", "LinearRing"):
        try:
            yield geom
        except Exception:
            return
        return
    if gtype == "Polygon":
        try:
            yield geom.exterior
            for ring in geom.interiors:
                yield ring
        except Exception:
            pass
        return
    if hasattr(geom, "geoms"):
        for part in geom.geoms:
            yield from _iter_line_geometries_for_clean_copy(part)


def _render_subdivision_clean_copy_pdf(
    db: Session,
    parent_plot_id: int,
    parent_poly_wgs84: Polygon,
    child_rows: list[dict],
    output_pdf_path: str,
    *,
    title_text: str,
    paper_size: str,
    scale_text: str,
    coordinate_system: str,
    epsg_code: int,
    station_names: list[str] | None,
    north_arrow_style: str,
    north_arrow_color: str,
    beacon_style: str,
    road_width_m: float | None,
    area_overrides: dict[str, str] | None = None,
    boundary_color: str | None = None,
    grid_color: str | None = None,
    text_color: str | None = None,
    road_color: str | None = None,
    river_color: str | None = None,
    building_color: str | None = None,
    building_hatch_type: str | None = None,
    road_style: str | None = None,
    station_font: str | None = None,
    station_size: int | None = None,
    bearing_font: str | None = None,
    bearing_size: int | None = None,
):
    # None means "not overridden" - fall back to this clean-copy template's own established
    # look so omitting these params leaves existing exports unchanged. grid_color is unused
    # here (this template has no coordinate grid), kept only for signature symmetry.
    boundary_color = boundary_color or "black"
    text_color = text_color or "black"
    road_color = road_color or "#f97316"
    river_color = river_color or "#1d4ed8"
    building_color = building_color or "black"
    building_hatch_type = building_hatch_type or "diagonal"
    paper_name = str(paper_size or "A4").upper()
    if paper_name not in {"A4", "A3", "A2", "A1", "A0"}:
        paper_name = "A4"

    display_epsg = int(epsg_code or 4326)
    if str(coordinate_system or "").strip().lower() == "wgs84" or display_epsg == 4326:
        display_epsg = _metric_epsg_for_wgs84_polygon(parent_poly_wgs84)

    parent_metric = gpd.GeoDataFrame(geometry=[parent_poly_wgs84], crs="EPSG:4326").to_crs(epsg=display_epsg).geometry.iloc[0]

    child_metric_rows: list[dict] = []
    for row in child_rows:
        geom_geojson_raw = row.get("geom_geojson")
        if not geom_geojson_raw:
            continue
        try:
            geom_wgs = _clean_single_polygon(shape(json.loads(geom_geojson_raw)))
        except Exception:
            geom_wgs = None
        if geom_wgs is None or geom_wgs.is_empty:
            continue
        geom_metric = gpd.GeoDataFrame(geometry=[geom_wgs], crs="EPSG:4326").to_crs(epsg=display_epsg).geometry.iloc[0]
        child_metric_rows.append(
            {
                "child_plot_id": int(row.get("child_plot_id") or 0),
                "lot_no": str(row.get("lot_no") or "").strip() or f"LOT-{int(row.get('child_plot_id') or 0)}",
                "area_m2": float(row.get("area_m2") or 0.0),
                "geometry": geom_metric,
            }
        )

    if not child_metric_rows:
        raise HTTPException(status_code=404, detail="Subdivision batch has no valid lot geometries.")

    detected_rows = db.execute(
        text("SELECT geom, feature_type FROM detected_features WHERE plot_id=:id"),
        {"id": int(parent_plot_id)},
    ).fetchall()
    override_rows = db.execute(
        text(
            """
            SELECT feature_type, action, name, width_m, ST_AsGeoJSON(geom) AS geojson
            FROM plot_feature_overrides
            WHERE plot_id = :id
            """
        ),
        {"id": int(parent_plot_id)},
    ).fetchall()
    roads_auto_rows = db.execute(
        text(
            """
            WITH roads AS (
                SELECT
                    CASE
                        WHEN ST_SRID(r.geom) = 4326 THEN r.geom
                        WHEN ST_SRID(r.geom) = 0 THEN ST_SetSRID(r.geom, 4326)
                        ELSE ST_Transform(r.geom, 4326)
                    END AS geom,
                    r.name AS name
                FROM lines r
                WHERE r.highway IS NOT NULL
            )
            SELECT roads.geom, roads.name
            FROM roads
            JOIN plot_buffers b ON b.plot_id = :plot_id
            WHERE ST_Intersects(roads.geom, b.geom)
            """
        ),
        {"plot_id": int(parent_plot_id)},
    ).fetchall()

    buildings_wgs: list[Any] = []
    rivers_wgs: list[Any] = []
    fences_wgs: list[Any] = []
    detected_roads_wgs: list[Any] = []
    for row in detected_rows:
        try:
            geom = wkb.loads(row.geom)
        except Exception:
            continue
        feature_type = str(row.feature_type or "").strip().lower()
        if feature_type == "building":
            buildings_wgs.append(geom)
        elif feature_type == "river":
            rivers_wgs.append(geom)
        elif feature_type == "fence":
            fences_wgs.append(geom)
        elif feature_type == "road":
            detected_roads_wgs.append(geom)

    overrides: list[dict[str, Any]] = []
    for row in override_rows:
        geom = None
        if row.geojson:
            try:
                geom = shape(json.loads(row.geojson))
            except Exception:
                geom = None
        overrides.append(
            {
                "feature_type": str(row.feature_type or "").strip().lower(),
                "action": str(row.action or "").strip().lower(),
                "name": row.name,
                "width_m": row.width_m,
                "geom": geom,
            }
        )

    def apply_overrides(base_list: list[Any], feature_type: str):
        result = list(base_list)
        added: list[Any] = []
        delete_geoms: list[Any] = []
        use_coverage_match = feature_type in ("road", "river", "fence")
        for ov in overrides:
            if ov.get("feature_type") != feature_type:
                continue
            geom = ov.get("geom")
            if geom is None:
                continue
            try:
                if hasattr(geom, "is_valid") and not geom.is_valid:
                    geom = geom.buffer(0)
            except Exception:
                pass
            if ov.get("action") in ("delete", "update"):
                if use_coverage_match:
                    result = [g for g in result if not _feature_geom_replaced_by_type(g, geom, feature_type)]
                else:
                    result = [g for g in result if not g.intersects(geom)]
                delete_geoms.append(geom)
            if ov.get("action") in ("add", "update"):
                result.append(geom)
                added.append(geom)
        if delete_geoms:
            if use_coverage_match:
                added = [
                    g for g in added
                    if not any(_feature_geom_replaced_by_type(g, dg, feature_type) for dg in delete_geoms)
                ]
            else:
                added = [g for g in added if not any(g.intersects(dg) for dg in delete_geoms)]
        return result, added, delete_geoms

    buildings_wgs, added_buildings_wgs, _ = apply_overrides(buildings_wgs, "building")
    rivers_wgs, _, _ = apply_overrides(rivers_wgs, "river")
    fences_wgs, added_fences_wgs, _ = apply_overrides(fences_wgs, "fence")

    roads_wgs: list[dict[str, Any]] = []
    for row in roads_auto_rows:
        try:
            roads_wgs.append(
                {
                    "geom": wkb.loads(row.geom),
                    "name": str(getattr(row, "name", "") or "").strip(),
                    "width_m": None,
                }
            )
        except Exception:
            continue
    for geom in detected_roads_wgs:
        roads_wgs.append(
            {
                "geom": geom,
                "name": "",
                "width_m": None,
            }
        )
    road_delete_geoms: list[Any] = []
    added_roads_wgs: list[dict[str, Any]] = []
    for ov in overrides:
        if ov.get("feature_type") != "road":
            continue
        geom = ov.get("geom")
        if geom is None:
            continue
        try:
            if hasattr(geom, "is_valid") and not geom.is_valid:
                geom = geom.buffer(0)
        except Exception:
            pass
        if ov.get("action") in ("delete", "update"):
            road_delete_geoms.append(geom)
        if ov.get("action") in ("add", "update"):
            width_value = None
            try:
                width_parsed = float(ov.get("width_m") or 0.0)
                if width_parsed > 0:
                    width_value = width_parsed
            except Exception:
                width_value = None
            added_roads_wgs.append(
                {
                    "geom": geom,
                    "name": str(ov.get("name") or "").strip(),
                    "width_m": width_value,
                }
            )

    roads_wgs = [
        r for r in roads_wgs
        if not any(
            _feature_geom_replaced_by_type(r.get("geom"), dg, "road")
            for dg in road_delete_geoms
            if r.get("geom") is not None
        )
    ]
    roads_wgs.extend(added_roads_wgs)

    paper_config = get_paper_config(paper_name)
    font_scale = float(paper_config.get("scale", 1.0))
    dpi = 220 if paper_name in {"A4", "A3"} else (170 if paper_name == "A2" else 130)
    scale_ratio = parse_scale_ratio(scale_text)
    parent_area_m2 = float(parent_metric.area or 0.0)
    old_font_family = matplotlib.rcParams.get("font.family")
    fig = None
    try:
        matplotlib.rcParams["font.family"] = "DejaVu Serif"
        fig = plt.figure(figsize=(paper_config["width"], paper_config["height"]), dpi=dpi)
        fig.patch.set_facecolor("white")

        # Clean-copy page frame (outer + inner) to match the template feel.
        fig.add_artist(
            patches.Rectangle(
                (0.02, 0.02),
                0.96,
                0.96,
                transform=fig.transFigure,
                fill=False,
                lw=1.0,
                edgecolor="black",
                zorder=300,
                clip_on=False,
            )
        )
        fig.add_artist(
            patches.Rectangle(
                (0.03, 0.03),
                0.94,
                0.94,
                transform=fig.transFigure,
                fill=False,
                lw=0.85,
                edgecolor="black",
                zorder=300,
                clip_on=False,
            )
        )

        raw_title = str(title_text or "").strip()
        wrap_width = 54 if paper_name in {"A0", "A1", "A2"} else 48
        title_lines: list[str] = []
        if raw_title:
            for segment in [s.strip() for s in str(raw_title).upper().splitlines() if s.strip()]:
                seg_clean = re.sub(r"\s+", " ", segment)
                title_lines.extend(textwrap.wrap(seg_clean, width=wrap_width, break_long_words=False))
        if not title_lines:
            title_lines = ["SURVEY PLAN CLEAN COPY PLAN"]
        title_lines = title_lines[:5]

        y_title = 0.955
        line_gap = 0.03 if paper_name in {"A0", "A1", "A2"} else 0.028
        title_size = max(9, int(10.6 * font_scale))
        for idx, line in enumerate(title_lines):
            fig.text(
                0.5,
                y_title - idx * line_gap,
                line,
                ha="center",
                va="top",
                fontsize=title_size,
                weight="bold",
                color="black",
                fontfamily="DejaVu Serif",
            )

        if parent_area_m2 > 0:
            area_y = y_title - len(title_lines) * line_gap - 0.006
            fig.text(
                0.5,
                area_y,
                f"AREA={format_area_display(parent_area_m2)}",
                ha="center",
                va="top",
                fontsize=max(8, int(9.4 * font_scale)),
                weight="bold",
                color="black",
                fontfamily="DejaVu Serif",
            )

        map_left, map_bottom, map_right = 0.055, 0.075, 0.945
        header_clearance = 0.085 + (len(title_lines) * line_gap)
        map_top = max(0.66, 0.96 - header_clearance)
        map_width = map_right - map_left
        map_height = map_top - map_bottom
        fig.add_artist(
            patches.Rectangle(
                (map_left, map_bottom),
                map_width,
                map_height,
                transform=fig.transFigure,
                fill=False,
                lw=0.9,
                edgecolor="black",
                zorder=250,
                clip_on=False,
            )
        )
        ax = fig.add_axes([map_left + 0.004, map_bottom + 0.004, map_width - 0.008, map_height - 0.008])
        ax.set_aspect("equal", adjustable="box")
        ax.set_anchor("C")
        ax.set_facecolor("white")

        # Fit map to parent parcel bounds; lock limits so later plots cannot stretch it.
        minx, miny, maxx, maxy = parent_metric.bounds
        span_x = max(maxx - minx, 1.0)
        span_y = max(maxy - miny, 1.0)
        pad_x = max(1.0, span_x * 0.1)
        pad_y = max(1.0, span_y * 0.1)
        target_xlim = (minx - pad_x, maxx + pad_x)
        target_ylim = (miny - pad_y, maxy + pad_y)
        ax.set_xlim(*target_xlim)
        ax.set_ylim(*target_ylim)
        ax.set_autoscale_on(False)
        extent_poly = box(target_xlim[0], target_ylim[0], target_xlim[1], target_ylim[1])
        clip_buffer = max(1.0, min(span_x, span_y) * 0.05)

        if rivers_wgs:
            gpd.GeoDataFrame(geometry=rivers_wgs, crs="EPSG:4326").to_crs(epsg=display_epsg).plot(
                ax=ax, color=river_color, lw=max(0.8, 1.0 * font_scale), zorder=5
            )

        road_edge_lw = max(0.75, 0.9 * font_scale)
        road_edge_color = road_color
        # This clean-copy renderer draws roads with its own hand-rolled ax.plot calls rather than
        # the shared _draw_road_edges - an explicit "solid" choice is still honored here for
        # consistency, but the dashed default's tick-symbol embellishment isn't duplicated into
        # this secondary export path.
        road_edge_linestyle = "-" if road_style == "solid" else (0, (7, 4))
        road_label_size = max(7, int(7.2 * font_scale))
        road_snap_tol = max(1.0, (5.0 / 1000.0) * scale_ratio)
        road_label_features: list[tuple[Any, str]] = []
        road_line_zorder = 21
        road_label_zorder = 23

        def _line_length_total(geom_obj: Any) -> float:
            total = 0.0
            for part in _iter_line_geometries_for_clean_copy(geom_obj):
                if part is None or part.is_empty:
                    continue
                try:
                    total += float(getattr(part, "length", 0.0))
                except Exception:
                    continue
            return total

        for road_item in roads_wgs:
            road_geom = road_item.get("geom")
            if road_geom is None:
                continue
            road_name = str(road_item.get("name") or "").strip()
            width_value = None
            try:
                parsed_width = float(road_item.get("width_m") or 0.0)
                if parsed_width > 0:
                    width_value = parsed_width
            except Exception:
                width_value = None
            if width_value is None:
                try:
                    global_width = float(road_width_m or 0.0)
                    width_value = global_width if global_width > 0 else 10.0
                except Exception:
                    width_value = 10.0
            half_width = max(1.0, width_value / 2.0)

            try:
                projected = gpd.GeoSeries([road_geom], crs="EPSG:4326").to_crs(epsg=display_epsg).iloc[0]
            except Exception:
                continue
            clipped = projected.intersection(extent_poly.buffer(road_snap_tol))
            if clipped.is_empty:
                continue
            snapped_clipped = snap(clipped, extent_poly.boundary, road_snap_tol)
            edge_lines = _collect_connected_road_edge_lines([(snapped_clipped, half_width)], snap_tol_m=road_snap_tol)
            source_len = _line_length_total(snapped_clipped)
            drawn_edge_len = 0.0
            if not edge_lines:
                # Fallback: draw per-part offset lines so road still appears for edge cases.
                for part in _iter_line_geometries_for_clean_copy(snapped_clipped):
                    if part is None or part.is_empty:
                        continue
                    try:
                        for side in ("left", "right"):
                            edge_geom = part.parallel_offset(half_width, side, join_style=2)
                            for edge_part in _iter_line_geometries_for_clean_copy(edge_geom):
                                if edge_part is None or edge_part.is_empty:
                                    continue
                                x_vals, y_vals = edge_part.xy
                                ax.plot(
                                    x_vals,
                                    y_vals,
                                    color=road_edge_color,
                                    lw=road_edge_lw,
                                    linestyle=road_edge_linestyle,
                                    zorder=road_line_zorder,
                                )
                                try:
                                    drawn_edge_len += float(getattr(edge_part, "length", 0.0))
                                except Exception:
                                    pass
                    except Exception:
                        continue
            else:
                for seg in edge_lines:
                    try:
                        x_vals, y_vals = seg.xy
                        ax.plot(
                            x_vals,
                            y_vals,
                            color=road_edge_color,
                            lw=road_edge_lw,
                            linestyle=road_edge_linestyle,
                            zorder=road_line_zorder,
                        )
                        drawn_edge_len += float(getattr(seg, "length", 0.0))
                    except Exception:
                        continue
            # If edge geometry is too weak (or effectively invisible), force a centerline fallback.
            if source_len > 0 and drawn_edge_len < (0.35 * source_len):
                for part in _iter_line_geometries_for_clean_copy(snapped_clipped):
                    if part is None or part.is_empty:
                        continue
                    try:
                        x_vals, y_vals = part.xy
                        ax.plot(
                            x_vals,
                            y_vals,
                            color=road_edge_color,
                            lw=max(0.7, 0.85 * road_edge_lw),
                            linestyle=road_edge_linestyle,
                            zorder=road_line_zorder,
                        )
                    except Exception:
                        continue
            if road_name:
                road_label_features.append((snapped_clipped, road_name))

        for geom, road_name in road_label_features:
            try:
                if geom is None or geom.is_empty:
                    continue
                label_line = max(
                    [lp for lp in _iter_line_geometries_for_clean_copy(geom) if lp is not None and not lp.is_empty],
                    key=lambda ln: float(getattr(ln, "length", 0.0)),
                    default=None,
                )
                if label_line is None:
                    continue
                mid = label_line.interpolate(0.5, normalized=True)
                p1 = label_line.interpolate(0.45, normalized=True)
                p2 = label_line.interpolate(0.55, normalized=True)
                angle = math.degrees(math.atan2(p2.y - p1.y, p2.x - p1.x))
                if angle < -90 or angle > 90:
                    angle += 180
                # Shrink road-name text to available road length so labels fit.
                label_len_m = max(1.0, float(getattr(label_line, "length", 1.0)))
                char_count = max(1, len(road_name))
                # Approx text-on-ground length model: chars * 0.62 * pt * 0.3528mm * scale_ratio
                max_pt_fit = (0.74 * label_len_m * 1000.0) / (char_count * 0.62 * 0.3528 * max(1, scale_ratio))
                fitted_size = max(5.0, min(float(road_label_size), float(max_pt_fit)))
                ax.text(
                    mid.x,
                    mid.y,
                    road_name.upper(),
                    fontsize=fitted_size,
                    color=road_color,
                    ha="center",
                    va="center",
                    rotation=angle,
                    weight="normal",
                    zorder=road_label_zorder,
                )
            except Exception:
                continue

        all_buildings = list(buildings_wgs) + list(added_buildings_wgs or [])
        if all_buildings:
            draw_building_hatch(
                ax,
                all_buildings,
                display_epsg,
                scale_ratio=scale_ratio,
                font_scale=font_scale,
                color=building_color,
                hatch_type=building_hatch_type,
            )
            try:
                gpd.GeoDataFrame(geometry=all_buildings, crs="EPSG:4326").to_crs(epsg=display_epsg).plot(
                    ax=ax, facecolor="none", edgecolor=building_color, lw=max(0.8, 0.9 * font_scale), zorder=8
                )
            except Exception:
                pass

        all_fences = list(fences_wgs or []) + list(added_fences_wgs or [])
        if all_fences:
            draw_fences(
                ax,
                all_fences,
                display_epsg=display_epsg,
                scale_ratio=scale_ratio,
                font_scale=font_scale,
            )
        fence_avoid_geom = build_fence_avoid_geom(all_fences, display_epsg=display_epsg, scale_ratio=scale_ratio)

        for row in child_metric_rows:
            geom_metric = _clean_single_polygon(row["geometry"])
            if geom_metric is None or geom_metric.is_empty:
                continue
            try:
                x_vals, y_vals = geom_metric.exterior.xy
                ax.plot(x_vals, y_vals, color=boundary_color, linewidth=max(0.85, 0.95 * font_scale), zorder=17)
            except Exception:
                continue

        boundary_mm = 0.7 if paper_name in ["A0"] else 0.5 if paper_name in ["A1"] else 0.35
        boundary_lw_pts = boundary_mm * 72.0 / 25.4
        gpd.GeoDataFrame(geometry=[parent_metric], crs=f"EPSG:{display_epsg}").plot(
            ax=ax, facecolor="none", edgecolor=boundary_color, lw=boundary_lw_pts, zorder=20
        )
        ax.set_xlim(target_xlim)
        ax.set_ylim(target_ylim)

        min_label_mm = 12
        min_label_length_m = (min_label_mm / 1000.0) * scale_ratio
        annotate_vertices(
            ax,
            parent_metric,
            int(parent_plot_id),
            station_names=station_names if station_names else None,
            font_scale=font_scale,
            min_label_length_m=min_label_length_m,
            avoid_geom=fence_avoid_geom,
            scale_ratio=scale_ratio,
            boundary_poly=parent_metric,
            beacon_style=beacon_style,
            show_station_names=False,
            show_beacons=False,
            text_color=text_color,
            boundary_color=boundary_color,
            station_font=station_font,
            station_size=station_size,
            bearing_font=bearing_font,
            bearing_size=bearing_size,
        )

        for row in child_metric_rows:
            geom_metric = _clean_single_polygon(row["geometry"])
            if geom_metric is None or geom_metric.is_empty:
                continue
            lot_no = str(row.get("lot_no") or "").strip() or "LOT"
            child_plot_id = int(row.get("child_plot_id") or 0)
            area_m2 = float(row.get("area_m2") or 0.0)
            area_label = _resolve_clean_copy_area_label(lot_no, child_plot_id, area_m2, area_overrides)
            label_pt = geom_metric.representative_point()
            ax.text(
                label_pt.x,
                label_pt.y,
                f"{lot_no}\n{area_label}",
                ha="center",
                va="center",
                fontsize=max(7, int(6.8 * font_scale)),
                color=text_color,
                zorder=22,
                fontfamily="DejaVu Serif",
                weight="bold",
            )

        # Place clean-copy north arrow above the map frame and tight to right page edge,
        # matching the requested template look.
        clean_anchor_x = 0.948  # shifted slightly left from right inner border
        clean_anchor_y = min(0.915, max(map_top + 0.055, 0.86))
        add_north_arrow(
            ax,
            font_scale=font_scale,
            style=str(north_arrow_style or "one_side_stem"),
            color=str(north_arrow_color or "blue"),
            anchor_x=clean_anchor_x,
            anchor_y=clean_anchor_y,
        )

        ax.set_aspect("equal")
        ax.axis("off")
        fig.savefig(output_pdf_path, format="pdf", dpi=dpi, facecolor=fig.get_facecolor())
    finally:
        if fig is not None:
            plt.close(fig)
        matplotlib.rcParams["font.family"] = old_font_family


def _compose_child_title(parent_meta: dict, lot_no: str, estate_name: str | None) -> str:
    estate = str(estate_name or parent_meta.get("estate_name") or "").strip()
    base_title = str(parent_meta.get("title_text") or "SURVEY PLAN").strip()
    if estate:
        return f"{estate} - {lot_no}".strip()
    if base_title:
        return f"{base_title} - {lot_no}".strip()
    return f"SURVEY PLAN - {lot_no}".strip()


def _render_survey_plan_pdf_for_plot(db: Session, plot_id: int, output_pdf_path: str):
    meta = get_plot_meta(db, plot_id)
    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom_wgs84,
        meta["coordinate_system"],
        meta.get("survey_input_coordinates"),
    )
    render_coordinate_system = measurement_context["render_coordinate_system"]
    epsg_code = measurement_context["epsg_code"]
    crs_name = measurement_context["crs_name"]

    tmp_map = tempfile.NamedTemporaryFile(suffix="_map.png", delete=False)
    map_path = tmp_map.name
    tmp_map.close()

    render_plot_map_layout(
        db=db,
        plot_id=plot_id,
        output_path=map_path,
        title_text=meta["title_text"],
        location_text=meta["location_text"],
        lga_text=meta["lga_text"],
        state_text=meta["state_text"],
        scale_text=meta["scale_text"],
        surveyor_name=meta["surveyor_name"],
        surveyor_rank=meta["surveyor_rank"],
        certification_statement=meta.get("certification_statement"),
        station_names=None,
        coordinate_system=render_coordinate_system,
        epsg_code=epsg_code,
        crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
        measurement_polygon=measurement_context["measurement_polygon"],
        measurement_area_m2=measurement_context["measurement_area_m2"],
        paper_size=meta["paper_size"],
        north_arrow_style="one_side_stem",
        north_arrow_color="blue",
        beacon_style="cross",
        road_width_m=None,
        road_width_override_m=None,
        template_name=meta.get("template_name") or DEFAULT_TEMPLATE_NAME,
        adamawa_rof_no=meta.get("adamawa_rof_no") or "",
        adamawa_owner_name=meta.get("adamawa_owner_name") or "",
        adamawa_authority_title=meta.get("adamawa_authority_title") or DEFAULT_ADAMAWA_AUTHORITY_TITLE,
        adamawa_authority_date_text=meta.get("adamawa_authority_date_text") or DEFAULT_ADAMAWA_AUTHORITY_DATE,
        adamawa_control_point_name=meta.get("adamawa_control_point_name") or "",
        adamawa_northing=meta.get("adamawa_northing") or "",
        adamawa_easting=meta.get("adamawa_easting") or "",
        adamawa_elevation=meta.get("adamawa_elevation") or "",
        adamawa_origin_text=meta.get("adamawa_origin_text") or DEFAULT_ADAMAWA_ORIGIN_TEXT,
        adamawa_topo_sheet_text=meta.get("adamawa_topo_sheet_text") or DEFAULT_ADAMAWA_TOPO_SHEET_TEXT,
        adamawa_computation_no=meta.get("adamawa_computation_no") or "",
        adamawa_cadastral_sheet_no=meta.get("adamawa_cadastral_sheet_no") or "",
        adamawa_plan_no=meta.get("adamawa_plan_no") or "",
        adamawa_surveyed_by_text=meta.get("adamawa_surveyed_by_text") or "",
        adamawa_disclaimer_text=meta.get("adamawa_disclaimer_text") or DEFAULT_ADAMAWA_DISCLAIMER_TEXT,
        cadastral_plan_no=meta.get("cadastral_plan_no") or "",
        cadastral_area_name=meta.get("cadastral_area_name") or "",
        cadastral_datum_text=meta.get("cadastral_datum_text") or "",
        cadastral_firm_block_text=meta.get("cadastral_firm_block_text") or "",
        fct_file_no=meta.get("fct_file_no") or "",
        fct_district=meta.get("fct_district") or "",
        fct_cadastral_zone=meta.get("fct_cadastral_zone") or "",
        fct_origin_beacon_text=meta.get("fct_origin_beacon_text") or "",
        fct_cadastral_map_ref=meta.get("fct_cadastral_map_ref") or "",
        fct_title_prefix=meta.get("fct_title_prefix") or "",
    )
    report = get_plot_report(plot_id, db)
    generate_plot_report_pdf(report, output_pdf_path, map_path, paper_size=meta["paper_size"])
    safe_remove(map_path)


@router.post("/{plot_id}/subdivision/preview")
def preview_plot_subdivision(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    method: str = Body("by_count"),
    split_count: int | None = Body(None),
    target_area_m2: float | None = Body(None),
    orientation_deg: float = Body(0.0),
    lot_prefix: str = Body("LOT"),
    estate_name: str = Body(""),
    fraction_weights: list[float] | None = Body(None),
    fraction_breaks: list[float] | None = Body(None),
    custom_areas_m2: list[float] | None = Body(None),
    lot_names: list[str] | None = Body(None),
    exclude_road: bool = Body(False),
    road_width_m: float = Body(10.0),
    lot_width: float | None = Body(None),
    lot_height: float | None = Body(None),
    dimension_unit: str = Body("m"),
):
    parent_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    payload = _compute_subdivision_payload(
        plot_id,
        parent_geom_wgs84,
        db,
        method=method,
        split_count=split_count,
        target_area_m2=target_area_m2,
        orientation_deg=_coerce_float(orientation_deg, 0.0),
        lot_prefix=lot_prefix,
        fraction_weights=fraction_weights,
        fraction_breaks=fraction_breaks,
        custom_areas_m2=custom_areas_m2,
        lot_names=lot_names,
        exclude_road=exclude_road,
        road_width_m=road_width_m,
        lot_width=lot_width,
        lot_height=lot_height,
        dimension_unit=dimension_unit,
    )
    payload["estate_name"] = str(estate_name or "").strip()
    log_survey_activity(
        db,
        event_type="preview_completed",
        workflow="subdivision",
        request=request,
        plot_id=plot_id,
        details={"method": method, "lot_count": payload.get("resolved_count")},
    )
    return payload


@router.post("/{plot_id}/subdivision/apply")
def apply_plot_subdivision(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    method: str = Body("by_count"),
    split_count: int | None = Body(None),
    target_area_m2: float | None = Body(None),
    orientation_deg: float = Body(0.0),
    lot_prefix: str = Body("LOT"),
    estate_name: str = Body(""),
    fraction_weights: list[float] | None = Body(None),
    fraction_breaks: list[float] | None = Body(None),
    custom_areas_m2: list[float] | None = Body(None),
    lot_names: list[str] | None = Body(None),
    include_feature_detection: bool = Body(False),
    exclude_road: bool = Body(False),
    road_width_m: float = Body(10.0),
    lot_width: float | None = Body(None),
    lot_height: float | None = Body(None),
    dimension_unit: str = Body("m"),
):
    parent_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    safe_estate_name = str(estate_name or "").strip()
    payload = _compute_subdivision_payload(
        plot_id,
        parent_geom_wgs84,
        db,
        method=method,
        split_count=split_count,
        target_area_m2=target_area_m2,
        orientation_deg=_coerce_float(orientation_deg, 0.0),
        lot_prefix=lot_prefix,
        fraction_weights=fraction_weights,
        fraction_breaks=fraction_breaks,
        custom_areas_m2=custom_areas_m2,
        lot_names=lot_names,
        exclude_road=exclude_road,
        road_width_m=road_width_m,
        lot_width=lot_width,
        lot_height=lot_height,
        dimension_unit=dimension_unit,
    )
    parent_meta = get_plot_meta(db, plot_id)

    # Child plots inherit the parent's ownership if it's already claimed; otherwise fall back to
    # whatever Survey session this request carries right now. Without this, `Plot(geom=...)` below
    # would leave every subdivision lot permanently unowned (owner_user_id NULL) even after the
    # surveyor logs in - unlike regular plots, which stamp ownership at creation time.
    parent_owner_id = db.execute(
        text("SELECT owner_user_id FROM plots WHERE id = :id"), {"id": plot_id}
    ).scalar()
    child_owner_id = parent_owner_id
    if child_owner_id is None:
        try:
            survey_session = resolve_survey_session(db, request)
            child_owner_id = survey_session.user_id if survey_session else None
        except Exception:
            child_owner_id = None

    batch_id = db.execute(
        text(
            """
            INSERT INTO plot_subdivision_batches (
                parent_plot_id, estate_name, method, requested_count, target_area_m2,
                orientation_deg, generated_count, total_area_m2, status,
                exclude_road, road_width_m, lot_width_m, lot_height_m, dimension_unit,
                leftover_area_m2, excluded_area_m2, leftover_geojson, excluded_geojson
            )
            VALUES (
                :parent_plot_id, :estate_name, :method, :requested_count, :target_area_m2,
                :orientation_deg, 0, 0, 'processing',
                :exclude_road, :road_width_m, :lot_width_m, :lot_height_m, :dimension_unit,
                :leftover_area_m2, :excluded_area_m2,
                CAST(:leftover_geojson AS JSONB), CAST(:excluded_geojson AS JSONB)
            )
            RETURNING id
            """
        ),
        {
            "parent_plot_id": int(plot_id),
            "estate_name": safe_estate_name or None,
            "method": payload["method"],
            "requested_count": payload.get("requested_count"),
            "target_area_m2": payload.get("target_area_m2"),
            "orientation_deg": payload.get("orientation_deg") or 0.0,
            "exclude_road": bool(payload.get("exclude_road")),
            "road_width_m": payload.get("road_width_m"),
            "lot_width_m": payload.get("lot_width_m"),
            "lot_height_m": payload.get("lot_height_m"),
            "dimension_unit": payload.get("dimension_unit"),
            "leftover_area_m2": payload.get("leftover_area_m2"),
            "excluded_area_m2": payload.get("excluded_area_m2"),
            "leftover_geojson": json.dumps(payload.get("leftover_geojson")) if payload.get("leftover_geojson") else None,
            "excluded_geojson": json.dumps(payload.get("excluded_geojson")) if payload.get("excluded_geojson") else None,
        },
    ).scalar()
    if batch_id is None:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to create subdivision batch.")

    created_items: list[dict] = []
    try:
        for row in payload["plots"]:
            raw_geom = shape(row["geometry"])
            if raw_geom.geom_type != "Polygon":
                # A lot combining both sides of an excluded road (see
                # _subdivide_polygon_multi_part_equal_count) is a MultiPolygon in the interactive
                # preview - but every saved plot is a single contiguous parcel at the database
                # level, so silently keeping only the larger piece here (which
                # _clean_single_polygon would otherwise do) would quietly save less land than the
                # preview showed and promised. Fail loudly instead.
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Lot {row['lot_no']} combines area from both sides of the excluded road and can't be "
                        "saved as a single plot. Increase the plot count so every lot fits on one side, or turn "
                        "off 'Exclude access road' for this split."
                    ),
                )
            geom_obj = _clean_single_polygon(raw_geom)
            if geom_obj is None:
                raise HTTPException(status_code=400, detail=f"Invalid generated geometry for lot {row['lot_no']}.")

            child_plot = Plot(geom=from_shape(geom_obj, srid=4326), owner_user_id=child_owner_id)
            db.add(child_plot)
            db.flush()
            child_plot_id = int(child_plot.id)
            lot_no = str(row.get("lot_no") or f"LOT-{child_plot_id}")

            _apply_child_plot_meta(
                db,
                plot_id=child_plot_id,
                title_text=_compose_child_title(parent_meta, lot_no, safe_estate_name),
                parent_meta=parent_meta,
                adamawa_owner_name=lot_no,
            )
            _set_plot_subdivision_meta(
                db,
                child_plot_id=child_plot_id,
                parent_plot_id=int(plot_id),
                batch_id=int(batch_id),
                lot_no=lot_no,
                estate_name=safe_estate_name,
            )
            db.execute(
                text(
                    """
                    INSERT INTO plot_subdivision_items (batch_id, child_plot_id, lot_no, area_m2)
                    VALUES (:batch_id, :child_plot_id, :lot_no, :area_m2)
                    """
                ),
                {
                    "batch_id": int(batch_id),
                    "child_plot_id": child_plot_id,
                    "lot_no": lot_no,
                    "area_m2": float(row.get("area_m2") or 0.0),
                },
            )

            if include_feature_detection:
                _run_plot_feature_detection(db, child_plot_id)

            created_items.append(
                {
                    "child_plot_id": child_plot_id,
                    "lot_no": lot_no,
                    "area_m2": float(row.get("area_m2") or 0.0),
                    "area_hectares": float(row.get("area_hectares") or 0.0),
                }
            )

        db.execute(
            text(
                """
                UPDATE plot_subdivision_batches
                SET generated_count = :generated_count,
                    total_area_m2 = :total_area_m2,
                    status = 'completed',
                    updated_at = NOW()
                WHERE id = :batch_id
                """
            ),
            {
                "batch_id": int(batch_id),
                "generated_count": len(created_items),
                "total_area_m2": float(payload.get("derived_total_area_m2") or 0.0),
            },
        )
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except Exception as exc:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Subdivision apply failed: {exc}")

    return {
        "batch_id": int(batch_id),
        "parent_plot_id": int(plot_id),
        "estate_name": safe_estate_name,
        "method": payload["method"],
        "generated_count": len(created_items),
        "total_area_m2": payload.get("derived_total_area_m2"),
        "plots": created_items,
    }


@router.get("/{plot_id}/subdivision/batches")
def list_plot_subdivision_batches(plot_id: int, db: Session = Depends(get_db)):
    rows = db.execute(
        text(
            """
            SELECT b.id,
                   b.parent_plot_id,
                   b.estate_name,
                   b.method,
                   b.requested_count,
                   b.target_area_m2,
                   b.orientation_deg,
                   b.generated_count,
                   b.total_area_m2,
                   b.status,
                   b.created_at,
                   b.updated_at,
                   COALESCE(COUNT(i.id), 0) AS item_count
            FROM plot_subdivision_batches b
            LEFT JOIN plot_subdivision_items i ON i.batch_id = b.id
            WHERE b.parent_plot_id = :plot_id
            GROUP BY b.id
            ORDER BY b.created_at DESC, b.id DESC
            """
        ),
        {"plot_id": int(plot_id)},
    ).mappings().all()
    return [dict(r) for r in rows]


@router.get("/subdivision/batches/{batch_id}")
def get_plot_subdivision_batch(batch_id: int, include_geojson: bool = Query(False), db: Session = Depends(get_db)):
    batch_row = db.execute(
        text(
            """
            SELECT id, parent_plot_id, estate_name, method, requested_count, target_area_m2,
                   orientation_deg, generated_count, total_area_m2, status, created_at, updated_at
            FROM plot_subdivision_batches
            WHERE id = :batch_id
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().first()
    if not batch_row:
        raise HTTPException(status_code=404, detail="Subdivision batch not found.")

    if include_geojson:
        item_rows = db.execute(
            text(
                """
                SELECT i.id,
                       i.batch_id,
                       i.child_plot_id,
                       i.lot_no,
                       i.area_m2,
                       i.created_at,
                       ST_AsGeoJSON(p.geom) AS geojson
                FROM plot_subdivision_items i
                JOIN plots p ON p.id = i.child_plot_id
                WHERE i.batch_id = :batch_id
                ORDER BY i.id ASC
                """
            ),
            {"batch_id": int(batch_id)},
        ).mappings().all()
    else:
        item_rows = db.execute(
            text(
                """
                SELECT id, batch_id, child_plot_id, lot_no, area_m2, created_at
                FROM plot_subdivision_items
                WHERE batch_id = :batch_id
                ORDER BY id ASC
                """
            ),
            {"batch_id": int(batch_id)},
        ).mappings().all()

    return {
        "batch": dict(batch_row),
        "items": [dict(r) for r in item_rows],
    }


def _get_subdivision_batch_export_context(db: Session, batch_id: int) -> dict[str, Any]:
    batch_row = db.execute(
        text(
            """
            SELECT id, parent_plot_id, estate_name, method, created_at
            FROM plot_subdivision_batches
            WHERE id = :batch_id
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().first()
    if not batch_row:
        raise HTTPException(status_code=404, detail="Subdivision batch not found.")

    items = db.execute(
        text(
            """
            SELECT child_plot_id, lot_no, area_m2
            FROM plot_subdivision_items
            WHERE batch_id = :batch_id
            ORDER BY id ASC
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().all()
    if not items:
        raise HTTPException(status_code=404, detail="Subdivision batch has no generated plots.")

    estate_tag = _safe_filename_fragment(str(batch_row.get("estate_name") or ""), f"batch_{batch_id}")
    zip_name = f"{estate_tag}_survey_plans_batch_{batch_id}.zip"
    cache_dir = os.path.join(REPORTS_DIR, "subdivision_batches")
    os.makedirs(cache_dir, exist_ok=True)
    cached_zip_path = os.path.join(cache_dir, zip_name)
    return {
        "batch": dict(batch_row),
        "items": [dict(item) for item in items],
        "zip_name": zip_name,
        "cached_zip_path": cached_zip_path,
    }


def _generate_subdivision_batch_zip(
    db: Session,
    *,
    batch_id: int,
    items: list[dict[str, Any]],
    zip_name: str,
    cached_zip_path: str,
    require_cached_copy: bool = False,
) -> tuple[str, str | None]:
    if os.path.isfile(cached_zip_path):
        return cached_zip_path, None

    tmp_dir = tempfile.mkdtemp(prefix=f"subdivision_batch_{batch_id}_")
    # No "sep=," Excel hint line here - GIS/DGPS CSV readers treat it as a malformed data row
    # (wrong column count) rather than the delimiter directive Excel understands it as.
    export_rows: list[list[str]] = [["lot_no", "child_plot_id", "area_m2"]]
    setting_out_rows: list[list[str]] = [[
        "lot_no",
        "child_plot_id",
        "point_index",
        "station",
        "longitude",
        "latitude",
        "easting",
        "northing",
        "utm_epsg",
    ]]
    pdf_files: list[str] = []

    try:
        for item in items:
            child_plot_id = int(item["child_plot_id"])
            lot_no = str(item.get("lot_no") or f"LOT-{child_plot_id}")
            safe_lot = _safe_filename_fragment(lot_no, f"LOT_{child_plot_id}")
            pdf_name = f"{safe_lot}_survey_plan.pdf"
            pdf_path = os.path.join(tmp_dir, pdf_name)
            _render_survey_plan_pdf_for_plot(db, child_plot_id, pdf_path)
            pdf_files.append(pdf_path)
            export_rows.append([lot_no, str(child_plot_id), f"{float(item.get('area_m2') or 0.0):.2f}"])

            try:
                geom_geojson_raw = db.execute(
                    text("SELECT ST_AsGeoJSON(geom) FROM plots WHERE id = :id"),
                    {"id": child_plot_id},
                ).scalar()
                if geom_geojson_raw:
                    geom_obj = _clean_single_polygon(shape(json.loads(geom_geojson_raw)))
                    if geom_obj is not None:
                        coords_wgs = list(geom_obj.exterior.coords)
                        if len(coords_wgs) >= 2 and coords_wgs[0] == coords_wgs[-1]:
                            coords_wgs = coords_wgs[:-1]
                        utm_epsg = _metric_epsg_for_wgs84_polygon(geom_obj)
                        metric_poly = (
                            gpd.GeoDataFrame(geometry=[geom_obj], crs="EPSG:4326")
                            .to_crs(epsg=utm_epsg)
                            .geometry.iloc[0]
                        )
                        coords_metric = list(metric_poly.exterior.coords)
                        if len(coords_metric) >= 2 and coords_metric[0] == coords_metric[-1]:
                            coords_metric = coords_metric[:-1]
                        for idx, (wgs_pt, metric_pt) in enumerate(zip(coords_wgs, coords_metric), start=1):
                            lng_val = float(wgs_pt[0])
                            lat_val = float(wgs_pt[1])
                            easting_val = float(metric_pt[0])
                            northing_val = float(metric_pt[1])
                            setting_out_rows.append([
                                lot_no,
                                str(child_plot_id),
                                str(idx),
                                _station_name(idx - 1),
                                _format_coordinate_number(lng_val, 8),
                                _format_coordinate_number(lat_val, 8),
                                _format_coordinate_number(easting_val, 3),
                                _format_coordinate_number(northing_val, 3),
                                str(int(utm_epsg)),
                            ])
            except Exception:
                pass

        manifest_path = os.path.join(tmp_dir, "batch_manifest.csv")
        with open(manifest_path, "w", encoding="utf-8-sig", newline="") as f:
            # "sep=," as the literal first line is a Microsoft-documented Excel hint that forces
            # comma-delimited parsing on double-click, regardless of the machine's regional list
            # separator (many Windows locales default to semicolon) - safe here since this file is
            # a human-facing summary, not a GIS/DGPS ingestion target.
            f.write("sep=,\n")
            writer = csv.writer(
                f,
                delimiter=",",
                quotechar='"',
                quoting=csv.QUOTE_MINIMAL,
                lineterminator="\n",
            )
            writer.writerows(export_rows)

        setting_out_path = os.path.join(tmp_dir, "setting_out_points_dgps.csv")
        with open(setting_out_path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(
                f,
                delimiter=",",
                quotechar='"',
                quoting=csv.QUOTE_MINIMAL,
                lineterminator="\n",
            )
            writer.writerows(setting_out_rows)

        # Excel-friendly companion of the same setting-out data, with the "sep=," hint. Kept as a
        # separate file rather than added to setting_out_points_dgps.csv above, because QGIS/Civil3D/
        # DGPS-receiver CSV importers treat a "sep=," first line as a malformed data row (wrong
        # column count), not the delimiter directive Excel understands it as - the DGPS file has to
        # stay strict for those tools, so Excel users get their own copy instead.
        setting_out_excel_path = os.path.join(tmp_dir, "setting_out_points.csv")
        with open(setting_out_excel_path, "w", encoding="utf-8-sig", newline="") as f:
            f.write("sep=,\n")
            writer = csv.writer(
                f,
                delimiter=",",
                quotechar='"',
                quoting=csv.QUOTE_MINIMAL,
                lineterminator="\n",
            )
            writer.writerows(setting_out_rows)

        zip_path = os.path.join(tmp_dir, zip_name)
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(manifest_path, arcname="batch_manifest.csv")
            zf.write(setting_out_path, arcname="setting_out_points_dgps.csv")
            zf.write(setting_out_excel_path, arcname="setting_out_points.csv")
            for fp in pdf_files:
                if os.path.isfile(fp):
                    zf.write(fp, arcname=os.path.basename(fp))

        try:
            shutil.copyfile(zip_path, cached_zip_path)
            safe_rmtree(tmp_dir)
            return cached_zip_path, None
        except Exception as exc:
            if require_cached_copy:
                safe_rmtree(tmp_dir)
                raise HTTPException(status_code=500, detail=f"Failed to persist subdivision batch export: {exc}") from exc
            return zip_path, tmp_dir
    except HTTPException:
        safe_rmtree(tmp_dir)
        raise
    except Exception as exc:
        safe_rmtree(tmp_dir)
        raise HTTPException(status_code=500, detail=f"Failed to export subdivision batch: {exc}") from exc


def _get_plot_export_job(db: Session, job_id: str) -> dict[str, Any] | None:
    row = db.execute(
        text(
            """
            SELECT
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                cache_key,
                status,
                file_name,
                local_path,
                request_payload,
                error_text,
                started_at,
                completed_at,
                created_at,
                updated_at
            FROM plot_export_jobs
            WHERE id = :job_id
            LIMIT 1
            """
        ),
        {"job_id": str(job_id)},
    ).mappings().first()
    return dict(row) if row else None


def _set_plot_export_job_status(
    db: Session,
    job_id: str,
    *,
    status: str,
    local_path: str | None = None,
    file_name: str | None = None,
    error_text: str | None = None,
    started: bool = False,
    completed: bool = False,
):
    if status not in PLOT_EXPORT_JOB_STATUS_VALUES:
        raise ValueError(f"Unsupported plot export job status: {status}")
    updates = ["status = :status", "updated_at = NOW()"]
    params: dict[str, Any] = {"job_id": str(job_id), "status": status}
    if started:
        updates.append("started_at = COALESCE(started_at, NOW())")
    if completed:
        updates.append("completed_at = NOW()")
    if local_path is not None:
        updates.append("local_path = :local_path")
        params["local_path"] = local_path
    if file_name is not None:
        updates.append("file_name = :file_name")
        params["file_name"] = file_name
    if error_text is not None:
        updates.append("error_text = :error_text")
        params["error_text"] = error_text
    db.execute(text(f"UPDATE plot_export_jobs SET {', '.join(updates)} WHERE id = :job_id"), params)
    db.commit()


def sweep_stale_plot_export_jobs(db: Session, *, stale_after_minutes: int = 10) -> int:
    """Marks any export job stuck in 'running' for longer than stale_after_minutes as failed.

    Background export jobs run in daemon threads with no server-side timeout of their own - if
    one hangs (e.g. inside a stalled basemap tile fetch), it would otherwise sit in
    status='running' forever with no way for the client to know it's dead. Python can't safely
    force-cancel a running thread, so this doesn't stop the runaway thread itself, but it does
    free the job's cache-key slot (_find_plot_export_job_by_cache_key already excludes 'failed'
    jobs) so the next request for the same export starts a clean new attempt instead of polling
    a job that will never complete.
    """
    result = db.execute(
        text(
            """
            UPDATE plot_export_jobs
            SET status = 'failed',
                error_text = 'Export timed out',
                completed_at = NOW(),
                updated_at = NOW()
            WHERE status = 'running'
              AND started_at IS NOT NULL
              AND started_at < NOW() - make_interval(mins => :stale_after_minutes)
            RETURNING id
            """
        ),
        {"stale_after_minutes": int(stale_after_minutes)},
    )
    stale_ids = [row[0] for row in result.fetchall()]
    db.commit()
    return len(stale_ids)


def _build_plot_export_public_base_url(request: Request | None = None) -> str:
    # request.url_for(...) trusts request.url.scheme, which resolves to "http" behind a reverse
    # proxy/tunnel (Cloudflare) that terminates TLS before forwarding to this app - producing a
    # download_url like "http://api.landcheck.online/..." that browsers block as mixed content on
    # an https:// page. Prefer an explicit public URL env var, then the X-Forwarded-Proto header
    # the proxy actually sets, matching the same pattern already used for other public URLs in
    # this codebase (see _build_public_api_base_url in routers/green.py).
    configured = str(os.getenv("LANDCHECK_API_PUBLIC_URL") or "").strip().rstrip("/")
    if configured:
        return configured
    if request is not None:
        forwarded_proto = str(request.headers.get("x-forwarded-proto") or request.url.scheme or "https").split(",")[0].strip()
        forwarded_host = str(request.headers.get("x-forwarded-host") or request.headers.get("host") or request.url.netloc).split(",")[0].strip()
        if forwarded_proto and forwarded_host:
            return f"{forwarded_proto}://{forwarded_host}".rstrip("/")
        return str(request.base_url).rstrip("/")
    return ""


def _serialize_plot_export_job(job: dict[str, Any], request: Request | None = None) -> dict[str, Any]:
    item = dict(job)
    download_url: str | None = None
    if str(item.get("status") or "").strip().lower() == "completed":
        job_id = str(item.get("id") or "").strip()
        base_url = _build_plot_export_public_base_url(request)
        download_url = f"{base_url}/plots/export-jobs/{job_id}/download" if base_url else f"/plots/export-jobs/{job_id}/download"
    item["download_url"] = download_url
    return item


def _normalize_plot_export_job_payload(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return dict(value)
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            return {}
    return {}


def _build_plot_export_cache_key(
    db: Session,
    *,
    plot_id: int,
    export_type: str,
    payload: dict[str, Any] | None = None,
) -> str:
    normalized_payload = dict(payload or {})
    revision_token = (
        build_preview_revision_token(db, plot_id)
        if str(export_type or "").strip().lower() == "survey-plan.pdf"
        else build_plot_geom_revision_token(db, plot_id)
    )
    signature_payload = {
        "schema_version": "plot_export_job_v1",
        "plot_id": int(plot_id),
        "export_type": str(export_type or "").strip().lower(),
        "render_options": normalized_payload,
        "revision": revision_token,
    }
    signature_text = json.dumps(signature_payload, sort_keys=True, separators=(",", ":"), default=str)
    return hashlib.sha256(signature_text.encode("utf-8")).hexdigest()


def _find_plot_export_job_by_cache_key(
    db: Session,
    *,
    plot_id: int,
    export_type: str,
    cache_key: str,
) -> dict[str, Any] | None:
    row = db.execute(
        text(
            """
            SELECT
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                cache_key,
                status,
                file_name,
                local_path,
                request_payload,
                error_text,
                started_at,
                completed_at,
                created_at,
                updated_at
            FROM plot_export_jobs
            WHERE plot_id = :plot_id
              AND export_type = :export_type
              AND cache_key = :cache_key
              AND status != 'failed'
            ORDER BY
                CASE status
                    WHEN 'completed' THEN 0
                    WHEN 'running' THEN 1
                    WHEN 'queued' THEN 2
                    ELSE 3
                END,
                created_at DESC
            LIMIT 1
            """
        ),
        {
            "plot_id": int(plot_id),
            "export_type": str(export_type or "").strip(),
            "cache_key": str(cache_key or "").strip(),
        },
    ).mappings().first()
    return dict(row) if row else None


def _get_subdivision_batch_clean_copy_context(
    db: Session,
    *,
    batch_id: int,
    title_text: str = "",
    area_labels: list[dict] | None = None,
    paper_size: str | None = None,
    scale_text: str | None = None,
    coordinate_system: str | None = None,
    station_names: list[str] | None = None,
    north_arrow_style: str = "one_side_stem",
    north_arrow_color: str = "blue",
    beacon_style: str = "cross",
    road_width_m: float | None = None,
    boundary_color: str | None = None,
    grid_color: str | None = None,
    text_color: str | None = None,
    road_color: str | None = None,
    river_color: str | None = None,
    building_color: str | None = None,
    building_hatch_type: str | None = None,
    road_style: str | None = None,
    station_font: str | None = None,
    station_size: int | None = None,
    bearing_font: str | None = None,
    bearing_size: int | None = None,
) -> dict[str, Any]:
    batch_row = db.execute(
        text(
            """
            SELECT id, parent_plot_id, estate_name, method, created_at
            FROM plot_subdivision_batches
            WHERE id = :batch_id
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().first()
    if not batch_row:
        raise HTTPException(status_code=404, detail="Subdivision batch not found.")

    child_rows = db.execute(
        text(
            """
            SELECT
                i.child_plot_id,
                i.lot_no,
                i.area_m2,
                ST_AsGeoJSON(p.geom) AS geom_geojson
            FROM plot_subdivision_items i
            JOIN plots p ON p.id = i.child_plot_id
            WHERE i.batch_id = :batch_id
            ORDER BY i.id ASC
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().all()
    if not child_rows:
        raise HTTPException(status_code=404, detail="Subdivision batch has no generated plots.")

    parent_plot_id = int(batch_row.get("parent_plot_id") or 0)
    if parent_plot_id <= 0:
        raise HTTPException(status_code=400, detail="Subdivision batch parent parcel is invalid.")

    parent_meta = get_plot_meta(db, parent_plot_id)
    effective_paper_size = str(paper_size or parent_meta.get("paper_size") or "A4").upper()
    if effective_paper_size not in {"A4", "A3", "A2", "A1", "A0"}:
        effective_paper_size = "A4"
    effective_scale_text = str(scale_text or parent_meta.get("scale_text") or "1 : 1000")
    effective_coordinate_system = str(coordinate_system or parent_meta.get("coordinate_system") or "wgs84")
    parent_plot_geom_wgs84 = _load_plot_polygon_wgs84(db, parent_plot_id)
    effective_render_coordinate_system, effective_epsg, _ = _resolve_survey_render_crs(
        effective_coordinate_system,
        parent_plot_geom_wgs84,
    )

    clean_title = str(title_text or "").strip()
    if not clean_title:
        estate_name = str(batch_row.get("estate_name") or "").strip()
        clean_title = f"{estate_name} CLEAN COPY PLAN" if estate_name else "SURVEY PLAN"

    area_override_map = _extract_clean_copy_area_overrides(area_labels)
    cache_key_payload = {
        "render_version": CLEAN_COPY_RENDER_VERSION,
        "batch_id": int(batch_id),
        "title_text": clean_title,
        "paper_size": effective_paper_size,
        "scale_text": effective_scale_text,
        "coordinate_system": effective_coordinate_system,
        "station_names": list(station_names or []),
        "north_arrow_style": str(north_arrow_style or "one_side_stem"),
        "north_arrow_color": str(north_arrow_color or "blue"),
        "beacon_style": str(beacon_style or "cross"),
        "road_width_m": float(road_width_m or 0.0),
        "boundary_color": boundary_color or "",
        "grid_color": grid_color or "",
        "text_color": text_color or "",
        "road_color": road_color or "",
        "river_color": river_color or "",
        "building_color": building_color or "",
        "building_hatch_type": building_hatch_type or "",
        "road_style": road_style or "",
        "station_font": station_font or "",
        "station_size": station_size or 0,
        "bearing_font": bearing_font or "",
        "bearing_size": bearing_size or 0,
        "area_overrides": sorted(area_override_map.items()),
    }
    cache_hash = hashlib.sha1(json.dumps(cache_key_payload, sort_keys=True).encode("utf-8")).hexdigest()[:12]
    estate_tag = _safe_filename_fragment(str(batch_row.get("estate_name") or ""), f"batch_{batch_id}")
    pdf_name = f"{estate_tag}_clean_copy_batch_{batch_id}_{cache_hash}.pdf"
    cache_dir = os.path.join(REPORTS_DIR, "subdivision_clean_copy")
    os.makedirs(cache_dir, exist_ok=True)
    cached_pdf_path = os.path.join(cache_dir, pdf_name)
    return {
        "batch": dict(batch_row),
        "child_rows": [dict(row) for row in child_rows],
        "parent_plot_id": parent_plot_id,
        "pdf_name": pdf_name,
        "cached_pdf_path": cached_pdf_path,
        "title_text": clean_title,
        "paper_size": effective_paper_size,
        "scale_text": effective_scale_text,
        "coordinate_system": effective_render_coordinate_system,
        "epsg_code": int(effective_epsg),
        "station_names": list(station_names or []),
        "north_arrow_style": str(north_arrow_style or "one_side_stem"),
        "north_arrow_color": str(north_arrow_color or "blue"),
        "beacon_style": str(beacon_style or "cross"),
        "road_width_m": road_width_m,
        "boundary_color": boundary_color,
        "grid_color": grid_color,
        "text_color": text_color,
        "road_color": road_color,
        "river_color": river_color,
        "building_color": building_color,
        "building_hatch_type": building_hatch_type,
        "road_style": road_style,
        "station_font": station_font,
        "station_size": station_size,
        "bearing_font": bearing_font,
        "bearing_size": bearing_size,
        "area_overrides": area_override_map,
    }


def _generate_subdivision_batch_clean_copy_pdf(
    db: Session,
    *,
    context: dict[str, Any],
) -> str:
    cached_pdf_path = str(context.get("cached_pdf_path") or "").strip()
    if not cached_pdf_path:
        raise HTTPException(status_code=500, detail="Clean copy export cache path is missing.")
    if os.path.isfile(cached_pdf_path):
        return cached_pdf_path

    parent_plot_id = int(context.get("parent_plot_id") or 0)
    parent_poly_wgs84 = _load_plot_polygon_wgs84(db, parent_plot_id)
    try:
        _render_subdivision_clean_copy_pdf(
            db=db,
            parent_plot_id=parent_plot_id,
            parent_poly_wgs84=parent_poly_wgs84,
            child_rows=list(context.get("child_rows") or []),
            output_pdf_path=cached_pdf_path,
            title_text=str(context.get("title_text") or ""),
            paper_size=str(context.get("paper_size") or "A4"),
            scale_text=str(context.get("scale_text") or "1 : 1000"),
            coordinate_system=str(context.get("coordinate_system") or "wgs84"),
            epsg_code=int(context.get("epsg_code") or 4326),
            station_names=list(context.get("station_names") or []),
            north_arrow_style=str(context.get("north_arrow_style") or "one_side_stem"),
            north_arrow_color=str(context.get("north_arrow_color") or "blue"),
            beacon_style=str(context.get("beacon_style") or "cross"),
            road_width_m=context.get("road_width_m"),
            boundary_color=context.get("boundary_color") or None,
            grid_color=context.get("grid_color") or None,
            text_color=context.get("text_color") or None,
            road_color=context.get("road_color") or None,
            river_color=context.get("river_color") or None,
            building_color=context.get("building_color") or None,
            building_hatch_type=context.get("building_hatch_type") or None,
            road_style=context.get("road_style") or None,
            station_font=context.get("station_font") or None,
            station_size=context.get("station_size") or None,
            bearing_font=context.get("bearing_font") or None,
            bearing_size=context.get("bearing_size") or None,
            area_overrides=dict(context.get("area_overrides") or {}),
        )
    except HTTPException:
        raise
    except Exception as exc:
        safe_remove(cached_pdf_path)
        raise HTTPException(status_code=500, detail=f"Failed to export clean copy PDF: {exc}") from exc

    return cached_pdf_path


def _run_subdivision_batch_export_job(job_id: str):
    db = SessionLocal()
    try:
        job = _get_plot_export_job(db, job_id)
        if not job:
            return
        batch_id = int(job.get("subdivision_batch_id") or 0)
        if batch_id <= 0:
            raise RuntimeError("Subdivision batch export job is missing a batch id.")
        _set_plot_export_job_status(db, job_id, status="running", started=True, error_text="")
        context = _get_subdivision_batch_export_context(db, batch_id)
        export_path, _cleanup_dir = _generate_subdivision_batch_zip(
            db,
            batch_id=batch_id,
            items=context["items"],
            zip_name=context["zip_name"],
            cached_zip_path=context["cached_zip_path"],
            require_cached_copy=True,
        )
        _set_plot_export_job_status(
            db,
            job_id,
            status="completed",
            local_path=export_path,
            file_name=context["zip_name"],
            error_text="",
            completed=True,
        )
    except Exception as exc:
        try:
            _set_plot_export_job_status(db, job_id, status="failed", error_text=str(exc), completed=True)
        except Exception:
            db.rollback()
    finally:
        db.close()


def _run_subdivision_batch_clean_copy_export_job(job_id: str):
    db = SessionLocal()
    try:
        job = _get_plot_export_job(db, job_id)
        if not job:
            return
        batch_id = int(job.get("subdivision_batch_id") or 0)
        if batch_id <= 0:
            raise RuntimeError("Subdivision clean copy export job is missing a batch id.")
        payload = _normalize_plot_export_job_payload(job.get("request_payload"))
        _set_plot_export_job_status(db, job_id, status="running", started=True, error_text="")
        context = _get_subdivision_batch_clean_copy_context(
            db,
            batch_id=batch_id,
            title_text=str(payload.get("title_text") or ""),
            area_labels=payload.get("area_labels") if isinstance(payload.get("area_labels"), list) else None,
            paper_size=str(payload.get("paper_size") or "") or None,
            scale_text=str(payload.get("scale_text") or "") or None,
            coordinate_system=str(payload.get("coordinate_system") or "") or None,
            station_names=payload.get("station_names") if isinstance(payload.get("station_names"), list) else None,
            north_arrow_style=str(payload.get("north_arrow_style") or "one_side_stem"),
            north_arrow_color=str(payload.get("north_arrow_color") or "blue"),
            beacon_style=str(payload.get("beacon_style") or "cross"),
            road_width_m=float(payload.get("road_width_m") or 0.0) if payload.get("road_width_m") is not None else None,
            boundary_color=payload.get("boundary_color") or None,
            grid_color=payload.get("grid_color") or None,
            text_color=payload.get("text_color") or None,
            road_color=payload.get("road_color") or None,
            river_color=payload.get("river_color") or None,
            building_color=payload.get("building_color") or None,
            building_hatch_type=payload.get("building_hatch_type") or None,
            road_style=payload.get("road_style") or None,
            station_font=payload.get("station_font") or None,
            station_size=payload.get("station_size") or None,
            bearing_font=payload.get("bearing_font") or None,
            bearing_size=payload.get("bearing_size") or None,
        )
        export_path = _generate_subdivision_batch_clean_copy_pdf(db, context=context)
        _set_plot_export_job_status(
            db,
            job_id,
            status="completed",
            local_path=export_path,
            file_name=str(context.get("pdf_name") or "subdivision_clean_copy.pdf"),
            error_text="",
            completed=True,
        )
    except Exception as exc:
        try:
            _set_plot_export_job_status(db, job_id, status="failed", error_text=str(exc), completed=True)
        except Exception:
            db.rollback()
    finally:
        db.close()


@router.get("/subdivision/batches/{batch_id}/export/survey-plans.zip")
def export_subdivision_batch_survey_plans(
    batch_id: int,
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = None,
):
    context = _get_subdivision_batch_export_context(db, batch_id)
    served_path, cleanup_dir = _generate_subdivision_batch_zip(
        db,
        batch_id=batch_id,
        items=context["items"],
        zip_name=context["zip_name"],
        cached_zip_path=context["cached_zip_path"],
        require_cached_copy=False,
    )
    served_background = background_tasks
    if cleanup_dir:
        if served_background is None:
            served_background = BackgroundTasks()
        served_background.add_task(safe_rmtree, cleanup_dir)
    return FileResponse(
        served_path,
        media_type="application/zip",
        filename=context["zip_name"],
        background=served_background,
    )


@router.post("/subdivision/batches/{batch_id}/export-jobs/survey-plans")
def create_subdivision_batch_export_job(
    batch_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    context = _get_subdivision_batch_export_context(db, batch_id)
    existing = db.execute(
        text(
            """
            SELECT
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                status,
                file_name,
                local_path,
                error_text,
                started_at,
                completed_at,
                created_at,
                updated_at
            FROM plot_export_jobs
            WHERE export_type = 'subdivision_batch_survey_plans'
              AND subdivision_batch_id = :batch_id
              AND status IN ('queued', 'running')
            ORDER BY created_at DESC
            LIMIT 1
            """
        ),
        {"batch_id": int(batch_id)},
    ).mappings().first()
    if existing:
        return _serialize_plot_export_job(dict(existing), request=request)

    job_id = uuid.uuid4().hex
    if os.path.isfile(context["cached_zip_path"]):
        db.execute(
            text(
                """
                INSERT INTO plot_export_jobs (
                    id,
                    export_type,
                    plot_id,
                    subdivision_batch_id,
                    status,
                    file_name,
                    local_path,
                    started_at,
                    completed_at,
                    created_at,
                    updated_at
                )
                VALUES (
                    :id,
                    'subdivision_batch_survey_plans',
                    :plot_id,
                    :batch_id,
                    'completed',
                    :file_name,
                    :local_path,
                    NOW(),
                    NOW(),
                    NOW(),
                    NOW()
                )
                """
            ),
            {
                "id": job_id,
                "plot_id": int(context["batch"].get("parent_plot_id") or 0) or None,
                "batch_id": int(batch_id),
                "file_name": context["zip_name"],
                "local_path": context["cached_zip_path"],
            },
        )
        db.commit()
        job = _get_plot_export_job(db, job_id)
        return _serialize_plot_export_job(job or {"id": job_id, "status": "completed"}, request=request)

    db.execute(
        text(
            """
            INSERT INTO plot_export_jobs (
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                status,
                file_name,
                created_at,
                updated_at
            )
            VALUES (
                :id,
                'subdivision_batch_survey_plans',
                :plot_id,
                :batch_id,
                'queued',
                :file_name,
                NOW(),
                NOW()
            )
            """
        ),
        {
            "id": job_id,
            "plot_id": int(context["batch"].get("parent_plot_id") or 0) or None,
            "batch_id": int(batch_id),
            "file_name": context["zip_name"],
        },
    )
    db.commit()
    threading.Thread(target=_run_subdivision_batch_export_job, args=(job_id,), daemon=True).start()
    job = _get_plot_export_job(db, job_id)
    return _serialize_plot_export_job(job or {"id": job_id, "status": "queued"}, request=request)


@router.get("/export-jobs/{job_id}")
def get_plot_export_job_status(
    job_id: str,
    request: Request,
    db: Session = Depends(get_db),
):
    job = _get_plot_export_job(db, job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Plot export job not found.")
    return _serialize_plot_export_job(job, request=request)


@router.get("/export-jobs/{job_id}/download", name="download_plot_export_job")
def download_plot_export_job(
    job_id: str,
    request: Request,
    db: Session = Depends(get_db),
):
    job = _get_plot_export_job(db, job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Plot export job not found.")
    if str(job.get("status") or "").strip().lower() != "completed":
        raise HTTPException(status_code=409, detail="Plot export job is not finished yet.")

    local_path = str(job.get("local_path") or "").strip()
    file_name = str(job.get("file_name") or "").strip() or "plot_export.zip"
    if (not local_path or not os.path.isfile(local_path)) and int(job.get("subdivision_batch_id") or 0) > 0:
        context = _get_subdivision_batch_export_context(db, int(job.get("subdivision_batch_id") or 0))
        if os.path.isfile(context["cached_zip_path"]):
            local_path = context["cached_zip_path"]
            file_name = context["zip_name"]
            _set_plot_export_job_status(
                db,
                job_id,
                status="completed",
                local_path=local_path,
                file_name=file_name,
            )
    if not local_path or not os.path.isfile(local_path):
        raise HTTPException(status_code=404, detail="Export file is no longer available.")
    log_survey_activity(
        db,
        event_type="export_downloaded",
        workflow="subdivision" if int(job.get("subdivision_batch_id") or 0) else "survey_plan",
        request=request,
        plot_id=int(job["plot_id"]) if job.get("plot_id") else None,
        subdivision_batch_id=int(job["subdivision_batch_id"]) if job.get("subdivision_batch_id") else None,
        details={"export_type": job.get("export_type"), "file_name": file_name, "job_id": job_id},
    )
    media_type = mimetypes.guess_type(file_name)[0] or "application/octet-stream"
    return FileResponse(
        local_path,
        media_type=media_type,
        filename=file_name,
    )


def _cache_single_plot_export_file(
    *,
    plot_id: int,
    export_type: str,
    cache_key: str,
    source_path: str,
    file_name: str,
) -> str:
    cache_dir = os.path.join(REPORTS_DIR, "export_jobs", str(int(plot_id)))
    os.makedirs(cache_dir, exist_ok=True)
    safe_export_type = re.sub(r"[^a-z0-9]+", "_", str(export_type or "").strip().lower()).strip("_") or "export"
    ext = os.path.splitext(str(file_name or source_path))[1] or os.path.splitext(source_path)[1] or ".bin"
    cache_path = os.path.join(cache_dir, f"{safe_export_type}_{str(cache_key or '')[:12]}{ext}")
    if os.path.abspath(source_path) != os.path.abspath(cache_path):
        shutil.copyfile(source_path, cache_path)
    return cache_path


def _run_single_plot_export_job(job_id: str):
    db = SessionLocal()
    source_path = ""
    cleanup_dir = ""
    try:
        job = _get_plot_export_job(db, job_id)
        if not job:
            return
        plot_id = int(job.get("plot_id") or 0)
        if plot_id <= 0:
            raise RuntimeError("Plot export job is missing a plot id.")
        export_type = str(job.get("export_type") or "").strip().lower()
        cache_key = str(job.get("cache_key") or "").strip()
        payload = _normalize_plot_export_job_payload(job.get("request_payload"))
        _set_plot_export_job_status(db, job_id, status="running", started=True, error_text="")
        response = None
        if export_type == "survey-plan.pdf":
            response = download_plot_report_pdf(plot_id=plot_id, db=db, background_tasks=None, **payload)
        elif export_type == "orthophoto.pdf":
            response = orthophoto_pdf(plot_id=plot_id, db=db, background_tasks=None, **payload)
        elif export_type == "topomap.pdf":
            response = orthophoto_pdf(
                plot_id=plot_id,
                db=db,
                background_tasks=None,
                **{
                    **payload,
                    "use_topo_map": True,
                    "title_text": str(payload.get("title_text") or "TOPO MAP"),
                },
            )
        elif export_type == "survey-plan.dxf":
            response = download_survey_plan_dwg(plot_id=plot_id, db=db)
        elif export_type == "survey-plan.shapefile":
            response = download_survey_plan_shapefile(plot_id=plot_id, db=db, background_tasks=None)
            source_path = str(getattr(response, "path", "") or "").strip()
            cleanup_dir = os.path.dirname(source_path) if source_path else ""
        elif export_type == "technical-report.docx":
            response = download_plot_technical_report_docx(plot_id=plot_id, db=db, **payload)
        else:
            raise RuntimeError(f"Unsupported plot export job type: {export_type}")

        if not source_path:
            source_path = str(getattr(response, "path", "") or "").strip()
        if not source_path or not os.path.isfile(source_path):
            raise RuntimeError("Export generated without a readable output file.")
        file_name = str(getattr(response, "filename", "") or "").strip() or str(job.get("file_name") or "").strip() or os.path.basename(source_path)
        local_path = source_path
        if export_type == "survey-plan.shapefile":
            local_path = _cache_single_plot_export_file(
                plot_id=plot_id,
                export_type=export_type,
                cache_key=cache_key,
                source_path=source_path,
                file_name=file_name,
            )
        _set_plot_export_job_status(
            db,
            job_id,
            status="completed",
            local_path=local_path,
            file_name=file_name,
            error_text="",
            completed=True,
        )
    except Exception as exc:
        try:
            _set_plot_export_job_status(db, job_id, status="failed", error_text=str(exc), completed=True)
        except Exception:
            db.rollback()
    finally:
        if cleanup_dir:
            safe_rmtree(cleanup_dir)
        db.close()


def _insert_single_plot_export_job(
    db: Session,
    *,
    plot_id: int,
    export_type: str,
    cache_key: str,
    file_name: str,
    request_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    job_id = uuid.uuid4().hex
    db.execute(
        text(
            """
            INSERT INTO plot_export_jobs (
                id,
                export_type,
                plot_id,
                cache_key,
                status,
                file_name,
                request_payload,
                created_at,
                updated_at
            )
            VALUES (
                :id,
                :export_type,
                :plot_id,
                :cache_key,
                'queued',
                :file_name,
                CAST(:request_payload AS JSONB),
                NOW(),
                NOW()
            )
            """
        ),
        {
            "id": job_id,
            "export_type": export_type,
            "plot_id": int(plot_id),
            "cache_key": cache_key,
            "file_name": file_name,
            "request_payload": json.dumps(request_payload or {}),
        },
    )
    db.commit()
    threading.Thread(target=_run_single_plot_export_job, args=(job_id,), daemon=True).start()
    return _get_plot_export_job(db, job_id) or {"id": job_id, "status": "queued"}


@router.post("/{plot_id}/export-jobs/survey-plan.pdf")
def create_plot_survey_report_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    title_text: str = Body("SURVEY PLAN"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    certification_statement: str = Body(DEFAULT_CERTIFICATION_STATEMENT),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
    beacon_style: str = Body("cross"),
    road_width_m: float | None = Body(None),
    road_width_override_m: float | None = Body(None),
    boundary_color: str | None = Body(None),
    grid_color: str | None = Body(None),
    text_color: str | None = Body(None),
    road_color: str | None = Body(None),
    river_color: str | None = Body(None),
    building_color: str | None = Body(None),
    building_hatch_type: str | None = Body(None),
    road_style: str | None = Body(None),
    title_font: str | None = Body(None),
    title_size: int | None = Body(None),
    grid_font: str | None = Body(None),
    grid_size: int | None = Body(None),
    station_font: str | None = Body(None),
    station_size: int | None = Body(None),
    bearing_font: str | None = Body(None),
    bearing_size: int | None = Body(None),
    area_font: str | None = Body(None),
    area_size: int | None = Body(None),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    adamawa_rof_no: str = Body(""),
    adamawa_owner_name: str = Body(""),
    adamawa_authority_title: str = Body(DEFAULT_ADAMAWA_AUTHORITY_TITLE),
    adamawa_authority_date_text: str = Body(DEFAULT_ADAMAWA_AUTHORITY_DATE),
    adamawa_control_point_name: str = Body(""),
    adamawa_northing: str = Body(""),
    adamawa_easting: str = Body(""),
    adamawa_elevation: str = Body(""),
    adamawa_origin_text: str = Body(DEFAULT_ADAMAWA_ORIGIN_TEXT),
    adamawa_topo_sheet_text: str = Body(DEFAULT_ADAMAWA_TOPO_SHEET_TEXT),
    adamawa_computation_no: str = Body(""),
    adamawa_cadastral_sheet_no: str = Body(""),
    adamawa_plan_no: str = Body(""),
    adamawa_surveyed_by_text: str = Body(""),
    adamawa_disclaimer_text: str = Body(DEFAULT_ADAMAWA_DISCLAIMER_TEXT),
    cadastral_plan_no: str = Body(""),
    cadastral_area_name: str = Body(""),
    cadastral_datum_text: str = Body(""),
    cadastral_firm_block_text: str = Body(""),
    fct_file_no: str = Body(""),
    fct_district: str = Body(""),
    fct_cadastral_zone: str = Body(""),
    fct_origin_beacon_text: str = Body(""),
    fct_cadastral_map_ref: str = Body(""),
    fct_title_prefix: str = Body(""),
    survey_input_coordinates: Optional[list] = Body(default=None),
):
    request_payload = {
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "scale_text": scale_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "certification_statement": certification_statement,
        "station_names": station_names or [],
        "coordinate_system": coordinate_system,
        "paper_size": paper_size,
        "north_arrow_style": north_arrow_style,
        "north_arrow_color": north_arrow_color,
        "beacon_style": beacon_style,
        "road_width_m": road_width_m,
        "road_width_override_m": road_width_override_m,
        "boundary_color": boundary_color,
        "grid_color": grid_color,
        "text_color": text_color,
        "road_color": road_color,
        "river_color": river_color,
        "building_color": building_color,
        "building_hatch_type": building_hatch_type,
        "road_style": road_style,
        "title_font": title_font,
        "title_size": title_size,
        "grid_font": grid_font,
        "grid_size": grid_size,
        "station_font": station_font,
        "station_size": station_size,
        "bearing_font": bearing_font,
        "bearing_size": bearing_size,
        "area_font": area_font,
        "area_size": area_size,
        "template_name": template_name,
        "adamawa_rof_no": adamawa_rof_no,
        "adamawa_owner_name": adamawa_owner_name,
        "adamawa_authority_title": adamawa_authority_title,
        "adamawa_authority_date_text": adamawa_authority_date_text,
        "adamawa_control_point_name": adamawa_control_point_name,
        "adamawa_northing": adamawa_northing,
        "adamawa_easting": adamawa_easting,
        "adamawa_elevation": adamawa_elevation,
        "adamawa_origin_text": adamawa_origin_text,
        "adamawa_topo_sheet_text": adamawa_topo_sheet_text,
        "adamawa_computation_no": adamawa_computation_no,
        "adamawa_cadastral_sheet_no": adamawa_cadastral_sheet_no,
        "adamawa_plan_no": adamawa_plan_no,
        "adamawa_surveyed_by_text": adamawa_surveyed_by_text,
        "adamawa_disclaimer_text": adamawa_disclaimer_text,
        "cadastral_plan_no": cadastral_plan_no,
        "cadastral_area_name": cadastral_area_name,
        "cadastral_datum_text": cadastral_datum_text,
        "cadastral_firm_block_text": cadastral_firm_block_text,
        "fct_file_no": fct_file_no,
        "fct_district": fct_district,
        "fct_cadastral_zone": fct_cadastral_zone,
        "fct_origin_beacon_text": fct_origin_beacon_text,
        "fct_cadastral_map_ref": fct_cadastral_map_ref,
        "fct_title_prefix": fct_title_prefix,
        "survey_input_coordinates": survey_input_coordinates or [],
    }
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.pdf",
        payload=request_payload,
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.pdf",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.pdf",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, request_payload, "Survey_Plan", "pdf"),
        request_payload=request_payload,
    )
    return _serialize_plot_export_job(job, request=request)


@router.post("/{plot_id}/export-jobs/orthophoto.pdf")
def create_plot_orthophoto_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    title_text: str = Body("ORTHOPHOTO"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    use_topo_map: bool = Body(False),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
):
    request_payload = {
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "scale_text": scale_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "station_names": station_names or [],
        "coordinate_system": coordinate_system,
        "paper_size": paper_size,
        "use_topo_map": False,
        "north_arrow_style": north_arrow_style,
        "north_arrow_color": north_arrow_color,
    }
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="orthophoto.pdf",
        payload=request_payload,
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="orthophoto.pdf",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="orthophoto.pdf",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, request_payload, "Orthophoto", "pdf"),
        request_payload=request_payload,
    )
    return _serialize_plot_export_job(job, request=request)


@router.post("/{plot_id}/export-jobs/topomap.pdf")
def create_plot_topomap_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    title_text: str = Body("TOPO MAP"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    topo_source: str = Body("opentopomap"),
    contour_interval: float | None = Body(None),
    building_hatch_type: str = Body("solid"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
):
    request_payload = {
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "scale_text": scale_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "station_names": station_names or [],
        "coordinate_system": coordinate_system,
        "paper_size": paper_size,
        "use_topo_map": True,
        "topo_source": topo_source or "opentopomap",
        "contour_interval": contour_interval,
        "building_hatch_type": building_hatch_type or "solid",
        "north_arrow_style": north_arrow_style,
        "north_arrow_color": north_arrow_color,
    }
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="topomap.pdf",
        payload=request_payload,
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="topomap.pdf",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="topomap.pdf",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, request_payload, "Topo_Map", "pdf"),
        request_payload=request_payload,
    )
    return _serialize_plot_export_job(job, request=request)


@router.post("/{plot_id}/export-jobs/survey-plan.dxf")
def create_plot_dxf_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.dxf",
        payload={},
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.dxf",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.dxf",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, get_plot_meta(db, plot_id), "Survey_Plan", "dxf"),
        request_payload={},
    )
    return _serialize_plot_export_job(job, request=request)


@router.post("/{plot_id}/export-jobs/survey-plan.shapefile")
def create_plot_shapefile_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.shapefile",
        payload={},
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.shapefile",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="survey-plan.shapefile",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, get_plot_meta(db, plot_id), "Survey_Plan_Shapefile", "zip"),
        request_payload={},
    )
    return _serialize_plot_export_job(job, request=request)


@router.post("/subdivision/batches/{batch_id}/export-jobs/clean-copy.pdf")
def create_subdivision_batch_clean_copy_export_job(
    batch_id: int,
    request: Request,
    db: Session = Depends(get_db),
    title_text: str = Body(""),
    area_labels: list[dict] | None = Body(None),
    paper_size: str | None = Body(None),
    scale_text: str | None = Body(None),
    coordinate_system: str | None = Body(None),
    station_names: list[str] | None = Body(None),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
    beacon_style: str = Body("cross"),
    road_width_m: float | None = Body(None),
    boundary_color: str | None = Body(None),
    grid_color: str | None = Body(None),
    text_color: str | None = Body(None),
    road_color: str | None = Body(None),
    river_color: str | None = Body(None),
    building_color: str | None = Body(None),
    building_hatch_type: str | None = Body(None),
    road_style: str | None = Body(None),
    station_font: str | None = Body(None),
    station_size: int | None = Body(None),
    bearing_font: str | None = Body(None),
    bearing_size: int | None = Body(None),
):
    payload = {
        "title_text": str(title_text or ""),
        "area_labels": area_labels if isinstance(area_labels, list) else [],
        "paper_size": paper_size,
        "scale_text": scale_text,
        "coordinate_system": coordinate_system,
        "station_names": station_names if isinstance(station_names, list) else [],
        "north_arrow_style": str(north_arrow_style or "one_side_stem"),
        "north_arrow_color": str(north_arrow_color or "blue"),
        "beacon_style": str(beacon_style or "cross"),
        "road_width_m": road_width_m,
        "boundary_color": boundary_color,
        "grid_color": grid_color,
        "text_color": text_color,
        "road_color": road_color,
        "river_color": river_color,
        "building_color": building_color,
        "building_hatch_type": building_hatch_type,
        "road_style": road_style,
        "station_font": station_font,
        "station_size": station_size,
        "bearing_font": bearing_font,
        "bearing_size": bearing_size,
    }
    context = _get_subdivision_batch_clean_copy_context(
        db,
        batch_id=int(batch_id),
        title_text=str(payload.get("title_text") or ""),
        area_labels=payload["area_labels"],
        paper_size=paper_size,
        scale_text=scale_text,
        coordinate_system=coordinate_system,
        station_names=payload["station_names"],
        north_arrow_style=str(payload.get("north_arrow_style") or "one_side_stem"),
        north_arrow_color=str(payload.get("north_arrow_color") or "blue"),
        beacon_style=str(payload.get("beacon_style") or "cross"),
        road_width_m=road_width_m,
        boundary_color=boundary_color,
        grid_color=grid_color,
        text_color=text_color,
        road_color=road_color,
        river_color=river_color,
        building_color=building_color,
        building_hatch_type=building_hatch_type,
        road_style=road_style,
        station_font=station_font,
        station_size=station_size,
        bearing_font=bearing_font,
        bearing_size=bearing_size,
    )
    existing = db.execute(
        text(
            """
            SELECT
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                status,
                file_name,
                local_path,
                request_payload,
                error_text,
                started_at,
                completed_at,
                created_at,
                updated_at
            FROM plot_export_jobs
            WHERE export_type = 'subdivision_batch_clean_copy_pdf'
              AND subdivision_batch_id = :batch_id
              AND file_name = :file_name
              AND status IN ('queued', 'running')
            ORDER BY created_at DESC
            LIMIT 1
            """
        ),
        {"batch_id": int(batch_id), "file_name": str(context.get("pdf_name") or "")},
    ).mappings().first()
    if existing:
        return _serialize_plot_export_job(dict(existing), request=request)

    job_id = uuid.uuid4().hex
    if os.path.isfile(str(context.get("cached_pdf_path") or "").strip()):
        db.execute(
            text(
                """
                INSERT INTO plot_export_jobs (
                    id,
                    export_type,
                    plot_id,
                    subdivision_batch_id,
                    status,
                    file_name,
                    local_path,
                    request_payload,
                    started_at,
                    completed_at,
                    created_at,
                    updated_at
                )
                VALUES (
                    :id,
                    'subdivision_batch_clean_copy_pdf',
                    :plot_id,
                    :batch_id,
                    'completed',
                    :file_name,
                    :local_path,
                    CAST(:request_payload AS JSONB),
                    NOW(),
                    NOW(),
                    NOW(),
                    NOW()
                )
                """
            ),
            {
                "id": job_id,
                "plot_id": int(context.get("parent_plot_id") or 0) or None,
                "batch_id": int(batch_id),
                "file_name": str(context.get("pdf_name") or "subdivision_clean_copy.pdf"),
                "local_path": str(context.get("cached_pdf_path") or ""),
                "request_payload": json.dumps(payload),
            },
        )
        db.commit()
        job = _get_plot_export_job(db, job_id)
        return _serialize_plot_export_job(job or {"id": job_id, "status": "completed"}, request=request)

    db.execute(
        text(
            """
            INSERT INTO plot_export_jobs (
                id,
                export_type,
                plot_id,
                subdivision_batch_id,
                status,
                file_name,
                request_payload,
                created_at,
                updated_at
            )
            VALUES (
                :id,
                'subdivision_batch_clean_copy_pdf',
                :plot_id,
                :batch_id,
                'queued',
                :file_name,
                CAST(:request_payload AS JSONB),
                NOW(),
                NOW()
            )
            """
        ),
        {
            "id": job_id,
            "plot_id": int(context.get("parent_plot_id") or 0) or None,
            "batch_id": int(batch_id),
            "file_name": str(context.get("pdf_name") or "subdivision_clean_copy.pdf"),
            "request_payload": json.dumps(payload),
        },
    )
    db.commit()
    threading.Thread(target=_run_subdivision_batch_clean_copy_export_job, args=(job_id,), daemon=True).start()
    job = _get_plot_export_job(db, job_id)
    return _serialize_plot_export_job(job or {"id": job_id, "status": "queued"}, request=request)


@router.post("/subdivision/batches/{batch_id}/export/clean-copy.pdf")
def export_subdivision_batch_clean_copy_pdf(
    batch_id: int,
    db: Session = Depends(get_db),
    title_text: str = Body(""),
    area_labels: list[dict] | None = Body(None),
    paper_size: str | None = Body(None),
    scale_text: str | None = Body(None),
    coordinate_system: str | None = Body(None),
    station_names: list[str] | None = Body(None),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
    beacon_style: str = Body("cross"),
    road_width_m: float | None = Body(None),
    boundary_color: str | None = Body(None),
    grid_color: str | None = Body(None),
    text_color: str | None = Body(None),
    road_color: str | None = Body(None),
    river_color: str | None = Body(None),
    building_color: str | None = Body(None),
    building_hatch_type: str | None = Body(None),
    road_style: str | None = Body(None),
    station_font: str | None = Body(None),
    station_size: int | None = Body(None),
    bearing_font: str | None = Body(None),
    bearing_size: int | None = Body(None),
):
    context = _get_subdivision_batch_clean_copy_context(
        db,
        batch_id=int(batch_id),
        title_text=title_text,
        area_labels=area_labels,
        paper_size=paper_size,
        scale_text=scale_text,
        coordinate_system=coordinate_system,
        station_names=station_names,
        north_arrow_style=north_arrow_style,
        north_arrow_color=north_arrow_color,
        beacon_style=beacon_style,
        road_width_m=road_width_m,
        boundary_color=boundary_color,
        grid_color=grid_color,
        text_color=text_color,
        road_color=road_color,
        river_color=river_color,
        building_color=building_color,
        building_hatch_type=building_hatch_type,
        road_style=road_style,
        station_font=station_font,
        station_size=station_size,
        bearing_font=bearing_font,
        bearing_size=bearing_size,
    )
    cached_pdf_path = _generate_subdivision_batch_clean_copy_pdf(db, context=context)
    return _pdf_response_with_r2(
        cached_pdf_path,
        str(context.get("pdf_name") or "subdivision_clean_copy.pdf"),
        category="survey_subdivision_clean_copy",
    )


# ---------------- CREATE PLOT ----------------

@router.post("")
def create_plot(
    payload: Union[PlotCreateRequest, List[List[float]]],
    request: Request,
    db: Session = Depends(get_db),
):

    coords = payload.coordinates if isinstance(payload, PlotCreateRequest) else payload
    meta = payload.meta if isinstance(payload, PlotCreateRequest) else None
    client_request_id = (
        str(payload.client_request_id or "").strip()
        if isinstance(payload, PlotCreateRequest) and payload.client_request_id
        else ""
    )

    if len(coords) < 3:
        raise HTTPException(status_code=400, detail="Polygon requires at least 3 points")

    ensure_plots_created_at(db)

    if client_request_id:
        # A resend of the same client-generated attempt id (e.g. after a dropped response on a
        # flaky connection) returns the plot already created for it instead of creating a
        # duplicate - see ensure_plot_idempotency_columns for the backing unique index.
        existing_plot = db.execute(
            text("SELECT id FROM plots WHERE client_request_id = :client_request_id LIMIT 1"),
            {"client_request_id": client_request_id},
        ).mappings().first()
        if existing_plot:
            return {"plot_id": int(existing_plot["id"])}

    polygon = Polygon(coords)
    geom = from_shape(polygon, srid=4326)

    # Anonymous creation stays fully supported (the "value first" gate-free flow) - this only
    # stamps ownership when the request happens to already carry a signed-in Survey session.
    owner_user_id = None
    try:
        survey_session = resolve_survey_session(db, request)
        owner_user_id = survey_session.user_id if survey_session else None
    except Exception:
        owner_user_id = None

    plot = Plot(geom=geom, client_request_id=client_request_id or None, owner_user_id=owner_user_id)
    db.add(plot)
    db.commit()
    db.refresh(plot)

    # Ensure meta table exists and create a base row for this plot
    ensure_plot_meta_table(db)
    db.execute(text("""
        INSERT INTO plot_meta (plot_id)
        VALUES (:plot_id)
        ON CONFLICT (plot_id) DO NOTHING
    """), {"plot_id": plot.id})
    db.commit()

    if meta:
        upsert_plot_meta(
            db=db,
            plot_id=plot.id,
            title_text=meta.title or None,
            location_text=meta.location or None,
            lga_text=meta.lga or None,
            state_text=meta.state or None,
            surveyor_name=meta.surveyor or None,
            surveyor_rank=meta.rank or None,
            scale_text=meta.scale or None,
            survey_input_coordinates=getattr(meta, "survey_input_coordinates", None),
        )

    # Buffer + buildings/roads/rivers detection (Nigeria: local-table query; elsewhere: Overpass-
    # backed regional cache, run synchronously so it's already complete by the time this request
    # returns) - shared with subdivision child plots, see _run_plot_feature_detection.
    _run_plot_feature_detection(db, plot.id)

    db.commit()

    return {"plot_id": plot.id}


# ---------------- OWNERSHIP (Survey auth) ----------------

@router.post("/claim")
def claim_plots(request: Request, plot_ids: List[int] = Body(..., embed=True), db: Session = Depends(get_db)):
    session = require_survey_session(db, request)
    ensure_survey_activity_table(db)
    if not plot_ids:
        return {"claimed": []}
    # Only ever claims currently-unowned plots - never reassigns a plot someone else already owns.
    rows = db.execute(
        text(
            """
            UPDATE plots
            SET owner_user_id = :user_id
            WHERE id = ANY(:plot_ids) AND owner_user_id IS NULL
            RETURNING id
            """
        ),
        {"user_id": session.user_id, "plot_ids": plot_ids},
    ).mappings().all()

    # Subdivision lots generated from a parent plot before the surveyor logged in never got their
    # own owner_user_id stamped (only the parent plot id is tracked in the browser's draft list -
    # see savePlotToStorage in SurveyPlan.tsx). Claiming the parent should transitively claim every
    # lot generated from it too, so they show up in the admin/"my plots" listing the same way.
    child_rows = db.execute(
        text(
            """
            UPDATE plots child
            SET owner_user_id = :user_id
            FROM plot_meta pm
            WHERE child.id = pm.plot_id
              AND child.owner_user_id IS NULL
              AND pm.parent_plot_id IS NOT NULL
              AND pm.parent_plot_id IN (
                  SELECT id FROM plots WHERE owner_user_id = :user_id
              )
            RETURNING child.id
            """
        ),
        {"user_id": session.user_id},
    ).mappings().all()
    claimed_ids = [int(row["id"]) for row in rows] + [int(row["id"]) for row in child_rows]
    if claimed_ids:
        # Attribute pre-sign-in previews/exports to the account that claimed this browser's plots.
        # Only anonymous events are updated; activity belonging to another account is immutable.
        db.execute(
            text("""
                UPDATE survey_activity_events
                SET actor_user_id = :user_id
                WHERE actor_user_id IS NULL AND plot_id = ANY(:plot_ids)
            """),
            {"user_id": session.user_id, "plot_ids": claimed_ids},
        )
    db.commit()
    return {"claimed": claimed_ids}


@router.get("/mine")
def list_my_plots(request: Request, db: Session = Depends(get_db)):
    session = require_survey_session(db, request)
    ensure_plot_meta_table(db)
    rows = db.execute(
        text(
            """
            SELECT p.id, p.created_at, m.title_text, m.location_text, m.scale_text,
                   m.parent_plot_id, m.subdivision_batch_id, m.subdivision_lot_no, m.estate_name
            FROM plots p
            LEFT JOIN plot_meta m ON m.plot_id = p.id
            WHERE p.owner_user_id = :user_id
            ORDER BY p.created_at DESC
            """
        ),
        {"user_id": session.user_id},
    ).mappings().all()
    return {
        "plots": [
            {
                "plot_id": int(row["id"]),
                "created_at": row["created_at"].isoformat() if row["created_at"] else None,
                "title": row["title_text"],
                "location": row["location_text"],
                "scale": row["scale_text"],
                "status": "completed" if (row["title_text"] or "").strip() else "draft",
                "workflow_type": "subdivision" if row["parent_plot_id"] else "survey_plan",
                "parent_plot_id": row["parent_plot_id"],
                "subdivision_batch_id": row["subdivision_batch_id"],
                "subdivision_lot_no": row["subdivision_lot_no"],
                "estate_name": row["estate_name"],
            }
            for row in rows
        ]
    }


@router.post("/bulk-delete")
def bulk_delete_plots(request: Request, plot_ids: List[int] = Body(..., embed=True), db: Session = Depends(get_db)):
    session = require_survey_session(db, request)
    if not plot_ids:
        return {"deleted": []}

    # Only ever deletes plots the caller actually owns - never someone else's, and never an
    # anonymous (unclaimed) one, even if its id is guessed/passed in.
    owned_rows = db.execute(
        text("SELECT id FROM plots WHERE id = ANY(:plot_ids) AND owner_user_id = :user_id"),
        {"plot_ids": plot_ids, "user_id": session.user_id},
    ).mappings().all()
    owned_ids = [int(row["id"]) for row in owned_rows]
    if not owned_ids:
        return {"deleted": []}

    # plot_meta / plot_feature_overrides / plot_export_jobs / subdivision tables all have
    # ON DELETE CASCADE, but detected_features and plot_buffers are plain SQLAlchemy ForeignKeys
    # without it - deleting plots first would fail on those with a foreign-key violation, so
    # clear them explicitly first.
    db.execute(text("DELETE FROM detected_features WHERE plot_id = ANY(:ids)"), {"ids": owned_ids})
    db.execute(text("DELETE FROM plot_buffers WHERE plot_id = ANY(:ids)"), {"ids": owned_ids})
    db.execute(text("DELETE FROM plots WHERE id = ANY(:ids)"), {"ids": owned_ids})
    db.commit()
    return {"deleted": owned_ids}


# ---------------- FEATURES SUMMARY ----------------

@router.get("/{plot_id}/features")
def get_plot_features(plot_id: int, db: Session = Depends(get_db)):

    sql = text("""
        SELECT feature_type, location, COUNT(*) as count
        FROM detected_features
        WHERE plot_id = :plot_id
        GROUP BY feature_type, location
    """)

    rows = db.execute(sql, {"plot_id": plot_id}).fetchall()

    response = {"plot_id": plot_id, "inside": {}, "buffer": {}}

    for r in rows:
        response[r.location][r.feature_type] = r.count

    return response


@router.get("/{plot_id}/features/geojson")
def get_plot_features_geojson(
    plot_id: int,
    db: Session = Depends(get_db),
    scale_text: str | None = Query(default=None),
    paper_size: str | None = Query(default=None),
    template_name: str | None = Query(default=None),
):
    plot_geom = None
    display_epsg: int | None = None
    extent_wgs84 = None

    # Buildings and rivers from detected_features
    feature_rows = db.execute(text("""
        SELECT feature_type, ST_AsGeoJSON(geom) AS geojson
        FROM detected_features
        WHERE plot_id = :plot_id
    """), {"plot_id": plot_id}).fetchall()

    # Roads from lines (same logic as renderer), clipped to the plot buffer so a long
    # real-world road doesn't ship in full over a slow connection. ST_Intersection can turn a
    # road that crosses the buffer more than once into a MultiLineString (or a bare Point on a
    # tangent touch) - ST_Dump explodes that into individual rows first, and the geometry-type
    # filter drops degenerate points, so the CAD editor (which only understands plain
    # LineString features) never receives anything it can't render.
    road_rows = db.execute(text("""
        WITH roads AS (
            SELECT
                CASE
                    WHEN ST_SRID(r.geom) = 4326 THEN r.geom
                    WHEN ST_SRID(r.geom) = 0 THEN ST_SetSRID(r.geom, 4326)
                    ELSE ST_Transform(r.geom, 4326)
                END AS geom,
                r.highway,
                r.name
            FROM lines r
            WHERE r.highway IS NOT NULL
        ),
        clipped AS (
            SELECT
                roads.name,
                (ST_Dump(ST_Intersection(roads.geom, b.geom))).geom AS geom
            FROM roads
            JOIN plot_buffers b ON b.plot_id = :plot_id
            WHERE ST_Intersects(roads.geom, b.geom)
        )
        SELECT ST_AsGeoJSON(geom) AS geojson, name
        FROM clipped
        WHERE ST_GeometryType(geom) = 'ST_LineString'
    """), {"plot_id": plot_id}).fetchall()

    # Overrides
    override_rows = db.execute(text("""
        SELECT feature_type, action, name, width_m, ST_AsGeoJSON(geom) AS geojson
        FROM plot_feature_overrides
        WHERE plot_id = :plot_id
    """), {"plot_id": plot_id}).fetchall()

    def to_feature(geojson_str, props):
        return {
            "type": "Feature",
            "geometry": json.loads(geojson_str) if geojson_str else None,
            "properties": props,
        }

    def _iter_line_parts(geom):
        if geom is None:
            return []
        geom_type = getattr(geom, "geom_type", "")
        if geom_type == "LineString":
            return [geom] if getattr(geom, "length", 0.0) > 0 else []
        if geom_type == "MultiLineString":
            return [part for part in getattr(geom, "geoms", []) if getattr(part, "length", 0.0) > 0]
        parts = []
        for part in getattr(geom, "geoms", []):
            part_type = getattr(part, "geom_type", "")
            if part_type == "LineString" and getattr(part, "length", 0.0) > 0:
                parts.append(part)
            elif part_type == "MultiLineString":
                parts.extend([sub for sub in getattr(part, "geoms", []) if getattr(sub, "length", 0.0) > 0])
        return parts

    def _metric_length_m(geom):
        if geom is None:
            return None
        try:
            metric_epsg = display_epsg or (_metric_epsg_for_wgs84_polygon(plot_geom) if plot_geom is not None else None)
            if not metric_epsg:
                return None
            projected = gpd.GeoSeries([geom], crs="EPSG:4326").to_crs(epsg=metric_epsg).iloc[0]
            return float(projected.length)
        except Exception:
            return None

    def _line_anchor_point(geom):
        candidate = geom
        if extent_wgs84 is not None:
            try:
                clipped = geom.intersection(extent_wgs84)
                clipped_parts = _iter_line_parts(clipped)
                if clipped_parts:
                    candidate = max(clipped_parts, key=lambda part: getattr(part, "length", 0.0))
            except Exception:
                pass
        parts = _iter_line_parts(candidate)
        if not parts:
            parts = _iter_line_parts(geom)
        if not parts:
            return None
        line = max(parts, key=lambda part: getattr(part, "length", 0.0))
        try:
            return line.interpolate(line.length / 2.0)
        except Exception:
            try:
                return line.representative_point()
            except Exception:
                return None

    def _position_hint(geom):
        if extent_wgs84 is None:
            return ""
        anchor = _line_anchor_point(geom)
        if anchor is None:
            return ""
        minx, miny, maxx, maxy = extent_wgs84.bounds
        width = max(maxx - minx, 1e-9)
        height = max(maxy - miny, 1e-9)
        x_rel = (anchor.x - minx) / width
        y_rel = (anchor.y - miny) / height

        if x_rel < 0.33:
            horiz = "west"
        elif x_rel > 0.67:
            horiz = "east"
        else:
            horiz = "center"

        if y_rel < 0.33:
            vert = "south"
        elif y_rel > 0.67:
            vert = "north"
        else:
            vert = "center"

        if vert == "center" and horiz == "center":
            return "center"
        if vert == "center":
            return horiz
        if horiz == "center":
            return vert
        return f"{vert} {horiz}"

    def _segment_key(kind: str, geom) -> str:
        try:
            digest = hashlib.sha1(bytes(geom.wkb)).hexdigest()[:12]
        except Exception:
            digest = hashlib.sha1(str(geom).encode("utf-8")).hexdigest()[:12]
        return f"{kind}-{digest}"

    def _line_sort_key(pair):
        geom = pair[0]
        anchor = _line_anchor_point(geom)
        if anchor is None:
            return (999.0, 999.0)
        return (-float(anchor.y), float(anchor.x))

    def _decorate_line_pairs(pairs: list[tuple[Any, dict]], kind: str) -> list[tuple[Any, dict]]:
        sorted_pairs = sorted(pairs, key=_line_sort_key)
        out: list[tuple[Any, dict]] = []
        for idx, (geom, feature) in enumerate(sorted_pairs, start=1):
            props = dict((feature or {}).get("properties") or {})
            props["segment_key"] = _segment_key(kind, geom)
            props["segment_index"] = idx
            props["position_hint"] = _position_hint(geom)
            length_m = _metric_length_m(geom)
            if length_m is not None:
                props["length_m"] = round(length_m, 2)
            feature["properties"] = props
            out.append((geom, feature))
        return out

    # Each entry is (shapely geometry, feature dict) so overrides can be applied by spatial
    # intersection - same proven pattern already used in map_renderer_layout.py's apply_overrides
    # (and a third copy in this file's _render_subdivision_clean_copy_pdf), just not previously
    # wired into this endpoint. Without this, a "delete" override was silently ignored on reload
    # and an "update" override just appended a duplicate instead of replacing the original.
    buildings: list[tuple[Any, dict]] = []
    rivers: list[tuple[Any, dict]] = []
    fences: list[tuple[Any, dict]] = []
    for r in feature_rows:
        if not r.geojson:
            continue
        try:
            geom = shape(json.loads(r.geojson))
        except Exception:
            continue
        feat = to_feature(r.geojson, {"source": "detected"})
        if r.feature_type == "building":
            buildings.append((geom, feat))
        elif r.feature_type == "river":
            rivers.append((geom, feat))
        elif r.feature_type == "fence":
            fences.append((geom, feat))

    roads: list[tuple[Any, dict]] = []
    for r in road_rows:
        if not r.geojson:
            continue
        try:
            geom = shape(json.loads(r.geojson))
        except Exception:
            continue
        roads.append((geom, to_feature(r.geojson, {"source": "detected", "name": r.name})))

    # The live `lines` query above depends on a plot_buffers row existing for this plot (created
    # by _run_plot_feature_detection). Some plots never got one - e.g. older plots, or subdivision
    # child lots that inherit detected_features straight from the parent without re-running full
    # detection - so it can come back empty even though this plot's roads were captured into
    # detected_features at creation time (which is exactly what the survey plan preview/render
    # falls back to in preview mode). Without this fallback the CAD editor would show no roads at
    # all for such a plot while the preview still does.
    if not roads:
        for r in feature_rows:
            if r.feature_type != "road" or not r.geojson:
                continue
            try:
                geom = shape(json.loads(r.geojson))
            except Exception:
                continue
            roads.append((geom, to_feature(r.geojson, {"source": "detected", "name": None})))

    overrides: list[dict[str, Any]] = []
    for r in override_rows:
        if not r.geojson:
            continue
        try:
            geom = shape(json.loads(r.geojson))
        except Exception:
            continue
        overrides.append({
            "feature_type": str(r.feature_type or "").strip().lower(),
            "action": str(r.action or "").strip().lower(),
            "name": r.name,
            "width_m": r.width_m,
            "geom": geom,
            "geojson": r.geojson,
        })

    def _line_replaced_by(geom, override_geom, tol_deg: float = 0.00001) -> bool:
        # A plain intersects() can't tell "this is the segment the override replaces" apart from
        # "this is a different road/river that merely touches it at a junction" - two distinct
        # roads meeting at a junction legitimately intersect at that single point. This checks how
        # much of `geom`'s own length lies within a small buffer of the override instead: a
        # segment actually being replaced is covered almost along its whole length, while a
        # merely-crossing/touching one only has a single point (~0 length) inside the buffer.
        # Without this, naming one road (an "update" override) could make an unrelated road it
        # happens to join vanish from the CAD editor entirely.
        try:
            total_len = max(getattr(geom, "length", 0.0), 1e-9)
            uncovered = geom.difference(override_geom.buffer(tol_deg))
            uncovered_len = getattr(uncovered, "length", 0.0)
            return uncovered_len < total_len * 0.1
        except Exception:
            return geom.intersects(override_geom)

    def apply_overrides(base_pairs: list[tuple[Any, dict]], feature_type: str) -> list[tuple[Any, dict]]:
        result = list(base_pairs)
        use_line_test = feature_type in ("road", "river")
        for ov in overrides:
            if ov["feature_type"] != feature_type:
                continue
            geom = ov["geom"]
            try:
                if not geom.is_valid:
                    geom = geom.buffer(0)
            except Exception:
                pass
            if ov["action"] in ("delete", "update"):
                if use_line_test:
                    result = [(g, f) for (g, f) in result if not _line_replaced_by(g, geom)]
                else:
                    result = [(g, f) for (g, f) in result if not g.intersects(geom)]
            if ov["action"] in ("add", "update"):
                feat = to_feature(ov["geojson"], {"source": "override", "name": ov["name"], "width_m": ov["width_m"]})
                result.append((geom, feat))
        return result

    buildings = apply_overrides(buildings, "building")
    rivers = apply_overrides(rivers, "river")
    fences = apply_overrides(fences, "fence")
    roads = apply_overrides(roads, "road")

    # When the caller (the Road Names panel) knows the current scale/paper size, only surface
    # roads/rivers that will actually print on the sheet at that scale - the plot_buffers-clipped
    # data above can extend well beyond what a given scale + paper size will actually show.
    if scale_text:
        try:
            plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id=:id"), {"id": plot_id}).scalar()
            if plot_wkb:
                plot_geom = wkb.loads(plot_wkb)
                display_epsg = _metric_epsg_for_wgs84_polygon(plot_geom)

                proj_geom = gpd.GeoSeries([plot_geom], crs="EPSG:4326").to_crs(epsg=display_epsg).iloc[0]

                paper_config = get_paper_config(paper_size or "A4")
                map_width_frac, map_height_frac = _survey_template_map_frame(template_name)

                scale_ratio = parse_scale_ratio(scale_text)
                minx, miny, maxx, maxy = proj_geom.bounds
                cx, cy = (minx + maxx) / 2.0, (miny + maxy) / 2.0
                inch_to_m = 0.0254
                real_w = paper_config["width"] * map_width_frac * inch_to_m * scale_ratio
                real_h = paper_config["height"] * map_height_frac * inch_to_m * scale_ratio
                extent_proj = box(cx - real_w / 2.0, cy - real_h / 2.0, cx + real_w / 2.0, cy + real_h / 2.0)
                extent_wgs84 = gpd.GeoSeries([extent_proj], crs=f"EPSG:{display_epsg}").to_crs(epsg=4326).iloc[0]

                roads = [(g, f) for (g, f) in roads if g.intersects(extent_wgs84)]
                rivers = [(g, f) for (g, f) in rivers if g.intersects(extent_wgs84)]
        except Exception:
            pass

    if plot_geom is None:
        try:
            plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id=:id"), {"id": plot_id}).scalar()
            if plot_wkb:
                plot_geom = wkb.loads(plot_wkb)
                display_epsg = display_epsg or _metric_epsg_for_wgs84_polygon(plot_geom)
        except Exception:
            plot_geom = None

    roads = _decorate_line_pairs(roads, "road")
    rivers = _decorate_line_pairs(rivers, "river")

    return {
        "roads": {"type": "FeatureCollection", "features": [f for _, f in roads]},
        "buildings": {"type": "FeatureCollection", "features": [f for _, f in buildings]},
        "rivers": {"type": "FeatureCollection", "features": [f for _, f in rivers]},
        "fences": {"type": "FeatureCollection", "features": [f for _, f in fences]},
    }


@router.post("/{plot_id}/geometry")
def update_plot_geometry(
    plot_id: int,
    payload: PlotGeometryUpdateRequest,
    db: Session = Depends(get_db),
):
    coords = payload.coordinates or []
    if len(coords) < 3:
        raise HTTPException(status_code=400, detail="Polygon requires at least 3 points")

    plot = db.query(Plot).filter(Plot.id == plot_id).first()
    if not plot:
        raise HTTPException(status_code=404, detail="Plot not found")

    polygon = Polygon(coords)
    plot.geom = from_shape(polygon, srid=4326)
    db.add(plot)
    db.commit()

    return {"ok": True, "plot_id": plot_id}


# ---------------- FEATURE OVERRIDES ----------------

@router.get("/{plot_id}/feature-overrides")
def get_feature_overrides(plot_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT id, feature_type, action, name, width_m, ST_AsGeoJSON(geom) AS geojson, created_at, updated_at
        FROM plot_feature_overrides
        WHERE plot_id = :plot_id
        ORDER BY id DESC
    """), {"plot_id": plot_id}).fetchall()

    return [
        {
            "id": r.id,
            "feature_type": r.feature_type,
            "action": r.action,
            "name": r.name,
            "width_m": r.width_m,
            "geojson": r.geojson,
            "created_at": r.created_at,
            "updated_at": r.updated_at,
        }
        for r in rows
    ]


def _feature_geom_replaced_by(candidate_geom, new_geom, tol_deg: float = 0.00001) -> bool:
    """Whether `candidate_geom` (an existing override or detected feature) is genuinely the same
    real-world feature `new_geom` is replacing, rather than a different feature that merely
    touches or crosses it (e.g. two roads meeting at a junction, or a fence corner grazing a
    building). A plain ST_Intersects/.intersects() test can't tell those apart - it's true for a
    single shared point just as much as for two nearly-identical lines - so using it to decide
    what to delete would destroy unrelated edits (e.g. wiping out a previous "delete this other
    road" override just because the newly-saved geometry happens to touch it).

    Coverage is measured proportionally to `candidate_geom`'s own size (length for lines, area for
    polygons): the real match is covered almost entirely by a small buffer around `new_geom`,
    while a merely-touching feature only has a single point/edge (~0 length or area) inside it.
    """
    try:
        geom_type = candidate_geom.geom_type
        buffered = new_geom.buffer(tol_deg)
        if geom_type in ("LineString", "MultiLineString"):
            total = max(getattr(candidate_geom, "length", 0.0), 1e-9)
            uncovered = candidate_geom.difference(buffered)
            return getattr(uncovered, "length", 0.0) < total * 0.1
        if geom_type in ("Polygon", "MultiPolygon"):
            total = max(getattr(candidate_geom, "area", 0.0), 1e-12)
            uncovered = candidate_geom.difference(buffered)
            return getattr(uncovered, "area", 0.0) < total * 0.1
        return candidate_geom.distance(new_geom) < tol_deg
    except Exception:
        return candidate_geom.intersects(new_geom)


def _feature_geom_replaced_by_type(candidate_geom, new_geom, feature_type: str, tol_deg: float = 0.00001) -> bool:
    normalized = str(feature_type or "").strip().lower()
    try:
        buffered = new_geom.buffer(tol_deg)
        if normalized in ("road", "river", "fence"):
            total = max(getattr(candidate_geom, "length", 0.0), 1e-9)
            uncovered = candidate_geom.difference(buffered)
            return getattr(uncovered, "length", 0.0) < total * 0.1
        if normalized in ("building",):
            total = max(getattr(candidate_geom, "area", 0.0), 1e-12)
            uncovered = candidate_geom.difference(buffered)
            return getattr(uncovered, "area", 0.0) < total * 0.1
        return candidate_geom.distance(new_geom) < tol_deg
    except Exception:
        return candidate_geom.intersects(new_geom)


@router.post("/{plot_id}/feature-overrides")
def add_feature_override(
    plot_id: int,
    db: Session = Depends(get_db),
    feature_type: str = Body(...),
    action: str = Body(...),
    name: str = Body(default=""),
    width_m: float | None = Body(default=None),
    wkt: str | None = Body(default=None),
    geojson: dict | None = Body(default=None),
    client_request_id: str | None = Body(default=None),
):
    if feature_type not in {"road", "building", "river", "fence"}:
        raise HTTPException(status_code=400, detail="Invalid feature_type")
    if action not in {"add", "delete", "update"}:
        raise HTTPException(status_code=400, detail="Invalid action")

    client_request_id = str(client_request_id or "").strip() or None
    if client_request_id:
        # See create_plot's matching check - lets a resend of the same attempt id (after a
        # dropped response) return the existing row instead of inserting a duplicate override.
        existing_override = db.execute(
            text("SELECT id FROM plot_feature_overrides WHERE client_request_id = :client_request_id LIMIT 1"),
            {"client_request_id": client_request_id},
        ).mappings().first()
        if existing_override:
            return {"status": "ok"}

    geom_wkt = None
    geom_geojson = None
    if wkt:
        geom_wkt = wkt
    elif geojson:
        try:
            geom_geojson = geojson
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid geojson")
    if not geom_wkt:
        if not geom_geojson:
            raise HTTPException(status_code=400, detail="Geometry is required")

    if action == "update":
        # An "update" replaces whatever's at this geometry - without this, repeated edits (e.g.
        # renaming a road more than once from the Road Names panel) pile up as separate override
        # rows that all intersect the same road. At render time those stale rows could still
        # contribute their own (now-superseded) name, or their slightly-different geometry could
        # linger alongside the new one and corrupt the road's shape - deleting them here keeps
        # exactly one override per logical feature.
        #
        # This must only delete overrides that are genuinely the SAME feature, not merely nearby
        # ones - a coarse ST_Intersects DELETE here was destructive: saving a new "update" (e.g.
        # naming a road) would wipe out ANY prior override that happened to touch it at even a
        # single point, including a completely unrelated "delete this other road" edit from
        # earlier in the session, silently undoing a user's previous work. So candidates are found
        # with a generous but bounded spatial filter, then only the ones _feature_geom_replaced_by
        # confirms as the same feature are actually removed.
        new_geom_obj = shape(geom_geojson) if geom_geojson else shapely_wkt.loads(geom_wkt)
        new_geom_json = json.dumps(geom_geojson) if geom_geojson else None
        candidate_rows = db.execute(text(f"""
            SELECT id, ST_AsGeoJSON(geom) AS geojson
            FROM plot_feature_overrides
            WHERE plot_id = :plot_id AND feature_type = :feature_type
              AND ST_DWithin(
                    geom::geography,
                    {"ST_SetSRID(ST_GeomFromGeoJSON(:new_geojson), 4326)" if new_geom_json else "ST_SetSRID(ST_GeomFromText(:new_wkt), 4326)"}::geography,
                    3
              )
        """), {
            "plot_id": plot_id,
            "feature_type": feature_type,
            **({"new_geojson": new_geom_json} if new_geom_json else {"new_wkt": geom_wkt}),
        }).fetchall()

        ids_to_delete = []
        for row in candidate_rows:
            try:
                candidate_geom = shape(json.loads(row.geojson))
            except Exception:
                continue
            if _feature_geom_replaced_by(candidate_geom, new_geom_obj):
                ids_to_delete.append(row.id)
        if ids_to_delete:
            db.execute(
                text("DELETE FROM plot_feature_overrides WHERE id = ANY(:ids)"),
                {"ids": ids_to_delete},
            )

    if geom_geojson:
        db.execute(text("""
            INSERT INTO plot_feature_overrides (plot_id, feature_type, action, name, width_m, geom, client_request_id)
            VALUES (:plot_id, :feature_type, :action, :name, :width_m, ST_SetSRID(ST_GeomFromGeoJSON(:geojson), 4326), :client_request_id)
        """), {
            "plot_id": plot_id,
            "feature_type": feature_type,
            "action": action,
            "name": name or None,
            "width_m": width_m,
            "geojson": json.dumps(geom_geojson),
            "client_request_id": client_request_id,
        })
    else:
        db.execute(text("""
            INSERT INTO plot_feature_overrides (plot_id, feature_type, action, name, width_m, geom, client_request_id)
            VALUES (:plot_id, :feature_type, :action, :name, :width_m, ST_SetSRID(ST_GeomFromText(:wkt), 4326), :client_request_id)
        """), {
            "plot_id": plot_id,
            "feature_type": feature_type,
            "action": action,
            "name": name or None,
            "width_m": width_m,
            "wkt": geom_wkt,
            "client_request_id": client_request_id,
        })
    db.commit()
    return {"status": "ok"}


@router.delete("/{plot_id}/feature-overrides/{override_id}")
def delete_feature_override(plot_id: int, override_id: int, db: Session = Depends(get_db)):
    res = db.execute(text("""
        DELETE FROM plot_feature_overrides
        WHERE id = :id AND plot_id = :plot_id
    """), {"id": override_id, "plot_id": plot_id})
    db.commit()
    if res.rowcount == 0:
        raise HTTPException(status_code=404, detail="Override not found")
    return {"status": "ok"}


# ---------------- REPORT DATA ----------------

@router.get("/{plot_id}/report")
def get_plot_report(plot_id: int, db: Session = Depends(get_db)):
    meta = get_plot_meta(db, plot_id)
    plot_geom = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom,
        meta.get("coordinate_system"),
        meta.get("survey_input_coordinates"),
    )
    area = measurement_context["measurement_area_m2"]

    rows = db.execute(text("""
        SELECT feature_type, location, COUNT(*) as count
        FROM detected_features
        WHERE plot_id = :plot_id
        GROUP BY feature_type, location
    """), {"plot_id": plot_id}).fetchall()

    inside = {}
    buffer = {}

    for r in rows:
        (inside if r.location == "inside" else buffer)[r.feature_type] = r.count

    return {
        "plot_id": plot_id,
        "area_m2": round(area, 2) if area else None,
        "features": {"inside": inside, "buffer": buffer}
    }


# ---------------- SURVEY PLAN PDF ----------------

@router.post("/{plot_id}/meta")
def save_plot_metadata(
    plot_id: int,
    db: Session = Depends(get_db),
    title_text: str = Body("SURVEY PLAN"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    certification_statement: str = Body(DEFAULT_CERTIFICATION_STATEMENT),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    adamawa_rof_no: str = Body(""),
    adamawa_owner_name: str = Body(""),
    adamawa_authority_title: str = Body(DEFAULT_ADAMAWA_AUTHORITY_TITLE),
    adamawa_authority_date_text: str = Body(DEFAULT_ADAMAWA_AUTHORITY_DATE),
    adamawa_control_point_name: str = Body(""),
    adamawa_northing: str = Body(""),
    adamawa_easting: str = Body(""),
    adamawa_elevation: str = Body(""),
    adamawa_origin_text: str = Body(DEFAULT_ADAMAWA_ORIGIN_TEXT),
    adamawa_topo_sheet_text: str = Body(DEFAULT_ADAMAWA_TOPO_SHEET_TEXT),
    adamawa_computation_no: str = Body(""),
    adamawa_cadastral_sheet_no: str = Body(""),
    adamawa_plan_no: str = Body(""),
    adamawa_surveyed_by_text: str = Body(""),
    adamawa_disclaimer_text: str = Body(DEFAULT_ADAMAWA_DISCLAIMER_TEXT),
    cadastral_plan_no: str = Body(""),
    cadastral_area_name: str = Body(""),
    cadastral_datum_text: str = Body(""),
    cadastral_firm_block_text: str = Body(""),
    fct_file_no: str = Body(""),
    fct_district: str = Body(""),
    fct_cadastral_zone: str = Body(""),
    fct_origin_beacon_text: str = Body(""),
    fct_cadastral_map_ref: str = Body(""),
    fct_title_prefix: str = Body(""),
    survey_input_coordinates: Optional[list] = Body(default=None),
):
    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        certification_statement=certification_statement,
        scale_text=scale_text,
        paper_size=paper_size,
        coordinate_system=coordinate_system,
        template_name=template_name,
        adamawa_rof_no=adamawa_rof_no,
        adamawa_owner_name=adamawa_owner_name,
        adamawa_authority_title=adamawa_authority_title,
        adamawa_authority_date_text=adamawa_authority_date_text,
        adamawa_control_point_name=adamawa_control_point_name,
        adamawa_northing=adamawa_northing,
        adamawa_easting=adamawa_easting,
        adamawa_elevation=adamawa_elevation,
        adamawa_origin_text=adamawa_origin_text,
        adamawa_topo_sheet_text=adamawa_topo_sheet_text,
        adamawa_computation_no=adamawa_computation_no,
        adamawa_cadastral_sheet_no=adamawa_cadastral_sheet_no,
        adamawa_plan_no=adamawa_plan_no,
        adamawa_surveyed_by_text=adamawa_surveyed_by_text,
        adamawa_disclaimer_text=adamawa_disclaimer_text,
        cadastral_plan_no=cadastral_plan_no,
        cadastral_area_name=cadastral_area_name,
        cadastral_datum_text=cadastral_datum_text,
        cadastral_firm_block_text=cadastral_firm_block_text,
        fct_file_no=fct_file_no,
        fct_district=fct_district,
        fct_cadastral_zone=fct_cadastral_zone,
        fct_origin_beacon_text=fct_origin_beacon_text,
        fct_cadastral_map_ref=fct_cadastral_map_ref,
        fct_title_prefix=fct_title_prefix,
        survey_input_coordinates=survey_input_coordinates,
    )
    return {"ok": True, "plot_id": int(plot_id)}


@router.post("/{plot_id}/report/pdf")
def download_plot_report_pdf(plot_id: int, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None,
    title_text: str = Body("SURVEY PLAN"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    certification_statement: str = Body(DEFAULT_CERTIFICATION_STATEMENT),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
    beacon_style: str = Body("cross"),
    road_width_m: float | None = Body(None),
    road_width_override_m: float | None = Body(None),
    boundary_color: str | None = Body(None),
    grid_color: str | None = Body(None),
    text_color: str | None = Body(None),
    road_color: str | None = Body(None),
    river_color: str | None = Body(None),
    building_color: str | None = Body(None),
    building_hatch_type: str | None = Body(None),
    road_style: str | None = Body(None),
    title_font: str | None = Body(None),
    title_size: int | None = Body(None),
    grid_font: str | None = Body(None),
    grid_size: int | None = Body(None),
    station_font: str | None = Body(None),
    station_size: int | None = Body(None),
    bearing_font: str | None = Body(None),
    bearing_size: int | None = Body(None),
    area_font: str | None = Body(None),
    area_size: int | None = Body(None),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    adamawa_rof_no: str = Body(""),
    adamawa_owner_name: str = Body(""),
    adamawa_authority_title: str = Body(DEFAULT_ADAMAWA_AUTHORITY_TITLE),
    adamawa_authority_date_text: str = Body(DEFAULT_ADAMAWA_AUTHORITY_DATE),
    adamawa_control_point_name: str = Body(""),
    adamawa_northing: str = Body(""),
    adamawa_easting: str = Body(""),
    adamawa_elevation: str = Body(""),
    adamawa_origin_text: str = Body(DEFAULT_ADAMAWA_ORIGIN_TEXT),
    adamawa_topo_sheet_text: str = Body(DEFAULT_ADAMAWA_TOPO_SHEET_TEXT),
    adamawa_computation_no: str = Body(""),
    adamawa_cadastral_sheet_no: str = Body(""),
    adamawa_plan_no: str = Body(""),
    adamawa_surveyed_by_text: str = Body(""),
    adamawa_disclaimer_text: str = Body(DEFAULT_ADAMAWA_DISCLAIMER_TEXT),
    cadastral_plan_no: str = Body(""),
    cadastral_area_name: str = Body(""),
    cadastral_datum_text: str = Body(""),
    cadastral_firm_block_text: str = Body(""),
    fct_file_no: str = Body(""),
    fct_district: str = Body(""),
    fct_cadastral_zone: str = Body(""),
    fct_origin_beacon_text: str = Body(""),
    fct_cadastral_map_ref: str = Body(""),
    fct_title_prefix: str = Body(""),
    survey_input_coordinates: Optional[list] = Body(default=None)):

    reports_dir = REPORTS_DIR
    maps_dir = os.path.join(REPORTS_DIR, "maps")
    os.makedirs(reports_dir, exist_ok=True)
    os.makedirs(maps_dir, exist_ok=True)

    pdf_path = f"{reports_dir}/plot_{plot_id}_report.pdf"

    tmp_map = tempfile.NamedTemporaryFile(suffix="_map.png", delete=False)
    map_path = tmp_map.name
    tmp_map.close()

    # Save/refresh plot metadata
    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        certification_statement=certification_statement,
        scale_text=scale_text,
        paper_size=paper_size,
        coordinate_system=coordinate_system,
        template_name=template_name,
        adamawa_rof_no=adamawa_rof_no,
        adamawa_owner_name=adamawa_owner_name,
        adamawa_authority_title=adamawa_authority_title,
        adamawa_authority_date_text=adamawa_authority_date_text,
        adamawa_control_point_name=adamawa_control_point_name,
        adamawa_northing=adamawa_northing,
        adamawa_easting=adamawa_easting,
        adamawa_elevation=adamawa_elevation,
        adamawa_origin_text=adamawa_origin_text,
        adamawa_topo_sheet_text=adamawa_topo_sheet_text,
        adamawa_computation_no=adamawa_computation_no,
        adamawa_cadastral_sheet_no=adamawa_cadastral_sheet_no,
        adamawa_plan_no=adamawa_plan_no,
        adamawa_surveyed_by_text=adamawa_surveyed_by_text,
        adamawa_disclaimer_text=adamawa_disclaimer_text,
        cadastral_plan_no=cadastral_plan_no,
        cadastral_area_name=cadastral_area_name,
        cadastral_datum_text=cadastral_datum_text,
        cadastral_firm_block_text=cadastral_firm_block_text,
        fct_file_no=fct_file_no,
        fct_district=fct_district,
        fct_cadastral_zone=fct_cadastral_zone,
        fct_origin_beacon_text=fct_origin_beacon_text,
        fct_cadastral_map_ref=fct_cadastral_map_ref,
        fct_title_prefix=fct_title_prefix,
        survey_input_coordinates=survey_input_coordinates,
    )

    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom_wgs84,
        coordinate_system,
        survey_input_coordinates=survey_input_coordinates,
    )
    render_coordinate_system = measurement_context["render_coordinate_system"]
    epsg_code = measurement_context["epsg_code"]
    crs_name = measurement_context["crs_name"]
    measurement_polygon = measurement_context["measurement_polygon"]
    measurement_area_m2 = measurement_context["measurement_area_m2"]

    render_plot_map_layout(
        db=db,
        plot_id=plot_id,
        output_path=map_path,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        scale_text=scale_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        certification_statement=certification_statement,
        station_names=station_names if station_names else None,
        coordinate_system=render_coordinate_system,
        epsg_code=epsg_code,
        crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
        measurement_polygon=measurement_polygon,
        measurement_area_m2=measurement_area_m2,
        paper_size=paper_size,
        north_arrow_style=north_arrow_style,
        north_arrow_color=north_arrow_color,
        beacon_style=beacon_style,
        road_width_m=road_width_m,
        road_width_override_m=road_width_override_m,
        boundary_color=boundary_color,
        grid_color=grid_color,
        text_color=text_color,
        road_color=road_color,
        river_color=river_color,
        building_color=building_color,
        building_hatch_type=building_hatch_type,
        road_style=road_style,
        title_font=title_font,
        title_size=title_size,
        grid_font=grid_font,
        grid_size=grid_size,
        station_font=station_font,
        station_size=station_size,
        bearing_font=bearing_font,
        bearing_size=bearing_size,
        area_font=area_font,
        area_size=area_size,
        template_name=template_name,
        adamawa_rof_no=adamawa_rof_no,
        adamawa_owner_name=adamawa_owner_name,
        adamawa_authority_title=adamawa_authority_title,
        adamawa_authority_date_text=adamawa_authority_date_text,
        adamawa_control_point_name=adamawa_control_point_name,
        adamawa_northing=adamawa_northing,
        adamawa_easting=adamawa_easting,
        adamawa_elevation=adamawa_elevation,
        adamawa_origin_text=adamawa_origin_text,
        adamawa_topo_sheet_text=adamawa_topo_sheet_text,
        adamawa_computation_no=adamawa_computation_no,
        adamawa_cadastral_sheet_no=adamawa_cadastral_sheet_no,
        adamawa_plan_no=adamawa_plan_no,
        adamawa_surveyed_by_text=adamawa_surveyed_by_text,
        adamawa_disclaimer_text=adamawa_disclaimer_text,
        cadastral_plan_no=cadastral_plan_no,
        cadastral_area_name=cadastral_area_name,
        cadastral_datum_text=cadastral_datum_text,
        cadastral_firm_block_text=cadastral_firm_block_text,
        fct_file_no=fct_file_no,
        fct_district=fct_district,
        fct_cadastral_zone=fct_cadastral_zone,
        fct_origin_beacon_text=fct_origin_beacon_text,
        fct_cadastral_map_ref=fct_cadastral_map_ref,
        fct_title_prefix=fct_title_prefix,
    )

    report = get_plot_report(plot_id, db)
    generate_plot_report_pdf(report, pdf_path, map_path, paper_size=paper_size)

    safe_remove(map_path)

    filename = _build_plot_export_filename(
        plot_id, {"title_text": title_text, "location_text": location_text}, "Survey_Plan", "pdf",
    )
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="survey-plan",
        project_id=plot_id,
    )


def download_plot_technical_report_docx(plot_id: int, db: Session = Depends(get_db), **payload):
    """Generates the Adamawa OSG "Survey Technical Report" .docx, auto-filled from plot_meta
    (persisted via the same upsert used by every other export) plus the live computed area.
    """
    reports_dir = REPORTS_DIR
    os.makedirs(reports_dir, exist_ok=True)
    docx_path = f"{reports_dir}/plot_{plot_id}_technical_report.docx"

    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        title_text=payload.get("title_text"),
        location_text=payload.get("location_text"),
        lga_text=payload.get("lga_text"),
        state_text=payload.get("state_text"),
        surveyor_name=payload.get("surveyor_name"),
        surveyor_rank=payload.get("surveyor_rank"),
        template_name=payload.get("template_name"),
        adamawa_rof_no=payload.get("adamawa_rof_no"),
        adamawa_owner_name=payload.get("adamawa_owner_name"),
        adamawa_authority_title=payload.get("adamawa_authority_title"),
        adamawa_authority_date_text=payload.get("adamawa_authority_date_text"),
        adamawa_control_point_name=payload.get("adamawa_control_point_name"),
        adamawa_northing=payload.get("adamawa_northing"),
        adamawa_easting=payload.get("adamawa_easting"),
        adamawa_elevation=payload.get("adamawa_elevation"),
        adamawa_origin_text=payload.get("adamawa_origin_text"),
        adamawa_topo_sheet_text=payload.get("adamawa_topo_sheet_text"),
        adamawa_computation_no=payload.get("adamawa_computation_no"),
        adamawa_cadastral_sheet_no=payload.get("adamawa_cadastral_sheet_no"),
        adamawa_plan_no=payload.get("adamawa_plan_no"),
        adamawa_surveyed_by_text=payload.get("adamawa_surveyed_by_text"),
        adamawa_disclaimer_text=payload.get("adamawa_disclaimer_text"),
        technical_report_instruments=payload.get("technical_report_instruments"),
        technical_report_dgps_type=payload.get("technical_report_dgps_type"),
        technical_report_num_surveyors=payload.get("technical_report_num_surveyors"),
        technical_report_num_technical_officers=payload.get("technical_report_num_technical_officers"),
        technical_report_num_labourers=payload.get("technical_report_num_labourers"),
        technical_report_recce_text=payload.get("technical_report_recce_text"),
        technical_report_demarcation_text=payload.get("technical_report_demarcation_text"),
        technical_report_computation_software_text=payload.get("technical_report_computation_software_text"),
        technical_report_plotting_software_text=payload.get("technical_report_plotting_software_text"),
        technical_report_general_observation_text=payload.get("technical_report_general_observation_text"),
    )

    meta = get_plot_meta(db, plot_id)
    plot_geom = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom,
        meta.get("coordinate_system"),
        meta.get("survey_input_coordinates"),
    )
    area_m2 = measurement_context["measurement_area_m2"] or 0

    _render_technical_report_docx(meta, float(area_m2), docx_path)

    filename = _build_plot_export_filename(plot_id, meta, "Technical_Report", "docx")
    return _docx_response_with_r2(
        docx_path,
        filename,
        category="technical-report",
        project_id=plot_id,
    )


@router.post("/{plot_id}/export-jobs/technical-report.docx")
def create_plot_technical_report_export_job(
    plot_id: int,
    request: Request,
    db: Session = Depends(get_db),
    title_text: str = Body("SURVEY PLAN"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    adamawa_rof_no: str = Body(""),
    adamawa_owner_name: str = Body(""),
    adamawa_authority_title: str = Body(DEFAULT_ADAMAWA_AUTHORITY_TITLE),
    adamawa_authority_date_text: str = Body(DEFAULT_ADAMAWA_AUTHORITY_DATE),
    adamawa_control_point_name: str = Body(""),
    adamawa_northing: str = Body(""),
    adamawa_easting: str = Body(""),
    adamawa_elevation: str = Body(""),
    adamawa_origin_text: str = Body(DEFAULT_ADAMAWA_ORIGIN_TEXT),
    adamawa_topo_sheet_text: str = Body(DEFAULT_ADAMAWA_TOPO_SHEET_TEXT),
    adamawa_computation_no: str = Body(""),
    adamawa_cadastral_sheet_no: str = Body(""),
    adamawa_plan_no: str = Body(""),
    adamawa_surveyed_by_text: str = Body(""),
    adamawa_disclaimer_text: str = Body(DEFAULT_ADAMAWA_DISCLAIMER_TEXT),
    technical_report_instruments: list[str] = Body(default=[]),
    technical_report_dgps_type: str = Body(""),
    technical_report_num_surveyors: int | None = Body(None),
    technical_report_num_technical_officers: int | None = Body(None),
    technical_report_num_labourers: int | None = Body(None),
    technical_report_recce_text: str = Body(""),
    technical_report_demarcation_text: str = Body(""),
    technical_report_computation_software_text: str = Body(""),
    technical_report_plotting_software_text: str = Body(""),
    technical_report_general_observation_text: str = Body(""),
):
    request_payload = {
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "template_name": template_name,
        "adamawa_rof_no": adamawa_rof_no,
        "adamawa_owner_name": adamawa_owner_name,
        "adamawa_authority_title": adamawa_authority_title,
        "adamawa_authority_date_text": adamawa_authority_date_text,
        "adamawa_control_point_name": adamawa_control_point_name,
        "adamawa_northing": adamawa_northing,
        "adamawa_easting": adamawa_easting,
        "adamawa_elevation": adamawa_elevation,
        "adamawa_origin_text": adamawa_origin_text,
        "adamawa_topo_sheet_text": adamawa_topo_sheet_text,
        "adamawa_computation_no": adamawa_computation_no,
        "adamawa_cadastral_sheet_no": adamawa_cadastral_sheet_no,
        "adamawa_plan_no": adamawa_plan_no,
        "adamawa_surveyed_by_text": adamawa_surveyed_by_text,
        "adamawa_disclaimer_text": adamawa_disclaimer_text,
        "technical_report_instruments": technical_report_instruments or [],
        "technical_report_dgps_type": technical_report_dgps_type,
        "technical_report_num_surveyors": technical_report_num_surveyors,
        "technical_report_num_technical_officers": technical_report_num_technical_officers,
        "technical_report_num_labourers": technical_report_num_labourers,
        "technical_report_recce_text": technical_report_recce_text,
        "technical_report_demarcation_text": technical_report_demarcation_text,
        "technical_report_computation_software_text": technical_report_computation_software_text,
        "technical_report_plotting_software_text": technical_report_plotting_software_text,
        "technical_report_general_observation_text": technical_report_general_observation_text,
    }
    cache_key = _build_plot_export_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="technical-report.docx",
        payload=request_payload,
    )
    existing = _find_plot_export_job_by_cache_key(
        db,
        plot_id=int(plot_id),
        export_type="technical-report.docx",
        cache_key=cache_key,
    )
    if existing:
        return _serialize_plot_export_job(existing, request=request)
    job = _insert_single_plot_export_job(
        db,
        plot_id=int(plot_id),
        export_type="technical-report.docx",
        cache_key=cache_key,
        file_name=_build_plot_export_filename(plot_id, request_payload, "Technical_Report", "docx"),
        request_payload=request_payload,
    )
    return _serialize_plot_export_job(job, request=request)


# ---------------- SIMPLE PDF DOWNLOAD (GET) ----------------

@router.get("/{plot_id}/download/pdf")
def simple_download_pdf(plot_id: int, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None):
    """Simple GET endpoint for basic PDF download from dashboard"""
    reports_dir = REPORTS_DIR
    maps_dir = os.path.join(REPORTS_DIR, "maps")
    os.makedirs(reports_dir, exist_ok=True)
    os.makedirs(maps_dir, exist_ok=True)

    pdf_path = f"{reports_dir}/plot_{plot_id}_report.pdf"

    tmp_map = tempfile.NamedTemporaryFile(suffix="_map.png", delete=False)
    map_path = tmp_map.name
    tmp_map.close()

    # Use default values
    render_plot_map_layout(
        db=db,
        plot_id=plot_id,
        output_path=map_path,
        title_text="SURVEY PLAN",
        location_text="",
        lga_text="",
        state_text="",
        scale_text="1 : 1000",
        surveyor_name="",
        surveyor_rank="",
        station_names=None,
        coordinate_system="wgs84",
        epsg_code=4326,
        crs_footer_text="COORDINATE SYSTEM: WGS84"
    )

    report = get_plot_report(plot_id, db)
    generate_plot_report_pdf(report, pdf_path, map_path)

    safe_remove(map_path)

    filename = _build_plot_export_filename(plot_id, get_plot_meta(db, plot_id), "Survey_Plan", "pdf")
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="survey-plan",
        project_id=plot_id,
    )


# ---------------- SURVEY PLAN PREVIEW ----------------

@router.post("/{plot_id}/scale-recommendation")
def get_plot_scale_recommendation(
    plot_id: int,
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    survey_input_coordinates: Optional[list] = Body(default=None),
    db: Session = Depends(get_db),
):
    """Recommend a standard scale before the first plan render.

    The result uses the same metric measurement polygon and template map frame as the renderer.
    The selected denominator is always rounded upward, never down, so a standard scale cannot
    crop the parcel merely to look tidier on the title block.
    """
    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom_wgs84,
        coordinate_system,
        survey_input_coordinates=survey_input_coordinates,
    )
    plot_metric = measurement_context["measurement_polygon"]
    map_width_fraction, map_height_fraction = _survey_template_map_frame(template_name)

    def recommendation_for(size: str) -> tuple[int, int]:
        config = get_paper_config(size)
        fitted = compute_fit_scale_ratio(
            plot_metric,
            config["width"] * map_width_fraction,
            config["height"] * map_height_fraction,
        )
        return fitted, _select_standard_survey_scale(fitted)

    requested_paper = str(paper_size or "A4").upper()
    if requested_paper not in {"A4", "A3", "A2", "A1", "A0"}:
        requested_paper = "A4"
    fitted_ratio, standard_ratio = recommendation_for(requested_paper)
    recommended_paper = requested_paper

    # A3 is recommended only where an A4 plan would need a small-scale 1:5000+ sheet. This
    # keeps ordinary plots on A4 while making large plots more readable without forcing A3.
    if requested_paper == "A4" and standard_ratio >= 5000:
        a3_fitted, a3_standard = recommendation_for("A3")
        if a3_standard < standard_ratio:
            recommended_paper = "A3"
            fitted_ratio, standard_ratio = a3_fitted, a3_standard

    return {
        "paper_size": recommended_paper,
        "scale_text": f"1 : {standard_ratio}",
        "scale_denominator": standard_ratio,
        "fitted_scale_denominator": fitted_ratio,
        "template_name": str(template_name or DEFAULT_TEMPLATE_NAME),
        "reason": "The selected standard scale contains the full parcel in the template map frame.",
    }


@router.post("/{plot_id}/report/preview")
def preview_plot_map(plot_id: int, request: Request, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None,
    title_text: str = Body("SURVEY PLAN"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("auto"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    certification_statement: str = Body(DEFAULT_CERTIFICATION_STATEMENT),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue"),
    beacon_style: str = Body("cross"),
    road_width_m: float | None = Body(None),
    road_width_override_m: float | None = Body(None),
    boundary_color: str | None = Body(None),
    grid_color: str | None = Body(None),
    text_color: str | None = Body(None),
    road_color: str | None = Body(None),
    river_color: str | None = Body(None),
    building_color: str | None = Body(None),
    building_hatch_type: str | None = Body(None),
    road_style: str | None = Body(None),
    title_font: str | None = Body(None),
    title_size: int | None = Body(None),
    grid_font: str | None = Body(None),
    grid_size: int | None = Body(None),
    station_font: str | None = Body(None),
    station_size: int | None = Body(None),
    bearing_font: str | None = Body(None),
    bearing_size: int | None = Body(None),
    area_font: str | None = Body(None),
    area_size: int | None = Body(None),
    template_name: str = Body(DEFAULT_TEMPLATE_NAME),
    adamawa_rof_no: str = Body(""),
    adamawa_owner_name: str = Body(""),
    adamawa_authority_title: str = Body(DEFAULT_ADAMAWA_AUTHORITY_TITLE),
    adamawa_authority_date_text: str = Body(DEFAULT_ADAMAWA_AUTHORITY_DATE),
    adamawa_control_point_name: str = Body(""),
    adamawa_northing: str = Body(""),
    adamawa_easting: str = Body(""),
    adamawa_elevation: str = Body(""),
    adamawa_origin_text: str = Body(DEFAULT_ADAMAWA_ORIGIN_TEXT),
    adamawa_topo_sheet_text: str = Body(DEFAULT_ADAMAWA_TOPO_SHEET_TEXT),
    adamawa_computation_no: str = Body(""),
    adamawa_cadastral_sheet_no: str = Body(""),
    adamawa_plan_no: str = Body(""),
    adamawa_surveyed_by_text: str = Body(""),
    adamawa_disclaimer_text: str = Body(DEFAULT_ADAMAWA_DISCLAIMER_TEXT),
    cadastral_plan_no: str = Body(""),
    cadastral_area_name: str = Body(""),
    cadastral_datum_text: str = Body(""),
    cadastral_firm_block_text: str = Body(""),
    fct_file_no: str = Body(""),
    fct_district: str = Body(""),
    fct_cadastral_zone: str = Body(""),
    fct_origin_beacon_text: str = Body(""),
    fct_cadastral_map_ref: str = Body(""),
    fct_title_prefix: str = Body(""),
    survey_input_coordinates: Optional[list] = Body(default=None)):

    effective_scale_text = str(scale_text or "").strip() or "auto"
    resolved_scale_text = effective_scale_text

    plot_geom_hex = db.execute(
        text("SELECT encode(ST_AsBinary(geom), 'hex') FROM plots WHERE id=:id"),
        {"id": plot_id},
    ).scalar()
    if not plot_geom_hex:
        raise HTTPException(status_code=404, detail="Plot not found")
    try:
        plot_geom_wgs84 = wkb.loads(plot_geom_hex, hex=True)
    except GEOSException:
        raise HTTPException(
            status_code=400,
            detail="Plot geometry is invalid. Re-save the parcel boundary and try preview again.",
        )
    if plot_geom_wgs84 is None or plot_geom_wgs84.is_empty:
        raise HTTPException(status_code=400, detail="Plot geometry is empty.")

    measurement_context = _resolve_measurement_polygon_context(
        plot_geom_wgs84,
        coordinate_system,
        survey_input_coordinates=survey_input_coordinates,
    )
    render_coordinate_system = measurement_context["render_coordinate_system"]
    effective_epsg = measurement_context["epsg_code"]
    crs_name = measurement_context["crs_name"]
    plot_metric = measurement_context["measurement_polygon"]
    measurement_area_m2 = measurement_context["measurement_area_m2"]

    paper_config = get_paper_config(paper_size)
    map_width_fraction, map_height_fraction = _survey_template_map_frame(template_name)
    if is_auto_scale_text(effective_scale_text):
        fitted_ratio = compute_fit_scale_ratio(
            plot_metric,
            paper_config["width"] * map_width_fraction,
            paper_config["height"] * map_height_fraction,
        )
        resolved_scale_text = f"1 : {fitted_ratio}"
    else:
        resolved_scale_text = f"1 : {parse_scale_ratio(effective_scale_text)}"

    payload_for_cache = {
        "_layout_version": PREVIEW_LAYOUT_VERSION,
        "title_text": title_text,
        "location_text": location_text,
        "lga_text": lga_text,
        "state_text": state_text,
        "scale_text": resolved_scale_text,
        "surveyor_name": surveyor_name,
        "surveyor_rank": surveyor_rank,
        "certification_statement": certification_statement,
        "station_names": station_names or [],
        "coordinate_system": coordinate_system,
        "paper_size": paper_size,
        "north_arrow_style": north_arrow_style,
        "north_arrow_color": north_arrow_color,
        "beacon_style": beacon_style,
        "road_width_m": road_width_m,
        "road_width_override_m": road_width_override_m,
        "boundary_color": boundary_color,
        "grid_color": grid_color,
        "text_color": text_color,
        "road_color": road_color,
        "river_color": river_color,
        "building_color": building_color,
        "building_hatch_type": building_hatch_type,
        "road_style": road_style,
        "title_font": title_font,
        "title_size": title_size,
        "grid_font": grid_font,
        "grid_size": grid_size,
        "station_font": station_font,
        "station_size": station_size,
        "bearing_font": bearing_font,
        "bearing_size": bearing_size,
        "area_font": area_font,
        "area_size": area_size,
        "template_name": template_name,
        "adamawa_rof_no": adamawa_rof_no,
        "adamawa_owner_name": adamawa_owner_name,
        "adamawa_authority_title": adamawa_authority_title,
        "adamawa_authority_date_text": adamawa_authority_date_text,
        "adamawa_control_point_name": adamawa_control_point_name,
        "adamawa_northing": adamawa_northing,
        "adamawa_easting": adamawa_easting,
        "adamawa_elevation": adamawa_elevation,
        "adamawa_origin_text": adamawa_origin_text,
        "adamawa_topo_sheet_text": adamawa_topo_sheet_text,
        "adamawa_computation_no": adamawa_computation_no,
        "adamawa_cadastral_sheet_no": adamawa_cadastral_sheet_no,
        "adamawa_plan_no": adamawa_plan_no,
        "adamawa_surveyed_by_text": adamawa_surveyed_by_text,
        "adamawa_disclaimer_text": adamawa_disclaimer_text,
        "cadastral_plan_no": cadastral_plan_no,
        "cadastral_area_name": cadastral_area_name,
        "cadastral_datum_text": cadastral_datum_text,
        "cadastral_firm_block_text": cadastral_firm_block_text,
        "fct_file_no": fct_file_no,
        "fct_district": fct_district,
        "fct_cadastral_zone": fct_cadastral_zone,
        "fct_origin_beacon_text": fct_origin_beacon_text,
        "fct_cadastral_map_ref": fct_cadastral_map_ref,
        "fct_title_prefix": fct_title_prefix,
        "survey_input_coordinates": survey_input_coordinates or [],
    }
    revision_token = build_preview_revision_token(db, plot_id)
    cache_key = build_preview_cache_key(plot_id, payload_for_cache, revision_token)
    prune_preview_cache(plot_id, variant="survey")
    cached_path = get_cached_preview_path(plot_id, cache_key, variant="survey")
    if cached_path:
        return FileResponse(
            cached_path,
            media_type="image/png",
            headers={
                "Cache-Control": "no-store",
                "X-LandCheck-Resolved-Scale": resolved_scale_text,
            },
        )

    cleanup_preview_files(plot_id)
    tmp_map = tempfile.NamedTemporaryFile(suffix="_preview.png", delete=False)
    map_path = tmp_map.name
    tmp_map.close()

    # Save/refresh plot metadata
    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        certification_statement=certification_statement,
        scale_text=resolved_scale_text,
        paper_size=paper_size,
        coordinate_system=coordinate_system,
        template_name=template_name,
        adamawa_rof_no=adamawa_rof_no,
        adamawa_owner_name=adamawa_owner_name,
        adamawa_authority_title=adamawa_authority_title,
        adamawa_authority_date_text=adamawa_authority_date_text,
        adamawa_control_point_name=adamawa_control_point_name,
        adamawa_northing=adamawa_northing,
        adamawa_easting=adamawa_easting,
        adamawa_elevation=adamawa_elevation,
        adamawa_origin_text=adamawa_origin_text,
        adamawa_topo_sheet_text=adamawa_topo_sheet_text,
        adamawa_computation_no=adamawa_computation_no,
        adamawa_cadastral_sheet_no=adamawa_cadastral_sheet_no,
        adamawa_plan_no=adamawa_plan_no,
        adamawa_surveyed_by_text=adamawa_surveyed_by_text,
        adamawa_disclaimer_text=adamawa_disclaimer_text,
        cadastral_plan_no=cadastral_plan_no,
        cadastral_area_name=cadastral_area_name,
        cadastral_datum_text=cadastral_datum_text,
        cadastral_firm_block_text=cadastral_firm_block_text,
        fct_file_no=fct_file_no,
        fct_district=fct_district,
        fct_cadastral_zone=fct_cadastral_zone,
        fct_origin_beacon_text=fct_origin_beacon_text,
        fct_cadastral_map_ref=fct_cadastral_map_ref,
        fct_title_prefix=fct_title_prefix,
        survey_input_coordinates=survey_input_coordinates,
    )

    epsg_code = effective_epsg

    render_plot_map_layout(
        db=db,
        plot_id=plot_id,
        output_path=map_path,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        scale_text=resolved_scale_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        certification_statement=certification_statement,
        station_names=station_names if station_names else None,
        coordinate_system=render_coordinate_system,
        epsg_code=epsg_code,
        crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
        measurement_polygon=plot_metric,
        measurement_area_m2=measurement_area_m2,
        paper_size=paper_size,
        north_arrow_style=north_arrow_style,
        north_arrow_color=north_arrow_color,
        beacon_style=beacon_style,
        road_width_m=road_width_m,
        road_width_override_m=road_width_override_m,
        boundary_color=boundary_color,
        grid_color=grid_color,
        text_color=text_color,
        road_color=road_color,
        river_color=river_color,
        building_color=building_color,
        building_hatch_type=building_hatch_type,
        road_style=road_style,
        title_font=title_font,
        title_size=title_size,
        grid_font=grid_font,
        grid_size=grid_size,
        station_font=station_font,
        station_size=station_size,
        bearing_font=bearing_font,
        bearing_size=bearing_size,
        area_font=area_font,
        area_size=area_size,
        preview_mode=True,
        template_name=template_name,
        adamawa_rof_no=adamawa_rof_no,
        adamawa_owner_name=adamawa_owner_name,
        adamawa_authority_title=adamawa_authority_title,
        adamawa_authority_date_text=adamawa_authority_date_text,
        adamawa_control_point_name=adamawa_control_point_name,
        adamawa_northing=adamawa_northing,
        adamawa_easting=adamawa_easting,
        adamawa_elevation=adamawa_elevation,
        adamawa_origin_text=adamawa_origin_text,
        adamawa_topo_sheet_text=adamawa_topo_sheet_text,
        adamawa_computation_no=adamawa_computation_no,
        adamawa_cadastral_sheet_no=adamawa_cadastral_sheet_no,
        adamawa_plan_no=adamawa_plan_no,
        adamawa_surveyed_by_text=adamawa_surveyed_by_text,
        adamawa_disclaimer_text=adamawa_disclaimer_text,
        cadastral_plan_no=cadastral_plan_no,
        cadastral_area_name=cadastral_area_name,
        cadastral_datum_text=cadastral_datum_text,
        cadastral_firm_block_text=cadastral_firm_block_text,
        fct_file_no=fct_file_no,
        fct_district=fct_district,
        fct_cadastral_zone=fct_cadastral_zone,
        fct_origin_beacon_text=fct_origin_beacon_text,
        fct_cadastral_map_ref=fct_cadastral_map_ref,
        fct_title_prefix=fct_title_prefix,
    )

    cache_path = preview_cache_path(plot_id, cache_key, variant="survey")
    served_path = map_path
    served_background = background_tasks
    try:
        shutil.copyfile(map_path, cache_path)
        safe_remove(map_path)
        served_path = cache_path
        served_background = None
    except Exception:
        pass

    if served_path == map_path:
        if served_background is None:
            served_background = BackgroundTasks()
        served_background.add_task(safe_remove, map_path)

    log_survey_activity(
        db,
        event_type="preview_completed",
        workflow="survey_plan",
        request=request,
        plot_id=plot_id,
        details={"template_name": template_name, "paper_size": paper_size, "scale": resolved_scale_text},
    )

    return FileResponse(
        served_path,
        media_type="image/png",
        headers={
            "Cache-Control": "no-store",
            "X-LandCheck-Resolved-Scale": resolved_scale_text,
        },
        background=served_background,
    )


# ---------------- BACK COMPUTATION ----------------

@router.post("/{plot_id}/back-computation/pdf")
def download_back_computation_pdf(plot_id: int, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None,
    coordinate_system: str = Body("wgs84"),
    station_names: list[str] = Body(default=[]),
    survey_input_coordinates: Optional[list] = Body(default=None)):

    reports_dir = REPORTS_DIR
    os.makedirs(reports_dir, exist_ok=True)

    pdf_path = f"{reports_dir}/plot_{plot_id}_back_computation.pdf"

    meta = get_plot_meta(db, plot_id)

    plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id=:id"), {"id": plot_id}).scalar()
    plot_geom = wkb.loads(plot_wkb)

    measurement_context = _resolve_measurement_polygon_context(
        plot_geom,
        coordinate_system or meta.get("coordinate_system"),
        survey_input_coordinates=(
            survey_input_coordinates
            if survey_input_coordinates is not None
            else meta.get("survey_input_coordinates")
        ),
    )
    poly = measurement_context["measurement_polygon"]
    area_m2 = measurement_context["measurement_area_m2"]
    crs_name = measurement_context["crs_name"]

    # Use custom station names if provided
    labels = station_names if station_names else None

    rows, sum_de, sum_dn = compute_back_computation(poly, labels)

    render_back_computation_pdf(rows, sum_de, sum_dn, area_m2, plot_id, pdf_path, crs_name)

    filename = _build_plot_export_filename(plot_id, meta, "Back_Computation", "pdf")
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="back-computation",
        project_id=plot_id,
    )


# ---------------- ORTHOPHOTO ----------------

@router.post("/{plot_id}/orthophoto/preview")
def orthophoto_preview(plot_id: int, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None,
    scale_text: str = Body("1 : 1000"),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    use_topo_map: bool = Body(False),
    topo_source: str = Body("opentopomap"),
    elevation_points: list = Body(default=[]),
    contour_interval: float | None = Body(None),
    building_hatch_type: str = Body("solid"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue")):

    topo_source = topo_source or "opentopomap"
    building_hatch_type = building_hatch_type or "solid"
    payload_for_cache = {
        "scale_text": scale_text,
        "station_names": station_names or [],
        "coordinate_system": coordinate_system,
        "paper_size": paper_size,
        "use_topo_map": bool(use_topo_map),
        "topo_source": topo_source,
        "elevation_points": elevation_points if topo_source == "userdata" else [],
        "contour_interval": contour_interval,
        "building_hatch_type": building_hatch_type,
        "north_arrow_style": north_arrow_style,
        "north_arrow_color": north_arrow_color,
    }
    revision_token = build_plot_geom_revision_token(db, plot_id)
    cache_key = build_preview_cache_key(plot_id, payload_for_cache, revision_token)
    cache_variant = "topomap" if use_topo_map else "orthophoto"
    prune_preview_cache(plot_id, variant=cache_variant)
    cached_path = get_cached_preview_path(plot_id, cache_key, variant=cache_variant, extension="jpg")
    if cached_path:
        return FileResponse(
            cached_path,
            media_type="image/jpeg",
            headers={"Cache-Control": "no-store"},
        )

    cleanup_preview_files(plot_id)
    tmp_png = tempfile.NamedTemporaryFile(suffix="_orthophoto_preview.jpg", delete=False)
    png_path = tmp_png.name
    tmp_png.close()

    # Save/refresh plot metadata (scale, paper size, coord system). Elevation points are only
    # persisted when the caller actually sent some - an empty/omitted list must never overwrite
    # previously-uploaded height data (same NULL-means-"don't touch" convention as every other
    # plot_meta field here).
    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        scale_text=scale_text,
        paper_size=paper_size,
        coordinate_system=coordinate_system,
        elevation_points=elevation_points if (topo_source == "userdata" and elevation_points) else None,
    )

    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    render_coordinate_system, epsg_code, crs_name = _resolve_survey_render_crs(
        coordinate_system,
        plot_geom_wgs84,
    )

    persisted_elevation_points = elevation_points
    if use_topo_map and topo_source == "userdata" and not persisted_elevation_points:
        persisted_elevation_points = get_plot_meta(db, plot_id).get("elevation_points") or []

    if use_topo_map:
        topo_source_footer = "SOURCE: Your uploaded elevation data" if topo_source == "userdata" else "SOURCE: Global satellite elevation model (Copernicus/NASA DEM)"
    else:
        topo_source_footer = "SOURCE: Satellite Imagery"

    render_orthophoto_png(
        db=db,
        plot_id=plot_id,
        output_path=png_path,
        title_text="TOPO MAP" if use_topo_map else "ORTHOPHOTO",
        scale_text=scale_text,
        station_names=station_names if station_names else None,
        coordinate_system=render_coordinate_system,
        epsg_code=epsg_code,
        crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
        source_footer_text=topo_source_footer,
        use_topo_map=use_topo_map,
        topo_source=topo_source,
        elevation_points=persisted_elevation_points,
        contour_interval=contour_interval,
        building_hatch_type=building_hatch_type,
        paper_size=paper_size,
        north_arrow_style=north_arrow_style,
        north_arrow_color=north_arrow_color,
        preview_mode=True,
    )

    cache_path = preview_cache_path(plot_id, cache_key, variant=cache_variant, extension="jpg")
    served_path = png_path
    served_background = background_tasks
    try:
        shutil.copyfile(png_path, cache_path)
        safe_remove(png_path)
        served_path = cache_path
        served_background = None
    except Exception:
        pass

    if served_path == png_path:
        if served_background is None:
            served_background = BackgroundTasks()
        served_background.add_task(safe_remove, png_path)

    return FileResponse(
        served_path,
        media_type="image/jpeg",
        headers={"Cache-Control": "no-store"},
        background=served_background,
    )


@router.post("/{plot_id}/orthophoto/pdf")
def orthophoto_pdf(plot_id: int, db: Session = Depends(get_db), background_tasks: BackgroundTasks = None,
    title_text: str = Body("ORTHOPHOTO"),
    location_text: str = Body(""),
    lga_text: str = Body(""),
    state_text: str = Body(""),
    scale_text: str = Body("1 : 1000"),
    surveyor_name: str = Body(""),
    surveyor_rank: str = Body(""),
    station_names: list[str] = Body(default=[]),
    coordinate_system: str = Body("wgs84"),
    paper_size: str = Body("A4"),
    use_topo_map: bool = Body(False),
    topo_source: str = Body("opentopomap"),
    contour_interval: float | None = Body(None),
    building_hatch_type: str = Body("solid"),
    north_arrow_style: str = Body("one_side_stem"),
    north_arrow_color: str = Body("blue")):

    topo_source = topo_source or "opentopomap"
    building_hatch_type = building_hatch_type or "solid"
    out_dir = os.path.join(REPORTS_DIR, "orthophoto")
    os.makedirs(out_dir, exist_ok=True)

    # Use different filename for topo vs satellite (response name only)
    map_type = "topo" if use_topo_map else "satellite"
    pdf_path = f"{out_dir}/plot_{plot_id}_orthophoto_{map_type}.pdf"

    tmp_png = tempfile.NamedTemporaryFile(suffix=f"_{map_type}.png", delete=False)
    png_path = tmp_png.name
    tmp_png.close()

    # Save/refresh plot metadata (scale, paper size, coord system)
    upsert_plot_meta(
        db=db,
        plot_id=plot_id,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        scale_text=scale_text,
        paper_size=paper_size,
        coordinate_system=coordinate_system,
    )

    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    render_coordinate_system, epsg_code, crs_name = _resolve_survey_render_crs(
        coordinate_system,
        plot_geom_wgs84,
    )

    elevation_points = None
    if use_topo_map and topo_source == "userdata":
        elevation_points = get_plot_meta(db, plot_id).get("elevation_points") or None

    if use_topo_map:
        topo_source_footer = "SOURCE: Your uploaded elevation data" if topo_source == "userdata" else "SOURCE: Global satellite elevation model (Copernicus/NASA DEM)"
    else:
        topo_source_footer = "SOURCE: Satellite Imagery"

    render_orthophoto_png(
        db=db,
        plot_id=plot_id,
        output_path=png_path,
        title_text=title_text,
        location_text=location_text,
        lga_text=lga_text,
        state_text=state_text,
        scale_text=scale_text,
        surveyor_name=surveyor_name,
        surveyor_rank=surveyor_rank,
        source_footer_text=topo_source_footer,
        station_names=station_names if station_names else None,
        coordinate_system=render_coordinate_system,
        epsg_code=epsg_code,
        crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
        use_topo_map=use_topo_map,
        topo_source=topo_source,
        elevation_points=elevation_points,
        contour_interval=contour_interval,
        building_hatch_type=building_hatch_type,
        paper_size=paper_size,
        north_arrow_style=north_arrow_style,
        north_arrow_color=north_arrow_color,
    )

    render_orthophoto_pdf_from_png(png_path, pdf_path, paper_size=paper_size)

    filename = _build_plot_export_filename(
        plot_id, {"title_text": title_text, "location_text": location_text},
        "Topo_Map" if use_topo_map else "Orthophoto", "pdf",
    )
    safe_remove(png_path)

    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="orthophoto" if not use_topo_map else "topo-map",
        project_id=plot_id,
    )
@router.get("/{plot_id}/survey-plan/dwg")
def download_survey_plan_dwg(plot_id: int, db: Session = Depends(get_db)):
    meta = get_plot_meta(db, plot_id)

    out_dir = os.path.join(REPORTS_DIR, "dwg")
    os.makedirs(out_dir, exist_ok=True)

    dxf_path = f"{out_dir}/plot_{plot_id}_survey_plan.dxf"

    # Build the boundary the same way the PDF/back-computation exports do - directly from the
    # surveyor's original input coordinates, reprojected once - so the DXF's bearings/distances/
    # area match the PDF (and AutoCAD, since that's what a surveyor would drive Civil 3D from)
    # instead of round-tripping through the plot's stored WGS84 geometry a second time.
    plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
    measurement_context = _resolve_measurement_polygon_context(
        plot_geom_wgs84,
        meta.get("coordinate_system") or "wgs84",
        meta.get("survey_input_coordinates"),
    )

    export_survey_plan_to_dxf(
        db,
        plot_id,
        dxf_path,
        coordinate_system=meta.get("coordinate_system") or "wgs84",
        measurement_polygon=measurement_context["measurement_polygon"],
        export_epsg_override=measurement_context["epsg_code"],
    )

    return FileResponse(
        dxf_path,
        media_type="application/dxf",
        filename=_build_plot_export_filename(plot_id, meta, "Survey_Plan", "dxf"),
    )


@router.get("/{plot_id}/survey-plan/shapefile")
def download_survey_plan_shapefile(
    plot_id: int,
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = None,
):
    plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id=:id"), {"id": plot_id}).scalar()
    if not plot_wkb:
        raise HTTPException(status_code=404, detail="Plot not found")

    meta = get_plot_meta(db, plot_id)
    plot_geom = wkb.loads(plot_wkb)
    feature_sets = get_plot_features_geojson(plot_id, db)

    tmp_dir = tempfile.mkdtemp(prefix=f"plot_{plot_id}_shp_")
    export_filename = _build_plot_export_filename(plot_id, meta, "Survey_Plan_Shapefile", "zip")
    zip_path = os.path.join(tmp_dir, export_filename)

    try:
        def station_name(idx: int) -> str:
            name = ""
            num = idx
            while True:
                name = chr(65 + (num % 26)) + name
                num = (num // 26) - 1
                if num < 0:
                    break
            return name

        def write_layer(layer_suffix: str, records: list[dict]):
            if not records:
                return
            layer_base = f"plot_{plot_id}_{layer_suffix}"
            layer_path = os.path.join(tmp_dir, f"{layer_base}.shp")
            gdf = gpd.GeoDataFrame(records, crs="EPSG:4326")
            if gdf.empty:
                return
            gdf.to_file(layer_path, driver="ESRI Shapefile")

        # 1) Plot boundary polygon (main parcel geometry)
        write_layer("boundary", [{
            "plot_id": int(plot_id),
            "title": (meta.get("title_text") or "SURVEY PLAN")[:80],
            "loc_text": (meta.get("location_text") or "")[:80],
            "lga_text": (meta.get("lga_text") or "")[:80],
            "state_txt": (meta.get("state_text") or "")[:80],
            "surveyr": (meta.get("surveyor_name") or "")[:80],
            "rank_txt": (meta.get("surveyor_rank") or "")[:80],
            "scale_txt": (meta.get("scale_text") or "")[:40],
            "geometry": plot_geom,
        }])

        # 2) Station/vertex points (A, B, C...)
        vertex_records = []
        coords = list(getattr(plot_geom, "exterior", plot_geom).coords)
        if coords and len(coords) > 1 and coords[0] == coords[-1]:
            coords = coords[:-1]
        for idx, (lng, lat, *_) in enumerate(coords):
            vertex_records.append({
                "plot_id": int(plot_id),
                "stn": station_name(idx),
                "ord_no": idx + 1,
                "lng": float(lng),
                "lat": float(lat),
                "geometry": Point(float(lng), float(lat)),
            })
        write_layer("stations", vertex_records)

        # 3) Map features shown in plan (roads/buildings/rivers/fences)
        layer_map = [
            ("roads", feature_sets.get("roads", {}), "road"),
            ("buildings", feature_sets.get("buildings", {}), "building"),
            ("rivers", feature_sets.get("rivers", {}), "river"),
            ("fences", feature_sets.get("fences", {}), "fence"),
        ]
        for suffix, collection, default_type in layer_map:
            feat_records = []
            for feat in (collection or {}).get("features", []):
                geom_json = feat.get("geometry")
                if not geom_json:
                    continue
                try:
                    geom_obj = shape(geom_json)
                except Exception:
                    continue
                props = feat.get("properties") or {}
                feat_records.append({
                    "plot_id": int(plot_id),
                    "ftype": default_type[:10],
                    "src": str(props.get("source") or "")[:20],
                    "name": str(props.get("name") or "")[:80],
                    "width_m": float(props.get("width_m")) if props.get("width_m") is not None else None,
                    "geometry": geom_obj,
                })
            write_layer(suffix, feat_records)
    except Exception as e:
        safe_rmtree(tmp_dir)
        raise HTTPException(status_code=500, detail=f"Failed to generate shapefile: {e}")

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zip_name = os.path.basename(zip_path)
        for fname in os.listdir(tmp_dir):
            if fname == zip_name:
                continue
            fp = os.path.join(tmp_dir, fname)
            if os.path.isfile(fp):
                zf.write(fp, arcname=fname)

    if background_tasks is None:
        background_tasks = BackgroundTasks()
    background_tasks.add_task(safe_rmtree, tmp_dir)

    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=export_filename,
        background=background_tasks,
    )


@router.get("/{plot_id}/reports/survey-plan")
def get_saved_survey_plan_pdf(plot_id: int, refresh: bool = False, db: Session = Depends(get_db)):
    pdf_path = resolve_existing_path([
        os.path.join(REPORTS_DIR, f"plot_{plot_id}_report_{SURVEY_REPORT_RENDER_VERSION}.pdf"),
        f"app/reports/plot_{plot_id}_report_{SURVEY_REPORT_RENDER_VERSION}.pdf",
    ])
    if refresh or not pdf_path:
        meta = get_plot_meta(db, plot_id)
        maps_dir = os.path.join(REPORTS_DIR, "maps")
        os.makedirs(REPORTS_DIR, exist_ok=True)
        os.makedirs(maps_dir, exist_ok=True)
        pdf_path = os.path.join(REPORTS_DIR, f"plot_{plot_id}_report_{SURVEY_REPORT_RENDER_VERSION}.pdf")
        tmp_map = tempfile.NamedTemporaryFile(suffix="_map.png", delete=False)
        map_path = tmp_map.name
        tmp_map.close()
        plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
        measurement_context = _resolve_measurement_polygon_context(
            plot_geom_wgs84,
            meta["coordinate_system"],
            meta.get("survey_input_coordinates"),
        )
        render_coordinate_system = measurement_context["render_coordinate_system"]
        epsg_code = measurement_context["epsg_code"]
        crs_name = measurement_context["crs_name"]
        render_plot_map_layout(
            db=db,
            plot_id=plot_id,
            output_path=map_path,
            title_text=meta["title_text"],
            location_text=meta["location_text"],
            lga_text=meta["lga_text"],
            state_text=meta["state_text"],
            scale_text=meta["scale_text"],
            surveyor_name=meta["surveyor_name"],
            surveyor_rank=meta["surveyor_rank"],
            certification_statement=meta.get("certification_statement"),
            station_names=None,
            coordinate_system=render_coordinate_system,
            epsg_code=epsg_code,
            crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
            measurement_polygon=measurement_context["measurement_polygon"],
            measurement_area_m2=measurement_context["measurement_area_m2"],
            paper_size=meta["paper_size"],
            north_arrow_style="one_side_stem",
            north_arrow_color="blue",
            beacon_style="cross",
            road_width_m=None,
            road_width_override_m=None,
            template_name=meta.get("template_name") or DEFAULT_TEMPLATE_NAME,
            adamawa_rof_no=meta.get("adamawa_rof_no") or "",
            adamawa_owner_name=meta.get("adamawa_owner_name") or "",
            adamawa_authority_title=meta.get("adamawa_authority_title") or DEFAULT_ADAMAWA_AUTHORITY_TITLE,
            adamawa_authority_date_text=meta.get("adamawa_authority_date_text") or DEFAULT_ADAMAWA_AUTHORITY_DATE,
            adamawa_control_point_name=meta.get("adamawa_control_point_name") or "",
            adamawa_northing=meta.get("adamawa_northing") or "",
            adamawa_easting=meta.get("adamawa_easting") or "",
            adamawa_elevation=meta.get("adamawa_elevation") or "",
            adamawa_origin_text=meta.get("adamawa_origin_text") or DEFAULT_ADAMAWA_ORIGIN_TEXT,
            adamawa_topo_sheet_text=meta.get("adamawa_topo_sheet_text") or DEFAULT_ADAMAWA_TOPO_SHEET_TEXT,
            adamawa_computation_no=meta.get("adamawa_computation_no") or "",
            adamawa_cadastral_sheet_no=meta.get("adamawa_cadastral_sheet_no") or "",
            adamawa_plan_no=meta.get("adamawa_plan_no") or "",
            adamawa_surveyed_by_text=meta.get("adamawa_surveyed_by_text") or "",
            adamawa_disclaimer_text=meta.get("adamawa_disclaimer_text") or DEFAULT_ADAMAWA_DISCLAIMER_TEXT,
            cadastral_plan_no=meta.get("cadastral_plan_no") or "",
            cadastral_area_name=meta.get("cadastral_area_name") or "",
            cadastral_datum_text=meta.get("cadastral_datum_text") or "",
            cadastral_firm_block_text=meta.get("cadastral_firm_block_text") or "",
            fct_file_no=meta.get("fct_file_no") or "",
            fct_district=meta.get("fct_district") or "",
            fct_cadastral_zone=meta.get("fct_cadastral_zone") or "",
            fct_origin_beacon_text=meta.get("fct_origin_beacon_text") or "",
            fct_cadastral_map_ref=meta.get("fct_cadastral_map_ref") or "",
            fct_title_prefix=meta.get("fct_title_prefix") or "",
        )
        report = get_plot_report(plot_id, db)
        generate_plot_report_pdf(report, pdf_path, map_path, paper_size=meta["paper_size"])
        safe_remove(map_path)
    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="Survey plan PDF not found")
    # meta is only loaded above when the cache was actually (re)built - fetch it fresh here too so
    # the filename is still identity-based on a cache hit that skipped that block.
    filename = _build_plot_export_filename(plot_id, get_plot_meta(db, plot_id), "Survey_Plan", "pdf")
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="survey-plan",
        project_id=plot_id,
    )


@router.get("/{plot_id}/reports/orthophoto")
def get_saved_orthophoto_pdf(plot_id: int, map_type: str = "satellite", refresh: bool = False, db: Session = Depends(get_db)):
    safe_type = "topo" if str(map_type).lower() in ["topo", "topomap", "topo_map"] else "satellite"
    pdf_path = resolve_existing_path([
        os.path.join(REPORTS_DIR, "orthophoto", f"plot_{plot_id}_orthophoto_{safe_type}.pdf"),
        f"app/reports/orthophoto/plot_{plot_id}_orthophoto_{safe_type}.pdf",
    ])
    if refresh or not pdf_path:
        meta = get_plot_meta(db, plot_id)
        out_dir = os.path.join(REPORTS_DIR, "orthophoto")
        os.makedirs(out_dir, exist_ok=True)
        pdf_path = os.path.join(out_dir, f"plot_{plot_id}_orthophoto_{safe_type}.pdf")
        tmp_png = tempfile.NamedTemporaryFile(suffix=f"_{safe_type}.png", delete=False)
        png_path = tmp_png.name
        tmp_png.close()
        plot_geom_wgs84 = _load_plot_polygon_wgs84(db, plot_id)
        render_coordinate_system, epsg_code, crs_name = _resolve_survey_render_crs(
            meta["coordinate_system"],
            plot_geom_wgs84,
        )
        render_orthophoto_png(
            db=db,
            plot_id=plot_id,
            output_path=png_path,
            title_text=meta["title_text"] if safe_type == "satellite" else "TOPO MAP",
            location_text=meta["location_text"],
            lga_text=meta["lga_text"],
            state_text=meta["state_text"],
            scale_text=meta["scale_text"],
            surveyor_name=meta["surveyor_name"],
            surveyor_rank=meta["surveyor_rank"],
            station_names=None,
            coordinate_system=render_coordinate_system,
            epsg_code=epsg_code,
            crs_footer_text=f"COORDINATE SYSTEM: {crs_name}",
            source_footer_text=(
                "SOURCE: Your uploaded elevation data" if safe_type == "topo" and meta.get("elevation_points")
                else "SOURCE: Global satellite elevation model (Copernicus/NASA DEM)" if safe_type == "topo"
                else "SOURCE: Satellite Imagery"
            ),
            use_topo_map=(safe_type == "topo"),
            topo_source="userdata" if meta.get("elevation_points") else "opentopomap",
            elevation_points=meta.get("elevation_points"),
            contour_interval=meta.get("contour_interval"),
            building_hatch_type="solid",
            paper_size=meta["paper_size"],
            north_arrow_style="one_side_stem",
            north_arrow_color="blue",
        )
        render_orthophoto_pdf_from_png(png_path, pdf_path, paper_size=meta["paper_size"])
        safe_remove(png_path)
    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="Orthophoto PDF not found")
    filename = _build_plot_export_filename(
        plot_id, get_plot_meta(db, plot_id), "Topo_Map" if safe_type == "topo" else "Orthophoto", "pdf",
    )
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="orthophoto" if safe_type != "topo" else "topo-map",
        project_id=plot_id,
    )


@router.get("/{plot_id}/reports/back-computation")
def get_saved_back_computation_pdf(plot_id: int, refresh: bool = False, db: Session = Depends(get_db)):
    pdf_path = resolve_existing_path([
        os.path.join(REPORTS_DIR, f"plot_{plot_id}_back_computation.pdf"),
        f"app/reports/plot_{plot_id}_back_computation.pdf",
    ])
    if refresh or not pdf_path:
        meta = get_plot_meta(db, plot_id)
        pdf_path = os.path.join(REPORTS_DIR, f"plot_{plot_id}_back_computation.pdf")
        os.makedirs(REPORTS_DIR, exist_ok=True)
        # Get plot geometry
        plot_wkb = db.execute(text("SELECT geom FROM plots WHERE id=:id"), {"id": plot_id}).scalar()
        plot_geom = wkb.loads(plot_wkb)
        measurement_context = _resolve_measurement_polygon_context(
            plot_geom,
            meta["coordinate_system"],
            meta.get("survey_input_coordinates"),
        )
        poly = measurement_context["measurement_polygon"]
        area_m2 = measurement_context["measurement_area_m2"]
        crs_name = measurement_context["crs_name"]
        rows, sum_de, sum_dn = compute_back_computation(poly, None)
        render_back_computation_pdf(rows, sum_de, sum_dn, area_m2, plot_id, pdf_path, crs_name)
    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="Back computation PDF not found")
    filename = _build_plot_export_filename(plot_id, get_plot_meta(db, plot_id), "Back_Computation", "pdf")
    return _pdf_response_with_r2(
        pdf_path,
        filename,
        category="back-computation",
        project_id=plot_id,
    )
