from fastapi import APIRouter, Body, Depends, HTTPException, Query, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import text
from datetime import datetime, date, timedelta, timezone
import json
import os
import tempfile
import csv
import io
import uuid
import zipfile
import re
import hashlib
import hmac
import secrets
import smtplib
from pathlib import Path
from threading import Lock
from urllib.parse import quote, unquote, urlparse
from email.message import EmailMessage

import boto3
from botocore.exceptions import ClientError

from app.db import SessionLocal
from app.utils.green_pdf import (
    render_green_report_pdf,
    render_green_work_report_pdf,
    render_green_custodian_report_pdf,
    render_green_existing_trees_report_pdf,
    render_green_org_credentials_pdf,
)
from app.utils.carbon import (
    compute_project_carbon,
    generate_co2_projection_table,
    estimate_tree_co2_kg,
    estimate_annual_co2_kg,
    estimate_lifetime_co2_kg,
    list_known_species,
    _normalize_species_key,
    _get_species_params,
    _infer_tree_reference_date,
    tree_age_years,
    project_dbh,
    calculate_agb_chave,
)

router = APIRouter(prefix="/green", tags=["green"])

_GREEN_SCHEMA_BOOTSTRAP_LOCK = Lock()
_GREEN_SCHEMA_READY = False
_GREEN_SCHEMA_ADVISORY_LOCK_ID = 903670421

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports", "green")
LIVE_SOURCE_REFERENCES = [
    {
        "label": "FAO - Forest restoration monitoring and maintenance sequence",
        "url": "https://www.fao.org/sustainable-forest-management-toolbox/modules/forest-restoration/en",
    },
    {
        "label": "FAO - Post-planting operations (watering, protection, replacement)",
        "url": "https://www.fao.org/4/u2247e/u2247e0a.htm",
    },
    {
        "label": "FAO - Savanna plantation field maintenance practices (Nigeria-relevant context)",
        "url": "https://www.fao.org/4/93269e/93269e03.htm",
    },
    {
        "label": "NiMet seasonal outlook context for local onset/dry-period planning",
        "url": "https://www.nimet.gov.ng/news?id=94",
    },
]
VERRA_VCS_REFERENCES = [
    {
        "label": "Verra - Verified Carbon Standard (program overview)",
        "url": "https://verra.org/programs/verified-carbon-standard/",
    },
    {
        "label": "Verra - VCS program rules and requirements",
        "url": "https://verra.org/project/vcs-program-rules-and-requirements/",
    },
    {
        "label": "Verra - Monitoring report templates and guidance (official resources page)",
        "url": "https://verra.org/project/vcs-program-rules-and-requirements/",
    },
]

MAINTENANCE_ACTIVITY_ORDER = ("watering", "weeding", "protection", "inspection", "replacement")
SUPERVISION_TASK_TYPE = "supervision"
ASSIGNABLE_TASK_TYPES = set(MAINTENANCE_ACTIVITY_ORDER) | {SUPERVISION_TASK_TYPE}
AGE_SURVIVAL_CHECKPOINTS_DAYS = (30, 90, 180)
SEASON_VALUES = {"rainy", "dry"}
TASK_STATUS_VALUES = {"pending", "done", "overdue"}
REVIEW_STATE_VALUES = {"none", "submitted", "approved", "rejected", "reopened", "metadata_edit"}
PLANTING_MODEL_VALUES = {"direct", "community_distributed", "mixed"}
DEFAULT_PLANTING_MODEL = "direct"
EXISTING_TREE_SCOPE_VALUES = {"exclude_from_planting_kpi", "include_in_planting_kpi"}
DEFAULT_EXISTING_TREE_SCOPE = "exclude_from_planting_kpi"
TREE_ORIGIN_VALUES = {"new_planting", "existing_inventory", "natural_regeneration"}
TREE_ATTRIBUTION_SCOPE_VALUES = {"full", "monitor_only"}
CUSTODIAN_TYPE_VALUES = {"household", "school", "community_group"}
TREE_PROJECT_LINK_TYPE_VALUES = {"owner", "reference"}
TREE_STATUS_VALUES = {
    "alive",
    "healthy",
    "dead",
    "needs_attention",
    "pending_planting",
    "pest",
    "disease",
    "need_replacement",
    "needs_replacement",
    "damaged",
    "removed",
    "need_watering",
    "need_protection",
}
REPLACEMENT_TRIGGER_STATUSES = {"dead", "damaged", "removed", "need_replacement", "needs_replacement"}
TREE_STATUS_ALIASES = {
    "needreplacement": "need_replacement",
    "need_replacement": "need_replacement",
    "needsreplacement": "needs_replacement",
    "needs_replacement": "needs_replacement",
    "need replacement": "need_replacement",
    "needs replacement": "needs_replacement",
    "deseas": "disease",
    "diseased": "disease",
    "needsattention": "needs_attention",
    "need_attention": "needs_attention",
    "need_watering": "need_watering",
    "needwatering": "need_watering",
    "need watering": "need_watering",
    "need_protection": "need_protection",
    "needprotection": "need_protection",
    "need protection": "need_protection",
}
HEALTHY_TREE_STATUSES = {"alive", "healthy"}
DEAD_TREE_STATUSES = {"dead", "removed"}
ATTENTION_TREE_STATUSES = {
    "needs_attention",
    "pest",
    "disease",
    "need_replacement",
    "needs_replacement",
    "damaged",
    "need_watering",
    "need_protection",
}
TREE_STATUS_COLOR_HEX = {
    "alive": "22c55e",
    "healthy": "16a34a",
    "pest": "eab308",
    "disease": "f97316",
    "need_replacement": "ef4444",
    "needs_replacement": "ef4444",
    "damaged": "dc2626",
    "dead": "b91c1c",
    "removed": "7f1d1d",
    "needs_attention": "f59e0b",
    "pending_planting": "3b82f6",
    "need_watering": "0ea5e9",
    "need_protection": "a855f7",
}


def get_db():
    db = SessionLocal()
    try:
        _ensure_green_schema_ready(db)
        yield db
    finally:
        db.close()


def _ensure_green_schema_ready(db: Session):
    global _GREEN_SCHEMA_READY
    if _GREEN_SCHEMA_READY:
        return

    with _GREEN_SCHEMA_BOOTSTRAP_LOCK:
        if _GREEN_SCHEMA_READY:
            return

        advisory_lock_acquired = False
        try:
            # Serialize bootstrap across multiple API workers/processes.
            try:
                db.execute(
                    text("SELECT pg_advisory_lock(:lock_id)"),
                    {"lock_id": _GREEN_SCHEMA_ADVISORY_LOCK_ID},
                )
                advisory_lock_acquired = True
            except Exception:
                # If advisory lock is unavailable, continue with best effort.
                db.rollback()

            ensure_green_tables(db)
            _GREEN_SCHEMA_READY = True
        finally:
            if advisory_lock_acquired:
                try:
                    db.execute(
                        text("SELECT pg_advisory_unlock(:lock_id)"),
                        {"lock_id": _GREEN_SCHEMA_ADVISORY_LOCK_ID},
                    )
                    db.commit()
                except Exception:
                    db.rollback()


def _normalize_name(value: str | None) -> str:
    return (value or "").strip().lower()


def _slugify_text(value: str | None, fallback: str = "organization") -> str:
    raw = (value or "").strip().lower()
    raw = re.sub(r"[^a-z0-9]+", "-", raw).strip("-")
    return raw or fallback


def _ensure_unique_org_slug(db: Session, slug_base: str, exclude_org_id: int | None = None) -> str:
    base = _slugify_text(slug_base)
    candidate = base
    suffix = 2
    while True:
        params = {"slug": candidate}
        if exclude_org_id is None:
            row = db.execute(
                text("SELECT id FROM green_organizations WHERE LOWER(slug) = LOWER(:slug) LIMIT 1"),
                params,
            ).first()
        else:
            row = db.execute(
                text(
                    """
                    SELECT id FROM green_organizations
                    WHERE LOWER(slug) = LOWER(:slug)
                      AND id <> :exclude_org_id
                    LIMIT 1
                    """
                ),
                {**params, "exclude_org_id": int(exclude_org_id)},
            ).first()
        if not row:
            return candidate
        candidate = f"{base}-{suffix}"
        suffix += 1


def _generate_prefixed_uid(prefix: str) -> str:
    token = uuid.uuid4().hex[:8].upper()
    return f"{prefix}-{token}"


def _ensure_unique_user_uid(db: Session, candidate: str | None = None, exclude_user_id: int | None = None) -> str:
    base = (candidate or "").strip().upper()
    if not base:
        base = _generate_prefixed_uid("USR")
    if not re.fullmatch(r"[A-Z0-9][A-Z0-9._-]{2,63}", base):
        base = re.sub(r"[^A-Z0-9._-]+", "-", base.upper()).strip("-")
        if not base:
            base = _generate_prefixed_uid("USR")
    value = base
    while True:
        if exclude_user_id is None:
            found = db.execute(
                text("SELECT id FROM green_users WHERE UPPER(COALESCE(user_uid, '')) = :uid LIMIT 1"),
                {"uid": value},
            ).first()
        else:
            found = db.execute(
                text(
                    """
                    SELECT id FROM green_users
                    WHERE UPPER(COALESCE(user_uid, '')) = :uid
                      AND id <> :user_id
                    LIMIT 1
                    """
                ),
                {"uid": value, "user_id": int(exclude_user_id)},
            ).first()
        if not found:
            return value
        value = _generate_prefixed_uid("USR")


def _ensure_unique_role_uid(db: Session, candidate: str | None = None, exclude_role_id: int | None = None) -> str:
    base = (candidate or "").strip().upper()
    if not base:
        base = _generate_prefixed_uid("ROL")
    if not re.fullmatch(r"[A-Z0-9][A-Z0-9._-]{2,63}", base):
        base = re.sub(r"[^A-Z0-9._-]+", "-", base.upper()).strip("-")
        if not base:
            base = _generate_prefixed_uid("ROL")
    value = base
    while True:
        if exclude_role_id is None:
            found = db.execute(
                text("SELECT id FROM green_roles WHERE UPPER(COALESCE(role_uid, '')) = :uid LIMIT 1"),
                {"uid": value},
            ).first()
        else:
            found = db.execute(
                text(
                    """
                    SELECT id FROM green_roles
                    WHERE UPPER(COALESCE(role_uid, '')) = :uid
                      AND id <> :role_id
                    LIMIT 1
                    """
                ),
                {"uid": value, "role_id": int(exclude_role_id)},
            ).first()
        if not found:
            return value
        value = _generate_prefixed_uid("ROL")


def _ensure_unique_role_key(db: Session, key_or_name: str | None, exclude_role_id: int | None = None) -> str:
    base = _slugify_text(key_or_name or "role", fallback="role").replace("-", "_")
    candidate = base
    suffix = 2
    while True:
        if exclude_role_id is None:
            found = db.execute(
                text("SELECT id FROM green_roles WHERE LOWER(COALESCE(role_key, '')) = LOWER(:role_key) LIMIT 1"),
                {"role_key": candidate},
            ).first()
        else:
            found = db.execute(
                text(
                    """
                    SELECT id FROM green_roles
                    WHERE LOWER(COALESCE(role_key, '')) = LOWER(:role_key)
                      AND id <> :role_id
                    LIMIT 1
                    """
                ),
                {"role_key": candidate, "role_id": int(exclude_role_id)},
            ).first()
        if not found:
            return candidate
        candidate = f"{base}_{suffix}"
        suffix += 1


def _hash_password_value(password: str) -> str:
    raw = str(password or "")
    if len(raw) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    iterations = 260000
    salt = secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", raw.encode("utf-8"), salt.encode("utf-8"), iterations)
    return f"pbkdf2_sha256${iterations}${salt}${digest.hex()}"


def _verify_password_value(password: str, encoded: str | None) -> bool:
    if not encoded:
        return False
    try:
        algo, iter_str, salt, digest_hex = str(encoded).split("$", 3)
        if algo != "pbkdf2_sha256":
            return False
        iterations = int(iter_str)
        derived = hashlib.pbkdf2_hmac("sha256", str(password or "").encode("utf-8"), salt.encode("utf-8"), iterations)
        return hmac.compare_digest(derived.hex(), digest_hex)
    except Exception:
        return False


def _generate_temporary_login_password(length: int = 12) -> str:
    alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnopqrstuvwxyz23456789"
    size = max(8, int(length or 12))
    return "".join(secrets.choice(alphabet) for _ in range(size))


def _env_bool(name: str, default: bool = False) -> bool:
    raw = str(os.getenv(name) or "").strip().lower()
    if not raw:
        return bool(default)
    return raw in {"1", "true", "yes", "on"}


def _send_new_user_credentials_email(
    *,
    to_email: str,
    full_name: str,
    organization_name: str | None,
    username: str,
    password: str,
    allow_green: bool,
    allow_work: bool,
):
    smtp_host = str(os.getenv("SMTP_HOST") or "").strip()
    if not smtp_host:
        raise RuntimeError("SMTP_HOST is not configured")
    smtp_port = int(str(os.getenv("SMTP_PORT") or "587").strip() or "587")
    smtp_user = str(os.getenv("SMTP_USERNAME") or "").strip()
    smtp_pass = str(os.getenv("SMTP_PASSWORD") or "").strip()
    smtp_from_email = str(os.getenv("SMTP_FROM_EMAIL") or smtp_user or "").strip()
    smtp_from_name = str(os.getenv("SMTP_FROM_NAME") or "LandCheck").strip()
    if not smtp_from_email:
        raise RuntimeError("SMTP_FROM_EMAIL (or SMTP_USERNAME) is not configured")
    use_ssl = _env_bool("SMTP_USE_SSL", False)
    use_tls = _env_bool("SMTP_USE_TLS", not use_ssl)
    green_url = str(os.getenv("LANDCHECK_GREEN_URL") or "").strip() or "https://landcheck.online/green/login"
    work_url = str(os.getenv("LANDCHECK_WORK_URL") or "").strip() or "https://landcheck.online/green-work/login"

    access_lines = []
    if allow_green:
        access_lines.append(f"- LandCheck Green: {green_url}")
    if allow_work:
        access_lines.append(f"- LandCheck Work: {work_url}")
    if not access_lines:
        access_lines.append("- No app access enabled")

    recipient_name = (full_name or "").strip() or "User"
    org_line = f"Organization: {organization_name}\n" if organization_name else ""
    body = (
        f"Hello {recipient_name},\n\n"
        "Your LandCheck account has been created.\n\n"
        f"{org_line}"
        "Login details:\n"
        f"Username: {username}\n"
        f"Temporary Password: {password}\n\n"
        "Access:\n"
        f"{chr(10).join(access_lines)}\n\n"
        "Please log in and change/reset your password through your administrator after first access.\n\n"
        "Regards,\n"
        "LandCheck"
    )

    msg = EmailMessage()
    msg["Subject"] = "Your LandCheck Login Credentials"
    msg["From"] = f"{smtp_from_name} <{smtp_from_email}>" if smtp_from_name else smtp_from_email
    msg["To"] = to_email
    msg.set_content(body)

    if use_ssl:
        with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=20) as server:
            if smtp_user:
                server.login(smtp_user, smtp_pass)
            server.send_message(msg)
        return

    with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as server:
        try:
            server.ehlo()
        except Exception:
            pass
        if use_tls:
            server.starttls()
            try:
                server.ehlo()
            except Exception:
                pass
        if smtp_user:
            server.login(smtp_user, smtp_pass)
        server.send_message(msg)


def _send_organization_welcome_email(
    *,
    to_email: str,
    organization_name: str,
    organization_slug: str | None = None,
    status: str | None = None,
    short_name: str | None = None,
    website_url: str | None = None,
    contact_phone: str | None = None,
    country: str | None = None,
    state_region: str | None = None,
    city: str | None = None,
):
    smtp_host = str(os.getenv("SMTP_HOST") or "").strip()
    if not smtp_host:
        raise RuntimeError("SMTP_HOST is not configured")
    smtp_port = int(str(os.getenv("SMTP_PORT") or "587").strip() or "587")
    smtp_user = str(os.getenv("SMTP_USERNAME") or "").strip()
    smtp_pass = str(os.getenv("SMTP_PASSWORD") or "").strip()
    smtp_from_email = str(os.getenv("SMTP_FROM_EMAIL") or smtp_user or "").strip()
    smtp_from_name = str(os.getenv("SMTP_FROM_NAME") or "LandCheck").strip()
    if not smtp_from_email:
        raise RuntimeError("SMTP_FROM_EMAIL (or SMTP_USERNAME) is not configured")
    use_ssl = _env_bool("SMTP_USE_SSL", False)
    use_tls = _env_bool("SMTP_USE_TLS", not use_ssl)
    green_url = str(os.getenv("LANDCHECK_GREEN_URL") or "").strip() or "https://landcheck.online/green/login"
    work_url = str(os.getenv("LANDCHECK_WORK_URL") or "").strip() or "https://landcheck.online/green-work/login"

    detail_lines: list[str] = []
    detail_lines.append(f"Organization: {organization_name}")
    if short_name:
        detail_lines.append(f"Short name: {short_name}")
    if organization_slug:
        detail_lines.append(f"Slug: {organization_slug}")
    if status:
        detail_lines.append(f"Partnership status: {status}")
    location_parts = [part for part in [city, state_region, country] if str(part or "").strip()]
    if location_parts:
        detail_lines.append(f"Location: {', '.join(location_parts)}")
    if contact_phone:
        detail_lines.append(f"Contact phone: {contact_phone}")
    if website_url:
        detail_lines.append(f"Website: {website_url}")

    body = (
        f"Hello {organization_name},\n\n"
        "Welcome to the LandCheck partnership.\n\n"
        "Your organization has been created on LandCheck and onboarding has been completed.\n\n"
        "Organization details:\n"
        f"{chr(10).join(f'- {line}' for line in detail_lines)}\n\n"
        "What LandCheck does:\n"
        "- LandCheck Green helps field teams capture tree planting and maintenance activities with GPS and photo evidence.\n"
        "- LandCheck Work helps supervisors assign tasks, review submissions, monitor progress, and export reports.\n"
        "- The platform supports project-based monitoring for planting, survival tracking, and operational reporting.\n\n"
        "We value your commitment to environmental stewardship and climate action. Welcome onboard!\n\n"
        "Regards,\n"
        "Samuel Yohanna\n"
        "Founder\n"
        "LandCheck Geospatial Technologies Limited\n"
        "RC: 9350241\n\n"
        f"LandCheck Green: {green_url}\n"
        f"LandCheck Work: {work_url}\n"
        "WhatsApp: +49 1776732638\n"
    )

    msg = EmailMessage()
    msg["Subject"] = f"Welcome to LandCheck Partnership - {organization_name}"
    msg["From"] = f"{smtp_from_name} <{smtp_from_email}>" if smtp_from_name else smtp_from_email
    msg["To"] = to_email
    msg.set_content(body)

    if use_ssl:
        with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=20) as server:
            if smtp_user:
                server.login(smtp_user, smtp_pass)
            server.send_message(msg)
        return

    with smtplib.SMTP(smtp_host, smtp_port, timeout=20) as server:
        try:
            server.ehlo()
        except Exception:
            pass
        if use_tls:
            server.starttls()
            try:
                server.ehlo()
            except Exception:
                pass
        if smtp_user:
            server.login(smtp_user, smtp_pass)
        server.send_message(msg)


def _next_project_tree_no(db: Session, project_id: int) -> int:
    # Lock the project row so concurrent inserts don't issue the same local tree number.
    db.execute(text("SELECT id FROM tree_projects WHERE id = :project_id FOR UPDATE"), {"project_id": int(project_id)})
    next_no = db.execute(
        text("SELECT COALESCE(MAX(project_tree_no), 0) + 1 FROM trees WHERE project_id = :project_id"),
        {"project_id": int(project_id)},
    ).scalar()
    try:
        return int(next_no or 1)
    except Exception:
        return 1


def _normalize_planting_model(value: str | None) -> str:
    normalized = _normalize_name(value)
    if normalized in PLANTING_MODEL_VALUES:
        return normalized
    return DEFAULT_PLANTING_MODEL


def _normalize_existing_tree_scope(value: str | None) -> str:
    normalized = _normalize_name(value)
    if normalized in EXISTING_TREE_SCOPE_VALUES:
        return normalized
    return DEFAULT_EXISTING_TREE_SCOPE


def _normalize_tree_origin(value: str | None) -> str:
    normalized = _normalize_name(value).replace("-", "_").replace(" ", "_")
    if normalized in TREE_ORIGIN_VALUES:
        return normalized
    return "new_planting"


def _normalize_tree_attribution_scope(value: str | None) -> str:
    normalized = _normalize_name(value).replace("-", "_").replace(" ", "_")
    if normalized in TREE_ATTRIBUTION_SCOPE_VALUES:
        return normalized
    return "full"


def _normalize_species_allocations(value: object) -> list[dict]:
    if not isinstance(value, list):
        return []
    merged: dict[str, dict] = {}
    for item in value:
        if not isinstance(item, dict):
            continue
        species_raw = str(item.get("species") or "").strip()
        if not species_raw:
            continue
        try:
            count_value = int(item.get("count") or 0)
        except Exception:
            continue
        if count_value <= 0:
            continue
        key = _normalize_name(species_raw)
        if not key:
            continue
        if key not in merged:
            merged[key] = {"species": species_raw, "count": 0}
        merged[key]["count"] = int(merged[key]["count"]) + count_value
    return list(merged.values())


def _normalize_custodian_type(value: str | None) -> str:
    normalized = _normalize_name(value).replace("-", "_").replace(" ", "_")
    if normalized in CUSTODIAN_TYPE_VALUES:
        return normalized
    return "household"


def _normalize_tree_link_type(value: str | None) -> str:
    normalized = _normalize_name(value)
    if normalized in TREE_PROJECT_LINK_TYPE_VALUES:
        return normalized
    return "reference"


def _normalize_tree_status(value: str | None) -> str:
    raw = _normalize_name(value).replace("-", "_")
    collapsed = raw.replace("_", "").replace(" ", "")
    if raw in TREE_STATUS_ALIASES:
        return TREE_STATUS_ALIASES[raw]
    if collapsed in TREE_STATUS_ALIASES:
        return TREE_STATUS_ALIASES[collapsed]
    if " " in raw:
        spaced = raw.replace(" ", "_")
        if spaced in TREE_STATUS_ALIASES:
            return TREE_STATUS_ALIASES[spaced]
    return raw.replace(" ", "_")


def _resolve_tree_scope_defaults(
    *,
    tree_origin: str,
    attribution_scope: str | None = None,
    count_in_planting_kpis: bool | None = None,
    count_in_carbon_scope: bool | None = None,
    project_existing_scope: str | None = None,
) -> tuple[str, bool, bool]:
    origin = _normalize_tree_origin(tree_origin)
    resolved_scope = _normalize_tree_attribution_scope(attribution_scope)
    resolved_existing_scope = _normalize_existing_tree_scope(project_existing_scope)

    if origin == "existing_inventory":
        default_scope = "monitor_only"
        if count_in_planting_kpis is None:
            default_planting_scope = resolved_existing_scope == "include_in_planting_kpi"
        else:
            default_planting_scope = bool(count_in_planting_kpis)
        if count_in_carbon_scope is None:
            default_carbon_scope = resolved_existing_scope == "include_in_planting_kpi"
        else:
            default_carbon_scope = bool(count_in_carbon_scope)
    else:
        default_scope = "full"
        default_planting_scope = bool(count_in_planting_kpis) if count_in_planting_kpis is not None else True
        default_carbon_scope = bool(count_in_carbon_scope) if count_in_carbon_scope is not None else True

    if not attribution_scope:
        resolved_scope = default_scope

    if resolved_scope == "monitor_only":
        if count_in_planting_kpis is None:
            default_planting_scope = False
        if count_in_carbon_scope is None:
            default_carbon_scope = False

    return resolved_scope, bool(default_planting_scope), bool(default_carbon_scope)


def _is_done_status(status: str | None) -> bool:
    return _normalize_name(status) in {"done", "completed", "closed"}


def _start_of_day(value: date | datetime) -> date:
    if isinstance(value, datetime):
        return value.date()
    return value


def _parse_date_value(value: str | datetime | date | None) -> date | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    raw = str(value).strip()
    if not raw:
        return None
    try:
        if len(raw) >= 10:
            return date.fromisoformat(raw[:10])
        return date.fromisoformat(raw)
    except Exception:
        return None


def _parse_datetime_value(value: str | datetime | None) -> datetime | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        dt_value = value
    else:
        raw = str(value).strip()
        if not raw:
            return None
        candidate = raw
        if candidate.endswith("Z"):
            candidate = candidate[:-1] + "+00:00"
        try:
            dt_value = datetime.fromisoformat(candidate)
        except Exception:
            return None
    if dt_value.tzinfo is not None:
        dt_value = dt_value.astimezone(timezone.utc).replace(tzinfo=None)
    return dt_value


def _to_date_input(value: date | None) -> str:
    return value.isoformat() if value else ""


def _add_days(value: date, days: int) -> date:
    return value + timedelta(days=days)


def _day_diff(target: date, reference: date) -> int:
    return (_start_of_day(target) - _start_of_day(reference)).days


def _safe_json(value: dict | list | None) -> str:
    try:
        return json.dumps(value or {}, default=str)
    except Exception:
        return "{}"


def _normalize_photo_urls(value: object) -> list[str]:
    raw: object = value
    if isinstance(raw, str):
        text_value = raw.strip()
        if not text_value:
            return []
        try:
            parsed = json.loads(text_value)
            raw = parsed
        except Exception:
            raw = [text_value]
    if isinstance(raw, tuple):
        raw = list(raw)
    if not isinstance(raw, list):
        return []
    seen: set[str] = set()
    items: list[str] = []
    for item in raw:
        url = str(item or "").strip()
        if not url or url in seen:
            continue
        seen.add(url)
        items.append(url)
    return items[:30]


def _merge_photo_evidence(photo_url: str | None, photo_urls: object) -> tuple[str | None, list[str]]:
    merged = _normalize_photo_urls(photo_urls)
    single = (photo_url or "").strip()
    if single and single not in merged:
        merged.append(single)
    primary = merged[-1] if merged else None
    return primary, merged


def _is_valid_linear_ring(ring: object) -> bool:
    if not isinstance(ring, list) or len(ring) < 4:
        return False
    for point in ring:
        if not isinstance(point, (list, tuple)) or len(point) < 2:
            return False
        try:
            float(point[0])
            float(point[1])
        except Exception:
            return False
    return True


def _normalize_work_area_geojson(value: dict | str | None) -> dict | None:
    if value is None:
        return None

    raw: object = value
    if isinstance(raw, str):
        raw = raw.strip()
        if not raw:
            return None
        try:
            raw = json.loads(raw)
        except Exception:
            return None

    if not isinstance(raw, dict):
        return None

    geometry = raw.get("geometry") if str(raw.get("type") or "").strip() == "Feature" else raw
    if not isinstance(geometry, dict):
        return None

    geom_type = str(geometry.get("type") or "").strip()
    coords = geometry.get("coordinates")

    if geom_type == "Polygon":
        if not isinstance(coords, list) or len(coords) == 0:
            return None
        normalized_rings: list[list[list[float]]] = []
        for ring in coords:
            if not _is_valid_linear_ring(ring):
                return None
            normalized_ring = [[float(point[0]), float(point[1])] for point in ring]
            normalized_rings.append(normalized_ring)
        return {"type": "Polygon", "coordinates": normalized_rings}

    if geom_type == "MultiPolygon":
        if not isinstance(coords, list) or len(coords) == 0:
            return None
        normalized_polygons: list[list[list[list[float]]]] = []
        for polygon in coords:
            if not isinstance(polygon, list) or len(polygon) == 0:
                return None
            normalized_rings: list[list[list[float]]] = []
            for ring in polygon:
                if not _is_valid_linear_ring(ring):
                    return None
                normalized_ring = [[float(point[0]), float(point[1])] for point in ring]
                normalized_rings.append(normalized_ring)
            normalized_polygons.append(normalized_rings)
        return {"type": "MultiPolygon", "coordinates": normalized_polygons}

    return None


def _get_maintenance_intervals(activity: str, tree_age_days: int, season: str) -> dict:
    activity_key = _normalize_name(activity)
    season_key = "dry" if _normalize_name(season) == "dry" else "rainy"
    age = max(int(tree_age_days or 0), 0)

    if activity_key == "watering":
        if season_key == "rainy":
            return {"first_days": 0, "repeat_days": 14 if age < 90 else 21}
        return {"first_days": 0, "repeat_days": 5 if age < 90 else 7}

    if activity_key == "weeding":
        if season_key == "rainy":
            if age < 365:
                return {"first_days": 21, "repeat_days": 45}
            if age < 730:
                return {"first_days": 30, "repeat_days": 90}
            return {"first_days": 30, "repeat_days": 150}
        if age < 365:
            return {"first_days": 35, "repeat_days": 90}
        if age < 730:
            return {"first_days": 45, "repeat_days": 150}
        return {"first_days": 45, "repeat_days": 210}

    if activity_key == "protection":
        if season_key == "rainy":
            return {"first_days": 0, "repeat_days": 45}
        return {"first_days": 0, "repeat_days": 21}

    if activity_key == "inspection":
        if season_key == "rainy":
            return {"first_days": 14, "repeat_days": 30 if age < 180 else 90}
        return {"first_days": 7, "repeat_days": 21 if age < 180 else 60}

    if activity_key == "replacement":
        if season_key == "rainy":
            return {"first_days": 42, "repeat_days": 180}
        return {"first_days": 56, "repeat_days": 210}

    return {"first_days": 30, "repeat_days": 90}


def _should_skip_existing_tree_routine_activity(
    activity: str,
    tree_origin: str,
    tree_age_days: int | None,
    tree_status: str,
    has_open_task: bool,
) -> bool:
    origin_key = _normalize_tree_origin(tree_origin)
    activity_key = _normalize_name(activity)
    status_key = _normalize_tree_status(tree_status)
    if origin_key != "existing_inventory" or has_open_task:
        return False
    if activity_key == "watering":
        if status_key == "need_watering":
            return False
        return tree_age_days is None or tree_age_days >= 365
    if activity_key == "weeding":
        return tree_age_days is None or tree_age_days >= 730
    return False


def _point_in_ring(lng: float, lat: float, ring: list[list[float]]) -> bool:
    if not isinstance(ring, list) or len(ring) < 4:
        return False
    inside = False
    j = len(ring) - 1
    for i in range(len(ring)):
        try:
            xi, yi = float(ring[i][0]), float(ring[i][1])
            xj, yj = float(ring[j][0]), float(ring[j][1])
        except Exception:
            j = i
            continue
        intersects = ((yi > lat) != (yj > lat)) and (
            lng < ((xj - xi) * (lat - yi) / ((yj - yi) or 1e-12) + xi)
        )
        if intersects:
            inside = not inside
        j = i
    return inside


def _point_in_polygon_geojson(lng: float, lat: float, geometry: dict | None) -> bool:
    if not isinstance(geometry, dict):
        return False
    geom_type = str(geometry.get("type") or "").strip()
    coordinates = geometry.get("coordinates")
    if geom_type == "Polygon":
        polygons = [coordinates]
    elif geom_type == "MultiPolygon":
        polygons = coordinates if isinstance(coordinates, list) else []
    else:
        return False
    for polygon in polygons:
        if not isinstance(polygon, list) or not polygon:
            continue
        outer = polygon[0]
        holes = polygon[1:] if len(polygon) > 1 else []
        if not _point_in_ring(lng, lat, outer):
            continue
        if any(_point_in_ring(lng, lat, hole) for hole in holes if isinstance(hole, list)):
            continue
        return True
    return False


def _find_matching_auto_first_cycle_work_order(
    db: Session,
    *,
    project_id: int,
    assignee_name: str,
    species: str | None,
    lng: float,
    lat: float,
) -> dict | None:
    assignee_clean = (assignee_name or "").strip()
    if not assignee_clean:
        return None
    rows = db.execute(
        text(
            """
            SELECT id, assignee_name, maintenance_schedule, species_allocations,
                   COALESCE(area_enabled, FALSE) AS area_enabled, area_geojson,
                   status, created_at
            FROM green_work_orders
            WHERE project_id = :project_id
              AND work_type = 'planting'
              AND COALESCE(auto_assign_first_cycle_maintenance, FALSE) = TRUE
              AND LOWER(TRIM(assignee_name)) = LOWER(TRIM(:assignee_name))
              AND LOWER(COALESCE(status, 'assigned')) NOT IN ('done', 'completed', 'closed', 'cancelled')
            ORDER BY created_at DESC, id DESC
            """
        ),
        {"project_id": int(project_id), "assignee_name": assignee_clean},
    ).mappings().all()
    species_key = _normalize_name(species)
    for raw in rows:
        item = dict(raw)
        raw_species_allocations = item.get("species_allocations")
        if isinstance(raw_species_allocations, str):
            try:
                raw_species_allocations = json.loads(raw_species_allocations)
            except Exception:
                raw_species_allocations = []
        normalized_species_allocations = _normalize_species_allocations(raw_species_allocations)
        if normalized_species_allocations:
            allowed_keys = {
                _normalize_name(entry.get("species"))
                for entry in normalized_species_allocations
                if str(entry.get("species") or "").strip()
            }
            if not species_key or species_key not in allowed_keys:
                continue

        area_enabled = bool(item.get("area_enabled"))
        if area_enabled:
            raw_geojson = item.get("area_geojson")
            if isinstance(raw_geojson, str):
                try:
                    raw_geojson = json.loads(raw_geojson)
                except Exception:
                    raw_geojson = None
            normalized_geojson = _normalize_work_area_geojson(raw_geojson) if raw_geojson else None
            if not normalized_geojson or not _point_in_polygon_geojson(lng, lat, normalized_geojson):
                continue
        return item
    return None


def _auto_assign_first_cycle_maintenance_from_order(
    db: Session,
    *,
    project_id: int,
    tree_id: int,
    assignee_name: str,
    planting_date_value: str | None,
    order_row: dict,
) -> list[int]:
    assignee_clean = (assignee_name or "").strip()
    if not assignee_clean:
        return []
    base_date = _parse_date_value(planting_date_value) or date.today()
    season_key = _normalize_name(order_row.get("maintenance_schedule"))
    if season_key not in SEASON_VALUES:
        season_key = "rainy"
    created_task_ids: list[int] = []
    for activity in ("watering", "weeding", "protection", "inspection"):
        existing_task_id = db.execute(
            text(
                """
                SELECT id
                FROM tree_tasks
                WHERE tree_id = :tree_id
                  AND LOWER(task_type) = :task_type
                LIMIT 1
                """
            ),
            {"tree_id": int(tree_id), "task_type": activity},
        ).scalar()
        if existing_task_id:
            continue
        intervals = _get_maintenance_intervals(activity, 0, season_key)
        due_date = _add_days(base_date, int(intervals.get("first_days") or 0))
        task_id = db.execute(
            text(
                """
                INSERT INTO tree_tasks (
                    tree_id, task_type, assignee_name, due_date, priority, status, notes,
                    review_state, auto_generated, model_season
                )
                VALUES (
                    :tree_id, :task_type, :assignee_name, :due_date, 'normal', 'pending', :notes,
                    'none', FALSE, :model_season
                )
                RETURNING id
                """
            ),
            {
                "tree_id": int(tree_id),
                "task_type": activity,
                "assignee_name": assignee_clean,
                "due_date": _to_date_input(due_date),
                "notes": f"Auto-assigned first-cycle maintenance from planting order #{int(order_row.get('id') or 0)}.",
                "model_season": season_key,
            },
        ).scalar()
        if not task_id:
            continue
        created_task_ids.append(int(task_id))
        _log_audit_event(
            db,
            project_id=int(project_id),
            entity_type="task",
            entity_id=int(task_id),
            action="first_cycle_task_auto_assigned",
            actor="system",
            details={
                "tree_id": int(tree_id),
                "source_work_order_id": int(order_row.get("id") or 0),
                "task_type": activity,
                "due_date": _to_date_input(due_date),
                "season": season_key,
            },
        )
    return created_task_ids


def _get_lifecycle_start_date(planting_date_obj: date | None, replacement_done_obj: date | None) -> date | None:
    if planting_date_obj and replacement_done_obj:
        return replacement_done_obj if replacement_done_obj > planting_date_obj else planting_date_obj
    return replacement_done_obj or planting_date_obj


def _get_project_id_for_tree(db: Session, tree_id: int) -> int | None:
    return db.execute(
        text("SELECT project_id FROM trees WHERE id = :tree_id"),
        {"tree_id": tree_id},
    ).scalar()


def _get_project_settings(db: Session, project_id: int) -> dict:
    row = db.execute(
        text(
            """
            SELECT id, planting_model, allow_existing_tree_link, default_existing_tree_scope
            FROM tree_projects
            WHERE id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Project not found")
    return {
        "id": int(row.get("id")),
        "planting_model": _normalize_planting_model(row.get("planting_model")),
        "allow_existing_tree_link": bool(row.get("allow_existing_tree_link")),
        "default_existing_tree_scope": _normalize_existing_tree_scope(row.get("default_existing_tree_scope")),
    }


def _record_tree_status_history(
    db: Session,
    tree_id: int,
    status: str,
    project_id: int | None = None,
    status_date: date | datetime | str | None = None,
    source: str = "manual",
    source_task_id: int | None = None,
    changed_by: str | None = None,
    notes: str | None = None,
):
    normalized_status = _normalize_tree_status(status)
    if normalized_status not in TREE_STATUS_VALUES:
        return

    resolved_project_id = int(project_id) if project_id is not None else _get_project_id_for_tree(db, tree_id)
    if resolved_project_id is None:
        return

    resolved_date = _parse_date_value(status_date) or date.today()
    existing = db.execute(
        text(
            """
            SELECT id
            FROM green_tree_status_history
            WHERE tree_id = :tree_id
              AND status = :status
              AND status_date = :status_date
              AND COALESCE(source, '') = COALESCE(:source, '')
              AND COALESCE(source_task_id, 0) = COALESCE(:source_task_id, 0)
            ORDER BY id DESC
            LIMIT 1
            """
        ),
        {
            "tree_id": int(tree_id),
            "status": normalized_status,
            "status_date": resolved_date,
            "source": source or "manual",
            "source_task_id": source_task_id,
        },
    ).scalar()
    if existing:
        return

    db.execute(
        text(
            """
            INSERT INTO green_tree_status_history (
                tree_id, project_id, status, status_date, source, source_task_id, changed_by, notes
            )
            VALUES (
                :tree_id, :project_id, :status, :status_date, :source, :source_task_id, :changed_by, :notes
            )
            """
        ),
        {
            "tree_id": int(tree_id),
            "project_id": int(resolved_project_id),
            "status": normalized_status,
            "status_date": resolved_date,
            "source": source or "manual",
            "source_task_id": source_task_id,
            "changed_by": (changed_by or "").strip() or None,
            "notes": notes or None,
        },
    )


def _get_project_id_for_task(db: Session, task_id: int) -> int | None:
    return db.execute(
        text("""
            SELECT tr.project_id
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE t.id = :task_id
        """),
        {"task_id": task_id},
    ).scalar()


def _log_audit_event(
    db: Session,
    project_id: int | None,
    entity_type: str,
    entity_id: int | None,
    action: str,
    actor: str | None = None,
    details: dict | None = None,
):
    db.execute(
        text("""
            INSERT INTO green_audit_events (
                project_id, entity_type, entity_id, action, actor, details
            )
            VALUES (:project_id, :entity_type, :entity_id, :action, :actor, CAST(:details AS JSONB))
        """),
        {
            "project_id": project_id,
            "entity_type": entity_type,
            "entity_id": entity_id,
            "action": action,
            "actor": actor,
            "details": _safe_json(details),
        },
    )


def _get_species_maturity_map(project_id: int, db: Session) -> dict[str, int]:
    rows = db.execute(
        text("""
            SELECT species_key, maturity_years
            FROM green_species_maturity
            WHERE project_id = :project_id
        """),
        {"project_id": project_id},
    ).mappings().all()
    result: dict[str, int] = {}
    for row in rows:
        key = _normalize_name(row.get("species_key"))
        years = int(row.get("maturity_years") or 0)
        if key and years > 0:
            result[key] = years
    return result


def _task_needs_evidence(task_type: str | None) -> dict:
    # Premium default: every maintenance completion requires note + photo proof.
    activity = _normalize_name(task_type)
    if activity in ASSIGNABLE_TASK_TYPES:
        return {"require_notes": True, "require_photo": True}
    return {"require_notes": False, "require_photo": False}


def _has_required_evidence(
    task_type: str | None,
    notes: str | None,
    photo_url: str | None,
    photo_urls: object = None,
) -> tuple[bool, str]:
    policy = _task_needs_evidence(task_type)
    notes_ok = bool((notes or "").strip())
    merged_primary, merged_urls = _merge_photo_evidence(photo_url, photo_urls)
    photo_ok = bool(merged_primary) or bool(merged_urls)
    if policy["require_notes"] and not notes_ok:
        return False, "Notes are required before submission."
    if policy["require_photo"] and not photo_ok:
        return False, "Photo proof is required before submission."
    return True, ""


def _is_replacement_trigger_status(status: str | None) -> bool:
    return _normalize_tree_status(status) in REPLACEMENT_TRIGGER_STATUSES


def _tree_status_color_hex(status: str | None) -> str:
    return TREE_STATUS_COLOR_HEX.get(_normalize_tree_status(status), "22c55e")


def _record_alert(
    db: Session,
    project_id: int,
    alert_type: str,
    severity: str,
    message: str,
    tree_id: int | None = None,
    task_id: int | None = None,
    payload: dict | None = None,
):
    existing = db.execute(
        text("""
            SELECT id
            FROM green_alerts
            WHERE project_id = :project_id
              AND COALESCE(tree_id, 0) = COALESCE(:tree_id, 0)
              AND COALESCE(task_id, 0) = COALESCE(:task_id, 0)
              AND alert_type = :alert_type
              AND status = 'open'
            LIMIT 1
        """),
        {
            "project_id": project_id,
            "tree_id": tree_id,
            "task_id": task_id,
            "alert_type": alert_type,
        },
    ).scalar()
    if existing:
        return existing
    return db.execute(
        text("""
            INSERT INTO green_alerts (
                project_id, tree_id, task_id, alert_type, severity, message, payload
            )
            VALUES (
                :project_id, :tree_id, :task_id, :alert_type, :severity, :message, CAST(:payload AS JSONB)
            )
            RETURNING id
        """),
        {
            "project_id": project_id,
            "tree_id": tree_id,
            "task_id": task_id,
            "alert_type": alert_type,
            "severity": severity,
            "message": message,
            "payload": _safe_json(payload),
        },
    ).scalar()


def _resolve_task_alerts(db: Session, task_id: int):
    db.execute(
        text("""
            UPDATE green_alerts
            SET status = 'resolved',
                resolved_at = NOW()
            WHERE task_id = :task_id
              AND status = 'open'
        """),
        {"task_id": task_id},
    )


def _refresh_project_alerts(db: Session, project_id: int):
    db.execute(
        text("""
            DELETE FROM green_alerts
            WHERE project_id = :project_id
              AND status = 'open'
              AND alert_type IN ('task_overdue', 'task_due_soon', 'task_submitted', 'missing_evidence')
        """),
        {"project_id": project_id},
    )

    rows = db.execute(
        text("""
            SELECT t.id AS task_id, t.tree_id, t.task_type, t.assignee_name, t.status, t.review_state,
                   t.due_date, t.notes, t.photo_url, t.completed_at,
                   tr.project_id
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE tr.project_id = :project_id
        """),
        {"project_id": project_id},
    ).mappings().all()

    today = date.today()
    due_soon_limit = today + timedelta(days=3)
    for row in rows:
        task_id = int(row["task_id"])
        tree_id = int(row["tree_id"])
        status = _normalize_name(row.get("status"))
        review_state = _normalize_name(row.get("review_state"))
        due_date = _parse_date_value(row.get("due_date"))
        evidence_ok, _ = _has_required_evidence(row.get("task_type"), row.get("notes"), row.get("photo_url"))

        if review_state == "submitted":
            _record_alert(
                db,
                project_id=project_id,
                alert_type="task_submitted",
                severity="warning",
                message=f"Task #{task_id} is awaiting supervisor review.",
                tree_id=tree_id,
                task_id=task_id,
                payload={"review_state": review_state},
            )

        if status != "done" and due_date:
            if due_date < today:
                days = abs(_day_diff(due_date, today))
                _record_alert(
                    db,
                    project_id=project_id,
                    alert_type="task_overdue",
                    severity="danger",
                    message=f"Task #{task_id} overdue by {days} day{'s' if days != 1 else ''}.",
                    tree_id=tree_id,
                    task_id=task_id,
                    payload={"due_date": _to_date_input(due_date)},
                )
            elif due_date <= due_soon_limit:
                left_days = max(_day_diff(due_date, today), 0)
                _record_alert(
                    db,
                    project_id=project_id,
                    alert_type="task_due_soon",
                    severity="warning",
                    message=f"Task #{task_id} due in {left_days} day{'s' if left_days != 1 else ''}.",
                    tree_id=tree_id,
                    task_id=task_id,
                    payload={"due_date": _to_date_input(due_date)},
                )

        if _is_done_status(status) and review_state == "submitted" and not evidence_ok:
            _record_alert(
                db,
                project_id=project_id,
                alert_type="missing_evidence",
                severity="danger",
                message=f"Task #{task_id} submitted without complete evidence.",
                tree_id=tree_id,
                task_id=task_id,
                payload={"task_type": row.get("task_type")},
            )


def _compute_live_maintenance_rows(
    db: Session,
    project_id: int,
    season_mode: str = "rainy",
    assignee_name: str | None = None,
    tree_scope: str = "new_planting",
) -> dict:
    season = "dry" if _normalize_name(season_mode) == "dry" else "rainy"
    assignee_key = _normalize_name(assignee_name)
    scope_key = _normalize_name(tree_scope or "new_planting")
    if scope_key not in {"new_planting", "existing_inventory", "all"}:
        scope_key = "new_planting"
    trees = db.execute(
        text("""
            SELECT
                id,
                created_by,
                status,
                species,
                planting_date,
                tree_origin,
                created_at,
                tree_age_months,
                COALESCE(count_in_planting_kpis, TRUE) AS count_in_planting_kpis
            FROM trees
            WHERE project_id = :project_id
            ORDER BY id ASC
        """),
        {"project_id": project_id},
    ).mappings().all()
    filtered_trees: list[dict] = []
    for row in trees:
        item = dict(row)
        origin_key = _normalize_tree_origin(item.get("tree_origin"))
        in_new_scope = origin_key == "new_planting" and bool(item.get("count_in_planting_kpis", True))
        in_existing_scope = origin_key == "existing_inventory"
        if scope_key == "new_planting" and not in_new_scope:
            continue
        if scope_key == "existing_inventory" and not in_existing_scope:
            continue
        filtered_trees.append(item)
    task_rows = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.status, t.review_state, t.due_date,
                   t.priority, t.notes, t.photo_url, t.created_at, t.completed_at, t.submitted_at
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE tr.project_id = :project_id
              AND (:assignee_name IS NULL OR t.assignee_name = :assignee_name)
            ORDER BY t.created_at DESC
        """),
        {"project_id": project_id, "assignee_name": assignee_name},
    ).mappings().all()

    species_maturity_map = _get_species_maturity_map(project_id, db)
    task_buckets: dict[str, list[dict]] = {}
    planting_task_buckets: dict[int, list[dict]] = {}
    for task in task_rows:
        task_type_key = _normalize_name(task.get("task_type"))
        tree_id = int(task.get("tree_id") or 0)
        task_item = dict(task)
        if task_type_key == "planting" and tree_id > 0:
            planting_task_buckets.setdefault(tree_id, []).append(task_item)
        if task_type_key not in MAINTENANCE_ACTIVITY_ORDER:
            continue
        key = f"{tree_id}:{task_type_key}"
        task_buckets.setdefault(key, []).append(task_item)

    today = date.today()
    rows: list[dict] = []
    for tree in filtered_trees:
        tree_id = int(tree["id"])
        tree_assignee = str(tree.get("created_by") or "-")
        origin_key = _normalize_tree_origin(tree.get("tree_origin"))
        if assignee_key:
            has_matching_task = any(
                _normalize_name(str(task.get("assignee_name") or "")) == assignee_key and int(task["tree_id"]) == tree_id
                for task in task_rows
            )
            if _normalize_name(tree_assignee) != assignee_key and not has_matching_task:
                continue

        tree_status = _normalize_tree_status(tree.get("status") or "alive")
        planting_submission_task = None
        if tree_status == "pending_planting":
            planting_rows = planting_task_buckets.get(tree_id, [])
            submitted_plantings = [
                task
                for task in planting_rows
                if _is_done_status(task.get("status")) and _normalize_name(task.get("review_state")) == "submitted"
            ]
            submitted_plantings.sort(
                key=lambda task: _parse_date_value(task.get("submitted_at") or task.get("created_at") or task.get("due_date")) or date.min,
                reverse=True,
            )
            planting_submission_task = submitted_plantings[0] if submitted_plantings else None
            if not planting_submission_task:
                # No submitted planting evidence yet, so no provisional maintenance preview.
                continue
        provisional_pending_approval = tree_status == "pending_planting" and planting_submission_task is not None
        replacement_required = _is_replacement_trigger_status(tree_status)
        planting_date_obj = _parse_date_value(tree.get("planting_date"))
        inferred_tree_age_days = _infer_tree_age_days_for_maintenance(tree, today)
        replacement_key = f"{tree_id}:replacement"
        replacement_done = sorted(
            [
                _parse_date_value(task.get("completed_at") or task.get("due_date") or task.get("created_at"))
                for task in task_buckets.get(replacement_key, [])
                if _is_done_status(task.get("status")) and _normalize_name(task.get("review_state")) in {"approved", "none"}
            ],
            reverse=True,
        )
        latest_replacement_date = replacement_done[0] if replacement_done else None
        base_lifecycle_start = planting_date_obj
        if base_lifecycle_start is None and inferred_tree_age_days is not None:
            base_lifecycle_start = _add_days(today, -int(inferred_tree_age_days))
        lifecycle_start = _get_lifecycle_start_date(base_lifecycle_start, latest_replacement_date)
        tree_age_days = _day_diff(today, lifecycle_start) if lifecycle_start else inferred_tree_age_days
        species_key = _normalize_name(tree.get("species"))
        maturity_years = species_maturity_map.get(species_key) if species_key else None
        maturity_reached = (
            tree_status in HEALTHY_TREE_STATUSES
            and maturity_years is not None
            and tree_age_days is not None
            and tree_age_days >= maturity_years * 365
        )

        for activity in MAINTENANCE_ACTIVITY_ORDER:
            if activity == "replacement" and not replacement_required:
                # Replacement is condition-triggered only; hide it unless current tree status requires it.
                continue
            bucket = task_buckets.get(f"{tree_id}:{activity}", [])
            done_tasks = [
                task
                for task in bucket
                if _is_done_status(task.get("status")) and _normalize_name(task.get("review_state")) in {"approved", "none"}
            ]
            open_tasks = [task for task in bucket if not (_is_done_status(task.get("status")) and _normalize_name(task.get("review_state")) in {"approved", "none"})]
            done_tasks.sort(
                key=lambda task: _parse_date_value(task.get("completed_at") or task.get("due_date") or task.get("created_at")) or date.min,
                reverse=True,
            )
            open_tasks.sort(
                key=lambda task: _parse_date_value(task.get("due_date") or task.get("created_at")) or date.max,
            )
            latest_done = done_tasks[0] if done_tasks else None
            active_task = open_tasks[0] if open_tasks else None

            if _should_skip_existing_tree_routine_activity(
                activity,
                origin_key,
                tree_age_days,
                tree_status,
                bool(active_task),
            ):
                continue

            status_text = "No open task"
            tone = "ok"
            indicator = "On schedule"
            model_due: date | None = None
            assigned_due = _parse_date_value(active_task.get("due_date")) if active_task else None

            if replacement_required and activity != "replacement":
                tone = "danger"
                status_text = "Paused"
                indicator = f"Tree status '{tree_status.replace('_', ' ')}' requires replacement first."
            elif activity == "replacement" and replacement_required:
                model_due = today
                if active_task:
                    tone = "warning"
                    status_text = f"Task #{active_task['id']} {active_task.get('status') or 'pending'}"
                    indicator = "Replacement assigned."
                else:
                    tone = "danger"
                    status_text = "Replacement required"
                    indicator = "Replacement due immediately."
            elif tree_status == "need_watering" and activity == "watering":
                model_due = today
                tone = "warning"
                status_text = f"Task #{active_task['id']} {active_task.get('status') or 'pending'}" if active_task else "Action required"
                indicator = "Inspection flagged need watering. Due immediately."
            elif tree_status == "need_protection" and activity == "protection":
                model_due = today
                tone = "warning"
                status_text = f"Task #{active_task['id']} {active_task.get('status') or 'pending'}" if active_task else "Action required"
                indicator = "Inspection flagged need protection. Due immediately."
            elif maturity_reached:
                tone = "info"
                status_text = "Lifecycle complete"
                indicator = f"Tree reached maturity (~{maturity_years} years)."
            else:
                intervals = _get_maintenance_intervals(activity, tree_age_days or 0, season)
                latest_done_date = _parse_date_value(
                    latest_done.get("completed_at") if latest_done else None
                ) or _parse_date_value(latest_done.get("due_date") if latest_done else None) or _parse_date_value(
                    latest_done.get("created_at") if latest_done else None
                )
                if latest_done_date:
                    model_due = _add_days(latest_done_date, intervals["repeat_days"])
                elif lifecycle_start:
                    model_due = _add_days(lifecycle_start, intervals["first_days"])

                if active_task:
                    rs = _normalize_name(active_task.get("review_state"))
                    if rs == "submitted":
                        status_text = f"Task #{active_task['id']} submitted"
                    elif rs == "rejected":
                        status_text = f"Task #{active_task['id']} rejected"
                    else:
                        status_text = f"Task #{active_task['id']} {active_task.get('status') or 'pending'}"

            effective_due: date | None = None
            if model_due and assigned_due:
                effective_due = model_due if model_due <= assigned_due else assigned_due
            else:
                effective_due = model_due or assigned_due

            countdown_days = _day_diff(effective_due, today) if effective_due else None
            if countdown_days is not None:
                if countdown_days < 0:
                    tone = "danger"
                    indicator = f"Not done, overdue by {abs(countdown_days)} day{'s' if abs(countdown_days) != 1 else ''}."
                elif countdown_days == 0:
                    tone = "warning"
                    indicator = "Due today."
                elif countdown_days <= 7:
                    tone = "warning"
                    indicator = f"Due in {countdown_days} day{'s' if countdown_days != 1 else ''}."

            overdue_open = 0
            for task in open_tasks:
                due = _parse_date_value(task.get("due_date"))
                if due and due < today:
                    overdue_open += 1

            intervals = _get_maintenance_intervals(activity, tree_age_days or 0, season)
            if activity == "replacement":
                rationale = (
                    "Replacement is condition-triggered (dead/damaged/removed/needs replacement) "
                    "and is not treated as a routine cyclical task."
                )
            else:
                rationale = f"{season.title()} season model: first {intervals['first_days']}d, repeat {intervals['repeat_days']}d."

            if provisional_pending_approval:
                planting_task_label = f"Planting task #{int(planting_submission_task['id'])}" if planting_submission_task else "Planting task"
                status_text = f"{planting_task_label} submitted (awaiting supervisor approval)"
                if countdown_days is None:
                    tone = "warning"
                    indicator = "Provisional preview while planting approval is pending."
                elif countdown_days < 0:
                    tone = "danger"
                    indicator = (
                        f"Provisional: due from planting date and now {abs(countdown_days)} day"
                        f"{'s' if abs(countdown_days) != 1 else ''} overdue while approval is pending."
                    )
                elif countdown_days == 0:
                    tone = "warning"
                    indicator = "Provisional: due today from planting date once approved."
                elif countdown_days <= 7:
                    tone = "warning"
                    indicator = f"Provisional: due in {countdown_days} day{'s' if countdown_days != 1 else ''} from planting date."
                else:
                    tone = "info"
                    indicator = f"Provisional: scheduled from planting date in {countdown_days} day{'s' if countdown_days != 1 else ''}."
                rationale = (
                    f"{rationale} Provisional preview only: planting submission is awaiting supervisor approval. "
                    "Rows are shown for planning visibility, but maintenance workflow becomes active after approval."
                )

            rows.append(
                {
                    "key": f"{tree_id}:{activity}",
                    "treeId": tree_id,
                    "treeOrigin": origin_key,
                    "assignee": tree_assignee,
                    "activity": activity,
                    "activityLabel": activity.replace("_", " ").title(),
                    "plantingDate": _to_date_input(planting_date_obj),
                    "treeAgeDays": tree_age_days,
                    "lastDoneAt": _to_date_input(_parse_date_value(latest_done.get("completed_at") if latest_done else None) or _parse_date_value(latest_done.get("due_date") if latest_done else None)),
                    "modelDueDate": _to_date_input(model_due),
                    "assignedDueDate": _to_date_input(assigned_due),
                    "effectiveDueDate": _to_date_input(effective_due),
                    "countdownDays": countdown_days,
                    "tone": tone,
                    "indicator": indicator,
                    "statusText": status_text,
                    "doneCount": len(done_tasks),
                    "pendingCount": len(open_tasks),
                    "overdueCount": overdue_open,
                    "openTaskId": int(active_task["id"]) if active_task else None,
                    "modelRationale": rationale,
                }
            )

    tone_order = {"danger": 0, "warning": 1, "info": 2, "ok": 3}
    rows.sort(
        key=lambda item: (
            tone_order.get(item.get("tone"), 3),
            item.get("countdownDays") if item.get("countdownDays") is not None else 999999,
            int(item.get("treeId") or 0),
        )
    )

    summary = {
        "total": len(rows),
        "danger": sum(1 for item in rows if item.get("tone") == "danger"),
        "warning": sum(1 for item in rows if item.get("tone") == "warning"),
        "ok": sum(1 for item in rows if item.get("tone") == "ok"),
        "info": sum(1 for item in rows if item.get("tone") == "info"),
        "dueSoon": sum(1 for item in rows if isinstance(item.get("countdownDays"), int) and 0 <= item["countdownDays"] <= 7),
    }
    return {"rows": rows, "summary": summary}


def _auto_schedule_next_cycle(db: Session, task_id: int, season_hint: str | None = None) -> int | None:
    task = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.priority, t.status, t.review_state,
                   t.completed_at, t.due_date, t.model_season,
                   tr.project_id, tr.status AS tree_status, tr.species, tr.planting_date
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE t.id = :task_id
        """),
        {"task_id": task_id},
    ).mappings().first()
    if not task:
        return None

    if _normalize_name(task.get("task_type")) not in MAINTENANCE_ACTIVITY_ORDER:
        return None
    if not _is_done_status(task.get("status")):
        return None
    if _normalize_name(task.get("review_state")) != "approved":
        return None

    tree_status = _normalize_tree_status(task.get("tree_status") or "alive")
    activity = _normalize_name(task.get("task_type"))
    if activity == "replacement":
        # Replacement is condition-triggered only; do not auto-generate recurring replacement cycles.
        return None
    if _is_replacement_trigger_status(tree_status):
        return None

    season = _normalize_name(season_hint or task.get("model_season") or "rainy")
    if season not in SEASON_VALUES:
        season = "rainy"

    project_id = int(task["project_id"])
    tree_id = int(task["tree_id"])
    today = date.today()
    completed_date = _parse_date_value(task.get("completed_at")) or _parse_date_value(task.get("due_date")) or today

    replacement_done = db.execute(
        text("""
            SELECT COALESCE(completed_at::date, due_date, created_at::date) AS stamp
            FROM tree_tasks
            WHERE tree_id = :tree_id
              AND LOWER(task_type) = 'replacement'
              AND LOWER(status) IN ('done', 'completed', 'closed')
              AND LOWER(review_state) = 'approved'
            ORDER BY stamp DESC
            LIMIT 1
        """),
        {"tree_id": tree_id},
    ).scalar()
    lifecycle_start = _get_lifecycle_start_date(
        _parse_date_value(task.get("planting_date")),
        _parse_date_value(replacement_done),
    )
    tree_age_days = _day_diff(today, lifecycle_start) if lifecycle_start else 0
    species_key = _normalize_name(task.get("species"))
    maturity_map = _get_species_maturity_map(project_id, db)
    maturity_years = maturity_map.get(species_key) if species_key else None
    if tree_status in HEALTHY_TREE_STATUSES and maturity_years and tree_age_days >= maturity_years * 365:
        return None

    open_exists = db.execute(
        text("""
            SELECT id
            FROM tree_tasks
            WHERE tree_id = :tree_id
              AND LOWER(task_type) = :task_type
              AND (
                LOWER(status) NOT IN ('done', 'completed', 'closed')
                OR LOWER(review_state) IN ('submitted', 'rejected', 'reopened')
              )
            LIMIT 1
        """),
        {"tree_id": tree_id, "task_type": activity},
    ).scalar()
    if open_exists:
        return None

    intervals = _get_maintenance_intervals(activity, tree_age_days, season)
    next_due = _add_days(completed_date, intervals["repeat_days"])
    new_task_id = db.execute(
        text("""
            INSERT INTO tree_tasks (
                tree_id, task_type, assignee_name, due_date, priority, status, notes,
                auto_generated, model_season, source_task_id, review_state
            )
            VALUES (
                :tree_id, :task_type, :assignee_name, :due_date, :priority, 'pending', :notes,
                TRUE, :model_season, :source_task_id, 'none'
            )
            RETURNING id
        """),
        {
            "tree_id": tree_id,
            "task_type": activity,
            "assignee_name": task.get("assignee_name"),
            "due_date": next_due,
            "priority": task.get("priority") or "normal",
            "notes": f"Auto-generated next cycle from Task #{task_id}.",
            "model_season": season,
            "source_task_id": task_id,
        },
    ).scalar()

    db.execute(
        text("""
            INSERT INTO green_maintenance_cycles (
                project_id, tree_id, task_type, model_season, source_task_id, generated_task_id, due_date, status
            )
            VALUES (
                :project_id, :tree_id, :task_type, :model_season, :source_task_id, :generated_task_id, :due_date, 'scheduled'
            )
        """),
        {
            "project_id": project_id,
            "tree_id": tree_id,
            "task_type": activity,
            "model_season": season,
            "source_task_id": task_id,
            "generated_task_id": new_task_id,
            "due_date": next_due,
        },
    )
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="task",
        entity_id=int(new_task_id),
        action="auto_cycle_generated",
        actor="system",
        details={
            "source_task_id": task_id,
            "season": season,
            "due_date": _to_date_input(next_due),
        },
    )
    return int(new_task_id)


def ensure_green_tables(db: Session):
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_organizations (
            id SERIAL PRIMARY KEY,
            name TEXT NOT NULL,
            slug TEXT NOT NULL UNIQUE,
            short_name TEXT,
            logo_url TEXT,
            status TEXT NOT NULL DEFAULT 'pilot',
            contact_email TEXT,
            contact_phone TEXT,
            website_url TEXT,
            country TEXT,
            state_region TEXT,
            city TEXT,
            address_text TEXT,
            notes TEXT,
            is_active BOOLEAN NOT NULL DEFAULT TRUE,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS tree_projects (
            id SERIAL PRIMARY KEY,
            organization_id INTEGER REFERENCES green_organizations(id) ON DELETE SET NULL,
            name TEXT NOT NULL,
            location_text TEXT,
            sponsor TEXT,
            planting_model TEXT NOT NULL DEFAULT 'direct',
            allow_existing_tree_link BOOLEAN NOT NULL DEFAULT FALSE,
            default_existing_tree_scope TEXT NOT NULL DEFAULT 'exclude_from_planting_kpi',
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS trees (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            project_tree_no INTEGER,
            geom GEOMETRY(POINT, 4326),
            species TEXT,
            planting_date DATE,
            status TEXT NOT NULL DEFAULT 'alive',
            notes TEXT,
            photo_url TEXT,
            photo_urls JSONB,
            created_by TEXT,
            tree_origin TEXT NOT NULL DEFAULT 'new_planting',
            custodian_id INTEGER,
            custody_started_at DATE,
            attribution_scope TEXT NOT NULL DEFAULT 'full',
            count_in_planting_kpis BOOLEAN NOT NULL DEFAULT TRUE,
            count_in_carbon_scope BOOLEAN NOT NULL DEFAULT TRUE,
            source_project_id INTEGER REFERENCES tree_projects(id) ON DELETE SET NULL,
            tree_height_m NUMERIC,
            tree_age_months NUMERIC,
            inventory_tree_count INTEGER NOT NULL DEFAULT 1,
            existing_area_geojson JSONB,
            existing_area_sqm NUMERIC,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS tree_visits (
            id SERIAL PRIMARY KEY,
            tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
            visit_date DATE NOT NULL,
            status TEXT NOT NULL,
            notes TEXT,
            photo_url TEXT,
            created_by TEXT,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_tree_status_history (
            id SERIAL PRIMARY KEY,
            tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            status TEXT NOT NULL,
            status_date DATE NOT NULL,
            source TEXT NOT NULL DEFAULT 'manual',
            source_task_id INTEGER,
            changed_by TEXT,
            notes TEXT,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_users (
            id SERIAL PRIMARY KEY,
            full_name TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'field_officer',
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_roles (
            id SERIAL PRIMARY KEY,
            role_uid TEXT NOT NULL,
            role_key TEXT NOT NULL,
            role_name TEXT NOT NULL,
            description TEXT,
            scope TEXT NOT NULL DEFAULT 'platform',
            is_system BOOLEAN NOT NULL DEFAULT FALSE,
            is_active BOOLEAN NOT NULL DEFAULT TRUE,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    try:
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS user_uid TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS email TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS phone TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS organization_id INTEGER"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS role_id INTEGER"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS allow_green BOOLEAN NOT NULL DEFAULT TRUE"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS allow_work BOOLEAN NOT NULL DEFAULT FALSE"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS work_username TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS work_password_hash TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS notes TEXT"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS is_active BOOLEAN NOT NULL DEFAULT TRUE"))
        db.execute(text("ALTER TABLE green_users ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT NOW()"))
        db.execute(text("ALTER TABLE green_organizations ADD COLUMN IF NOT EXISTS logo_url TEXT"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS role_uid TEXT"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS role_key TEXT"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS role_name TEXT"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS description TEXT"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS scope TEXT NOT NULL DEFAULT 'platform'"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS is_system BOOLEAN NOT NULL DEFAULT FALSE"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS is_active BOOLEAN NOT NULL DEFAULT TRUE"))
        db.execute(text("ALTER TABLE green_roles ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT NOW()"))
        db.execute(text("ALTER TABLE tree_projects ADD COLUMN IF NOT EXISTS organization_id INTEGER"))
        db.execute(text("ALTER TABLE tree_projects ADD COLUMN IF NOT EXISTS planting_model TEXT NOT NULL DEFAULT 'direct'"))
        db.execute(text("ALTER TABLE tree_projects ADD COLUMN IF NOT EXISTS allow_existing_tree_link BOOLEAN NOT NULL DEFAULT FALSE"))
        db.execute(
            text(
                "ALTER TABLE tree_projects ADD COLUMN IF NOT EXISTS default_existing_tree_scope TEXT NOT NULL DEFAULT 'exclude_from_planting_kpi'"
            )
        )
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS tree_origin TEXT NOT NULL DEFAULT 'new_planting'"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS custodian_id INTEGER"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS custody_started_at DATE"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS attribution_scope TEXT NOT NULL DEFAULT 'full'"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS count_in_planting_kpis BOOLEAN NOT NULL DEFAULT TRUE"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS count_in_carbon_scope BOOLEAN NOT NULL DEFAULT TRUE"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS source_project_id INTEGER"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS tree_height_m NUMERIC"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS tree_age_months NUMERIC"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS project_tree_no INTEGER"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS photo_urls JSONB"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS inventory_tree_count INTEGER NOT NULL DEFAULT 1"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS existing_area_geojson JSONB"))
        db.execute(text("ALTER TABLE trees ADD COLUMN IF NOT EXISTS existing_area_sqm NUMERIC"))
    except Exception:
        db.rollback()
    try:
        db.execute(
            text(
                """
                WITH ranked AS (
                    SELECT
                        id,
                        ROW_NUMBER() OVER (
                            PARTITION BY project_id
                            ORDER BY COALESCE(created_at, NOW()), id
                        ) AS rn
                    FROM trees
                )
                UPDATE trees t
                SET project_tree_no = ranked.rn
                FROM ranked
                WHERE t.id = ranked.id
                  AND (t.project_tree_no IS NULL OR t.project_tree_no <> ranked.rn)
                """
            )
        )
    except Exception:
        db.rollback()
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS green_custodians (
                id SERIAL PRIMARY KEY,
                project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
                custodian_type TEXT NOT NULL DEFAULT 'household',
                name TEXT NOT NULL,
                contact_person TEXT,
                phone TEXT,
                alt_phone TEXT,
                email TEXT,
                address_text TEXT,
                local_government TEXT,
                community_name TEXT,
                verification_status TEXT NOT NULL DEFAULT 'pending',
                notes TEXT,
                created_by TEXT,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
            """
        )
    )
    try:
        db.execute(text("ALTER TABLE green_custodians ADD COLUMN IF NOT EXISTS contact_person TEXT"))
        db.execute(text("ALTER TABLE green_custodians ADD COLUMN IF NOT EXISTS alt_phone TEXT"))
        db.execute(text("ALTER TABLE green_custodians ADD COLUMN IF NOT EXISTS email TEXT"))
        db.execute(text("ALTER TABLE green_custodians ADD COLUMN IF NOT EXISTS local_government TEXT"))
        db.execute(text("ALTER TABLE green_custodians ADD COLUMN IF NOT EXISTS notes TEXT"))
    except Exception:
        db.rollback()
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS green_distribution_events (
                id SERIAL PRIMARY KEY,
                project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
                event_date DATE NOT NULL,
                species TEXT,
                quantity INTEGER NOT NULL DEFAULT 0,
                source_batch_ref TEXT,
                distributed_by TEXT,
                notes TEXT,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW()
            )
            """
        )
    )
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS green_distribution_allocations (
                id SERIAL PRIMARY KEY,
                event_id INTEGER NOT NULL REFERENCES green_distribution_events(id) ON DELETE CASCADE,
                project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
                custodian_id INTEGER NOT NULL REFERENCES green_custodians(id) ON DELETE CASCADE,
                quantity_allocated INTEGER NOT NULL DEFAULT 0,
                supervision_target INTEGER NOT NULL DEFAULT 0,
                expected_planting_start DATE,
                expected_planting_end DATE,
                followup_cycle_days INTEGER NOT NULL DEFAULT 30,
                notes TEXT,
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(event_id, custodian_id)
            )
            """
        )
    )
    try:
        db.execute(text("ALTER TABLE green_distribution_allocations ADD COLUMN IF NOT EXISTS supervision_target INTEGER NOT NULL DEFAULT 0"))
    except Exception:
        db.rollback()
    db.execute(
        text(
            """
            CREATE TABLE IF NOT EXISTS tree_project_links (
                id SERIAL PRIMARY KEY,
                source_project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
                target_project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
                source_tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
                target_tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
                link_type TEXT NOT NULL DEFAULT 'reference',
                transfer_mode TEXT NOT NULL DEFAULT 'reference',
                created_by TEXT,
                created_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(source_project_id, target_project_id, source_tree_id, target_tree_id, link_type)
            )
            """
        )
    )
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS tree_tasks (
            id SERIAL PRIMARY KEY,
            tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
            task_type TEXT NOT NULL,
            assignee_name TEXT NOT NULL,
            due_date DATE,
            priority TEXT DEFAULT 'normal',
            status TEXT NOT NULL DEFAULT 'pending',
            notes TEXT,
            photo_url TEXT,
            created_at TIMESTAMP DEFAULT NOW(),
            completed_at TIMESTAMP
        )
    """))
    try:
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS priority TEXT DEFAULT 'normal'"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS review_state TEXT NOT NULL DEFAULT 'none'"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS submitted_at TIMESTAMP"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS reviewed_at TIMESTAMP"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS reviewed_by TEXT"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS review_notes TEXT"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS auto_generated BOOLEAN NOT NULL DEFAULT FALSE"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS model_season TEXT"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS source_task_id INTEGER"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS reported_tree_status TEXT"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS activity_lng DOUBLE PRECISION"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS activity_lat DOUBLE PRECISION"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS activity_recorded_at TIMESTAMP"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS photo_urls JSONB"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS custodian_id INTEGER"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS distribution_allocation_id INTEGER"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS supervision_visit_no INTEGER"))
        db.execute(text("ALTER TABLE tree_tasks ADD COLUMN IF NOT EXISTS supervision_total_visits INTEGER"))
    except Exception:
        db.rollback()
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_work_orders (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            assignee_name TEXT NOT NULL,
            work_type TEXT NOT NULL,
            target_trees INTEGER DEFAULT 0,
            species_allocations JSONB,
            maintenance_schedule TEXT,
            auto_assign_first_cycle_maintenance BOOLEAN NOT NULL DEFAULT FALSE,
            due_date DATE,
            area_enabled BOOLEAN NOT NULL DEFAULT FALSE,
            area_label TEXT,
            area_geojson JSONB,
            allow_existing_tree_area_reuse BOOLEAN NOT NULL DEFAULT FALSE,
            status TEXT NOT NULL DEFAULT 'assigned',
            planted_count INTEGER DEFAULT 0,
            visits_done INTEGER DEFAULT 0,
            last_update TIMESTAMP DEFAULT NOW(),
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    try:
        db.execute(text("ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS area_enabled BOOLEAN NOT NULL DEFAULT FALSE"))
        db.execute(text("ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS area_label TEXT"))
        db.execute(text("ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS area_geojson JSONB"))
        db.execute(text("ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS species_allocations JSONB"))
        db.execute(
            text(
                "ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS auto_assign_first_cycle_maintenance BOOLEAN NOT NULL DEFAULT FALSE"
            )
        )
        db.execute(
            text(
                "ALTER TABLE green_work_orders ADD COLUMN IF NOT EXISTS allow_existing_tree_area_reuse BOOLEAN NOT NULL DEFAULT FALSE"
            )
        )
    except Exception:
        db.rollback()
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_species_maturity (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            species_key TEXT NOT NULL,
            species_label TEXT,
            maturity_years INTEGER NOT NULL CHECK (maturity_years > 0),
            updated_at TIMESTAMP DEFAULT NOW(),
            created_at TIMESTAMP DEFAULT NOW(),
            UNIQUE(project_id, species_key)
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_task_reviews (
            id SERIAL PRIMARY KEY,
            task_id INTEGER NOT NULL REFERENCES tree_tasks(id) ON DELETE CASCADE,
            decision TEXT NOT NULL,
            reviewer_name TEXT,
            review_notes TEXT,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_audit_events (
            id SERIAL PRIMARY KEY,
            project_id INTEGER REFERENCES tree_projects(id) ON DELETE CASCADE,
            entity_type TEXT NOT NULL,
            entity_id INTEGER,
            action TEXT NOT NULL,
            actor TEXT,
            details JSONB,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_maintenance_cycles (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            tree_id INTEGER NOT NULL REFERENCES trees(id) ON DELETE CASCADE,
            task_type TEXT NOT NULL,
            model_season TEXT,
            source_task_id INTEGER REFERENCES tree_tasks(id) ON DELETE SET NULL,
            generated_task_id INTEGER REFERENCES tree_tasks(id) ON DELETE SET NULL,
            due_date DATE NOT NULL,
            status TEXT NOT NULL DEFAULT 'scheduled',
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_alerts (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            tree_id INTEGER REFERENCES trees(id) ON DELETE SET NULL,
            task_id INTEGER REFERENCES tree_tasks(id) ON DELETE SET NULL,
            alert_type TEXT NOT NULL,
            severity TEXT NOT NULL DEFAULT 'warning',
            message TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'open',
            payload JSONB,
            created_at TIMESTAMP DEFAULT NOW(),
            resolved_at TIMESTAMP
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_scheduled_reports (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            report_type TEXT NOT NULL DEFAULT 'donor',
            report_format TEXT NOT NULL DEFAULT 'pdf',
            recipients TEXT NOT NULL DEFAULT '',
            cron_expr TEXT,
            timezone TEXT NOT NULL DEFAULT 'Africa/Lagos',
            webhook_url TEXT,
            is_enabled BOOLEAN NOT NULL DEFAULT TRUE,
            created_by TEXT,
            last_run_at TIMESTAMP,
            next_run_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_kpi_snapshots (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            snapshot_at TIMESTAMP NOT NULL DEFAULT NOW(),
            metrics JSONB NOT NULL
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_alert_rules (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            rule_name TEXT NOT NULL,
            metric_key TEXT NOT NULL,
            comparator TEXT NOT NULL DEFAULT 'gte',
            threshold NUMERIC NOT NULL,
            severity TEXT NOT NULL DEFAULT 'warning',
            message_template TEXT,
            is_enabled BOOLEAN NOT NULL DEFAULT TRUE,
            created_by TEXT,
            created_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_alert_events (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            rule_id INTEGER REFERENCES green_alert_rules(id) ON DELETE SET NULL,
            severity TEXT NOT NULL DEFAULT 'warning',
            status TEXT NOT NULL DEFAULT 'open',
            metric_key TEXT,
            metric_value NUMERIC,
            threshold NUMERIC,
            message TEXT NOT NULL,
            payload JSONB,
            triggered_at TIMESTAMP DEFAULT NOW(),
            resolved_at TIMESTAMP
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_webhook_deliveries (
            id SERIAL PRIMARY KEY,
            event_id INTEGER REFERENCES green_alert_events(id) ON DELETE CASCADE,
            target_url TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            response_code INTEGER,
            response_body TEXT,
            attempt_count INTEGER NOT NULL DEFAULT 0,
            delivered_at TIMESTAMP,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("""
        CREATE TABLE IF NOT EXISTS green_verra_exports (
            id SERIAL PRIMARY KEY,
            project_id INTEGER NOT NULL REFERENCES tree_projects(id) ON DELETE CASCADE,
            season_mode TEXT NOT NULL DEFAULT 'rainy',
            assignee_name TEXT,
            output_format TEXT NOT NULL DEFAULT 'zip',
            monitoring_start DATE,
            monitoring_end DATE,
            methodology_id TEXT,
            verifier_notes TEXT,
            generated_by TEXT,
            file_name TEXT,
            payload_summary JSONB,
            created_at TIMESTAMP DEFAULT NOW()
        )
    """))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_project_id ON trees(project_id)"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_trees_project_tree_no ON trees(project_id, project_tree_no)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_geom ON trees USING GIST (geom)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_visits_tree_id ON tree_visits(tree_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_status_history_tree_date ON green_tree_status_history(tree_id, status_date DESC, id DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_status_history_project_date ON green_tree_status_history(project_id, status_date DESC, id DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_tasks_tree_id ON tree_tasks(tree_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_tasks_review_state ON tree_tasks(review_state)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_tasks_allocation ON tree_tasks(distribution_allocation_id, task_type, review_state)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_tasks_custodian ON tree_tasks(custodian_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_green_users_name ON green_users(full_name)"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_green_users_uid ON green_users(UPPER(user_uid))"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_green_users_work_username ON green_users(LOWER(work_username))"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_green_users_org ON green_users(organization_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_green_users_role_id ON green_users(role_id)"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_green_roles_uid ON green_roles(UPPER(role_uid))"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_green_roles_key ON green_roles(LOWER(role_key))"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_green_roles_active ON green_roles(is_active)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_green_orgs_name ON green_organizations(name)"))
    db.execute(text("CREATE UNIQUE INDEX IF NOT EXISTS idx_green_orgs_slug ON green_organizations(LOWER(slug))"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_work_orders_project_id ON green_work_orders(project_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_species_maturity_project_id ON green_species_maturity(project_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_task_reviews_task_id ON green_task_reviews(task_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_audit_project_created ON green_audit_events(project_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_cycles_project_due ON green_maintenance_cycles(project_id, due_date)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_alerts_project_status ON green_alerts(project_id, status, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_sched_reports_project ON green_scheduled_reports(project_id, is_enabled)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_kpi_snapshots_project_time ON green_kpi_snapshots(project_id, snapshot_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_alert_rules_project_enabled ON green_alert_rules(project_id, is_enabled)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_alert_events_project_time ON green_alert_events(project_id, triggered_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_webhook_event ON green_webhook_deliveries(event_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_verra_exports_project_time ON green_verra_exports(project_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_projects_model ON tree_projects(planting_model)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_projects_org ON tree_projects(organization_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_origin ON trees(tree_origin)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_scope_flags ON trees(project_id, count_in_planting_kpis, count_in_carbon_scope)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_custodian_id ON trees(custodian_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_trees_source_project_id ON trees(source_project_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_custodians_project ON green_custodians(project_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_dist_events_project ON green_distribution_events(project_id, event_date DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_dist_alloc_project ON green_distribution_allocations(project_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_dist_alloc_custodian ON green_distribution_allocations(custodian_id, created_at DESC)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_tree_project_links_target ON tree_project_links(target_project_id, created_at DESC)"))
    for role_key, role_name, is_system in [
        ("admin", "Admin", True),
        ("supervisor", "Supervisor", True),
        ("field_officer", "Field Officer", True),
        ("volunteer", "Volunteer", True),
        ("viewer", "Viewer", True),
    ]:
        existing_role = db.execute(
            text("SELECT id, role_uid FROM green_roles WHERE LOWER(role_key) = LOWER(:role_key) LIMIT 1"),
            {"role_key": role_key},
        ).mappings().first()
        if existing_role:
            if not existing_role.get("role_uid"):
                db.execute(
                    text(
                        """
                        UPDATE green_roles
                        SET role_uid = :role_uid,
                            role_name = COALESCE(NULLIF(role_name, ''), :role_name),
                            is_system = TRUE,
                            updated_at = NOW()
                        WHERE id = :id
                        """
                    ),
                    {
                        "id": int(existing_role["id"]),
                        "role_uid": _ensure_unique_role_uid(db),
                        "role_name": role_name,
                    },
                )
        else:
            db.execute(
                text(
                    """
                    INSERT INTO green_roles (role_uid, role_key, role_name, is_system, is_active)
                    VALUES (:role_uid, :role_key, :role_name, :is_system, TRUE)
                    """
                ),
                {
                    "role_uid": _ensure_unique_role_uid(db),
                    "role_key": role_key,
                    "role_name": role_name,
                    "is_system": bool(is_system),
                },
            )
    legacy_roles = db.execute(
        text(
            """
            SELECT DISTINCT TRIM(role) AS role_value
            FROM green_users
            WHERE COALESCE(TRIM(role), '') <> ''
            """
        )
    ).scalars().all()
    for role_value in legacy_roles:
        role_value_str = str(role_value or "").strip()
        if not role_value_str:
            continue
        role_key = _ensure_unique_role_key(db, role_value_str)
        existing = db.execute(
            text("SELECT id FROM green_roles WHERE LOWER(role_key) = LOWER(:role_key) LIMIT 1"),
            {"role_key": role_key},
        ).scalar()
        if not existing:
            db.execute(
                text(
                    """
                    INSERT INTO green_roles (role_uid, role_key, role_name, is_system, is_active)
                    VALUES (:role_uid, :role_key, :role_name, FALSE, TRUE)
                    """
                ),
                {
                    "role_uid": _ensure_unique_role_uid(db),
                    "role_key": role_key,
                    "role_name": role_value_str.replace("_", " ").title(),
                },
            )
    missing_user_rows = db.execute(
        text(
            """
            SELECT id
            FROM green_users
            WHERE COALESCE(TRIM(user_uid), '') = ''
            ORDER BY id ASC
            """
        )
    ).scalars().all()
    for user_id in missing_user_rows:
        db.execute(
            text("UPDATE green_users SET user_uid = :user_uid, updated_at = COALESCE(updated_at, NOW()) WHERE id = :id"),
            {"id": int(user_id), "user_uid": _ensure_unique_user_uid(db)},
        )
    users_missing_role_fk = db.execute(
        text(
            """
            SELECT u.id, u.role
            FROM green_users u
            WHERE u.role_id IS NULL
            """
        )
    ).mappings().all()
    for user_row in users_missing_role_fk:
        raw_role = str(user_row.get("role") or "").strip()
        if not raw_role:
            raw_role = "field_officer"
        direct_match = db.execute(
            text("SELECT id FROM green_roles WHERE LOWER(role_key) = LOWER(:rk) LIMIT 1"),
            {"rk": raw_role},
        ).scalar()
        role_id_val = direct_match
        if not role_id_val:
            normalized_key = _slugify_text(raw_role, fallback="role").replace("-", "_")
            role_id_val = db.execute(
                text("SELECT id FROM green_roles WHERE LOWER(role_key) = LOWER(:rk) LIMIT 1"),
                {"rk": normalized_key},
            ).scalar()
        if role_id_val:
            db.execute(
                text("UPDATE green_users SET role_id = :role_id, updated_at = COALESCE(updated_at, NOW()) WHERE id = :id"),
                {"id": int(user_row["id"]), "role_id": int(role_id_val)},
            )
    db.execute(
        text(
            """
            UPDATE green_users
            SET allow_work = TRUE
            WHERE allow_work = FALSE
              AND LOWER(COALESCE(role, '')) IN ('admin', 'supervisor')
            """
        )
    )
    db.execute(
        text(
            """
            INSERT INTO green_tree_status_history (tree_id, project_id, status, status_date, source, changed_by, notes)
            SELECT
                t.id,
                t.project_id,
                COALESCE(NULLIF(TRIM(t.status), ''), 'alive'),
                COALESCE(t.planting_date, t.created_at::date, CURRENT_DATE),
                'seed',
                t.created_by,
                'Auto-seeded baseline status from current tree row'
            FROM trees t
            WHERE NOT EXISTS (
                SELECT 1
                FROM green_tree_status_history h
                WHERE h.tree_id = t.id
            )
            """
        )
    )
    db.commit()


def _load_env_token() -> str | None:
    token_keys = ("MAPBOX_TOKEN", "MAPBOX_ACCESS_TOKEN", "MAPBOX_PUBLIC_TOKEN", "VITE_MAPBOX_TOKEN")
    for key in token_keys:
        token = os.environ.get(key)
        if token:
            return token.strip().strip('"').strip("'")
    env_path = os.path.join(BASE_DIR, "..", ".env")
    if not os.path.exists(env_path):
        return None
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip().startswith("#") or "=" not in line:
                    continue
                key, val = line.strip().split("=", 1)
                if key in {"MAPBOX_TOKEN", "MAPBOX_ACCESS_TOKEN", "MAPBOX_PUBLIC_TOKEN", "VITE_MAPBOX_TOKEN"}:
                    return val.strip().strip('"').strip("'")
    except Exception:
        return None
    return None


def _http_get_binary(url: str, timeout: int = 15) -> bytes | None:
    try:
        import requests

        resp = requests.get(url, timeout=timeout)
        if resp.status_code == 200 and resp.content:
            return resp.content
    except Exception:
        pass

    try:
        import urllib.request

        req = urllib.request.Request(url, headers={"User-Agent": "LandCheck/1.0"})
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = resp.read()
            if data:
                return data
    except Exception:
        return None
    return None


def _build_report_map_png(
    map_rows: list[dict],
    lng: float | None = None,
    lat: float | None = None,
    zoom: float | None = None,
    bearing: float | None = 0.0,
    pitch: float | None = 0.0,
) -> bytes | None:
    if not map_rows:
        return None

    def _coerce_optional_float(value: object) -> float | None:
        if value is None:
            return None
        try:
            if isinstance(value, bool):
                return None
            return float(value)
        except Exception:
            return None

    lng_value = _coerce_optional_float(lng)
    lat_value = _coerce_optional_float(lat)
    zoom_value = _coerce_optional_float(zoom)
    bearing_value = _coerce_optional_float(bearing)
    pitch_value = _coerce_optional_float(pitch)

    lats = [r.get("lat") for r in map_rows if r.get("lat") is not None]
    lngs = [r.get("lng") for r in map_rows if r.get("lng") is not None]
    if not lats or not lngs:
        return None

    center_lat = lat_value if lat_value is not None else sum(lats) / len(lats)
    center_lng = lng_value if lng_value is not None else sum(lngs) / len(lngs)
    z = zoom_value if zoom_value is not None else 13
    b = bearing_value if bearing_value is not None else 0
    p = pitch_value if pitch_value is not None else 0

    token = _load_env_token()
    if token:
        markers = []
        for r in map_rows[:60]:
            if r.get("lng") is None or r.get("lat") is None:
                continue
            color = _tree_status_color_hex(r.get("status"))
            markers.append(f"pin-s+{color}({r['lng']},{r['lat']})")
        overlay = ",".join(markers) if markers else ""
        overlay_part = f"{quote(overlay, safe='(),:+')}/" if overlay else ""
        for style in ("mapbox/satellite-streets-v12", "mapbox/satellite-v9"):
            if overlay:
                url = (
                    f"https://api.mapbox.com/styles/v1/{style}/static/"
                    f"{overlay_part}{center_lng},{center_lat},{z},{b},{p}/800x500@2x?access_token={token}"
                )
                map_png = _http_get_binary(url)
                if map_png:
                    return map_png
            url = (
                f"https://api.mapbox.com/styles/v1/{style}/static/"
                f"{center_lng},{center_lat},{z},{b},{p}/800x500@2x?access_token={token}"
            )
            map_png = _http_get_binary(url)
            if map_png:
                return map_png

    markers = "|".join([f"{r['lat']},{r['lng']},lightgreen1" for r in map_rows[:50] if r.get("lat") is not None and r.get("lng") is not None])
    marker_qs = quote(markers, safe="|,")
    osm_url = (
        "https://staticmap.openstreetmap.de/staticmap.php?"
        f"center={center_lat},{center_lng}&zoom={int(round(z))}&size=800x500&markers={marker_qs}"
    )
    return _http_get_binary(osm_url)


def _build_r2_settings() -> dict:
    default_endpoint = "https://751ea1abdb3fb6ff7f276b3753e4c6a1.r2.cloudflarestorage.com"
    default_bucket = "photosgreen"
    default_public_base = f"{default_endpoint}/{default_bucket}"

    endpoint_raw = (os.getenv("R2_ENDPOINT_URL") or default_endpoint).strip()
    public_base = (os.getenv("R2_PUBLIC_BASE_URL") or default_public_base).strip()
    bucket = (os.getenv("R2_BUCKET") or default_bucket).strip()
    access_key = (os.getenv("R2_ACCESS_KEY_ID") or "").strip()
    secret_key = (os.getenv("R2_SECRET_ACCESS_KEY") or "").strip()
    region = (os.getenv("R2_REGION") or "auto").strip()

    raw_for_parse = endpoint_raw or public_base
    if not raw_for_parse:
        raise HTTPException(
            status_code=500,
            detail="R2 is not configured. Set R2_ENDPOINT_URL (or R2_PUBLIC_BASE_URL), R2_BUCKET, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY.",
        )

    parsed = urlparse(raw_for_parse)
    if not parsed.scheme or not parsed.netloc:
        raise HTTPException(status_code=500, detail="Invalid R2 endpoint/public base URL format.")

    path_parts = [part for part in parsed.path.split("/") if part]
    if path_parts and not bucket:
        bucket = path_parts[0]

    endpoint_url = endpoint_raw.strip() if endpoint_raw else f"{parsed.scheme}://{parsed.netloc}"
    endpoint_parsed = urlparse(endpoint_url)
    endpoint_url = f"{endpoint_parsed.scheme}://{endpoint_parsed.netloc}"

    if not public_base:
        public_base = f"{endpoint_url.rstrip('/')}/{bucket}"

    if not bucket:
        raise HTTPException(status_code=500, detail="R2 bucket is not configured. Set R2_BUCKET.")
    if not access_key or not secret_key:
        raise HTTPException(
            status_code=500,
            detail="R2 credentials are not configured. Set R2_ACCESS_KEY_ID and R2_SECRET_ACCESS_KEY.",
        )

    return {
        "endpoint_url": endpoint_url,
        "public_base": public_base.rstrip("/"),
        "bucket": bucket,
        "access_key": access_key,
        "secret_key": secret_key,
        "region": region,
    }


def _make_r2_client(settings: dict):
    return boto3.client(
        "s3",
        endpoint_url=settings["endpoint_url"],
        aws_access_key_id=settings["access_key"],
        aws_secret_access_key=settings["secret_key"],
        region_name=settings["region"],
    )


def _normalize_object_key(raw_key: str, bucket: str) -> str:
    key = (raw_key or "").strip().lstrip("/")
    if not key:
        return ""

    # Handle keys that were encoded once/twice by older clients.
    for _ in range(3):
        decoded = unquote(key)
        if decoded == key:
            break
        key = decoded

    bucket_prefix = f"{bucket}/"
    if key.startswith(bucket_prefix):
        key = key[len(bucket_prefix):]
    return key


@router.get("/uploads/object/{object_key:path}")
def get_uploaded_photo(object_key: str):
    settings = _build_r2_settings()
    resolved_key = _normalize_object_key(object_key, settings["bucket"])
    if not resolved_key:
        raise HTTPException(status_code=400, detail="Invalid photo key.")
    try:
        client = _make_r2_client(settings)
        obj = client.get_object(Bucket=settings["bucket"], Key=resolved_key)
    except ClientError as exc:
        code = (exc.response.get("Error") or {}).get("Code", "")
        if code in {"NoSuchKey", "404", "NotFound"}:
            raise HTTPException(status_code=404, detail="Photo not found.")
        raise HTTPException(status_code=502, detail="Failed to read photo from storage.")
    except Exception:
        raise HTTPException(status_code=502, detail="Failed to read photo from storage.")

    content_type = obj.get("ContentType") or "application/octet-stream"
    cache_control = obj.get("CacheControl") or "public, max-age=86400"
    return StreamingResponse(
        obj["Body"].iter_chunks(),
        media_type=content_type,
        headers={"Cache-Control": cache_control},
    )


@router.post("/uploads/photo")
async def upload_photo_to_r2(
    request: Request,
    db: Session = Depends(get_db),
    file: UploadFile = File(...),
    folder: str = Form(default="trees"),
    tree_id: int | None = Form(default=None),
    task_id: int | None = Form(default=None),
):
    if tree_id is not None and task_id is not None:
        raise HTTPException(status_code=400, detail="Provide either tree_id or task_id, not both.")

    content_type = (file.content_type or "").strip().lower()
    if not content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="Only image uploads are allowed.")

    payload = await file.read()
    if not payload:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    max_bytes = 10 * 1024 * 1024
    if len(payload) > max_bytes:
        raise HTTPException(status_code=413, detail="Image too large. Max size is 10MB.")

    settings = _build_r2_settings()

    ext = Path(file.filename or "").suffix.lower()
    allowed_ext = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".heic", ".heif"}
    if ext not in allowed_ext:
        content_ext = {
            "image/jpeg": ".jpg",
            "image/png": ".png",
            "image/webp": ".webp",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/heic": ".heic",
            "image/heif": ".heif",
        }
        ext = content_ext.get(content_type, ".jpg")

    folder_parts = [part for part in (folder or "trees").split("/") if part and part not in {".", ".."}]
    safe_folder = "/".join(folder_parts) or "trees"
    date_path = datetime.utcnow().strftime("%Y/%m")
    object_key = f"{safe_folder}/{date_path}/{uuid.uuid4().hex}{ext}"

    try:
        client = _make_r2_client(settings)
        client.put_object(
            Bucket=settings["bucket"],
            Key=object_key,
            Body=payload,
            ContentType=content_type,
            CacheControl="public, max-age=31536000, immutable",
        )
    except Exception:
        raise HTTPException(status_code=502, detail="Photo upload failed.")

    public_url = f"{settings['public_base']}/{quote(object_key, safe='/')}"
    app_base = str(request.base_url).rstrip("/")
    proxy_url = f"{app_base}/green/uploads/object/{quote(object_key, safe='/')}"

    linked_tree_id = None
    linked_task_id = None

    if tree_id is not None:
        tree_row = db.execute(
            text(
                """
                SELECT id, photo_urls
                FROM trees
                WHERE id = :tree_id
                """
            ),
            {"tree_id": tree_id},
        ).mappings().first()
        if not tree_row:
            db.rollback()
            raise HTTPException(status_code=404, detail="Tree not found for photo link.")
        merged_tree_urls = _normalize_photo_urls(tree_row.get("photo_urls"))
        if proxy_url not in merged_tree_urls:
            merged_tree_urls.append(proxy_url)
        linked_tree_id = db.execute(text("""
            UPDATE trees
            SET photo_url = :photo_url,
                photo_urls = CAST(:photo_urls AS JSONB)
            WHERE id = :tree_id
            RETURNING id
        """), {"photo_url": proxy_url, "photo_urls": _safe_json(merged_tree_urls), "tree_id": tree_id}).scalar()
        if not linked_tree_id:
            db.rollback()
            raise HTTPException(status_code=404, detail="Tree not found for photo link.")
        linked_task_candidate = db.execute(
            text("""
                SELECT id, photo_urls
                FROM tree_tasks
                WHERE tree_id = :tree_id
                  AND LOWER(COALESCE(task_type, '')) = 'planting'
                  AND LOWER(COALESCE(review_state, 'none')) <> 'approved'
                ORDER BY created_at DESC, id DESC
                LIMIT 1
            """),
            {"tree_id": tree_id},
        ).mappings().first()
        if linked_task_candidate:
            merged_urls = _normalize_photo_urls(linked_task_candidate.get("photo_urls"))
            if proxy_url not in merged_urls:
                merged_urls.append(proxy_url)
            linked_task_row = db.execute(
                text(
                    """
                    UPDATE tree_tasks
                    SET photo_url = :photo_url,
                        photo_urls = CAST(:photo_urls AS JSONB)
                    WHERE id = :task_id
                    RETURNING id
                    """
                ),
                {
                    "photo_url": proxy_url,
                    "photo_urls": _safe_json(merged_urls),
                    "task_id": int(linked_task_candidate.get("id") or 0),
                },
            ).mappings().first()
            if linked_task_row:
                linked_task_id = linked_task_row["id"]
        db.commit()
    elif task_id is not None:
        locked = db.execute(
            text("""
                SELECT review_state, photo_urls
                FROM tree_tasks
                WHERE id = :task_id
            """),
            {"task_id": task_id},
        ).mappings().first()
        if locked and _normalize_name(str(locked.get("review_state"))) == "approved":
            raise HTTPException(status_code=409, detail="Task already approved and locked")
        merged_urls = _normalize_photo_urls((locked or {}).get("photo_urls"))
        if proxy_url not in merged_urls:
            merged_urls.append(proxy_url)
        task_row = db.execute(text("""
            UPDATE tree_tasks
            SET photo_url = :photo_url,
                photo_urls = CAST(:photo_urls AS JSONB)
            WHERE id = :task_id
            RETURNING id, tree_id
        """), {"photo_url": proxy_url, "photo_urls": _safe_json(merged_urls), "task_id": task_id}).mappings().first()
        if not task_row:
            db.rollback()
            raise HTTPException(status_code=404, detail="Task not found for photo link.")
        linked_task_id = task_row["id"]
        linked_tree_id = task_row["tree_id"]
        existing_tree_photo_urls = db.execute(
            text("SELECT photo_urls FROM trees WHERE id = :tree_id"),
            {"tree_id": linked_tree_id},
        ).scalar()
        merged_tree_urls = _normalize_photo_urls(existing_tree_photo_urls)
        if proxy_url not in merged_tree_urls:
            merged_tree_urls.append(proxy_url)
        db.execute(text("""
            UPDATE trees
            SET photo_url = :photo_url,
                photo_urls = CAST(:photo_urls AS JSONB)
            WHERE id = :tree_id
        """), {"photo_url": proxy_url, "photo_urls": _safe_json(merged_tree_urls), "tree_id": linked_tree_id})
        db.commit()

    return {
        "url": proxy_url,
        "key": object_key,
        "public_url": public_url,
        "linked_tree_id": linked_tree_id,
        "linked_task_id": linked_task_id,
    }


@router.post("/projects")
def create_project(
    db: Session = Depends(get_db),
    name: str = Body(...),
    location_text: str = Body(default=""),
    sponsor: str = Body(default=""),
    organization_id: int | None = Body(default=None),
    planting_model: str = Body(default=DEFAULT_PLANTING_MODEL),
    allow_existing_tree_link: bool = Body(default=False),
    default_existing_tree_scope: str = Body(default=DEFAULT_EXISTING_TREE_SCOPE),
):
    normalized_model = _normalize_planting_model(planting_model)
    normalized_existing_scope = _normalize_existing_tree_scope(default_existing_tree_scope)
    org_id_value = int(organization_id) if organization_id is not None else None
    if org_id_value is not None and org_id_value <= 0:
        org_id_value = None
    if org_id_value is not None:
        org_exists = db.execute(
            text("SELECT id FROM green_organizations WHERE id = :org_id AND COALESCE(is_active, TRUE) = TRUE"),
            {"org_id": org_id_value},
        ).scalar()
        if not org_exists:
            raise HTTPException(status_code=400, detail="Selected organization not found or inactive")
    row = db.execute(
        text("""
            INSERT INTO tree_projects (
                organization_id, name, location_text, sponsor, planting_model, allow_existing_tree_link, default_existing_tree_scope
            )
            VALUES (
                :organization_id, :name, :location_text, :sponsor, :planting_model, :allow_existing_tree_link, :default_existing_tree_scope
            )
            RETURNING id, organization_id, name, location_text, sponsor, planting_model, allow_existing_tree_link, default_existing_tree_scope, created_at
        """),
        {
            "organization_id": org_id_value,
            "name": name,
            "location_text": location_text,
            "sponsor": sponsor,
            "planting_model": normalized_model,
            "allow_existing_tree_link": bool(allow_existing_tree_link),
            "default_existing_tree_scope": normalized_existing_scope,
        },
    ).mappings().first()
    project = dict(row)
    _log_audit_event(
        db,
        project_id=project["id"],
        entity_type="project",
        entity_id=project["id"],
        action="project_created",
        details={
            "name": project.get("name"),
            "location_text": project.get("location_text"),
            "organization_id": project.get("organization_id"),
            "planting_model": project.get("planting_model"),
            "allow_existing_tree_link": bool(project.get("allow_existing_tree_link")),
            "default_existing_tree_scope": project.get("default_existing_tree_scope"),
        },
    )
    if project.get("organization_id"):
        org_meta = db.execute(
            text("SELECT name, slug, status, logo_url FROM green_organizations WHERE id = :org_id"),
            {"org_id": int(project["organization_id"])},
        ).mappings().first()
        if org_meta:
            project["organization_name"] = org_meta.get("name")
            project["organization_slug"] = org_meta.get("slug")
            project["organization_status"] = org_meta.get("status")
            project["organization_logo_url"] = org_meta.get("logo_url")
    db.commit()
    return project


@router.delete("/projects/{project_id}")
def delete_project(
    project_id: int,
    db: Session = Depends(get_db),
    confirm_name: str = Body(..., embed=True),
):
    project = db.execute(
        text(
            """
            SELECT id, name
            FROM tree_projects
            WHERE id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    expected_name = str(project.get("name") or "").strip()
    provided_name = str(confirm_name or "").strip()
    if not provided_name:
        raise HTTPException(status_code=400, detail="Project name confirmation is required")
    if provided_name != expected_name:
        raise HTTPException(status_code=400, detail="Project name confirmation does not match")

    tree_count = int(
        db.execute(
            text(
                """
                SELECT COUNT(*) AS total
                FROM trees
                WHERE project_id = :project_id
                """
            ),
            {"project_id": project_id},
        ).scalar()
        or 0
    )
    task_count = int(
        db.execute(
            text(
                """
                SELECT COUNT(*) AS total
                FROM tree_tasks
                WHERE tree_id IN (
                    SELECT id
                    FROM trees
                    WHERE project_id = :project_id
                )
                """
            ),
            {"project_id": project_id},
        ).scalar()
        or 0
    )
    work_order_count = int(
        db.execute(
            text(
                """
                SELECT COUNT(*) AS total
                FROM green_work_orders
                WHERE project_id = :project_id
                """
            ),
            {"project_id": project_id},
        ).scalar()
        or 0
    )

    try:
        deleted = db.execute(
            text(
                """
                DELETE FROM tree_projects
                WHERE id = :project_id
                RETURNING id
                """
            ),
            {"project_id": project_id},
        ).scalar()
        if not deleted:
            db.rollback()
            raise HTTPException(status_code=404, detail="Project not found")
        db.commit()
    except HTTPException:
        raise
    except Exception:
        db.rollback()
        raise HTTPException(status_code=409, detail="Project could not be deleted. Please retry.")

    return {
        "status": "ok",
        "deleted_project_id": project_id,
        "deleted_project_name": expected_name,
        "deleted_summary": {
            "trees": tree_count,
            "tasks": task_count,
            "work_orders": work_order_count,
        },
    }


@router.patch("/projects/{project_id}/settings")
def update_project_settings(
    project_id: int,
    db: Session = Depends(get_db),
    planting_model: str | None = Body(default=None),
    allow_existing_tree_link: bool | None = Body(default=None),
    default_existing_tree_scope: str | None = Body(default=None),
):
    existing = db.execute(
        text(
            """
            SELECT id, planting_model, allow_existing_tree_link, default_existing_tree_scope
            FROM tree_projects
            WHERE id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Project not found")

    next_model = _normalize_planting_model(planting_model) if planting_model is not None else _normalize_planting_model(existing.get("planting_model"))
    next_allow_existing = (
        bool(allow_existing_tree_link)
        if allow_existing_tree_link is not None
        else bool(existing.get("allow_existing_tree_link"))
    )
    next_scope = (
        _normalize_existing_tree_scope(default_existing_tree_scope)
        if default_existing_tree_scope is not None
        else _normalize_existing_tree_scope(existing.get("default_existing_tree_scope"))
    )

    row = db.execute(
        text(
            """
            UPDATE tree_projects
            SET planting_model = :planting_model,
                allow_existing_tree_link = :allow_existing_tree_link,
                default_existing_tree_scope = :default_existing_tree_scope
            WHERE id = :project_id
            RETURNING id, name, location_text, sponsor, planting_model, allow_existing_tree_link, default_existing_tree_scope, created_at
            """
        ),
        {
            "project_id": project_id,
            "planting_model": next_model,
            "allow_existing_tree_link": next_allow_existing,
            "default_existing_tree_scope": next_scope,
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="project",
        entity_id=project_id,
        action="project_settings_updated",
        details={
            "before": dict(existing),
            "after": {
                "planting_model": next_model,
                "allow_existing_tree_link": next_allow_existing,
                "default_existing_tree_scope": next_scope,
            },
        },
    )
    db.commit()
    return dict(row)


@router.patch("/projects/{project_id}/organization")
def assign_project_organization(
    project_id: int,
    db: Session = Depends(get_db),
    organization_id: int | None = Body(default=None),
):
    project_row = db.execute(
        text("SELECT id, name, organization_id FROM tree_projects WHERE id = :project_id"),
        {"project_id": project_id},
    ).mappings().first()
    if not project_row:
        raise HTTPException(status_code=404, detail="Project not found")
    org_id_value = int(organization_id) if organization_id is not None else None
    org_row = None
    if org_id_value is not None:
        org_row = db.execute(
            text("SELECT id, name, slug, status, logo_url FROM green_organizations WHERE id = :org_id"),
            {"org_id": org_id_value},
        ).mappings().first()
        if not org_row:
            raise HTTPException(status_code=404, detail="Organization not found")
    updated = db.execute(
        text(
            """
            UPDATE tree_projects
            SET organization_id = :organization_id
            WHERE id = :project_id
            RETURNING id, organization_id, name
            """
        ),
        {"organization_id": org_id_value, "project_id": project_id},
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="project",
        entity_id=project_id,
        action="project_organization_assigned",
        actor="super_admin",
        details={
            "before_organization_id": project_row.get("organization_id"),
            "after_organization_id": org_id_value,
            "after_organization_name": (org_row or {}).get("name") if org_row else None,
        },
    )
    db.commit()
    payload = dict(updated)
    if org_row:
        payload["organization_name"] = org_row.get("name")
        payload["organization_slug"] = org_row.get("slug")
        payload["organization_status"] = org_row.get("status")
        payload["organization_logo_url"] = org_row.get("logo_url")
    else:
        payload["organization_name"] = None
        payload["organization_slug"] = None
        payload["organization_status"] = None
        payload["organization_logo_url"] = None
    return payload


@router.get("/admin/organizations")
def list_admin_organizations(db: Session = Depends(get_db)):
    rows = db.execute(
        text(
            """
            SELECT
                o.id, o.name, o.slug, o.short_name, o.logo_url, o.status, o.contact_email, o.contact_phone, o.website_url,
                o.country, o.state_region, o.city, o.address_text, o.notes, COALESCE(o.is_active, TRUE) AS is_active,
                o.created_at, o.updated_at,
                COALESCE((
                    SELECT COUNT(*) FROM tree_projects p WHERE p.organization_id = o.id
                ), 0) AS projects_count,
                COALESCE((
                    SELECT COUNT(*) FROM trees t
                    JOIN tree_projects p ON p.id = t.project_id
                    WHERE p.organization_id = o.id
                ), 0) AS trees_count,
                COALESCE((
                    SELECT COUNT(*) FROM tree_tasks tt
                    JOIN trees t ON t.id = tt.tree_id
                    JOIN tree_projects p ON p.id = t.project_id
                    WHERE p.organization_id = o.id
                ), 0) AS tasks_count,
                COALESCE((
                    SELECT COUNT(*) FROM tree_tasks tt
                    JOIN trees t ON t.id = tt.tree_id
                    JOIN tree_projects p ON p.id = t.project_id
                    WHERE p.organization_id = o.id
                      AND LOWER(COALESCE(tt.review_state, 'none')) = 'submitted'
                ), 0) AS pending_review_count,
                COALESCE((
                    SELECT COUNT(*) FROM green_alerts ga
                    JOIN tree_projects p ON p.id = ga.project_id
                    WHERE p.organization_id = o.id
                      AND LOWER(COALESCE(ga.status, 'open')) = 'open'
                ), 0) AS open_alert_count,
                (
                    SELECT MAX(a.created_at)
                    FROM green_audit_events a
                    JOIN tree_projects p ON p.id = a.project_id
                    WHERE p.organization_id = o.id
                ) AS last_activity_at
            FROM green_organizations o
            ORDER BY COALESCE(o.updated_at, o.created_at) DESC, o.id DESC
            """
        )
    ).mappings().all()
    return [dict(r) for r in rows]


@router.post("/admin/organizations")
def create_admin_organization(
    db: Session = Depends(get_db),
    name: str = Body(...),
    slug: str | None = Body(default=None),
    short_name: str | None = Body(default=None),
    logo_url: str | None = Body(default=None),
    status: str = Body(default="pilot"),
    contact_email: str | None = Body(default=None),
    contact_phone: str | None = Body(default=None),
    website_url: str | None = Body(default=None),
    country: str | None = Body(default=None),
    state_region: str | None = Body(default=None),
    city: str | None = Body(default=None),
    address_text: str | None = Body(default=None),
    notes: str | None = Body(default=None),
    is_active: bool = Body(default=True),
):
    name_clean = (name or "").strip()
    if not name_clean:
        raise HTTPException(status_code=400, detail="Organization name is required")
    status_clean = _normalize_name(status) or "pilot"
    final_slug = _ensure_unique_org_slug(db, (slug or "").strip() or name_clean)
    row = db.execute(
        text(
            """
            INSERT INTO green_organizations (
                name, slug, short_name, status, contact_email, contact_phone, website_url,
                country, state_region, city, address_text, notes, logo_url, is_active
            )
            VALUES (
                :name, :slug, :short_name, :status, :contact_email, :contact_phone, :website_url,
                :country, :state_region, :city, :address_text, :notes, :logo_url, :is_active
            )
            RETURNING *
            """
        ),
        {
            "name": name_clean,
            "slug": final_slug,
            "short_name": (short_name or "").strip() or None,
            "status": status_clean,
            "contact_email": (contact_email or "").strip() or None,
            "contact_phone": (contact_phone or "").strip() or None,
            "website_url": (website_url or "").strip() or None,
            "country": (country or "").strip() or None,
            "state_region": (state_region or "").strip() or None,
            "city": (city or "").strip() or None,
            "address_text": (address_text or "").strip() or None,
            "notes": (notes or "").strip() or None,
            "logo_url": (logo_url or "").strip() or None,
            "is_active": bool(is_active),
        },
    ).mappings().first()
    org = dict(row)
    _log_audit_event(
        db,
        project_id=None,
        entity_type="organization",
        entity_id=int(org["id"]),
        action="organization_created",
        actor="super_admin",
        details={"name": org.get("name"), "slug": org.get("slug"), "status": org.get("status")},
    )
    db.commit()
    org["welcome_email_attempted"] = False
    org["welcome_email_sent"] = False
    org["welcome_email_error"] = None
    org_contact_email = str(org.get("contact_email") or "").strip()
    if org_contact_email:
        org["welcome_email_attempted"] = True
        try:
            _send_organization_welcome_email(
                to_email=org_contact_email,
                organization_name=str(org.get("name") or "Organization"),
                organization_slug=str(org.get("slug") or "").strip() or None,
                status=str(org.get("status") or "").strip() or None,
                short_name=str(org.get("short_name") or "").strip() or None,
                website_url=str(org.get("website_url") or "").strip() or None,
                contact_phone=str(org.get("contact_phone") or "").strip() or None,
                country=str(org.get("country") or "").strip() or None,
                state_region=str(org.get("state_region") or "").strip() or None,
                city=str(org.get("city") or "").strip() or None,
            )
            org["welcome_email_sent"] = True
        except Exception as exc:
            org["welcome_email_error"] = str(exc)
    return org


@router.patch("/admin/organizations/{org_id}")
def update_admin_organization(
    org_id: int,
    db: Session = Depends(get_db),
    name: str | None = Body(default=None),
    slug: str | None = Body(default=None),
    short_name: str | None = Body(default=None),
    logo_url: str | None = Body(default=None),
    status: str | None = Body(default=None),
    contact_email: str | None = Body(default=None),
    contact_phone: str | None = Body(default=None),
    website_url: str | None = Body(default=None),
    country: str | None = Body(default=None),
    state_region: str | None = Body(default=None),
    city: str | None = Body(default=None),
    address_text: str | None = Body(default=None),
    notes: str | None = Body(default=None),
    is_active: bool | None = Body(default=None),
):
    existing = db.execute(
        text("SELECT * FROM green_organizations WHERE id = :org_id"),
        {"org_id": org_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Organization not found")

    next_name = (name.strip() if isinstance(name, str) else str(existing.get("name") or "").strip())
    if not next_name:
        raise HTTPException(status_code=400, detail="Organization name is required")
    next_slug_source = (slug or "").strip() if slug is not None else str(existing.get("slug") or next_name)
    next_slug = _ensure_unique_org_slug(db, next_slug_source or next_name, exclude_org_id=org_id)
    next_status = _normalize_name(status) if status is not None else _normalize_name(existing.get("status"))
    if not next_status:
        next_status = "pilot"
    row = db.execute(
        text(
            """
            UPDATE green_organizations
            SET name = :name,
                slug = :slug,
                short_name = :short_name,
                logo_url = :logo_url,
                status = :status,
                contact_email = :contact_email,
                contact_phone = :contact_phone,
                website_url = :website_url,
                country = :country,
                state_region = :state_region,
                city = :city,
                address_text = :address_text,
                notes = :notes,
                is_active = :is_active,
                updated_at = NOW()
            WHERE id = :org_id
            RETURNING *
            """
        ),
        {
            "org_id": org_id,
            "name": next_name,
            "slug": next_slug,
            "short_name": (short_name.strip() if isinstance(short_name, str) else (existing.get("short_name") or None)) or None,
            "logo_url": (logo_url.strip() if isinstance(logo_url, str) else (existing.get("logo_url") or None)) or None,
            "status": next_status,
            "contact_email": (contact_email.strip() if isinstance(contact_email, str) else (existing.get("contact_email") or None)) or None,
            "contact_phone": (contact_phone.strip() if isinstance(contact_phone, str) else (existing.get("contact_phone") or None)) or None,
            "website_url": (website_url.strip() if isinstance(website_url, str) else (existing.get("website_url") or None)) or None,
            "country": (country.strip() if isinstance(country, str) else (existing.get("country") or None)) or None,
            "state_region": (state_region.strip() if isinstance(state_region, str) else (existing.get("state_region") or None)) or None,
            "city": (city.strip() if isinstance(city, str) else (existing.get("city") or None)) or None,
            "address_text": (address_text.strip() if isinstance(address_text, str) else (existing.get("address_text") or None)) or None,
            "notes": (notes.strip() if isinstance(notes, str) else (existing.get("notes") or None)) or None,
            "is_active": bool(is_active) if is_active is not None else bool(existing.get("is_active", True)),
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=None,
        entity_type="organization",
        entity_id=org_id,
        action="organization_updated",
        actor="super_admin",
        details={
            "before": {"name": existing.get("name"), "slug": existing.get("slug"), "status": existing.get("status")},
            "after": {"name": next_name, "slug": next_slug, "status": next_status},
        },
    )
    db.commit()
    return dict(row)


@router.get("/admin/roles")
def list_admin_roles(db: Session = Depends(get_db), include_inactive: bool = Query(default=True)):
    rows = db.execute(
        text(
            """
            SELECT id, role_uid, role_key, role_name, description, scope, is_system,
                   COALESCE(is_active, TRUE) AS is_active, created_at, updated_at
            FROM green_roles
            WHERE (:include_inactive = TRUE OR COALESCE(is_active, TRUE) = TRUE)
            ORDER BY COALESCE(is_system, FALSE) DESC, role_name ASC, id ASC
            """
        ),
        {"include_inactive": bool(include_inactive)},
    ).mappings().all()
    return [dict(r) for r in rows]


@router.post("/admin/roles")
def create_admin_role(
    db: Session = Depends(get_db),
    role_name: str = Body(...),
    role_key: str | None = Body(default=None),
    role_uid: str | None = Body(default=None),
    description: str | None = Body(default=None),
    scope: str = Body(default="platform"),
    is_active: bool = Body(default=True),
):
    role_name_clean = (role_name or "").strip()
    if not role_name_clean:
        raise HTTPException(status_code=400, detail="Role name is required")
    final_role_key = _ensure_unique_role_key(db, (role_key or "").strip() or role_name_clean)
    final_role_uid = _ensure_unique_role_uid(db, role_uid)
    row = db.execute(
        text(
            """
            INSERT INTO green_roles (role_uid, role_key, role_name, description, scope, is_system, is_active)
            VALUES (:role_uid, :role_key, :role_name, :description, :scope, FALSE, :is_active)
            RETURNING id, role_uid, role_key, role_name, description, scope, is_system, is_active, created_at, updated_at
            """
        ),
        {
            "role_uid": final_role_uid,
            "role_key": final_role_key,
            "role_name": role_name_clean,
            "description": (description or "").strip() or None,
            "scope": _normalize_name(scope) or "platform",
            "is_active": bool(is_active),
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=None,
        entity_type="role",
        entity_id=int(row["id"]),
        action="role_created",
        actor="super_admin",
        details={"role_key": row.get("role_key"), "role_uid": row.get("role_uid"), "role_name": row.get("role_name")},
    )
    db.commit()
    return dict(row)


@router.patch("/admin/roles/{role_id}")
def update_admin_role(
    role_id: int,
    db: Session = Depends(get_db),
    role_name: str | None = Body(default=None),
    role_key: str | None = Body(default=None),
    role_uid: str | None = Body(default=None),
    description: str | None = Body(default=None),
    scope: str | None = Body(default=None),
    is_active: bool | None = Body(default=None),
):
    existing = db.execute(
        text("SELECT * FROM green_roles WHERE id = :role_id"),
        {"role_id": role_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Role not found")
    if bool(existing.get("is_system")) and role_key is not None:
        # Keep system role keys stable to avoid breaking existing workflows.
        role_key = None
    next_role_name = (role_name.strip() if isinstance(role_name, str) else str(existing.get("role_name") or "").strip())
    if not next_role_name:
        raise HTTPException(status_code=400, detail="Role name is required")
    next_role_key = (
        _ensure_unique_role_key(db, role_key.strip(), exclude_role_id=role_id)
        if isinstance(role_key, str)
        else str(existing.get("role_key") or "")
    )
    next_role_uid = (
        _ensure_unique_role_uid(db, role_uid, exclude_role_id=role_id)
        if role_uid is not None
        else str(existing.get("role_uid") or _ensure_unique_role_uid(db))
    )
    next_scope = _normalize_name(scope) if scope is not None else _normalize_name(existing.get("scope"))
    if not next_scope:
        next_scope = "platform"
    row = db.execute(
        text(
            """
            UPDATE green_roles
            SET role_uid = :role_uid,
                role_key = :role_key,
                role_name = :role_name,
                description = :description,
                scope = :scope,
                is_active = :is_active,
                updated_at = NOW()
            WHERE id = :role_id
            RETURNING id, role_uid, role_key, role_name, description, scope, is_system, is_active, created_at, updated_at
            """
        ),
        {
            "role_id": role_id,
            "role_uid": next_role_uid,
            "role_key": next_role_key,
            "role_name": next_role_name,
            "description": (description.strip() if isinstance(description, str) else existing.get("description")) or None,
            "scope": next_scope,
            "is_active": bool(is_active) if is_active is not None else bool(existing.get("is_active", True)),
        },
    ).mappings().first()
    if not bool(existing.get("is_system")) and str(existing.get("role_key") or "") != str(row.get("role_key") or ""):
        db.execute(
            text(
                """
                UPDATE green_users
                SET role = :role_key, updated_at = NOW()
                WHERE role_id = :role_id
                """
            ),
            {"role_id": role_id, "role_key": row.get("role_key")},
        )
    _log_audit_event(
        db,
        project_id=None,
        entity_type="role",
        entity_id=role_id,
        action="role_updated",
        actor="super_admin",
        details={"before_role_key": existing.get("role_key"), "after_role_key": row.get("role_key")},
    )
    db.commit()
    return dict(row)


@router.get("/admin/overview")
def admin_overview(db: Session = Depends(get_db), recent_limit: int = Query(default=40, ge=1, le=200)):
    organization_count = int(db.execute(text("SELECT COUNT(*) FROM green_organizations")).scalar() or 0)
    active_organization_count = int(
        db.execute(text("SELECT COUNT(*) FROM green_organizations WHERE COALESCE(is_active, TRUE) = TRUE")).scalar() or 0
    )
    project_count = int(db.execute(text("SELECT COUNT(*) FROM tree_projects")).scalar() or 0)
    unassigned_project_count = int(
        db.execute(text("SELECT COUNT(*) FROM tree_projects WHERE organization_id IS NULL")).scalar() or 0
    )
    tree_count = int(db.execute(text("SELECT COUNT(*) FROM trees")).scalar() or 0)
    task_count = int(db.execute(text("SELECT COUNT(*) FROM tree_tasks")).scalar() or 0)
    pending_review_count = int(
        db.execute(text("SELECT COUNT(*) FROM tree_tasks WHERE LOWER(COALESCE(review_state, 'none')) = 'submitted'")).scalar() or 0
    )
    open_alert_count = int(
        db.execute(text("SELECT COUNT(*) FROM green_alerts WHERE LOWER(COALESCE(status, 'open')) = 'open'")).scalar() or 0
    )
    users_count = int(db.execute(text("SELECT COUNT(*) FROM green_users")).scalar() or 0)
    roles_count = int(db.execute(text("SELECT COUNT(*) FROM green_roles")).scalar() or 0)

    org_rows = list_admin_organizations(db)
    recent_rows = db.execute(
        text(
            """
            SELECT
                a.id, a.project_id, a.entity_type, a.entity_id, a.action, a.actor, a.created_at,
                p.name AS project_name, p.organization_id,
                o.name AS organization_name, o.slug AS organization_slug
            FROM green_audit_events a
            LEFT JOIN tree_projects p ON p.id = a.project_id
            LEFT JOIN green_organizations o ON o.id = p.organization_id
            ORDER BY a.created_at DESC, a.id DESC
            LIMIT :recent_limit
            """
        ),
        {"recent_limit": int(recent_limit)},
    ).mappings().all()

    return {
        "totals": {
            "organizations": organization_count,
            "active_organizations": active_organization_count,
            "projects": project_count,
            "unassigned_projects": unassigned_project_count,
            "trees": tree_count,
            "tasks": task_count,
            "pending_reviews": pending_review_count,
            "open_alerts": open_alert_count,
            "users": users_count,
            "roles": roles_count,
        },
        "organizations": org_rows,
        "recent_activity": [dict(r) for r in recent_rows],
    }


@router.get("/admin/organizations/{organization_id}/credentials/export/pdf")
def export_admin_org_credentials_pdf(
    organization_id: int,
    include_inactive: bool = Query(default=True),
    db: Session = Depends(get_db),
):
    org_row = db.execute(
        text(
            """
            SELECT id, name, slug, short_name, status
            FROM green_organizations
            WHERE id = :organization_id
            """
        ),
        {"organization_id": organization_id},
    ).mappings().first()
    if not org_row:
        raise HTTPException(status_code=404, detail="Organization not found")

    user_rows = db.execute(
        text(
            """
            SELECT
                u.id, u.user_uid, u.full_name, u.role,
                COALESCE(u.allow_green, TRUE) AS allow_green,
                COALESCE(u.allow_work, FALSE) AS allow_work,
                COALESCE(u.is_active, TRUE) AS is_active,
                u.work_username,
                r.role_name, r.role_key
            FROM green_users u
            LEFT JOIN green_roles r ON r.id = u.role_id
            WHERE u.organization_id = :organization_id
              AND (:include_inactive = TRUE OR COALESCE(u.is_active, TRUE) = TRUE)
            ORDER BY COALESCE(u.updated_at, u.created_at) DESC, u.id DESC
            """
        ),
        {"organization_id": organization_id, "include_inactive": bool(include_inactive)},
    ).mappings().all()

    content = render_green_org_credentials_pdf(dict(org_row), [dict(row) for row in user_rows])
    safe_slug = re.sub(r"[^a-z0-9_-]+", "-", str(org_row.get("slug") or org_row.get("name") or "organization").lower()).strip("-")
    safe_slug = safe_slug or f"organization-{organization_id}"
    filename = f"landcheck-user-credentials-{safe_slug}.pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(content)
        pdf_path = tmp.name
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


@router.get("/projects")
def list_projects(
    db: Session = Depends(get_db),
    organization_id: int | None = Query(default=None),
):
    rows = db.execute(text("""
        SELECT
            p.id, p.organization_id, p.name, p.location_text, p.sponsor,
            p.planting_model, p.allow_existing_tree_link, p.default_existing_tree_scope, p.created_at,
            o.name AS organization_name, o.slug AS organization_slug, o.status AS organization_status,
            o.logo_url AS organization_logo_url
        FROM tree_projects p
        LEFT JOIN green_organizations o ON o.id = p.organization_id
        WHERE (:organization_id IS NULL OR p.organization_id = :organization_id)
        ORDER BY p.created_at DESC
    """), {
        "organization_id": int(organization_id) if organization_id is not None else None,
    }).mappings().all()
    return [dict(r) for r in rows]


@router.get("/projects/{project_id}")
def get_project(project_id: int, db: Session = Depends(get_db)):
    project = db.execute(text("""
        SELECT
            p.id, p.organization_id, p.name, p.location_text, p.sponsor,
            p.planting_model, p.allow_existing_tree_link, p.default_existing_tree_scope, p.created_at,
            o.name AS organization_name, o.slug AS organization_slug, o.status AS organization_status,
            o.logo_url AS organization_logo_url
        FROM tree_projects p
        LEFT JOIN green_organizations o ON o.id = p.organization_id
        WHERE p.id = :project_id
    """), {"project_id": project_id}).mappings().first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    stats_rows = db.execute(
        text("""
            SELECT status, COALESCE(count_in_planting_kpis, TRUE) AS in_scope, COUNT(*) AS count
            FROM trees
            WHERE project_id = :project_id
            GROUP BY status, COALESCE(count_in_planting_kpis, TRUE)
        """),
        {"project_id": project_id},
    ).mappings().all()
    status_counts: dict[str, int] = {}
    status_counts_all: dict[str, int] = {}
    for row in stats_rows:
        status_key = _normalize_tree_status(row.get("status"))
        row_count = int(row.get("count") or 0)
        status_counts_all[status_key] = status_counts_all.get(status_key, 0) + row_count
        if bool(row.get("in_scope")):
            status_counts[status_key] = status_counts.get(status_key, 0) + row_count

    total = sum(status_counts.values())
    alive = sum(status_counts.get(status_key, 0) for status_key in HEALTHY_TREE_STATUSES)
    dead = sum(status_counts.get(status_key, 0) for status_key in DEAD_TREE_STATUSES)
    needs_attention = sum(status_counts.get(status_key, 0) for status_key in ATTENTION_TREE_STATUSES)
    survival_rate = round((alive / total) * 100, 1) if total else 0.0
    total_all = sum(status_counts_all.values())

    # Carbon summary
    tree_rows_for_carbon = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :project_id
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"project_id": project_id}).mappings().all()
    carbon = compute_project_carbon([dict(r) for r in tree_rows_for_carbon])

    custodian_summary = db.execute(
        text(
            """
            SELECT
                COUNT(*) AS total_custodians,
                SUM(CASE WHEN LOWER(COALESCE(custodian_type, '')) = 'household' THEN 1 ELSE 0 END) AS households,
                SUM(CASE WHEN LOWER(COALESCE(custodian_type, '')) = 'school' THEN 1 ELSE 0 END) AS schools,
                SUM(CASE WHEN LOWER(COALESCE(custodian_type, '')) = 'community_group' THEN 1 ELSE 0 END) AS community_groups
            FROM green_custodians
            WHERE project_id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first() or {}
    distribution_summary = db.execute(
        text(
            """
            SELECT
                COUNT(*) AS events_total,
                COALESCE(SUM(quantity), 0) AS seedlings_distributed
            FROM green_distribution_events
            WHERE project_id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first() or {}
    allocations_summary = db.execute(
        text(
            """
            SELECT
                COUNT(*) AS allocations_total,
                COALESCE(SUM(quantity_allocated), 0) AS allocated_seedlings
            FROM green_distribution_allocations
            WHERE project_id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().first() or {}

    return {
        **dict(project),
        "stats": {
            "total": total,
            "total_all": total_all,
            "alive": alive,
            "dead": dead,
            "needs_attention": needs_attention,
            "survival_rate": survival_rate,
        },
        "settings": {
            "planting_model": _normalize_planting_model(project.get("planting_model")),
            "allow_existing_tree_link": bool(project.get("allow_existing_tree_link")),
            "default_existing_tree_scope": _normalize_existing_tree_scope(project.get("default_existing_tree_scope")),
        },
        "community": {
            "custodians_total": int(custodian_summary.get("total_custodians") or 0),
            "households": int(custodian_summary.get("households") or 0),
            "schools": int(custodian_summary.get("schools") or 0),
            "community_groups": int(custodian_summary.get("community_groups") or 0),
            "distribution_events": int(distribution_summary.get("events_total") or 0),
            "seedlings_distributed": int(distribution_summary.get("seedlings_distributed") or 0),
            "distribution_allocations": int(allocations_summary.get("allocations_total") or 0),
            "seedlings_allocated": int(allocations_summary.get("allocated_seedlings") or 0),
        },
        "carbon": {
            "current_co2_kg": carbon["current_co2_kg"],
            "current_co2_tonnes": carbon["current_co2_tonnes"],
            "annual_co2_kg": carbon["annual_co2_kg"],
            "annual_co2_tonnes": carbon["annual_co2_tonnes"],
            "projected_lifetime_co2_tonnes": carbon["projected_lifetime_co2_tonnes"],
            "co2_per_tree_avg_kg": carbon["co2_per_tree_avg_kg"],
            "trees_missing_age_data": carbon.get("trees_missing_age_data", 0),
            "trees_with_fallback_age": carbon.get("trees_with_fallback_age", 0),
            "trees_pending_review": carbon.get("trees_pending_review", 0),
        },
    }


@router.get("/projects/{project_id}/trees")
def list_trees(project_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT
               t.id,
               t.project_id,
               t.project_tree_no,
               t.species,
               t.planting_date,
               t.status,
               t.notes,
               t.photo_url,
               t.photo_urls,
               t.created_by,
               t.created_at,
               t.tree_origin,
               t.custodian_id,
               c.name AS custodian_name,
               t.custody_started_at,
               t.attribution_scope,
               t.count_in_planting_kpis,
               t.count_in_carbon_scope,
               t.source_project_id,
               t.tree_height_m,
               t.tree_age_months,
               COALESCE(t.inventory_tree_count, 1) AS inventory_tree_count,
               t.existing_area_geojson,
               t.existing_area_sqm,
               ST_X(geom) AS lng, ST_Y(geom) AS lat
        FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.project_id = :project_id
        ORDER BY t.created_at DESC
    """), {"project_id": project_id}).mappings().all()
    return [dict(r) for r in rows]


@router.get("/projects/{project_id}/custodians")
def list_custodians(project_id: int, db: Session = Depends(get_db)):
    _get_project_settings(db, project_id)
    rows = db.execute(
        text(
            """
            SELECT
                id,
                project_id,
                custodian_type,
                name,
                contact_person,
                phone,
                alt_phone,
                email,
                address_text,
                local_government,
                community_name,
                verification_status,
                notes,
                created_by,
                created_at,
                updated_at
            FROM green_custodians
            WHERE project_id = :project_id
            ORDER BY created_at DESC, id DESC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/projects/{project_id}/custodians")
def create_custodian(
    project_id: int,
    db: Session = Depends(get_db),
    custodian_type: str = Body(default="household"),
    name: str = Body(...),
    contact_person: str | None = Body(default=None),
    phone: str | None = Body(default=None),
    alt_phone: str | None = Body(default=None),
    email: str | None = Body(default=None),
    address_text: str | None = Body(default=None),
    local_government: str | None = Body(default=None),
    community_name: str | None = Body(default=None),
    verification_status: str = Body(default="pending"),
    notes: str | None = Body(default=None),
    created_by: str | None = Body(default=None),
):
    _get_project_settings(db, project_id)
    clean_name = (name or "").strip()
    if not clean_name:
        raise HTTPException(status_code=400, detail="Custodian name is required")

    row = db.execute(
        text(
            """
            INSERT INTO green_custodians (
                project_id, custodian_type, name, contact_person, phone, alt_phone, email, address_text,
                local_government, community_name, verification_status, notes, created_by
            )
            VALUES (
                :project_id, :custodian_type, :name, :contact_person, :phone, :alt_phone, :email, :address_text,
                :local_government, :community_name, :verification_status, :notes, :created_by
            )
            RETURNING
                id,
                project_id,
                custodian_type,
                name,
                contact_person,
                phone,
                alt_phone,
                email,
                address_text,
                local_government,
                community_name,
                verification_status,
                notes,
                created_by,
                created_at,
                updated_at
            """
        ),
        {
            "project_id": project_id,
            "custodian_type": _normalize_custodian_type(custodian_type),
            "name": clean_name,
            "contact_person": (contact_person or "").strip() or None,
            "phone": (phone or "").strip() or None,
            "alt_phone": (alt_phone or "").strip() or None,
            "email": (email or "").strip() or None,
            "address_text": (address_text or "").strip() or None,
            "local_government": (local_government or "").strip() or None,
            "community_name": (community_name or "").strip() or None,
            "verification_status": (_normalize_name(verification_status) or "pending"),
            "notes": (notes or "").strip() or None,
            "created_by": (created_by or "").strip() or None,
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="custodian",
        entity_id=int(row.get("id")),
        action="custodian_created",
        actor=(created_by or "").strip() or None,
        details={
            "custodian_type": row.get("custodian_type"),
            "name": row.get("name"),
            "contact_person": row.get("contact_person"),
            "email": row.get("email"),
            "community_name": row.get("community_name"),
        },
    )
    db.commit()
    return dict(row)


@router.patch("/custodians/{custodian_id}")
def update_custodian(
    custodian_id: int,
    db: Session = Depends(get_db),
    custodian_type: str | None = Body(default=None),
    name: str | None = Body(default=None),
    contact_person: str | None = Body(default=None),
    phone: str | None = Body(default=None),
    alt_phone: str | None = Body(default=None),
    email: str | None = Body(default=None),
    address_text: str | None = Body(default=None),
    local_government: str | None = Body(default=None),
    community_name: str | None = Body(default=None),
    verification_status: str | None = Body(default=None),
    notes: str | None = Body(default=None),
):
    existing = db.execute(
        text(
            """
            SELECT
                id, project_id, custodian_type, name, contact_person, phone, alt_phone, email,
                address_text, local_government, community_name, verification_status, notes
            FROM green_custodians
            WHERE id = :custodian_id
            """
        ),
        {"custodian_id": custodian_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Custodian not found")

    next_name = (name or "").strip() if name is not None else str(existing.get("name") or "").strip()
    if not next_name:
        raise HTTPException(status_code=400, detail="Custodian name is required")

    row = db.execute(
        text(
            """
            UPDATE green_custodians
            SET
                custodian_type = COALESCE(:custodian_type, custodian_type),
                name = :name,
                contact_person = COALESCE(:contact_person, contact_person),
                phone = COALESCE(:phone, phone),
                alt_phone = COALESCE(:alt_phone, alt_phone),
                email = COALESCE(:email, email),
                address_text = COALESCE(:address_text, address_text),
                local_government = COALESCE(:local_government, local_government),
                community_name = COALESCE(:community_name, community_name),
                verification_status = COALESCE(:verification_status, verification_status),
                notes = COALESCE(:notes, notes),
                updated_at = NOW()
            WHERE id = :custodian_id
            RETURNING
                id,
                project_id,
                custodian_type,
                name,
                contact_person,
                phone,
                alt_phone,
                email,
                address_text,
                local_government,
                community_name,
                verification_status,
                notes,
                created_by,
                created_at,
                updated_at
            """
        ),
        {
            "custodian_id": custodian_id,
            "custodian_type": _normalize_custodian_type(custodian_type) if custodian_type is not None else None,
            "name": next_name,
            "contact_person": ((contact_person or "").strip() or None) if contact_person is not None else None,
            "phone": ((phone or "").strip() or None) if phone is not None else None,
            "alt_phone": ((alt_phone or "").strip() or None) if alt_phone is not None else None,
            "email": ((email or "").strip() or None) if email is not None else None,
            "address_text": ((address_text or "").strip() or None) if address_text is not None else None,
            "local_government": ((local_government or "").strip() or None) if local_government is not None else None,
            "community_name": ((community_name or "").strip() or None) if community_name is not None else None,
            "verification_status": (_normalize_name(verification_status) or "pending") if verification_status is not None else None,
            "notes": ((notes or "").strip() or None) if notes is not None else None,
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=int(existing.get("project_id")),
        entity_type="custodian",
        entity_id=custodian_id,
        action="custodian_updated",
        details={"before": dict(existing), "after": dict(row)},
    )
    db.commit()
    return dict(row)


@router.get("/projects/{project_id}/distribution-events")
def list_distribution_events(project_id: int, db: Session = Depends(get_db)):
    _get_project_settings(db, project_id)
    rows = db.execute(
        text(
            """
            SELECT
                id,
                project_id,
                event_date,
                species,
                quantity,
                source_batch_ref,
                distributed_by,
                notes,
                created_at,
                updated_at
            FROM green_distribution_events
            WHERE project_id = :project_id
            ORDER BY event_date DESC, id DESC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/projects/{project_id}/distribution-events")
def create_distribution_event(
    project_id: int,
    db: Session = Depends(get_db),
    event_date: str = Body(...),
    species: str | None = Body(default=None),
    quantity: int = Body(default=0),
    source_batch_ref: str | None = Body(default=None),
    distributed_by: str | None = Body(default=None),
    notes: str | None = Body(default=None),
):
    _get_project_settings(db, project_id)
    event_date_value = _parse_date_value(event_date)
    if event_date_value is None:
        raise HTTPException(status_code=400, detail="event_date is required")
    if int(quantity or 0) < 0:
        raise HTTPException(status_code=400, detail="quantity cannot be negative")

    row = db.execute(
        text(
            """
            INSERT INTO green_distribution_events (
                project_id, event_date, species, quantity, source_batch_ref, distributed_by, notes
            )
            VALUES (
                :project_id, :event_date, :species, :quantity, :source_batch_ref, :distributed_by, :notes
            )
            RETURNING
                id,
                project_id,
                event_date,
                species,
                quantity,
                source_batch_ref,
                distributed_by,
                notes,
                created_at,
                updated_at
            """
        ),
        {
            "project_id": project_id,
            "event_date": event_date_value,
            "species": (species or "").strip() or None,
            "quantity": int(quantity or 0),
            "source_batch_ref": (source_batch_ref or "").strip() or None,
            "distributed_by": (distributed_by or "").strip() or None,
            "notes": (notes or "").strip() or None,
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="distribution_event",
        entity_id=int(row.get("id")),
        action="distribution_event_created",
        actor=(distributed_by or "").strip() or None,
        details={
            "event_date": _to_date_input(event_date_value),
            "species": row.get("species"),
            "quantity": int(row.get("quantity") or 0),
        },
    )
    db.commit()
    return dict(row)


@router.get("/projects/{project_id}/distribution-allocations")
def list_distribution_allocations(project_id: int, db: Session = Depends(get_db)):
    _get_project_settings(db, project_id)
    rows = db.execute(
        text(
            """
            SELECT
                a.id,
                a.event_id,
                a.project_id,
                a.custodian_id,
                c.name AS custodian_name,
                c.custodian_type,
                e.event_date,
                e.species,
                e.quantity AS event_quantity,
                a.quantity_allocated,
                a.supervision_target,
                a.expected_planting_start,
                a.expected_planting_end,
                a.followup_cycle_days,
                a.notes,
                COALESCE(s.supervision_assigned, 0) AS supervision_assigned,
                COALESCE(s.supervision_done, 0) AS supervision_done,
                COALESCE(s.supervision_live, 0) AS supervision_live,
                GREATEST(COALESCE(a.supervision_target, 0) - COALESCE(s.supervision_done, 0), 0) AS supervision_remaining,
                a.created_at,
                a.updated_at
            FROM green_distribution_allocations a
            JOIN green_distribution_events e ON e.id = a.event_id
            JOIN green_custodians c ON c.id = a.custodian_id
            LEFT JOIN (
                SELECT
                    t.distribution_allocation_id,
                    COUNT(*) AS supervision_assigned,
                    SUM(
                        CASE
                            WHEN LOWER(COALESCE(t.review_state, 'none')) = 'approved'
                                 OR (
                                     LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                                     AND LOWER(COALESCE(t.review_state, 'none')) = 'none'
                                 )
                            THEN 1 ELSE 0
                        END
                    ) AS supervision_done,
                    SUM(
                        CASE
                            WHEN LOWER(COALESCE(t.review_state, 'none')) = 'approved'
                                 OR (
                                     LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                                     AND LOWER(COALESCE(t.review_state, 'none')) = 'none'
                                 )
                            THEN 0 ELSE 1
                        END
                    ) AS supervision_live
                FROM tree_tasks t
                WHERE LOWER(COALESCE(t.task_type, '')) = :supervision_type
                  AND t.distribution_allocation_id IS NOT NULL
                GROUP BY t.distribution_allocation_id
            ) s ON s.distribution_allocation_id = a.id
            WHERE a.project_id = :project_id
            ORDER BY e.event_date DESC, a.created_at DESC, a.id DESC
            """
        ),
        {"project_id": project_id, "supervision_type": SUPERVISION_TASK_TYPE},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/distribution-events/{event_id}/allocations")
def upsert_distribution_allocation(
    event_id: int,
    db: Session = Depends(get_db),
    custodian_id: int = Body(...),
    quantity_allocated: int = Body(default=0),
    supervision_target: int = Body(default=0),
    expected_planting_start: str | None = Body(default=None),
    expected_planting_end: str | None = Body(default=None),
    followup_cycle_days: int = Body(default=30),
    notes: str | None = Body(default=None),
):
    event_row = db.execute(
        text("SELECT id, project_id, event_date FROM green_distribution_events WHERE id = :event_id"),
        {"event_id": event_id},
    ).mappings().first()
    if not event_row:
        raise HTTPException(status_code=404, detail="Distribution event not found")

    custodian_row = db.execute(
        text("SELECT id, project_id, name FROM green_custodians WHERE id = :custodian_id"),
        {"custodian_id": custodian_id},
    ).mappings().first()
    if not custodian_row:
        raise HTTPException(status_code=404, detail="Custodian not found")
    if int(custodian_row.get("project_id") or 0) != int(event_row.get("project_id") or 0):
        raise HTTPException(status_code=400, detail="Custodian and event must belong to the same project")

    if int(quantity_allocated or 0) < 0:
        raise HTTPException(status_code=400, detail="quantity_allocated cannot be negative")
    if int(supervision_target or 0) < 0:
        raise HTTPException(status_code=400, detail="supervision_target cannot be negative")
    if int(supervision_target or 0) > 365:
        raise HTTPException(status_code=400, detail="supervision_target cannot be greater than 365")
    if int(followup_cycle_days or 0) < 1 or int(followup_cycle_days or 0) > 365:
        raise HTTPException(status_code=400, detail="followup_cycle_days must be between 1 and 365")

    start_date = _parse_date_value(expected_planting_start)
    end_date = _parse_date_value(expected_planting_end)
    if start_date and end_date and end_date < start_date:
        raise HTTPException(status_code=400, detail="expected_planting_end cannot be before expected_planting_start")

    row = db.execute(
        text(
            """
            INSERT INTO green_distribution_allocations (
                event_id, project_id, custodian_id, quantity_allocated,
                supervision_target, expected_planting_start, expected_planting_end, followup_cycle_days, notes
            )
            VALUES (
                :event_id, :project_id, :custodian_id, :quantity_allocated,
                :supervision_target, :expected_planting_start, :expected_planting_end, :followup_cycle_days, :notes
            )
            ON CONFLICT (event_id, custodian_id)
            DO UPDATE SET
                quantity_allocated = EXCLUDED.quantity_allocated,
                supervision_target = EXCLUDED.supervision_target,
                expected_planting_start = COALESCE(EXCLUDED.expected_planting_start, green_distribution_allocations.expected_planting_start),
                expected_planting_end = COALESCE(EXCLUDED.expected_planting_end, green_distribution_allocations.expected_planting_end),
                followup_cycle_days = EXCLUDED.followup_cycle_days,
                notes = COALESCE(EXCLUDED.notes, green_distribution_allocations.notes),
                updated_at = NOW()
            RETURNING
                id,
                event_id,
                project_id,
                custodian_id,
                quantity_allocated,
                supervision_target,
                expected_planting_start,
                expected_planting_end,
                followup_cycle_days,
                notes,
                created_at,
                updated_at
            """
        ),
        {
            "event_id": event_id,
            "project_id": int(event_row.get("project_id")),
            "custodian_id": int(custodian_id),
            "quantity_allocated": int(quantity_allocated or 0),
            "supervision_target": int(supervision_target or 0),
            "expected_planting_start": start_date,
            "expected_planting_end": end_date,
            "followup_cycle_days": int(followup_cycle_days or 30),
            "notes": (notes or "").strip() or None,
        },
    ).mappings().first()
    _log_audit_event(
        db,
        project_id=int(event_row.get("project_id")),
        entity_type="distribution_allocation",
        entity_id=int(row.get("id")),
        action="distribution_allocation_upserted",
        details={
            "event_id": int(event_id),
            "custodian_id": int(custodian_id),
            "quantity_allocated": int(row.get("quantity_allocated") or 0),
            "supervision_target": int(row.get("supervision_target") or 0),
            "followup_cycle_days": int(row.get("followup_cycle_days") or 0),
        },
    )
    db.commit()
    return dict(row)


@router.post("/distribution-allocations/{allocation_id}/assign-supervision")
def assign_distribution_supervision(
    allocation_id: int,
    db: Session = Depends(get_db),
    assignee_name: str = Body(...),
    visits_to_assign: int = Body(default=1),
    due_date: str | None = Body(default=None),
    priority: str = Body(default="normal"),
    actor_name: str | None = Body(default=None),
):
    assignee_clean = (assignee_name or "").strip()
    if not assignee_clean:
        raise HTTPException(status_code=400, detail="Assignee name is required")
    requested_visits = int(visits_to_assign or 0)
    if requested_visits <= 0:
        raise HTTPException(status_code=400, detail="visits_to_assign must be greater than 0")
    if requested_visits > 30:
        raise HTTPException(status_code=400, detail="visits_to_assign cannot be greater than 30 at once")

    allocation = db.execute(
        text(
            """
            SELECT
                a.id,
                a.project_id,
                a.custodian_id,
                a.supervision_target,
                a.followup_cycle_days,
                a.expected_planting_start,
                a.expected_planting_end,
                c.name AS custodian_name,
                c.custodian_type,
                c.contact_person,
                c.phone,
                c.email,
                c.community_name,
                e.event_date,
                e.species AS event_species
            FROM green_distribution_allocations a
            JOIN green_custodians c ON c.id = a.custodian_id
            JOIN green_distribution_events e ON e.id = a.event_id
            WHERE a.id = :allocation_id
            """
        ),
        {"allocation_id": allocation_id},
    ).mappings().first()
    if not allocation:
        raise HTTPException(status_code=404, detail="Allocation not found")

    supervision_target = int(allocation.get("supervision_target") or 0)
    if supervision_target <= 0:
        raise HTTPException(status_code=400, detail="Set supervision target in allocation before assigning visits")

    tree_rows = db.execute(
        text(
            """
            SELECT id, species, status, planting_date, ST_X(geom) AS lng, ST_Y(geom) AS lat
            FROM trees
            WHERE project_id = :project_id
              AND custodian_id = :custodian_id
            ORDER BY created_at ASC, id ASC
            """
        ),
        {
            "project_id": int(allocation.get("project_id") or 0),
            "custodian_id": int(allocation.get("custodian_id") or 0),
        },
    ).mappings().all()
    if not tree_rows:
        raise HTTPException(status_code=400, detail="No trees linked to this custodian yet. Capture at least one tree first.")

    task_counts = db.execute(
        text(
            """
            SELECT
                COUNT(*) AS assigned_count,
                SUM(
                    CASE
                        WHEN LOWER(COALESCE(review_state, 'none')) = 'approved'
                             OR (
                                 LOWER(COALESCE(status, '')) IN ('done', 'completed', 'closed')
                                 AND LOWER(COALESCE(review_state, 'none')) = 'none'
                             )
                        THEN 1 ELSE 0
                    END
                ) AS done_count
            FROM tree_tasks
            WHERE distribution_allocation_id = :allocation_id
              AND LOWER(COALESCE(task_type, '')) = :task_type
            """
        ),
        {"allocation_id": allocation_id, "task_type": SUPERVISION_TASK_TYPE},
    ).mappings().first()
    assigned_count = int((task_counts or {}).get("assigned_count") or 0)
    done_count = int((task_counts or {}).get("done_count") or 0)
    remaining_assignable = max(supervision_target - assigned_count, 0)
    if remaining_assignable <= 0:
        raise HTTPException(status_code=409, detail="All supervision visits for this allocation are already assigned")

    create_count = min(requested_visits, remaining_assignable)
    due_date_value = _parse_date_value(due_date)
    cycle_days = int(allocation.get("followup_cycle_days") or 14)
    baseline_due = (
        _parse_date_value(allocation.get("expected_planting_start"))
        or _parse_date_value(allocation.get("event_date"))
        or date.today()
    )
    custodian_name = str(allocation.get("custodian_name") or "").strip()
    contact_text = str(allocation.get("phone") or allocation.get("email") or allocation.get("contact_person") or "-").strip()
    community_text = str(allocation.get("community_name") or "-").strip()

    created_tasks: list[dict] = []
    for index in range(create_count):
        visit_no = assigned_count + index + 1
        target_tree = tree_rows[(assigned_count + index) % len(tree_rows)]
        task_due_date = due_date_value or (baseline_due + timedelta(days=max(cycle_days * visit_no, 1)))
        task_notes = (
            f"Custodian supervision visit {visit_no}/{supervision_target}. "
            f"Custodian: {custodian_name or '-'} ({str(allocation.get('custodian_type') or '-').replace('_', ' ')}). "
            f"Community: {community_text}. Contact: {contact_text}. "
            f"Check tree condition, capture GPS, upload visit photos, and document support actions."
        )
        row = db.execute(
            text(
                """
                INSERT INTO tree_tasks (
                    tree_id, task_type, assignee_name, due_date, priority, status, notes, photo_url, photo_urls,
                    review_state, submitted_at, completed_at, model_season,
                    custodian_id, distribution_allocation_id, supervision_visit_no, supervision_total_visits
                )
                VALUES (
                    :tree_id, :task_type, :assignee_name, :due_date, :priority, 'pending', :notes, NULL, CAST(:photo_urls AS JSONB),
                    'none', NULL, NULL, NULL,
                    :custodian_id, :distribution_allocation_id, :supervision_visit_no, :supervision_total_visits
                )
                RETURNING id, tree_id, due_date
                """
            ),
            {
                "tree_id": int(target_tree.get("id") or 0),
                "task_type": SUPERVISION_TASK_TYPE,
                "assignee_name": assignee_clean,
                "due_date": task_due_date,
                "priority": priority or "normal",
                "notes": task_notes,
                "photo_urls": _safe_json([]),
                "custodian_id": int(allocation.get("custodian_id") or 0),
                "distribution_allocation_id": allocation_id,
                "supervision_visit_no": visit_no,
                "supervision_total_visits": supervision_target,
            },
        ).mappings().first()
        created_tasks.append(
            {
                "task_id": int(row.get("id") or 0),
                "tree_id": int(row.get("tree_id") or 0),
                "tree_species": target_tree.get("species"),
                "tree_status": target_tree.get("status"),
                "tree_lng": target_tree.get("lng"),
                "tree_lat": target_tree.get("lat"),
                "due_date": _to_date_input(row.get("due_date")),
                "visit_no": visit_no,
                "supervision_total_visits": supervision_target,
                "custodian_name": custodian_name,
                "custodian_contact": contact_text,
                "custodian_community": community_text,
            }
        )

    _log_audit_event(
        db,
        project_id=int(allocation.get("project_id") or 0),
        entity_type="distribution_allocation",
        entity_id=allocation_id,
        action="custodian_supervision_assigned",
        actor=(actor_name or "").strip() or None,
        details={
            "assignee_name": assignee_clean,
            "created_count": len(created_tasks),
            "supervision_target": supervision_target,
            "previous_assigned": assigned_count,
            "supervision_done": done_count,
        },
    )
    db.commit()
    return {
        "allocation_id": allocation_id,
        "project_id": int(allocation.get("project_id") or 0),
        "custodian_id": int(allocation.get("custodian_id") or 0),
        "custodian_name": custodian_name,
        "supervision_target": supervision_target,
        "supervision_assigned": assigned_count + len(created_tasks),
        "supervision_done": done_count,
        "supervision_live": max((assigned_count + len(created_tasks)) - done_count, 0),
        "supervision_remaining": max(supervision_target - (assigned_count + len(created_tasks)), 0),
        "created_count": len(created_tasks),
        "tasks": created_tasks,
    }


@router.get("/projects/{project_id}/existing-tree-candidates")
def list_existing_tree_candidates(
    project_id: int,
    source_project_id: int = Query(...),
    limit: int = Query(default=500, ge=1, le=2000),
    db: Session = Depends(get_db),
):
    _get_project_settings(db, project_id)
    _get_project_settings(db, source_project_id)
    if int(project_id) == int(source_project_id):
        return {"source_project_id": source_project_id, "target_project_id": project_id, "items": []}

    rows = db.execute(
        text(
            """
            SELECT
                t.id,
                t.project_id,
                t.species,
                t.planting_date,
                t.status,
                t.tree_origin,
                t.tree_height_m,
                t.tree_age_months,
                t.created_by,
                t.created_at,
                t.source_project_id,
                ST_X(t.geom) AS lng,
                ST_Y(t.geom) AS lat
            FROM trees t
            WHERE t.project_id = :source_project_id
            ORDER BY t.created_at DESC, t.id DESC
            LIMIT :limit
            """
        ),
        {"source_project_id": source_project_id, "limit": int(limit)},
    ).mappings().all()
    linked_rows = db.execute(
        text(
            """
            SELECT source_tree_id
            FROM tree_project_links
            WHERE source_project_id = :source_project_id
              AND target_project_id = :target_project_id
              AND link_type = 'reference'
            """
        ),
        {"source_project_id": source_project_id, "target_project_id": project_id},
    ).mappings().all()
    linked_tree_ids = {int(row.get("source_tree_id") or 0) for row in linked_rows}
    items = []
    for row in rows:
        item = dict(row)
        item["already_referenced"] = int(item.get("id") or 0) in linked_tree_ids
        items.append(item)
    return {
        "source_project_id": source_project_id,
        "target_project_id": project_id,
        "items": items,
    }


@router.post("/projects/{project_id}/existing-trees/import")
def import_existing_trees(
    project_id: int,
    db: Session = Depends(get_db),
    source_project_id: int = Body(...),
    tree_ids: list[int] = Body(...),
    mode: str = Body(default="reference"),
    actor_name: str | None = Body(default=None),
    attribution_scope: str | None = Body(default=None),
    count_in_planting_kpis: bool | None = Body(default=None),
    count_in_carbon_scope: bool | None = Body(default=None),
):
    target_settings = _get_project_settings(db, project_id)
    _get_project_settings(db, source_project_id)
    if int(project_id) == int(source_project_id):
        raise HTTPException(status_code=400, detail="source_project_id must be different from target project")
    if not bool(target_settings.get("allow_existing_tree_link")):
        raise HTTPException(status_code=400, detail="Target project does not allow existing tree import/linking.")

    mode_key = _normalize_name(mode)
    if mode_key not in {"transfer", "reference"}:
        raise HTTPException(status_code=400, detail="mode must be 'transfer' or 'reference'")

    normalized_tree_ids = sorted({int(item) for item in (tree_ids or []) if int(item) > 0})
    if not normalized_tree_ids:
        raise HTTPException(status_code=400, detail="tree_ids is required")

    source_rows = db.execute(
        text(
            """
            SELECT
                id,
                project_id,
                geom,
                species,
                planting_date,
                status,
                notes,
                photo_url,
                created_by,
                created_at,
                tree_origin,
                tree_height_m,
                source_project_id
            FROM trees
            WHERE project_id = :source_project_id
              AND id = ANY(:tree_ids)
            ORDER BY id
            """
        ),
        {"source_project_id": source_project_id, "tree_ids": normalized_tree_ids},
    ).mappings().all()
    if not source_rows:
        raise HTTPException(status_code=404, detail="No matching trees found in source project")

    referenced_lookup = set()
    if mode_key == "reference":
        linked = db.execute(
            text(
                """
                SELECT source_tree_id
                FROM tree_project_links
                WHERE source_project_id = :source_project_id
                  AND target_project_id = :target_project_id
                  AND link_type = 'reference'
                """
            ),
            {"source_project_id": source_project_id, "target_project_id": project_id},
        ).mappings().all()
        referenced_lookup = {int(row.get("source_tree_id") or 0) for row in linked}

    imported_items: list[dict] = []
    moved_count = 0
    referenced_count = 0
    skipped_count = 0

    resolved_scope, default_planting_scope, default_carbon_scope = _resolve_tree_scope_defaults(
        tree_origin="existing_inventory",
        attribution_scope=attribution_scope,
        count_in_planting_kpis=count_in_planting_kpis,
        count_in_carbon_scope=count_in_carbon_scope,
        project_existing_scope=target_settings.get("default_existing_tree_scope"),
    )
    actor_clean = (actor_name or "").strip() or None

    for row in source_rows:
        source_tree_id = int(row.get("id") or 0)
        if source_tree_id <= 0:
            continue
        if mode_key == "reference" and source_tree_id in referenced_lookup:
            skipped_count += 1
            imported_items.append(
                {
                    "source_tree_id": source_tree_id,
                    "target_tree_id": None,
                    "mode": "reference",
                    "status": "skipped_already_referenced",
                }
            )
            continue

        target_project_tree_no = _next_project_tree_no(db, int(project_id))

        if mode_key == "transfer":
            target_tree_id = db.execute(
                text(
                    """
                    UPDATE trees
                    SET
                        project_id = :target_project_id,
                        project_tree_no = :project_tree_no,
                        tree_origin = 'existing_inventory',
                        custodian_id = NULL,
                        custody_started_at = COALESCE(custody_started_at, CURRENT_DATE),
                        attribution_scope = :attribution_scope,
                        count_in_planting_kpis = :count_in_planting_kpis,
                        count_in_carbon_scope = :count_in_carbon_scope,
                        source_project_id = COALESCE(source_project_id, :source_project_id)
                    WHERE id = :tree_id
                    RETURNING id
                    """
                ),
                {
                    "target_project_id": project_id,
                    "project_tree_no": target_project_tree_no,
                    "attribution_scope": resolved_scope,
                    "count_in_planting_kpis": default_planting_scope,
                    "count_in_carbon_scope": default_carbon_scope,
                    "source_project_id": source_project_id,
                    "tree_id": source_tree_id,
                },
            ).scalar()
            if not target_tree_id:
                skipped_count += 1
                continue
            moved_count += 1
            imported_items.append(
                {
                    "source_tree_id": source_tree_id,
                    "target_tree_id": int(target_tree_id),
                    "mode": "transfer",
                    "status": "moved",
                }
            )
            db.execute(
                text(
                    """
                    INSERT INTO tree_project_links (
                        source_project_id, target_project_id, source_tree_id, target_tree_id, link_type, transfer_mode, created_by
                    )
                    VALUES (
                        :source_project_id, :target_project_id, :source_tree_id, :target_tree_id, :link_type, 'transfer', :created_by
                    )
                    ON CONFLICT DO NOTHING
                    """
                ),
                {
                    "source_project_id": source_project_id,
                    "target_project_id": project_id,
                    "source_tree_id": source_tree_id,
                    "target_tree_id": int(target_tree_id),
                    "link_type": "owner",
                    "created_by": actor_clean,
                },
            )
            _record_tree_status_history(
                db,
                tree_id=int(target_tree_id),
                project_id=int(project_id),
                status=_normalize_tree_status(row.get("status")),
                status_date=date.today(),
                source="existing_transfer",
                changed_by=actor_clean,
                notes=f"Transferred from project {source_project_id}",
            )
        else:
            target_tree_id = db.execute(
                text(
                    """
                    INSERT INTO trees (
                        project_id, project_tree_no, geom, species, planting_date, status, notes, photo_url, created_by,
                        tree_origin, custodian_id, custody_started_at, attribution_scope,
                        count_in_planting_kpis, count_in_carbon_scope, source_project_id, tree_height_m, tree_age_months
                    )
                    SELECT
                        :target_project_id,
                        :project_tree_no,
                        geom,
                        species,
                        planting_date,
                        status,
                        notes,
                        photo_url,
                        created_by,
                        'existing_inventory',
                        NULL,
                        CURRENT_DATE,
                        :attribution_scope,
                        :count_in_planting_kpis,
                        :count_in_carbon_scope,
                        :source_project_id,
                        tree_height_m,
                        tree_age_months
                    FROM trees
                    WHERE id = :tree_id
                    RETURNING id
                    """
                ),
                {
                    "target_project_id": project_id,
                    "project_tree_no": target_project_tree_no,
                    "attribution_scope": resolved_scope,
                    "count_in_planting_kpis": default_planting_scope,
                    "count_in_carbon_scope": default_carbon_scope,
                    "source_project_id": source_project_id,
                    "tree_id": source_tree_id,
                },
            ).scalar()
            if not target_tree_id:
                skipped_count += 1
                continue
            referenced_count += 1
            imported_items.append(
                {
                    "source_tree_id": source_tree_id,
                    "target_tree_id": int(target_tree_id),
                    "mode": "reference",
                    "status": "referenced",
                }
            )
            db.execute(
                text(
                    """
                    INSERT INTO tree_project_links (
                        source_project_id, target_project_id, source_tree_id, target_tree_id, link_type, transfer_mode, created_by
                    )
                    VALUES (
                        :source_project_id, :target_project_id, :source_tree_id, :target_tree_id, 'reference', 'reference', :created_by
                    )
                    ON CONFLICT DO NOTHING
                    """
                ),
                {
                    "source_project_id": source_project_id,
                    "target_project_id": project_id,
                    "source_tree_id": source_tree_id,
                    "target_tree_id": int(target_tree_id),
                    "created_by": actor_clean,
                },
            )
            _record_tree_status_history(
                db,
                tree_id=int(target_tree_id),
                project_id=int(project_id),
                status=_normalize_tree_status(row.get("status")),
                status_date=_parse_date_value(row.get("planting_date")) or date.today(),
                source="existing_reference",
                changed_by=actor_clean,
                notes=f"Referenced from project {source_project_id}, tree {source_tree_id}",
            )

    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="tree_import",
        entity_id=None,
        action="existing_trees_imported",
        actor=actor_clean,
        details={
            "source_project_id": source_project_id,
            "mode": mode_key,
            "requested_count": len(normalized_tree_ids),
            "moved_count": moved_count,
            "referenced_count": referenced_count,
            "skipped_count": skipped_count,
            "attribution_scope": resolved_scope,
            "count_in_planting_kpis": default_planting_scope,
            "count_in_carbon_scope": default_carbon_scope,
        },
    )
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {
        "target_project_id": project_id,
        "source_project_id": source_project_id,
        "mode": mode_key,
        "moved_count": moved_count,
        "referenced_count": referenced_count,
        "skipped_count": skipped_count,
        "items": imported_items,
    }


@router.get("/projects/{project_id}/species-maturity")
def get_species_maturity(project_id: int, db: Session = Depends(get_db)):
    project_exists = db.execute(
        text("SELECT 1 FROM tree_projects WHERE id = :project_id"),
        {"project_id": project_id},
    ).scalar()
    if not project_exists:
        raise HTTPException(status_code=404, detail="Project not found")

    rows = db.execute(
        text("""
            SELECT species_key, species_label, maturity_years, updated_at
            FROM green_species_maturity
            WHERE project_id = :project_id
            ORDER BY COALESCE(species_label, species_key) ASC
        """),
        {"project_id": project_id},
    ).mappings().all()
    items = [dict(row) for row in rows]
    return {
        "project_id": project_id,
        "items": items,
        "map": {row["species_key"]: int(row["maturity_years"]) for row in rows},
    }


@router.put("/projects/{project_id}/species-maturity")
def upsert_species_maturity(
    project_id: int,
    db: Session = Depends(get_db),
    species_key: str = Body(...),
    maturity_years: int = Body(...),
    species_label: str | None = Body(default=None),
):
    normalized_key = (species_key or "").strip().lower()
    if not normalized_key:
        raise HTTPException(status_code=400, detail="species_key is required")
    if maturity_years < 1 or maturity_years > 50:
        raise HTTPException(status_code=400, detail="maturity_years must be between 1 and 50")

    project_exists = db.execute(
        text("SELECT 1 FROM tree_projects WHERE id = :project_id"),
        {"project_id": project_id},
    ).scalar()
    if not project_exists:
        raise HTTPException(status_code=404, detail="Project not found")

    cleaned_label = (species_label or "").strip() or None
    row = db.execute(
        text("""
            INSERT INTO green_species_maturity (project_id, species_key, species_label, maturity_years)
            VALUES (:project_id, :species_key, :species_label, :maturity_years)
            ON CONFLICT (project_id, species_key)
            DO UPDATE
            SET
                maturity_years = EXCLUDED.maturity_years,
                species_label = COALESCE(EXCLUDED.species_label, green_species_maturity.species_label),
                updated_at = NOW()
            RETURNING project_id, species_key, species_label, maturity_years, updated_at
        """),
        {
            "project_id": project_id,
            "species_key": normalized_key,
            "species_label": cleaned_label,
            "maturity_years": int(maturity_years),
        },
    ).mappings().first()
    db.commit()
    return dict(row)


# ---------------------------------------------------------------------------
# Carbon / CO2 Endpoints
# ---------------------------------------------------------------------------


@router.get("/projects/{project_id}/carbon-summary")
def carbon_summary(project_id: int, projection_years: int = Query(default=40), db: Session = Depends(get_db)):
    """Get CO2 sequestration summary for a project."""
    project = db.execute(
        text("SELECT id FROM tree_projects WHERE id = :pid"), {"pid": project_id},
    ).scalar()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    tree_rows = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :project_id
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"project_id": project_id}).mappings().all()
    trees = [dict(r) for r in tree_rows]
    summary = compute_project_carbon(trees, projection_years)
    summary["project_id"] = project_id
    return summary


@router.get("/projects/{project_id}/carbon-projection")
def carbon_projection(project_id: int, years: int = Query(default=30), db: Session = Depends(get_db)):
    """Get year-by-year CO2 projection for a project."""
    project = db.execute(
        text("SELECT id FROM tree_projects WHERE id = :pid"), {"pid": project_id},
    ).scalar()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    tree_rows = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :project_id
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"project_id": project_id}).mappings().all()
    trees = [dict(r) for r in tree_rows]
    projection = generate_co2_projection_table(trees, years)
    return {"project_id": project_id, "projection": projection}


@router.get("/trees/{tree_id}/carbon")
def tree_carbon(tree_id: int, db: Session = Depends(get_db)):
    """Get CO2 estimate for a single tree."""
    tree = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_height_m, tree_age_months, count_in_carbon_scope,
               COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE id = :tree_id
    """), {"tree_id": tree_id}).mappings().first()
    if not tree:
        raise HTTPException(status_code=404, detail="Tree not found")

    metrics = _build_tree_carbon_metrics(dict(tree), projection_years=40, enforce_carbon_scope=False)
    species = tree.get("species")

    return {
        "tree_id": tree_id,
        "species": species,
        "species_matched": metrics.get("species_matched", "Unknown"),
        "growth_class": metrics.get("growth_class", "medium"),
        "age_years": round(float(metrics.get("age_years") or 0.0), 1),
        "age_source": metrics.get("age_source", "none"),
        "current_co2_kg": metrics.get("current_co2_kg", 0.0),
        "annual_co2_kg": metrics.get("annual_co2_kg", 0.0),
        "lifetime_co2_kg": metrics.get("lifetime_co2_kg", 0.0),
        "lifetime_co2_tonnes": metrics.get("lifetime_co2_tonnes", 0.0),
        "height_used_for_co2": bool(metrics.get("height_used_for_co2")),
        "co2_height_source": metrics.get("co2_height_source", "modeled_height"),
        "tree_height_m": tree.get("tree_height_m"),
        "tree_age_months": tree.get("tree_age_months"),
        "inventory_tree_count": metrics.get("inventory_tree_count", 1),
        "methodology": metrics.get("co2_methodology", "Chave et al. (2014) pantropical allometric equation"),
    }


@router.get("/carbon/species-database")
def carbon_species_database():
    """List all species in the carbon estimation database."""
    return {"species": list_known_species()}


@router.post("/trees")
def add_tree(
    db: Session = Depends(get_db),
    project_id: int = Body(...),
    lng: float = Body(...),
    lat: float = Body(...),
    species: str = Body(default=""),
    planting_date: str | None = Body(default=None),
    status: str = Body(default="alive"),
    notes: str = Body(default=""),
    photo_url: str = Body(default=""),
    photo_urls: list[str] | None = Body(default=None),
    created_by: str = Body(default=""),
    tree_origin: str = Body(default="new_planting"),
    custodian_id: int | None = Body(default=None),
    custody_started_at: str | None = Body(default=None),
    attribution_scope: str | None = Body(default=None),
    count_in_planting_kpis: bool | None = Body(default=None),
    count_in_carbon_scope: bool | None = Body(default=None),
    source_project_id: int | None = Body(default=None),
    tree_height_m: float | None = Body(default=None),
    tree_age_months: float | None = Body(default=None),
    inventory_tree_count: int | None = Body(default=None),
    existing_area_geojson: dict | str | None = Body(default=None),
):
    project_settings = _get_project_settings(db, int(project_id))
    origin = _normalize_tree_origin(tree_origin)
    if origin not in TREE_ORIGIN_VALUES:
        raise HTTPException(status_code=400, detail="Invalid tree_origin")

    if tree_height_m is not None:
        try:
            tree_height_value = float(tree_height_m)
        except Exception:
            raise HTTPException(status_code=400, detail="tree_height_m must be numeric")
        if tree_height_value < 0 or tree_height_value > 120:
            raise HTTPException(status_code=400, detail="tree_height_m must be between 0 and 120")
    else:
        tree_height_value = None

    if tree_age_months is not None:
        try:
            tree_age_months_value = float(tree_age_months)
        except Exception:
            raise HTTPException(status_code=400, detail="tree_age_months must be numeric")
        if tree_age_months_value < 0 or tree_age_months_value > 2400:
            raise HTTPException(status_code=400, detail="tree_age_months must be between 0 and 2400")
    else:
        tree_age_months_value = None

    try:
        inventory_tree_count_value = int(inventory_tree_count) if inventory_tree_count is not None else 1
    except Exception:
        raise HTTPException(status_code=400, detail="inventory_tree_count must be a whole number")
    if inventory_tree_count_value < 1 or inventory_tree_count_value > 1000000:
        raise HTTPException(status_code=400, detail="inventory_tree_count must be between 1 and 1000000")

    normalized_existing_area_geojson = _normalize_work_area_geojson(existing_area_geojson)
    if origin != "existing_inventory":
        if normalized_existing_area_geojson is not None:
            raise HTTPException(status_code=400, detail="existing_area_geojson is only allowed for Existing Tree records")
        inventory_tree_count_value = 1
    else:
        if inventory_tree_count_value > 1 and normalized_existing_area_geojson is None:
            raise HTTPException(
                status_code=400,
                detail="Draw a polygon area when capturing more than one existing tree in a single record",
            )

    existing_area_sqm_value = None
    if normalized_existing_area_geojson is not None:
        try:
            existing_area_sqm_value = float(
                db.execute(
                    text(
                        """
                        SELECT ST_Area(
                            ST_Transform(
                                ST_SetSRID(ST_GeomFromGeoJSON(:geojson), 4326),
                                3857
                            )
                        )
                        """
                    ),
                    {"geojson": _safe_json(normalized_existing_area_geojson)},
                ).scalar()
                or 0.0
            )
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid existing_area_geojson")

    resolved_scope, planting_scope_flag, carbon_scope_flag = _resolve_tree_scope_defaults(
        tree_origin=origin,
        attribution_scope=attribution_scope,
        count_in_planting_kpis=count_in_planting_kpis,
        count_in_carbon_scope=count_in_carbon_scope,
        project_existing_scope=project_settings.get("default_existing_tree_scope"),
    )

    custodian_id_value = int(custodian_id) if custodian_id is not None else None
    if custodian_id_value is not None:
        custodian_row = db.execute(
            text(
                """
                SELECT id, project_id
                FROM green_custodians
                WHERE id = :custodian_id
                """
            ),
            {"custodian_id": custodian_id_value},
        ).mappings().first()
        if not custodian_row:
            raise HTTPException(status_code=404, detail="Custodian not found")
        if int(custodian_row.get("project_id") or 0) != int(project_id):
            raise HTTPException(status_code=400, detail="Custodian belongs to a different project")

    source_project_id_value = int(source_project_id) if source_project_id is not None else None
    if source_project_id_value is not None and source_project_id_value > 0:
        source_project_exists = db.execute(
            text("SELECT 1 FROM tree_projects WHERE id = :project_id"),
            {"project_id": source_project_id_value},
        ).scalar()
        if not source_project_exists:
            raise HTTPException(status_code=404, detail="source_project_id not found")

    requested_status = _normalize_tree_status(status or "alive")
    if requested_status not in TREE_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid status")
    # New planting is supervisor-reviewed first; tree remains pending until approval.
    normalized_status = "pending_planting" if origin == "new_planting" else requested_status
    created_by_clean = (created_by or "").strip()
    normalized_photo_url, normalized_photo_urls = _merge_photo_evidence(photo_url, photo_urls)
    reported_status = requested_status if requested_status != "pending_planting" else "alive"
    project_tree_no = _next_project_tree_no(db, int(project_id))
    row = db.execute(text("""
        INSERT INTO trees (
            project_id, project_tree_no, geom, species, planting_date, status, notes, photo_url, photo_urls, created_by,
            tree_origin, custodian_id, custody_started_at, attribution_scope,
            count_in_planting_kpis, count_in_carbon_scope, source_project_id, tree_height_m, tree_age_months,
            inventory_tree_count, existing_area_geojson, existing_area_sqm
        )
        VALUES (
            :project_id,
            :project_tree_no,
            ST_SetSRID(ST_MakePoint(:lng, :lat), 4326),
            :species,
            :planting_date,
            :status,
            :notes,
            :photo_url,
            CAST(:photo_urls AS JSONB),
            :created_by,
            :tree_origin,
            :custodian_id,
            :custody_started_at,
            :attribution_scope,
            :count_in_planting_kpis,
            :count_in_carbon_scope,
            :source_project_id,
            :tree_height_m,
            :tree_age_months,
            :inventory_tree_count,
            CAST(:existing_area_geojson AS JSONB),
            :existing_area_sqm
        )
        RETURNING id
    """), {
        "project_id": project_id,
        "project_tree_no": project_tree_no,
        "lng": lng,
        "lat": lat,
        "species": species or None,
        "planting_date": planting_date,
        "status": normalized_status,
        "notes": notes or None,
        "photo_url": normalized_photo_url,
        "photo_urls": _safe_json(normalized_photo_urls),
        "created_by": created_by_clean or None,
        "tree_origin": origin,
        "custodian_id": custodian_id_value,
        "custody_started_at": custody_started_at,
        "attribution_scope": resolved_scope,
        "count_in_planting_kpis": planting_scope_flag,
        "count_in_carbon_scope": carbon_scope_flag,
        "source_project_id": source_project_id_value,
        "tree_height_m": tree_height_value,
        "tree_age_months": tree_age_months_value,
        "inventory_tree_count": inventory_tree_count_value,
        "existing_area_geojson": _safe_json(normalized_existing_area_geojson),
        "existing_area_sqm": existing_area_sqm_value,
    }).scalar()

    _record_tree_status_history(
        db,
        tree_id=int(row),
        project_id=int(project_id),
        status=normalized_status,
        status_date=_parse_date_value(planting_date) or date.today(),
        source="tree_created_existing" if origin != "new_planting" else "tree_created",
        changed_by=created_by_clean or None,
        notes="Initial tree record created",
    )

    review_task_id = None
    auto_first_cycle_task_ids: list[int] = []
    initial_review_task_type: str | None = None
    if created_by_clean and origin in {"new_planting", "existing_inventory"}:
        initial_review_task_type = "planting" if origin == "new_planting" else "existing_inventory_intake"
        review_task_id = db.execute(
            text(
                """
                INSERT INTO tree_tasks (
                    tree_id, task_type, assignee_name, due_date, priority, status, notes, photo_url,
                    review_state, submitted_at, completed_at, reported_tree_status
                )
                VALUES (
                    :tree_id, :task_type, :assignee_name, :due_date, 'normal', 'done', :notes, :photo_url,
                    'submitted', NOW(), NOW(), :reported_tree_status
                )
                RETURNING id
                """
            ),
            {
                "tree_id": int(row),
                "task_type": initial_review_task_type,
                "assignee_name": created_by_clean,
                "due_date": planting_date,
                "notes": notes or None,
                "photo_url": photo_url or None,
                "reported_tree_status": reported_status,
            },
        ).scalar()
        db.execute(
            text(
                """
                INSERT INTO green_task_reviews (task_id, decision, reviewer_name, review_notes)
                VALUES (:task_id, 'submitted', :reviewer_name, :review_notes)
                """
            ),
            {"task_id": int(review_task_id), "reviewer_name": created_by_clean, "review_notes": notes or None},
        )
        _record_alert(
            db,
            project_id=project_id,
            alert_type="task_submitted",
            severity="warning",
            message=f"Task #{int(review_task_id)} is awaiting supervisor review.",
            tree_id=int(row),
            task_id=int(review_task_id),
        )
        _log_audit_event(
            db,
            project_id=project_id,
            entity_type="task",
            entity_id=int(review_task_id),
            action="task_submitted_for_review",
            actor=created_by_clean,
            details={"task_type": initial_review_task_type, "status": "done", "review_state": "submitted"},
        )
        if origin == "new_planting":
            matching_auto_order = _find_matching_auto_first_cycle_work_order(
                db,
                project_id=int(project_id),
                assignee_name=created_by_clean,
                species=species or None,
                lng=float(lng),
                lat=float(lat),
            )
            if matching_auto_order:
                auto_first_cycle_task_ids = _auto_assign_first_cycle_maintenance_from_order(
                    db,
                    project_id=int(project_id),
                    tree_id=int(row),
                    assignee_name=created_by_clean,
                    planting_date_value=planting_date,
                    order_row=matching_auto_order,
                )

    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="tree",
        entity_id=int(row),
        action="tree_created",
        actor=created_by_clean or None,
        details={
            "species": species or None,
            "status": normalized_status,
            "reported_status": reported_status,
            "planting_date": planting_date,
            "lng": lng,
            "lat": lat,
            "tree_origin": origin,
            "custodian_id": custodian_id_value,
            "custody_started_at": custody_started_at,
            "attribution_scope": resolved_scope,
            "count_in_planting_kpis": planting_scope_flag,
            "count_in_carbon_scope": carbon_scope_flag,
            "source_project_id": source_project_id_value,
            "tree_height_m": tree_height_value,
            "tree_age_months": tree_age_months_value,
            "inventory_tree_count": inventory_tree_count_value,
            "existing_area_geojson": normalized_existing_area_geojson,
            "existing_area_sqm": existing_area_sqm_value,
            "photo_urls_count": len(normalized_photo_urls),
            "project_tree_no": project_tree_no,
            "review_task_id": int(review_task_id) if review_task_id else None,
            "auto_first_cycle_task_ids": auto_first_cycle_task_ids,
        },
    )
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {
        "id": row,
        "project_tree_no": project_tree_no,
        "review_task_id": review_task_id,
        "auto_first_cycle_task_ids": auto_first_cycle_task_ids,
        "status": "submitted_for_review" if review_task_id else "created",
    }


@router.patch("/trees/{tree_id}")
def update_tree(
    tree_id: int,
    db: Session = Depends(get_db),
    lng: float | None = Body(default=None),
    lat: float | None = Body(default=None),
    species: str | None = Body(default=None),
    planting_date: str | None = Body(default=None),
    status: str | None = Body(default=None),
    notes: str | None = Body(default=None),
    photo_url: str | None = Body(default=None),
    photo_urls: list[str] | None = Body(default=None),
    actor_name: str | None = Body(default=None),
    tree_origin: str | None = Body(default=None),
    custodian_id: int | None = Body(default=None),
    custody_started_at: str | None = Body(default=None),
    attribution_scope: str | None = Body(default=None),
    count_in_planting_kpis: bool | None = Body(default=None),
    count_in_carbon_scope: bool | None = Body(default=None),
    source_project_id: int | None = Body(default=None),
    tree_height_m: float | None = Body(default=None),
    tree_age_months: float | None = Body(default=None),
    inventory_tree_count: int | None = Body(default=None),
    existing_area_geojson: dict | str | None = Body(default=None),
):
    if (lng is None) != (lat is None):
        raise HTTPException(status_code=400, detail="Both lng and lat are required together")
    if lng is not None and not (-180 <= float(lng) <= 180):
        raise HTTPException(status_code=400, detail="Invalid lng")
    if lat is not None and not (-90 <= float(lat) <= 90):
        raise HTTPException(status_code=400, detail="Invalid lat")
    normalized_status = _normalize_tree_status(status) if status is not None else None
    if normalized_status is not None and normalized_status not in TREE_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid status")
    normalized_origin = _normalize_tree_origin(tree_origin) if tree_origin is not None else None
    if normalized_origin is not None and normalized_origin not in TREE_ORIGIN_VALUES:
        raise HTTPException(status_code=400, detail="Invalid tree_origin")
    normalized_attribution_scope = (
        _normalize_tree_attribution_scope(attribution_scope) if attribution_scope is not None else None
    )
    if normalized_attribution_scope is not None and normalized_attribution_scope not in TREE_ATTRIBUTION_SCOPE_VALUES:
        raise HTTPException(status_code=400, detail="Invalid attribution_scope")

    tree_height_value = None
    if tree_height_m is not None:
        try:
            tree_height_value = float(tree_height_m)
        except Exception:
            raise HTTPException(status_code=400, detail="tree_height_m must be numeric")
        if tree_height_value < 0 or tree_height_value > 120:
            raise HTTPException(status_code=400, detail="tree_height_m must be between 0 and 120")

    tree_age_months_value = None
    if tree_age_months is not None:
        try:
            tree_age_months_value = float(tree_age_months)
        except Exception:
            raise HTTPException(status_code=400, detail="tree_age_months must be numeric")
        if tree_age_months_value < 0 or tree_age_months_value > 2400:
            raise HTTPException(status_code=400, detail="tree_age_months must be between 0 and 2400")

    inventory_tree_count_value = None
    if inventory_tree_count is not None:
        try:
            inventory_tree_count_value = int(inventory_tree_count)
        except Exception:
            raise HTTPException(status_code=400, detail="inventory_tree_count must be a whole number")
        if inventory_tree_count_value < 1 or inventory_tree_count_value > 1000000:
            raise HTTPException(status_code=400, detail="inventory_tree_count must be between 1 and 1000000")

    normalized_existing_area_geojson = _normalize_work_area_geojson(existing_area_geojson)

    source_project_id_value = None
    if source_project_id is not None:
        source_project_id_value = int(source_project_id)
        if source_project_id_value <= 0:
            raise HTTPException(status_code=400, detail="source_project_id must be positive")
        source_project_exists = db.execute(
            text("SELECT 1 FROM tree_projects WHERE id = :project_id"),
            {"project_id": source_project_id_value},
        ).scalar()
        if not source_project_exists:
            raise HTTPException(status_code=404, detail="source_project_id not found")

    existing = db.execute(
        text(
            """
            SELECT
                project_id,
                ST_X(geom) AS lng,
                ST_Y(geom) AS lat,
                species,
                planting_date,
                status,
                notes,
                photo_url,
                photo_urls,
                tree_origin,
                custodian_id,
                custody_started_at,
                attribution_scope,
                count_in_planting_kpis,
                count_in_carbon_scope,
                source_project_id,
                tree_height_m,
                tree_age_months,
                COALESCE(inventory_tree_count, 1) AS inventory_tree_count,
                existing_area_geojson,
                existing_area_sqm
            FROM trees
            WHERE id = :tree_id
            """
        ),
        {"tree_id": tree_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Tree not found")

    custodian_id_value = int(custodian_id) if custodian_id is not None else None
    if custodian_id_value is not None:
        custodian_row = db.execute(
            text("SELECT id, project_id FROM green_custodians WHERE id = :custodian_id"),
            {"custodian_id": custodian_id_value},
        ).mappings().first()
        if not custodian_row:
            raise HTTPException(status_code=404, detail="Custodian not found")
        if int(custodian_row.get("project_id") or 0) != int(existing["project_id"]):
            raise HTTPException(status_code=400, detail="Custodian belongs to a different project")

    origin_for_scope = normalized_origin if normalized_origin is not None else _normalize_tree_origin(existing.get("tree_origin"))
    scope_for_resolve = normalized_attribution_scope if normalized_attribution_scope is not None else existing.get("attribution_scope")
    project_settings = _get_project_settings(db, int(existing["project_id"]))
    resolved_scope, resolved_count_in_planting_kpis, resolved_count_in_carbon_scope = _resolve_tree_scope_defaults(
        tree_origin=origin_for_scope,
        attribution_scope=scope_for_resolve,
        count_in_planting_kpis=count_in_planting_kpis
        if count_in_planting_kpis is not None
        else (
            bool(existing.get("count_in_planting_kpis"))
            if normalized_origin is None and normalized_attribution_scope is None
            else None
        ),
        count_in_carbon_scope=count_in_carbon_scope
        if count_in_carbon_scope is not None
        else (
            bool(existing.get("count_in_carbon_scope"))
            if normalized_origin is None and normalized_attribution_scope is None
            else None
        ),
        project_existing_scope=project_settings.get("default_existing_tree_scope"),
    )

    if origin_for_scope != "existing_inventory":
        if inventory_tree_count_value is not None and inventory_tree_count_value != 1:
            raise HTTPException(status_code=400, detail="inventory_tree_count > 1 is only allowed for Existing Tree records")
        if normalized_existing_area_geojson is not None:
            raise HTTPException(status_code=400, detail="existing_area_geojson is only allowed for Existing Tree records")
        inventory_tree_count_value = 1 if inventory_tree_count is not None else None
    else:
        next_inventory_count = (
            inventory_tree_count_value
            if inventory_tree_count_value is not None
            else max(int(existing.get("inventory_tree_count") or 1), 1)
        )
        next_area_geojson = normalized_existing_area_geojson if existing_area_geojson is not None else existing.get("existing_area_geojson")
        if next_inventory_count > 1 and not next_area_geojson:
            raise HTTPException(
                status_code=400,
                detail="Draw a polygon area when capturing more than one existing tree in a single record",
            )

    effective_existing_area_geojson = (
        normalized_existing_area_geojson if existing_area_geojson is not None else existing.get("existing_area_geojson")
    )
    existing_area_sqm_value = None
    if existing_area_geojson is not None and normalized_existing_area_geojson is not None:
        try:
            existing_area_sqm_value = float(
                db.execute(
                    text(
                        """
                        SELECT ST_Area(
                            ST_Transform(
                                ST_SetSRID(ST_GeomFromGeoJSON(:geojson), 4326),
                                3857
                            )
                        )
                        """
                    ),
                    {"geojson": _safe_json(normalized_existing_area_geojson)},
                ).scalar()
                or 0.0
            )
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid existing_area_geojson")

    effective_photo_url = None
    effective_photo_urls_json = None
    if photo_url is not None or photo_urls is not None:
        base_photo_url = photo_url if photo_url is not None else str(existing.get("photo_url") or "")
        base_photo_urls = photo_urls if photo_urls is not None else existing.get("photo_urls")
        merged_photo_url, merged_photo_urls = _merge_photo_evidence(base_photo_url, base_photo_urls)
        effective_photo_url = merged_photo_url
        effective_photo_urls_json = _safe_json(merged_photo_urls)

    db.execute(text("""
        UPDATE trees
        SET geom = CASE
                WHEN :lng IS NOT NULL AND :lat IS NOT NULL THEN ST_SetSRID(ST_MakePoint(:lng, :lat), 4326)
                ELSE geom
            END,
            species = COALESCE(:species, species),
            planting_date = COALESCE(:planting_date, planting_date),
            status = COALESCE(:status, status),
            notes = COALESCE(:notes, notes),
            photo_url = COALESCE(:photo_url, photo_url),
            photo_urls = COALESCE(CAST(:photo_urls AS JSONB), photo_urls),
            tree_origin = COALESCE(:tree_origin, tree_origin),
            custodian_id = COALESCE(:custodian_id, custodian_id),
            custody_started_at = COALESCE(:custody_started_at, custody_started_at),
            attribution_scope = COALESCE(:attribution_scope, attribution_scope),
            count_in_planting_kpis = COALESCE(:count_in_planting_kpis, count_in_planting_kpis),
            count_in_carbon_scope = COALESCE(:count_in_carbon_scope, count_in_carbon_scope),
            source_project_id = COALESCE(:source_project_id, source_project_id),
            tree_height_m = COALESCE(:tree_height_m, tree_height_m),
            tree_age_months = COALESCE(:tree_age_months, tree_age_months),
            inventory_tree_count = COALESCE(:inventory_tree_count, inventory_tree_count),
            existing_area_geojson = COALESCE(CAST(:existing_area_geojson AS JSONB), existing_area_geojson),
            existing_area_sqm = COALESCE(:existing_area_sqm, existing_area_sqm)
        WHERE id = :tree_id
    """), {
        "lng": lng,
        "lat": lat,
        "species": species,
        "planting_date": planting_date,
        "status": normalized_status,
        "notes": notes,
        "photo_url": effective_photo_url,
        "photo_urls": effective_photo_urls_json,
        "tree_origin": normalized_origin,
        "custodian_id": custodian_id_value,
        "custody_started_at": custody_started_at,
        "attribution_scope": resolved_scope if (normalized_attribution_scope is not None or normalized_origin is not None) else None,
        "count_in_planting_kpis": resolved_count_in_planting_kpis if (count_in_planting_kpis is not None or normalized_origin is not None or normalized_attribution_scope is not None) else None,
        "count_in_carbon_scope": resolved_count_in_carbon_scope if (count_in_carbon_scope is not None or normalized_origin is not None or normalized_attribution_scope is not None) else None,
        "source_project_id": source_project_id_value,
        "tree_height_m": tree_height_value,
        "tree_age_months": tree_age_months_value,
        "inventory_tree_count": inventory_tree_count_value,
        "existing_area_geojson": _safe_json(normalized_existing_area_geojson) if existing_area_geojson is not None else None,
        "existing_area_sqm": existing_area_sqm_value,
        "tree_id": tree_id,
    })
    previous_status = _normalize_tree_status(existing.get("status"))
    next_status = normalized_status if normalized_status is not None else previous_status
    if normalized_status is not None and next_status != previous_status:
        _record_tree_status_history(
            db,
            tree_id=tree_id,
            project_id=int(existing["project_id"]),
            status=next_status,
            status_date=date.today(),
            source="tree_updated",
            changed_by=actor_name,
            notes="Tree status updated via tree patch endpoint",
        )
    _log_audit_event(
        db,
        project_id=int(existing["project_id"]),
        entity_type="tree",
        entity_id=tree_id,
        action="tree_updated",
        details={
            "before": dict(existing),
            "after": {
                "species": species if species is not None else existing.get("species"),
                "lng": lng if lng is not None else existing.get("lng"),
                "lat": lat if lat is not None else existing.get("lat"),
                "planting_date": planting_date if planting_date is not None else existing.get("planting_date"),
                "status": normalized_status if normalized_status is not None else existing.get("status"),
                "notes": notes if notes is not None else existing.get("notes"),
                "photo_url": effective_photo_url if (photo_url is not None or photo_urls is not None) else existing.get("photo_url"),
                "photo_urls": (
                    _normalize_photo_urls(effective_photo_urls_json)
                    if (photo_url is not None or photo_urls is not None)
                    else _normalize_photo_urls(existing.get("photo_urls"))
                ),
                "tree_origin": normalized_origin if normalized_origin is not None else existing.get("tree_origin"),
                "custodian_id": custodian_id_value if custodian_id is not None else existing.get("custodian_id"),
                "custody_started_at": custody_started_at if custody_started_at is not None else existing.get("custody_started_at"),
                "attribution_scope": (
                    resolved_scope
                    if (normalized_attribution_scope is not None or normalized_origin is not None)
                    else existing.get("attribution_scope")
                ),
                "count_in_planting_kpis": (
                    resolved_count_in_planting_kpis
                    if (count_in_planting_kpis is not None or normalized_origin is not None or normalized_attribution_scope is not None)
                    else existing.get("count_in_planting_kpis")
                ),
                "count_in_carbon_scope": (
                    resolved_count_in_carbon_scope
                    if (count_in_carbon_scope is not None or normalized_origin is not None or normalized_attribution_scope is not None)
                    else existing.get("count_in_carbon_scope")
                ),
                "source_project_id": source_project_id_value if source_project_id is not None else existing.get("source_project_id"),
                "tree_height_m": tree_height_value if tree_height_m is not None else existing.get("tree_height_m"),
                "tree_age_months": tree_age_months_value if tree_age_months is not None else existing.get("tree_age_months"),
                "inventory_tree_count": (
                    inventory_tree_count_value if inventory_tree_count is not None else existing.get("inventory_tree_count")
                ),
                "existing_area_geojson": (
                    effective_existing_area_geojson if existing_area_geojson is not None else existing.get("existing_area_geojson")
                ),
                "existing_area_sqm": (
                    existing_area_sqm_value if existing_area_geojson is not None else existing.get("existing_area_sqm")
                ),
            },
        },
    )
    db.commit()
    return {"status": "ok"}


@router.delete("/trees/{tree_id}")
def delete_tree(
    tree_id: int,
    confirm_tree_id: int | None = Query(default=None),
    project_id: int | None = Query(default=None),
    actor_name: str | None = Query(default=None),
    db: Session = Depends(get_db),
):
    row = db.execute(
        text(
            """
            SELECT id, project_id, species, status, created_by
            FROM trees
            WHERE id = :tree_id
            """
        ),
        {"tree_id": tree_id},
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Tree not found")

    project_id_value = int(row.get("project_id") or 0)
    if project_id is not None and int(project_id) != project_id_value:
        raise HTTPException(status_code=400, detail="Project mismatch for selected tree")
    if confirm_tree_id is not None and int(confirm_tree_id) != int(tree_id):
        raise HTTPException(status_code=400, detail="Confirmation tree id does not match")

    db.execute(text("DELETE FROM trees WHERE id = :tree_id"), {"tree_id": tree_id})
    _log_audit_event(
        db,
        project_id=project_id_value,
        entity_type="tree",
        entity_id=int(tree_id),
        action="tree_deleted",
        actor=(actor_name or row.get("created_by") or None),
        details={
            "species": row.get("species"),
            "status": row.get("status"),
        },
    )
    _refresh_project_alerts(db, project_id_value)
    db.commit()
    return {"status": "deleted", "id": int(tree_id), "project_id": project_id_value}


@router.post("/trees/{tree_id}/visits")
def add_visit(
    tree_id: int,
    db: Session = Depends(get_db),
    visit_date: str = Body(...),
    status: str = Body(...),
    notes: str = Body(default=""),
    photo_url: str = Body(default=""),
    created_by: str = Body(default=""),
):
    normalized_status = _normalize_tree_status(status)
    if normalized_status not in TREE_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid status")
    db.execute(text("""
        INSERT INTO tree_visits (tree_id, visit_date, status, notes, photo_url, created_by)
        VALUES (:tree_id, :visit_date, :status, :notes, :photo_url, :created_by)
    """), {
        "tree_id": tree_id,
        "visit_date": visit_date,
        "status": normalized_status,
        "notes": notes or None,
        "photo_url": photo_url or None,
        "created_by": created_by or None,
    })
    project_id = _get_project_id_for_tree(db, tree_id)
    _record_tree_status_history(
        db,
        tree_id=tree_id,
        project_id=project_id,
        status=normalized_status,
        status_date=_parse_date_value(visit_date) or date.today(),
        source="visit",
        changed_by=(created_by or "").strip() or None,
        notes=notes or None,
    )
    db.commit()
    return {"status": "ok"}


@router.post("/trees/{tree_id}/tasks")
def add_task(
    tree_id: int,
    db: Session = Depends(get_db),
    task_type: str = Body(...),
    assignee_name: str = Body(...),
    due_date: str | None = Body(default=None),
    priority: str = Body(default="normal"),
    status: str = Body(default="pending"),
    notes: str = Body(default=""),
    photo_url: str = Body(default=""),
    photo_urls: list[str] | None = Body(default=None),
    model_season: str | None = Body(default=None),
    custodian_id: int | None = Body(default=None),
    distribution_allocation_id: int | None = Body(default=None),
    supervision_visit_no: int | None = Body(default=None),
    supervision_total_visits: int | None = Body(default=None),
):
    if status not in {"pending", "done", "overdue"}:
        raise HTTPException(status_code=400, detail="Invalid status")
    activity = _normalize_name(task_type)
    if activity not in ASSIGNABLE_TASK_TYPES:
        raise HTTPException(status_code=400, detail="Invalid task type")
    tree_row = db.execute(
        text("SELECT project_id, status FROM trees WHERE id = :tree_id"),
        {"tree_id": tree_id},
    ).mappings().first()
    if not tree_row:
        raise HTTPException(status_code=404, detail="Tree not found")
    tree_status = _normalize_tree_status(tree_row.get("status") or "alive")
    replacement_required = _is_replacement_trigger_status(tree_status)
    if activity == "replacement" and not replacement_required:
        raise HTTPException(
            status_code=400,
            detail="Replacement can only be assigned when tree status is dead, damaged, removed, or needs replacement.",
        )
    if replacement_required and activity not in {"replacement", SUPERVISION_TASK_TYPE}:
        raise HTTPException(
            status_code=400,
            detail="This tree currently requires replacement. Assign and complete replacement first.",
        )

    if activity == SUPERVISION_TASK_TYPE:
        if distribution_allocation_id is None:
            raise HTTPException(status_code=400, detail="distribution_allocation_id is required for supervision tasks")
        allocation = db.execute(
            text(
                """
                SELECT id, project_id, custodian_id, supervision_target
                FROM green_distribution_allocations
                WHERE id = :allocation_id
                """
            ),
            {"allocation_id": int(distribution_allocation_id)},
        ).mappings().first()
        if not allocation:
            raise HTTPException(status_code=404, detail="Distribution allocation not found")
        if int(allocation.get("project_id") or 0) != int(tree_row.get("project_id") or 0):
            raise HTTPException(status_code=400, detail="Allocation and tree must belong to the same project")
        if custodian_id is None:
            custodian_id = int(allocation.get("custodian_id") or 0)
        if int(custodian_id or 0) != int(allocation.get("custodian_id") or 0):
            raise HTTPException(status_code=400, detail="custodian_id must match distribution allocation")
        if supervision_total_visits is None:
            supervision_total_visits = int(allocation.get("supervision_target") or 0) or None

    normalized_season = _normalize_name(model_season)
    if normalized_season and normalized_season not in SEASON_VALUES:
        normalized_season = "rainy"
    normalized_photo_url, normalized_photo_urls = _merge_photo_evidence(photo_url, photo_urls)

    review_state = "none"
    submitted_at = None
    completed_at = None
    if _is_done_status(status):
        evidence_ok, detail = _has_required_evidence(activity, notes, normalized_photo_url, normalized_photo_urls)
        if not evidence_ok:
            raise HTTPException(status_code=400, detail=detail)
        review_state = "submitted"
        submitted_at = datetime.utcnow()
        completed_at = datetime.utcnow()

    row = db.execute(text("""
        INSERT INTO tree_tasks (
            tree_id, task_type, assignee_name, due_date, priority, status, notes, photo_url,
            photo_urls, review_state, submitted_at, completed_at, model_season,
            custodian_id, distribution_allocation_id, supervision_visit_no, supervision_total_visits
        )
        VALUES (
            :tree_id, :task_type, :assignee_name, :due_date, :priority, :status, :notes, :photo_url,
            CAST(:photo_urls AS JSONB), :review_state, :submitted_at, :completed_at, :model_season,
            :custodian_id, :distribution_allocation_id, :supervision_visit_no, :supervision_total_visits
        )
        RETURNING id
    """), {
        "tree_id": tree_id,
        "task_type": activity,
        "assignee_name": assignee_name,
        "due_date": due_date,
        "priority": priority,
        "status": status,
        "notes": notes or None,
        "photo_url": normalized_photo_url,
        "photo_urls": _safe_json(normalized_photo_urls),
        "review_state": review_state,
        "submitted_at": submitted_at,
        "completed_at": completed_at,
        "model_season": normalized_season or None,
        "custodian_id": int(custodian_id) if custodian_id is not None else None,
        "distribution_allocation_id": int(distribution_allocation_id) if distribution_allocation_id is not None else None,
        "supervision_visit_no": int(supervision_visit_no) if supervision_visit_no is not None else None,
        "supervision_total_visits": int(supervision_total_visits) if supervision_total_visits is not None else None,
    }).scalar()
    _log_audit_event(
        db,
        project_id=int(tree_row["project_id"]),
        entity_type="task",
        entity_id=int(row),
        action="task_created",
        actor=assignee_name,
        details={
            "task_type": activity,
            "due_date": due_date,
            "priority": priority,
            "status": status,
            "review_state": review_state,
            "model_season": normalized_season or None,
        },
    )
    if review_state == "submitted":
        _record_alert(
            db,
            project_id=int(tree_row["project_id"]),
            alert_type="task_submitted",
            severity="warning",
            message=f"Task #{int(row)} is awaiting supervisor review.",
            tree_id=tree_id,
            task_id=int(row),
        )
    db.commit()
    return {"id": row}


@router.get("/trees/{tree_id}/tasks")
def list_tree_tasks(tree_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT id, tree_id, task_type, assignee_name, due_date, priority,
               status, notes, photo_url, photo_urls, created_at, completed_at, review_state,
               submitted_at, reviewed_at, reviewed_by, review_notes, auto_generated, model_season, source_task_id,
               reported_tree_status, activity_lng, activity_lat, activity_recorded_at,
               custodian_id, distribution_allocation_id, supervision_visit_no, supervision_total_visits
        FROM tree_tasks
        WHERE tree_id = :tree_id
          AND COALESCE(auto_generated, FALSE) = FALSE
        ORDER BY created_at DESC
    """), {"tree_id": tree_id}).mappings().all()
    return [dict(r) for r in rows]


@router.get("/tasks")
def list_tasks(
    project_id: int,
    assignee_name: str | None = None,
    db: Session = Depends(get_db),
):
    rows = db.execute(text("""
        SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.due_date, t.priority,
               t.status, t.notes, t.photo_url, t.photo_urls, t.created_at, t.completed_at, t.review_state,
               t.submitted_at, t.reviewed_at, t.reviewed_by, t.review_notes, t.auto_generated, t.model_season, t.source_task_id,
               t.reported_tree_status, t.activity_lng, t.activity_lat, t.activity_recorded_at,
               t.custodian_id, t.distribution_allocation_id, t.supervision_visit_no, t.supervision_total_visits,
               tr.status AS tree_status, tr.species AS tree_species,
               COALESCE(tr.planting_date, t.due_date, t.created_at::date) AS tree_planting_date,
               ST_X(tr.geom) AS lng, ST_Y(tr.geom) AS lat,
               c.name AS custodian_name, c.custodian_type, c.community_name AS custodian_community_name,
               c.contact_person AS custodian_contact_person, c.phone AS custodian_phone, c.email AS custodian_email,
               a.supervision_target, a.followup_cycle_days, e.event_date AS allocation_event_date, e.species AS allocation_species
        FROM tree_tasks t
        JOIN trees tr ON tr.id = t.tree_id
        LEFT JOIN green_custodians c ON c.id = COALESCE(t.custodian_id, tr.custodian_id)
        LEFT JOIN green_distribution_allocations a ON a.id = t.distribution_allocation_id
        LEFT JOIN green_distribution_events e ON e.id = a.event_id
        WHERE tr.project_id = :project_id
          AND COALESCE(t.auto_generated, FALSE) = FALSE
          AND (:assignee_name IS NULL OR t.assignee_name = :assignee_name)
        ORDER BY t.created_at DESC
    """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()
    return [dict(r) for r in rows]


@router.post("/users")
def create_user(
    db: Session = Depends(get_db),
    full_name: str = Body(...),
    role: str = Body(default="field_officer"),
    user_uid: str | None = Body(default=None),
    organization_id: int | None = Body(default=None),
    role_id: int | None = Body(default=None),
    email: str | None = Body(default=None),
    phone: str | None = Body(default=None),
    allow_green: bool = Body(default=True),
    allow_work: bool = Body(default=False),
    work_username: str | None = Body(default=None),
    work_password: str | None = Body(default=None),
    send_credentials_email: bool = Body(default=True),
    notes: str | None = Body(default=None),
    is_active: bool = Body(default=True),
):
    full_name_clean = (full_name or "").strip()
    if not full_name_clean:
        raise HTTPException(status_code=400, detail="Full name required")
    org_id_value = int(organization_id) if organization_id is not None else None
    org_name_value: str | None = None
    if org_id_value is not None:
        org_exists = db.execute(
            text("SELECT id, name FROM green_organizations WHERE id = :org_id"),
            {"org_id": org_id_value},
        ).mappings().first()
        if not org_exists:
            raise HTTPException(status_code=404, detail="Organization not found")
        org_name_value = str(org_exists.get("name") or "").strip() or None
    role_value = _normalize_name(role) or "field_officer"
    role_id_value = int(role_id) if role_id is not None else None
    if role_id_value is not None and role_id_value <= 0:
        role_id_value = None
    role_row = None
    if role_id_value is not None:
        role_row = db.execute(
            text("SELECT id, role_key, role_name FROM green_roles WHERE id = :role_id"),
            {"role_id": role_id_value},
        ).mappings().first()
        if not role_row:
            raise HTTPException(status_code=404, detail="Role not found")
        role_value = str(role_row.get("role_key") or role_value or "field_officer")
    else:
        role_row = db.execute(
            text("SELECT id, role_key, role_name FROM green_roles WHERE LOWER(role_key) = LOWER(:role_key) LIMIT 1"),
            {"role_key": role_value},
        ).mappings().first()
        role_id_value = int(role_row["id"]) if role_row else None
    user_uid_value = _ensure_unique_user_uid(db, user_uid)
    login_enabled = bool(allow_green) or bool(allow_work)
    work_username_clean = (work_username or "").strip().lower() or None
    if not work_username_clean and login_enabled:
        work_username_clean = str(user_uid_value or "").strip().lower() or None
    if work_username_clean:
        existing_work_username = db.execute(
            text("SELECT id FROM green_users WHERE LOWER(COALESCE(work_username, '')) = LOWER(:work_username) LIMIT 1"),
            {"work_username": work_username_clean},
        ).scalar()
        if existing_work_username:
            raise HTTPException(status_code=409, detail="Work username already exists")
    generated_password: str | None = None
    explicit_password_clean = str(work_password or "").strip()
    if login_enabled and not explicit_password_clean:
        generated_password = _generate_temporary_login_password()
    plain_login_password = generated_password or explicit_password_clean
    password_hash = _hash_password_value(plain_login_password or "") if login_enabled else None
    email_clean = (email or "").strip() or None
    if login_enabled and (not work_username_clean or not password_hash):
        raise HTTPException(
            status_code=400,
            detail="Green/Work-enabled users require login username and password",
        )
    if login_enabled and bool(send_credentials_email) and not email_clean:
        raise HTTPException(status_code=400, detail="Email is required to send login credentials")
    row = db.execute(text("""
        INSERT INTO green_users (
            full_name, role, user_uid, organization_id, role_id, email, phone,
            allow_green, allow_work, work_username, work_password_hash,
            notes, is_active, updated_at
        )
        VALUES (
            :full_name, :role, :user_uid, :organization_id, :role_id, :email, :phone,
            :allow_green, :allow_work, :work_username, :work_password_hash,
            :notes, :is_active, NOW()
        )
        RETURNING id, user_uid, full_name, role, organization_id, role_id, email, phone,
                  allow_green, allow_work, work_username, notes, is_active, created_at, updated_at
    """), {
        "full_name": full_name_clean,
        "role": role_value,
        "user_uid": user_uid_value,
        "organization_id": org_id_value,
        "role_id": role_id_value,
        "email": email_clean,
        "phone": (phone or "").strip() or None,
        "allow_green": bool(allow_green),
        "allow_work": bool(allow_work),
        "work_username": work_username_clean,
        "work_password_hash": password_hash,
        "notes": (notes or "").strip() or None,
        "is_active": bool(is_active),
    }).mappings().first()
    _log_audit_event(
        db,
        project_id=None,
        entity_type="user",
        entity_id=int(row["id"]),
        action="user_created",
        actor="super_admin",
        details={
            "user_uid": row.get("user_uid"),
            "role": row.get("role"),
            "organization_id": row.get("organization_id"),
            "allow_green": bool(row.get("allow_green")),
            "allow_work": bool(row.get("allow_work")),
        },
    )
    db.commit()
    payload = list_users(
        db=db,
        include_inactive=True,
        organization_id=None,
        role_id=None,
        user_id_filter=int(row["id"]),
    )[0]
    if generated_password:
        payload["generated_password"] = generated_password
        payload["generated_login_username"] = work_username_clean
    payload["credentials_email_attempted"] = False
    payload["credentials_email_sent"] = False
    payload["credentials_email_error"] = None
    if login_enabled and bool(send_credentials_email) and email_clean and work_username_clean and plain_login_password:
        payload["credentials_email_attempted"] = True
        try:
            _send_new_user_credentials_email(
                to_email=email_clean,
                full_name=full_name_clean,
                organization_name=org_name_value,
                username=work_username_clean,
                password=plain_login_password,
                allow_green=bool(allow_green),
                allow_work=bool(allow_work),
            )
            payload["credentials_email_sent"] = True
        except Exception as exc:
            payload["credentials_email_error"] = str(exc)
    return payload


@router.get("/users")
def list_users(
    db: Session = Depends(get_db),
    include_inactive: bool = Query(default=False),
    organization_id: int | None = Query(default=None),
    role_id: int | None = Query(default=None),
    user_id_filter: int | None = None,
):
    def _safe_int_or_none(value):
        if value is None:
            return None
        try:
            return int(value)
        except Exception:
            return None

    rows = db.execute(text("""
        SELECT
            u.id, u.user_uid, u.full_name, u.role, u.organization_id, u.role_id,
            u.email, u.phone,
            COALESCE(u.allow_green, TRUE) AS allow_green,
            COALESCE(u.allow_work, FALSE) AS allow_work,
            u.work_username,
            u.notes, COALESCE(u.is_active, TRUE) AS is_active,
            u.created_at, u.updated_at,
            o.name AS organization_name, o.slug AS organization_slug, o.logo_url AS organization_logo_url,
            r.role_uid, r.role_key, r.role_name
        FROM green_users u
        LEFT JOIN green_organizations o ON o.id = u.organization_id
        LEFT JOIN green_roles r ON r.id = u.role_id
        WHERE (:include_inactive = TRUE OR COALESCE(u.is_active, TRUE) = TRUE)
          AND (:organization_id IS NULL OR u.organization_id = :organization_id)
          AND (:role_id IS NULL OR u.role_id = :role_id)
          AND (:user_id_filter IS NULL OR u.id = :user_id_filter)
        ORDER BY COALESCE(u.updated_at, u.created_at) DESC, u.id DESC
    """), {
        "include_inactive": bool(include_inactive),
        "organization_id": _safe_int_or_none(organization_id),
        "role_id": _safe_int_or_none(role_id),
        "user_id_filter": _safe_int_or_none(user_id_filter),
    }).mappings().all()
    return [dict(r) for r in rows]


@router.patch("/users/{user_id}")
def update_user(
    user_id: int,
    db: Session = Depends(get_db),
    full_name: str | None = Body(default=None),
    role: str | None = Body(default=None),
    user_uid: str | None = Body(default=None),
    organization_id: int | None = Body(default=None),
    role_id: int | None = Body(default=None),
    email: str | None = Body(default=None),
    phone: str | None = Body(default=None),
    allow_green: bool | None = Body(default=None),
    allow_work: bool | None = Body(default=None),
    work_username: str | None = Body(default=None),
    work_password: str | None = Body(default=None),
    notes: str | None = Body(default=None),
    is_active: bool | None = Body(default=None),
):
    existing = db.execute(
        text(
            """
            SELECT id, user_uid, full_name, role, organization_id, role_id, email, phone,
                   COALESCE(allow_green, TRUE) AS allow_green,
                   COALESCE(allow_work, FALSE) AS allow_work,
                   work_username, work_password_hash,
                   notes, COALESCE(is_active, TRUE) AS is_active
            FROM green_users
            WHERE id = :user_id
            """
        ),
        {"user_id": user_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")
    org_id_value = None
    if organization_id is None:
        org_id_value = existing.get("organization_id")
    else:
        org_id_value = int(organization_id) if organization_id else None
        if org_id_value is not None:
            org_exists = db.execute(text("SELECT id FROM green_organizations WHERE id = :org_id"), {"org_id": org_id_value}).scalar()
            if not org_exists:
                raise HTTPException(status_code=404, detail="Organization not found")
    role_value = _normalize_name(role) if role is not None else _normalize_name(existing.get("role"))
    role_id_value = int(role_id) if role_id is not None else (int(existing["role_id"]) if existing.get("role_id") else None)
    if role_id is not None:
        if role_id_value:
            role_row = db.execute(
                text("SELECT id, role_key FROM green_roles WHERE id = :role_id"),
                {"role_id": role_id_value},
            ).mappings().first()
            if not role_row:
                raise HTTPException(status_code=404, detail="Role not found")
            role_value = str(role_row.get("role_key") or role_value or "field_officer")
        else:
            role_id_value = None
    elif role_value:
        role_row = db.execute(
            text("SELECT id, role_key FROM green_roles WHERE LOWER(role_key)=LOWER(:role_key) LIMIT 1"),
            {"role_key": role_value},
        ).mappings().first()
        if role_row:
            role_id_value = int(role_row["id"])
            role_value = str(role_row.get("role_key") or role_value)
    next_allow_work = bool(allow_work) if allow_work is not None else bool(existing.get("allow_work", False))
    next_allow_green = bool(allow_green) if allow_green is not None else bool(existing.get("allow_green", True))
    user_uid_value = (
        _ensure_unique_user_uid(db, user_uid, exclude_user_id=user_id)
        if user_uid is not None
        else str(existing.get("user_uid") or _ensure_unique_user_uid(db))
    )
    work_username_value = existing.get("work_username")
    if work_username is not None:
        proposed_work_username = str(work_username or "").strip().lower()
        work_username_value = proposed_work_username or None
    if (next_allow_green or next_allow_work) and not work_username_value:
        work_username_value = str(user_uid_value or "").strip().lower() or None
    if work_username_value:
        clash = db.execute(
            text(
                """
                SELECT id FROM green_users
                WHERE LOWER(COALESCE(work_username, '')) = LOWER(:work_username)
                  AND id <> :user_id
                LIMIT 1
                """
            ),
            {"work_username": work_username_value, "user_id": user_id},
        ).scalar()
        if clash:
            raise HTTPException(status_code=409, detail="Work username already exists")
    work_password_hash_value = existing.get("work_password_hash")
    if work_password is not None:
        work_password_hash_value = _hash_password_value(work_password) if str(work_password).strip() else None
    if (next_allow_green or next_allow_work) and (not work_username_value or not work_password_hash_value):
        raise HTTPException(
            status_code=400,
            detail="Green/Work-enabled users require login username and password",
        )
    next_full_name = (full_name.strip() if isinstance(full_name, str) else str(existing.get("full_name") or "").strip())
    if not next_full_name:
        raise HTTPException(status_code=400, detail="Full name required")
    row = db.execute(
        text(
            """
            UPDATE green_users
            SET user_uid = :user_uid,
                full_name = :full_name,
                role = :role,
                organization_id = :organization_id,
                role_id = :role_id,
                email = :email,
                phone = :phone,
                allow_green = :allow_green,
                allow_work = :allow_work,
                work_username = :work_username,
                work_password_hash = :work_password_hash,
                notes = :notes,
                is_active = :is_active,
                updated_at = NOW()
            WHERE id = :user_id
            RETURNING id
            """
        ),
        {
            "user_id": user_id,
            "user_uid": user_uid_value,
            "full_name": next_full_name,
            "role": role_value or "field_officer",
            "organization_id": org_id_value,
            "role_id": role_id_value,
            "email": (email.strip() if isinstance(email, str) else existing.get("email")) or None,
            "phone": (phone.strip() if isinstance(phone, str) else existing.get("phone")) or None,
            "allow_green": next_allow_green,
            "allow_work": next_allow_work,
            "work_username": work_username_value,
            "work_password_hash": work_password_hash_value,
            "notes": (notes.strip() if isinstance(notes, str) else existing.get("notes")) or None,
            "is_active": bool(is_active) if is_active is not None else bool(existing.get("is_active", True)),
        },
    ).scalar()
    _log_audit_event(
        db,
        project_id=None,
        entity_type="user",
        entity_id=int(row),
        action="user_updated",
        actor="super_admin",
        details={
            "user_uid": user_uid_value,
            "role": role_value or "field_officer",
            "organization_id": org_id_value,
            "allow_green": next_allow_green,
            "allow_work": next_allow_work,
            "is_active": bool(is_active) if is_active is not None else bool(existing.get("is_active", True)),
        },
    )
    db.commit()
    result = list_users(
        db=db,
        include_inactive=True,
        organization_id=None,
        role_id=None,
        user_id_filter=user_id,
    )
    return result[0] if result else {"id": user_id}


@router.delete("/users/{user_id}")
def delete_user(
    user_id: int,
    db: Session = Depends(get_db),
):
    existing = db.execute(
        text(
            """
            SELECT id, user_uid, full_name, role, organization_id, role_id,
                   email, phone,
                   COALESCE(allow_green, TRUE) AS allow_green,
                   COALESCE(allow_work, FALSE) AS allow_work,
                   work_username,
                   notes, COALESCE(is_active, TRUE) AS is_active
            FROM green_users
            WHERE id = :user_id
            """
        ),
        {"user_id": user_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    deleted_id = db.execute(
        text("DELETE FROM green_users WHERE id = :user_id RETURNING id"),
        {"user_id": user_id},
    ).scalar()
    if not deleted_id:
        raise HTTPException(status_code=404, detail="User not found")

    _log_audit_event(
        db,
        project_id=None,
        entity_type="user",
        entity_id=int(deleted_id),
        action="user_deleted",
        actor="super_admin",
        details={
            "user_uid": existing.get("user_uid"),
            "full_name": existing.get("full_name"),
            "role": existing.get("role"),
            "organization_id": existing.get("organization_id"),
            "allow_green": bool(existing.get("allow_green", True)),
            "allow_work": bool(existing.get("allow_work", False)),
            "work_username": existing.get("work_username"),
            "is_active": bool(existing.get("is_active", True)),
        },
    )
    db.commit()
    return {"ok": True, "id": int(deleted_id)}


@router.post("/users/{user_id}/reset-password")
def reset_user_password(
    user_id: int,
    db: Session = Depends(get_db),
    send_credentials_email: bool = Body(default=True),
    password_length: int = Body(default=12),
):
    existing = db.execute(
        text(
            """
            SELECT
                u.id, u.user_uid, u.full_name, u.role, u.organization_id, u.email,
                COALESCE(u.allow_green, TRUE) AS allow_green,
                COALESCE(u.allow_work, FALSE) AS allow_work,
                COALESCE(u.is_active, TRUE) AS is_active,
                u.work_username,
                o.name AS organization_name
            FROM green_users u
            LEFT JOIN green_organizations o ON o.id = u.organization_id
            WHERE u.id = :user_id
            """
        ),
        {"user_id": user_id},
    ).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="User not found")

    allow_green_value = bool(existing.get("allow_green", True))
    allow_work_value = bool(existing.get("allow_work", False))
    if not (allow_green_value or allow_work_value):
        raise HTTPException(status_code=400, detail="User has no Green/Work access enabled")

    user_uid_value = str(existing.get("user_uid") or "").strip()
    if not user_uid_value:
        user_uid_value = _ensure_unique_user_uid(db, None, exclude_user_id=user_id)

    login_username = str(existing.get("work_username") or "").strip().lower()
    if not login_username:
        login_username = user_uid_value.strip().lower()
    if not login_username:
        raise HTTPException(status_code=400, detail="Unable to determine login username")

    clash = db.execute(
        text(
            """
            SELECT id FROM green_users
            WHERE LOWER(COALESCE(work_username, '')) = LOWER(:work_username)
              AND id <> :user_id
            LIMIT 1
            """
        ),
        {"work_username": login_username, "user_id": user_id},
    ).scalar()
    if clash:
        raise HTTPException(status_code=409, detail="Login username already exists on another user")

    generated_password = _generate_temporary_login_password(password_length)
    password_hash = _hash_password_value(generated_password)

    db.execute(
        text(
            """
            UPDATE green_users
            SET user_uid = :user_uid,
                work_username = :work_username,
                work_password_hash = :work_password_hash,
                updated_at = NOW()
            WHERE id = :user_id
            """
        ),
        {
            "user_id": user_id,
            "user_uid": user_uid_value,
            "work_username": login_username,
            "work_password_hash": password_hash,
        },
    )

    _log_audit_event(
        db,
        project_id=None,
        entity_type="user",
        entity_id=int(user_id),
        action="user_password_reset",
        actor="super_admin",
        details={
            "user_uid": user_uid_value,
            "work_username": login_username,
            "allow_green": allow_green_value,
            "allow_work": allow_work_value,
            "send_credentials_email": bool(send_credentials_email),
        },
    )
    db.commit()

    payload = list_users(
        db=db,
        include_inactive=True,
        organization_id=None,
        role_id=None,
        user_id_filter=user_id,
    )[0]
    payload["generated_password"] = generated_password
    payload["generated_login_username"] = login_username
    payload["credentials_email_attempted"] = bool(send_credentials_email)
    payload["credentials_email_sent"] = False
    payload["credentials_email_error"] = None

    email_clean = str(existing.get("email") or "").strip()
    if send_credentials_email:
        if not email_clean:
            payload["credentials_email_error"] = "User has no email address"
        else:
            try:
                _send_new_user_credentials_email(
                    to_email=email_clean,
                    full_name=str(existing.get("full_name") or ""),
                    organization_name=str(existing.get("organization_name") or "").strip() or None,
                    username=login_username,
                    password=generated_password,
                    allow_green=allow_green_value,
                    allow_work=allow_work_value,
                )
                payload["credentials_email_sent"] = True
            except Exception as exc:
                payload["credentials_email_error"] = str(exc)

    return payload


@router.post("/auth/change-password")
def change_own_password(
    db: Session = Depends(get_db),
    user_id: int = Body(...),
    current_password: str = Body(...),
    new_password: str = Body(...),
    app: str | None = Body(default=None),
):
    try:
        user_id_value = int(user_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid user id")
    if user_id_value <= 0:
        raise HTTPException(status_code=400, detail="Password change is not available for this account")

    current_password_value = str(current_password or "")
    new_password_value = str(new_password or "")
    if not current_password_value:
        raise HTTPException(status_code=400, detail="Current password is required")
    if not new_password_value:
        raise HTTPException(status_code=400, detail="New password is required")

    app_scope = str(app or "").strip().lower()
    if app_scope not in {"", "green", "work"}:
        raise HTTPException(status_code=400, detail="Invalid app scope")

    user_row = db.execute(
        text(
            """
            SELECT
                u.id, u.user_uid, u.full_name, u.work_username, u.work_password_hash,
                COALESCE(u.allow_green, TRUE) AS allow_green,
                COALESCE(u.allow_work, FALSE) AS allow_work,
                COALESCE(u.is_active, TRUE) AS is_active,
                u.organization_id
            FROM green_users u
            WHERE u.id = :user_id
            LIMIT 1
            """
        ),
        {"user_id": user_id_value},
    ).mappings().first()
    if not user_row:
        raise HTTPException(status_code=404, detail="User not found")
    if not bool(user_row.get("is_active", True)):
        raise HTTPException(status_code=403, detail="User account is inactive")
    if user_row.get("organization_id") is None:
        raise HTTPException(status_code=403, detail="This user is not linked to an organization")
    if app_scope == "green" and not bool(user_row.get("allow_green", True)):
        raise HTTPException(status_code=403, detail="This user is not enabled for LandCheck Green")
    if app_scope == "work" and not bool(user_row.get("allow_work", False)):
        raise HTTPException(status_code=403, detail="This user is not enabled for LandCheck Work")

    stored_hash = user_row.get("work_password_hash")
    if not _verify_password_value(current_password_value, stored_hash):
        raise HTTPException(status_code=401, detail="Current password is incorrect")
    if _verify_password_value(new_password_value, stored_hash):
        raise HTTPException(status_code=400, detail="New password must be different from the current password")

    next_hash = _hash_password_value(new_password_value)
    db.execute(
        text(
            """
            UPDATE green_users
            SET work_password_hash = :work_password_hash,
                updated_at = NOW()
            WHERE id = :user_id
            """
        ),
        {"user_id": user_id_value, "work_password_hash": next_hash},
    )

    actor_label = str(user_row.get("user_uid") or user_row.get("work_username") or f"user:{user_id_value}")
    _log_audit_event(
        db,
        project_id=None,
        entity_type="user",
        entity_id=user_id_value,
        action="user_password_changed_self",
        actor=actor_label,
        details={
            "app": app_scope or None,
            "work_username": user_row.get("work_username"),
        },
    )
    db.commit()
    return {"ok": True}


@router.post("/work-auth/login")
def work_auth_login(
    db: Session = Depends(get_db),
    username: str = Body(...),
    password: str = Body(...),
    organization_id: int | None = Body(default=None),
):
    username_clean = str(username or "").strip()
    if not username_clean:
        raise HTTPException(status_code=400, detail="Username is required")

    # Preserve existing env-based admin login as a fallback.
    env_username = str(os.getenv("WORK_USERNAME") or os.getenv("VITE_WORK_USERNAME") or "admin").strip()
    env_password = str(os.getenv("WORK_PASSWORD") or os.getenv("VITE_WORK_PASSWORD") or "landcheckwork")
    if username_clean == env_username and str(password or "") == env_password:
        return {
            "ok": True,
            "auth_mode": "env_admin",
            "user": {
                "id": 0,
                "user_uid": "SYS-ADMIN",
                "full_name": "System Admin",
                "role": "super_admin",
                "role_key": "super_admin",
                "role_name": "Super Admin",
                "allow_work": True,
                "allow_green": True,
                "organization_id": None,
                "organization_name": None,
                "organization_logo_url": None,
            },
        }

    org_id_filter = int(organization_id) if organization_id is not None else None
    if org_id_filter is not None and org_id_filter <= 0:
        org_id_filter = None
    user_row = db.execute(
        text(
            """
            SELECT
                u.id, u.user_uid, u.full_name, u.role, u.role_id,
                COALESCE(u.allow_green, TRUE) AS allow_green,
                COALESCE(u.allow_work, FALSE) AS allow_work,
                COALESCE(u.is_active, TRUE) AS is_active,
                u.work_username, u.work_password_hash,
                u.organization_id,
                o.name AS organization_name, o.slug AS organization_slug, o.status AS organization_status,
                COALESCE(o.is_active, TRUE) AS organization_is_active,
                o.logo_url AS organization_logo_url,
                r.role_key, r.role_name, r.role_uid
            FROM green_users u
            LEFT JOIN green_organizations o ON o.id = u.organization_id
            LEFT JOIN green_roles r ON r.id = u.role_id
            WHERE (
                LOWER(COALESCE(u.work_username, '')) = LOWER(:username)
                OR LOWER(COALESCE(u.user_uid, '')) = LOWER(:username)
            )
              AND (:organization_id IS NULL OR u.organization_id = :organization_id)
            LIMIT 1
            """
        ),
        {"username": username_clean, "organization_id": org_id_filter},
    ).mappings().first()
    if not user_row:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    if not bool(user_row.get("is_active", True)):
        raise HTTPException(status_code=403, detail="User account is inactive")
    if not bool(user_row.get("allow_work", False)):
        raise HTTPException(status_code=403, detail="This user is not enabled for LandCheck Work")
    if user_row.get("organization_id") is None:
        raise HTTPException(status_code=403, detail="This user is not linked to an organization")
    if not bool(user_row.get("organization_is_active", True)):
        raise HTTPException(status_code=403, detail="This organization is inactive")
    org_status = _normalize_name(user_row.get("organization_status"))
    if org_status == "suspended":
        raise HTTPException(status_code=403, detail="This organization is suspended")
    if not _verify_password_value(password, user_row.get("work_password_hash")):
        raise HTTPException(status_code=401, detail="Invalid username or password")

    return {
        "ok": True,
        "auth_mode": "partner_user",
        "user": {
            "id": int(user_row["id"]),
            "user_uid": user_row.get("user_uid"),
            "full_name": user_row.get("full_name"),
            "role": user_row.get("role"),
            "role_id": user_row.get("role_id"),
            "role_uid": user_row.get("role_uid"),
            "role_key": user_row.get("role_key") or user_row.get("role"),
            "role_name": user_row.get("role_name") or user_row.get("role"),
            "allow_work": bool(user_row.get("allow_work")),
            "allow_green": bool(user_row.get("allow_green")),
            "organization_id": user_row.get("organization_id"),
            "organization_name": user_row.get("organization_name"),
            "organization_slug": user_row.get("organization_slug"),
            "organization_status": user_row.get("organization_status"),
            "organization_is_active": bool(user_row.get("organization_is_active", True)),
            "organization_logo_url": user_row.get("organization_logo_url"),
        },
    }


@router.post("/green-auth/login")
def green_auth_login(
    db: Session = Depends(get_db),
    username: str = Body(...),
    password: str = Body(...),
    organization_id: int | None = Body(default=None),
):
    username_clean = str(username or "").strip()
    if not username_clean:
        raise HTTPException(status_code=400, detail="Username is required")

    # Preserve env admin fallback for setup/support access.
    env_username = str(os.getenv("WORK_USERNAME") or os.getenv("VITE_WORK_USERNAME") or "admin").strip()
    env_password = str(os.getenv("WORK_PASSWORD") or os.getenv("VITE_WORK_PASSWORD") or "landcheckwork")
    if username_clean == env_username and str(password or "") == env_password:
        return {
            "ok": True,
            "auth_mode": "env_admin",
            "user": {
                "id": 0,
                "user_uid": "SYS-ADMIN",
                "full_name": "System Admin",
                "role": "super_admin",
                "role_key": "super_admin",
                "role_name": "Super Admin",
                "allow_work": True,
                "allow_green": True,
                "organization_id": None,
                "organization_name": None,
                "organization_logo_url": None,
            },
        }

    org_id_filter = int(organization_id) if organization_id is not None else None
    if org_id_filter is not None and org_id_filter <= 0:
        org_id_filter = None
    user_row = db.execute(
        text(
            """
            SELECT
                u.id, u.user_uid, u.full_name, u.role, u.role_id,
                COALESCE(u.allow_green, TRUE) AS allow_green,
                COALESCE(u.allow_work, FALSE) AS allow_work,
                COALESCE(u.is_active, TRUE) AS is_active,
                u.work_username, u.work_password_hash,
                u.organization_id,
                o.name AS organization_name, o.slug AS organization_slug, o.status AS organization_status,
                COALESCE(o.is_active, TRUE) AS organization_is_active,
                o.logo_url AS organization_logo_url,
                r.role_key, r.role_name, r.role_uid
            FROM green_users u
            LEFT JOIN green_organizations o ON o.id = u.organization_id
            LEFT JOIN green_roles r ON r.id = u.role_id
            WHERE (
                LOWER(COALESCE(u.work_username, '')) = LOWER(:username)
                OR LOWER(COALESCE(u.user_uid, '')) = LOWER(:username)
            )
              AND (:organization_id IS NULL OR u.organization_id = :organization_id)
            LIMIT 1
            """
        ),
        {"username": username_clean, "organization_id": org_id_filter},
    ).mappings().first()
    if not user_row:
        raise HTTPException(status_code=401, detail="Invalid username or password")
    if not bool(user_row.get("is_active", True)):
        raise HTTPException(status_code=403, detail="User account is inactive")
    if not bool(user_row.get("allow_green", True)):
        raise HTTPException(status_code=403, detail="This user is not enabled for LandCheck Green")
    if user_row.get("organization_id") is None:
        raise HTTPException(status_code=403, detail="This user is not linked to an organization")
    if not bool(user_row.get("organization_is_active", True)):
        raise HTTPException(status_code=403, detail="This organization is inactive")
    org_status = _normalize_name(user_row.get("organization_status"))
    if org_status == "suspended":
        raise HTTPException(status_code=403, detail="This organization is suspended")
    if not _verify_password_value(password, user_row.get("work_password_hash")):
        raise HTTPException(status_code=401, detail="Invalid username or password")

    return {
        "ok": True,
        "auth_mode": "partner_user",
        "user": {
            "id": int(user_row["id"]),
            "user_uid": user_row.get("user_uid"),
            "full_name": user_row.get("full_name"),
            "role": user_row.get("role"),
            "role_id": user_row.get("role_id"),
            "role_uid": user_row.get("role_uid"),
            "role_key": user_row.get("role_key") or user_row.get("role"),
            "role_name": user_row.get("role_name") or user_row.get("role"),
            "allow_work": bool(user_row.get("allow_work")),
            "allow_green": bool(user_row.get("allow_green")),
            "organization_id": user_row.get("organization_id"),
            "organization_name": user_row.get("organization_name"),
            "organization_slug": user_row.get("organization_slug"),
            "organization_status": user_row.get("organization_status"),
            "organization_is_active": bool(user_row.get("organization_is_active", True)),
            "organization_logo_url": user_row.get("organization_logo_url"),
        },
    }


@router.patch("/tasks/{task_id}")
def update_task(
    task_id: int,
    db: Session = Depends(get_db),
    status: str | None = Body(default=None),
    notes: str | None = Body(default=None),
    photo_url: str | None = Body(default=None),
    photo_urls: list[str] | None = Body(default=None),
    tree_status: str | None = Body(default=None),
    activity_lng: float | None = Body(default=None),
    activity_lat: float | None = Body(default=None),
    activity_recorded_at: str | None = Body(default=None),
    actor_name: str | None = Body(default=None),
):
    if status and status not in TASK_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid status")
    normalized_tree_status = _normalize_tree_status(tree_status) if tree_status is not None else None
    if normalized_tree_status is not None and normalized_tree_status not in TREE_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid tree status")
    if (activity_lng is None) != (activity_lat is None):
        raise HTTPException(status_code=400, detail="Both activity_lng and activity_lat are required together")
    if activity_lng is not None and not (-180 <= float(activity_lng) <= 180):
        raise HTTPException(status_code=400, detail="Invalid activity_lng")
    if activity_lat is not None and not (-90 <= float(activity_lat) <= 90):
        raise HTTPException(status_code=400, detail="Invalid activity_lat")
    activity_recorded_at_value = _parse_datetime_value(activity_recorded_at)
    if activity_recorded_at is not None and activity_recorded_at_value is None:
        raise HTTPException(status_code=400, detail="Invalid activity_recorded_at")
    if activity_lng is not None and activity_lat is not None and activity_recorded_at_value is None:
        activity_recorded_at_value = datetime.utcnow()
    existing = db.execute(text("""
        SELECT t.id, t.tree_id, t.task_type, t.status, t.review_state, t.notes, t.photo_url, t.photo_urls,
               t.completed_at, t.activity_lng, t.activity_lat, t.activity_recorded_at,
               t.submitted_at, t.reviewed_at, t.reviewed_by, t.review_notes, t.reported_tree_status,
               tr.project_id, tr.status AS tree_status
        FROM tree_tasks t
        JOIN trees tr ON tr.id = t.tree_id
        WHERE t.id = :task_id
    """), {"task_id": task_id}).mappings().first()
    if not existing:
        raise HTTPException(status_code=404, detail="Task not found")
    if _normalize_name(existing.get("review_state")) == "approved":
        raise HTTPException(status_code=409, detail="Task already approved and locked")
    next_status = status or existing.get("status")
    next_notes = notes if notes is not None else existing.get("notes")
    existing_photo, existing_photo_urls = _merge_photo_evidence(existing.get("photo_url"), existing.get("photo_urls"))
    photo_update_requested = photo_url is not None or photo_urls is not None
    next_photo_urls = list(existing_photo_urls)
    if photo_urls is not None:
        next_photo_urls = _normalize_photo_urls(photo_urls)
    if photo_url is not None:
        explicit_photo = str(photo_url or "").strip()
        if explicit_photo and explicit_photo not in next_photo_urls:
            next_photo_urls.append(explicit_photo)
    if photo_update_requested:
        next_photo = next_photo_urls[-1] if next_photo_urls else None
    else:
        next_photo = existing_photo
    next_review_state = existing.get("review_state") or "none"
    next_submitted_at = existing.get("submitted_at")
    next_completed_at = existing.get("completed_at")

    clear_review_fields = False
    if _is_done_status(next_status):
        evidence_ok, detail = _has_required_evidence(existing.get("task_type"), next_notes, next_photo, next_photo_urls)
        if not evidence_ok:
            raise HTTPException(status_code=400, detail=detail)
        next_review_state = "submitted"
        next_submitted_at = datetime.utcnow()
        next_completed_at = datetime.utcnow()
        clear_review_fields = True
    elif status is not None:
        # Task moved out of done-state; keep it editable.
        next_completed_at = None
        current_review_state = _normalize_name(existing.get("review_state"))
        if current_review_state == "rejected":
            # Keep rejected state visible until staff explicitly resubmits.
            next_review_state = "rejected"
            clear_review_fields = False
        else:
            if _normalize_name(next_review_state) in {"submitted", "rejected", "reopened"}:
                next_review_state = "none"
            next_submitted_at = None
            clear_review_fields = True

    row = db.execute(text("""
        UPDATE tree_tasks
        SET status = COALESCE(:status, status),
            notes = COALESCE(:notes, notes),
            photo_url = COALESCE(:photo_url, photo_url),
            photo_urls = COALESCE(CAST(:photo_urls AS JSONB), photo_urls),
            reported_tree_status = COALESCE(:reported_tree_status, reported_tree_status),
            activity_lng = COALESCE(:activity_lng, activity_lng),
            activity_lat = COALESCE(:activity_lat, activity_lat),
            activity_recorded_at = COALESCE(:activity_recorded_at, activity_recorded_at),
            review_state = :review_state,
            submitted_at = :submitted_at,
            reviewed_at = CASE WHEN :clear_review_fields THEN NULL ELSE reviewed_at END,
            reviewed_by = CASE WHEN :clear_review_fields THEN NULL ELSE reviewed_by END,
            review_notes = CASE WHEN :clear_review_fields THEN NULL ELSE review_notes END,
            completed_at = :completed_at
        WHERE id = :task_id
        RETURNING tree_id, photo_url, photo_urls, status, review_state, reported_tree_status, activity_lng, activity_lat, activity_recorded_at
    """), {
        "status": status,
        "notes": notes,
        "photo_url": next_photo if photo_update_requested else photo_url,
        "photo_urls": _safe_json(next_photo_urls) if photo_update_requested else None,
        "reported_tree_status": normalized_tree_status,
        "activity_lng": activity_lng,
        "activity_lat": activity_lat,
        "activity_recorded_at": activity_recorded_at_value,
        "review_state": next_review_state,
        "submitted_at": next_submitted_at,
        "clear_review_fields": clear_review_fields,
        "completed_at": next_completed_at,
        "task_id": task_id,
    }).mappings().first()
    resolved_photo = row.get("photo_url")
    if resolved_photo:
        db.execute(text("""
            UPDATE trees
            SET photo_url = COALESCE(:photo_url, photo_url)
            WHERE id = :tree_id
        """), {"photo_url": resolved_photo, "tree_id": row["tree_id"]})

    project_id = int(existing["project_id"])
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="task",
        entity_id=task_id,
        action="task_updated",
        actor=actor_name,
        details={
            "before": {
                "status": existing.get("status"),
                "review_state": existing.get("review_state"),
                "notes": existing.get("notes"),
                "photo_url": existing.get("photo_url"),
                "photo_urls": existing_photo_urls,
                "tree_status": existing.get("tree_status"),
                "reported_tree_status": existing.get("reported_tree_status"),
                "activity_lng": existing.get("activity_lng"),
                "activity_lat": existing.get("activity_lat"),
                "activity_recorded_at": _to_iso_text(existing.get("activity_recorded_at")),
            },
            "after": {
                "status": row.get("status"),
                "review_state": row.get("review_state"),
                "notes": next_notes,
                "photo_url": next_photo,
                "photo_urls": next_photo_urls,
                "tree_status": existing.get("tree_status"),
                "reported_tree_status": row.get("reported_tree_status"),
                "activity_lng": row.get("activity_lng"),
                "activity_lat": row.get("activity_lat"),
                "activity_recorded_at": _to_iso_text(row.get("activity_recorded_at")),
            },
        },
    )
    if _normalize_name(row.get("review_state")) == "submitted":
        _record_alert(
            db,
            project_id=project_id,
            alert_type="task_submitted",
            severity="warning",
            message=f"Task #{task_id} is awaiting supervisor review.",
            tree_id=int(row["tree_id"]),
            task_id=task_id,
        )
    else:
        _resolve_task_alerts(db, task_id)
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {"status": "ok"}


@router.get("/tasks/review-queue")
def task_review_queue(
    project_id: int,
    assignee_name: str | None = None,
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.status, t.review_state,
                   t.priority, t.due_date, t.notes, t.photo_url, t.photo_urls, t.submitted_at, t.created_at,
                   t.reported_tree_status, t.review_notes, t.activity_lng, t.activity_lat, t.activity_recorded_at,
                   t.custodian_id, t.distribution_allocation_id, t.supervision_visit_no, t.supervision_total_visits,
                   tr.project_id, tr.status AS tree_status, tr.species AS tree_species,
                   ST_X(tr.geom) AS tree_lng, ST_Y(tr.geom) AS tree_lat,
                   c.name AS custodian_name, c.custodian_type, c.community_name AS custodian_community_name,
                   c.contact_person AS custodian_contact_person, c.phone AS custodian_phone, c.email AS custodian_email
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            LEFT JOIN green_custodians c ON c.id = COALESCE(t.custodian_id, tr.custodian_id)
            WHERE tr.project_id = :project_id
              AND LOWER(t.review_state) = 'submitted'
              AND (:assignee_name IS NULL OR t.assignee_name = :assignee_name)
            ORDER BY COALESCE(t.submitted_at, t.created_at) DESC, t.id DESC
        """),
        {"project_id": project_id, "assignee_name": assignee_name},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/tasks/{task_id}/submit")
def submit_task_for_review(
    task_id: int,
    db: Session = Depends(get_db),
    notes: str | None = Body(default=None),
    photo_url: str | None = Body(default=None),
    photo_urls: list[str] | None = Body(default=None),
    tree_status: str | None = Body(default=None),
    activity_lng: float | None = Body(default=None),
    activity_lat: float | None = Body(default=None),
    activity_recorded_at: str | None = Body(default=None),
    actor_name: str | None = Body(default=None),
):
    normalized_tree_status = _normalize_tree_status(tree_status) if tree_status is not None else None
    if normalized_tree_status is not None and normalized_tree_status not in TREE_STATUS_VALUES:
        raise HTTPException(status_code=400, detail="Invalid tree status")
    if (activity_lng is None) != (activity_lat is None):
        raise HTTPException(status_code=400, detail="Both activity_lng and activity_lat are required together")
    if activity_lng is not None and not (-180 <= float(activity_lng) <= 180):
        raise HTTPException(status_code=400, detail="Invalid activity_lng")
    if activity_lat is not None and not (-90 <= float(activity_lat) <= 90):
        raise HTTPException(status_code=400, detail="Invalid activity_lat")
    activity_recorded_at_value = _parse_datetime_value(activity_recorded_at)
    if activity_recorded_at is not None and activity_recorded_at_value is None:
        raise HTTPException(status_code=400, detail="Invalid activity_recorded_at")
    if activity_lng is not None and activity_lat is not None and activity_recorded_at_value is None:
        activity_recorded_at_value = datetime.utcnow()
    task = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.task_type, t.status, t.review_state, t.notes, t.photo_url, t.photo_urls,
                   t.reported_tree_status, t.activity_lng, t.activity_lat, t.activity_recorded_at,
                   tr.project_id
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE t.id = :task_id
        """),
        {"task_id": task_id},
    ).mappings().first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if _normalize_name(task.get("review_state")) == "approved":
        raise HTTPException(status_code=409, detail="Task already approved and locked")
    merged_notes = notes if notes is not None else task.get("notes")
    existing_photo, existing_photo_urls = _merge_photo_evidence(task.get("photo_url"), task.get("photo_urls"))
    merged_photo_urls = list(existing_photo_urls)
    if photo_urls is not None:
        merged_photo_urls = _normalize_photo_urls(photo_urls)
    if photo_url is not None:
        explicit_photo = str(photo_url or "").strip()
        if explicit_photo and explicit_photo not in merged_photo_urls:
            merged_photo_urls.append(explicit_photo)
    merged_photo = merged_photo_urls[-1] if merged_photo_urls else existing_photo
    evidence_ok, detail = _has_required_evidence(task.get("task_type"), merged_notes, merged_photo, merged_photo_urls)
    if not evidence_ok:
        raise HTTPException(status_code=400, detail=detail)

    row = db.execute(
        text("""
            UPDATE tree_tasks
            SET status = 'done',
                notes = COALESCE(:notes, notes),
                photo_url = COALESCE(:photo_url, photo_url),
                photo_urls = COALESCE(CAST(:photo_urls AS JSONB), photo_urls),
                reported_tree_status = COALESCE(:reported_tree_status, reported_tree_status),
                activity_lng = COALESCE(:activity_lng, activity_lng),
                activity_lat = COALESCE(:activity_lat, activity_lat),
                activity_recorded_at = COALESCE(:activity_recorded_at, activity_recorded_at),
                review_state = 'submitted',
                submitted_at = NOW(),
                reviewed_at = NULL,
                reviewed_by = NULL,
                review_notes = NULL,
                completed_at = COALESCE(completed_at, NOW())
            WHERE id = :task_id
            RETURNING id, tree_id, status, review_state, reported_tree_status, activity_lng, activity_lat, activity_recorded_at
        """),
        {
            "task_id": task_id,
            "notes": notes,
            "photo_url": merged_photo,
            "photo_urls": _safe_json(merged_photo_urls),
            "reported_tree_status": normalized_tree_status,
            "activity_lng": activity_lng,
            "activity_lat": activity_lat,
            "activity_recorded_at": activity_recorded_at_value,
        },
    ).mappings().first()
    if merged_photo:
        db.execute(
            text("""
                UPDATE trees
                SET photo_url = COALESCE(:photo_url, photo_url)
                WHERE id = :tree_id
            """),
            {"photo_url": merged_photo or None, "tree_id": int(row["tree_id"])},
        )

    project_id = int(task["project_id"])
    db.execute(
        text("""
            INSERT INTO green_task_reviews (task_id, decision, reviewer_name, review_notes)
            VALUES (:task_id, 'submitted', :reviewer_name, :review_notes)
        """),
        {"task_id": task_id, "reviewer_name": actor_name, "review_notes": merged_notes},
    )
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="task",
        entity_id=task_id,
        action="task_submitted_for_review",
        actor=actor_name,
        details={
            "status": row.get("status"),
            "review_state": row.get("review_state"),
            "activity_lng": row.get("activity_lng"),
            "activity_lat": row.get("activity_lat"),
            "activity_recorded_at": _to_iso_text(row.get("activity_recorded_at")),
        },
    )
    _record_alert(
        db,
        project_id=project_id,
        alert_type="task_submitted",
        severity="warning",
        message=f"Task #{task_id} is awaiting supervisor review.",
        tree_id=int(row["tree_id"]),
        task_id=task_id,
    )
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {"status": "submitted", "task_id": task_id}


@router.post("/tasks/{task_id}/review")
def review_submitted_task(
    task_id: int,
    db: Session = Depends(get_db),
    decision: str = Body(...),
    reviewer_name: str = Body(default=""),
    review_notes: str = Body(default=""),
    season_mode: str | None = Body(default=None),
):
    decision_key = _normalize_name(decision)
    if decision_key not in {"approve", "reject", "metadata_edit"}:
        raise HTTPException(status_code=400, detail="Decision must be approve, reject, or metadata_edit")
    if decision_key in {"reject", "metadata_edit"} and not (review_notes or "").strip():
        raise HTTPException(status_code=400, detail="Review note is required for reject or metadata edit request")

    task = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.status, t.review_state, t.model_season,
                   t.reported_tree_status,
                   tr.project_id
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE t.id = :task_id
        """),
        {"task_id": task_id},
    ).mappings().first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if _normalize_name(task.get("review_state")) != "submitted":
        raise HTTPException(status_code=409, detail="Task is not in submitted state")

    project_id = int(task["project_id"])
    auto_generated_task_id = None
    if decision_key == "approve":
        db.execute(
            text("""
                UPDATE tree_tasks
                SET review_state = 'approved',
                    reviewed_at = NOW(),
                    reviewed_by = :reviewer_name,
                    review_notes = :review_notes,
                    completed_at = COALESCE(completed_at, NOW())
                WHERE id = :task_id
            """),
            {
                "task_id": task_id,
                "reviewer_name": reviewer_name or None,
                "review_notes": review_notes or None,
            },
        )
        task_type = _normalize_name(task.get("task_type"))
        reported_tree_status = _normalize_tree_status(task.get("reported_tree_status"))
        if task_type == "planting":
            approved_status = (
                reported_tree_status
                if reported_tree_status in TREE_STATUS_VALUES and reported_tree_status != "pending_planting"
                else "alive"
            )
            db.execute(
                text("""
                    UPDATE trees
                    SET status = :status
                    WHERE id = :tree_id
                """),
                {"status": approved_status, "tree_id": int(task["tree_id"])},
            )
            _record_tree_status_history(
                db,
                tree_id=int(task["tree_id"]),
                project_id=project_id,
                status=approved_status,
                status_date=date.today(),
                source="task_review_approved",
                source_task_id=task_id,
                changed_by=reviewer_name or None,
                notes=review_notes or None,
            )
        elif reported_tree_status in TREE_STATUS_VALUES:
            db.execute(
                text("""
                    UPDATE trees
                    SET status = :status
                    WHERE id = :tree_id
                """),
                {"status": reported_tree_status, "tree_id": int(task["tree_id"])},
            )
            _record_tree_status_history(
                db,
                tree_id=int(task["tree_id"]),
                project_id=project_id,
                status=reported_tree_status,
                status_date=date.today(),
                source="task_review_approved",
                source_task_id=task_id,
                changed_by=reviewer_name or None,
                notes=review_notes or None,
            )
        # Auto-maintenance generation disabled: supervisors assign maintenance manually.
        auto_generated_task_id = None
        _resolve_task_alerts(db, task_id)
        action_name = "task_review_approved"
    elif decision_key == "reject":
        db.execute(
            text("""
                UPDATE tree_tasks
                SET status = 'pending',
                    review_state = 'rejected',
                    submitted_at = NULL,
                    completed_at = NULL,
                    reviewed_at = NOW(),
                    reviewed_by = :reviewer_name,
                    review_notes = :review_notes
                WHERE id = :task_id
            """),
            {
                "task_id": task_id,
                "reviewer_name": reviewer_name or None,
                "review_notes": review_notes or None,
            },
        )
        action_name = "task_review_rejected"
    else:
        db.execute(
            text("""
                UPDATE tree_tasks
                SET status = 'pending',
                    review_state = 'metadata_edit',
                    submitted_at = NULL,
                    completed_at = NULL,
                    reviewed_at = NOW(),
                    reviewed_by = :reviewer_name,
                    review_notes = :review_notes
                WHERE id = :task_id
            """),
            {
                "task_id": task_id,
                "reviewer_name": reviewer_name or None,
                "review_notes": review_notes or None,
            },
        )
        action_name = "task_review_metadata_edit_requested"

    db.execute(
        text("""
            INSERT INTO green_task_reviews (task_id, decision, reviewer_name, review_notes)
            VALUES (:task_id, :decision, :reviewer_name, :review_notes)
        """),
        {
            "task_id": task_id,
            "decision": (
                "approved"
                if decision_key == "approve"
                else "rejected"
                if decision_key == "reject"
                else "metadata_edit"
            ),
            "reviewer_name": reviewer_name or None,
            "review_notes": review_notes or None,
        },
    )
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="task",
        entity_id=task_id,
        action=action_name,
        actor=reviewer_name or None,
        details={
            "decision": decision_key,
            "auto_generated_task_id": auto_generated_task_id,
            "review_notes": review_notes or None,
        },
    )
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {
        "status": "ok",
        "decision": (
            "approved"
            if decision_key == "approve"
            else "rejected"
            if decision_key == "reject"
            else "metadata_edit"
        ),
        "auto_generated_task_id": auto_generated_task_id,
    }


@router.post("/tasks/{task_id}/reopen")
def reopen_approved_task(
    task_id: int,
    db: Session = Depends(get_db),
    reviewer_name: str = Body(default=""),
    reason: str = Body(default=""),
):
    task = db.execute(
        text("""
            SELECT t.id, t.tree_id, t.review_state, tr.project_id
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE t.id = :task_id
        """),
        {"task_id": task_id},
    ).mappings().first()
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    if _normalize_name(task.get("review_state")) != "approved":
        raise HTTPException(status_code=409, detail="Only approved tasks can be reopened")

    db.execute(
        text("""
            UPDATE tree_tasks
            SET status = 'pending',
                review_state = 'reopened',
                submitted_at = NULL,
                completed_at = NULL,
                reviewed_at = NOW(),
                reviewed_by = :reviewer_name,
                review_notes = :review_notes
            WHERE id = :task_id
        """),
        {"task_id": task_id, "reviewer_name": reviewer_name or None, "review_notes": reason or None},
    )
    db.execute(
        text("""
            INSERT INTO green_task_reviews (task_id, decision, reviewer_name, review_notes)
            VALUES (:task_id, 'reopened', :reviewer_name, :review_notes)
        """),
        {"task_id": task_id, "reviewer_name": reviewer_name or None, "review_notes": reason or None},
    )
    project_id = int(task["project_id"])
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="task",
        entity_id=task_id,
        action="task_reopened",
        actor=reviewer_name or None,
        details={"reason": reason or None},
    )
    _refresh_project_alerts(db, project_id)
    db.commit()
    return {"status": "ok"}


@router.get("/projects/{project_id}/alerts")
def project_alerts(
    project_id: int,
    refresh: bool = Query(default=True),
    status: str = Query(default="open"),
    db: Session = Depends(get_db),
):
    if refresh:
        _refresh_project_alerts(db, project_id)
        db.commit()
    rows = db.execute(
        text("""
            SELECT id, project_id, tree_id, task_id, alert_type, severity, message, status,
                   payload, created_at, resolved_at
            FROM green_alerts
            WHERE project_id = :project_id
              AND (:status = 'all' OR status = :status)
            ORDER BY created_at DESC, id DESC
            LIMIT 300
        """),
        {"project_id": project_id, "status": status},
    ).mappings().all()
    items = [dict(row) for row in rows]
    summary = {
        "total": len(items),
        "danger": sum(1 for item in items if _normalize_name(item.get("severity")) == "danger"),
        "warning": sum(1 for item in items if _normalize_name(item.get("severity")) == "warning"),
        "info": sum(1 for item in items if _normalize_name(item.get("severity")) == "info"),
    }
    return {"project_id": project_id, "status_filter": status, "summary": summary, "items": items}


def _compare_metric(metric_value: float, comparator: str, threshold: float) -> bool:
    op = _normalize_name(comparator)
    if op == "gt":
        return metric_value > threshold
    if op == "gte":
        return metric_value >= threshold
    if op == "lt":
        return metric_value < threshold
    if op == "lte":
        return metric_value <= threshold
    if op == "eq":
        return metric_value == threshold
    return False


@router.get("/reports/kpi")
def reports_kpi(
    project_id: int = Query(...),
    days: int = Query(default=30, ge=1, le=365),
    snapshot: bool = Query(default=True),
    db: Session = Depends(get_db),
):
    metrics = _compute_kpi_snapshot(project_id, db)
    if snapshot:
        _store_kpi_snapshot(project_id, metrics, db)
        db.commit()

    trend = _build_kpi_trend_series(project_id, db, days=days)
    species_daily_survival = _build_species_daily_survival_series(project_id, db)
    return {
        "project_id": project_id,
        "current": metrics,
        "trend_days": days,
        "trend_basis": {
            "survival": "Monthly cumulative survival across planting cohorts using current tree statuses (starts from first planting_date).",
            "species_survival_daily": "Daily species survival from planting date using status history and live tree status.",
        },
        "trend": trend,
        "species_daily_survival": species_daily_survival,
    }


@router.get("/reports/schedule")
def list_report_schedules(
    project_id: int = Query(...),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text("""
            SELECT id, project_id, report_type, report_format, recipients, cron_expr, timezone, webhook_url,
                   is_enabled, created_by, last_run_at, next_run_at, created_at, updated_at
            FROM green_scheduled_reports
            WHERE project_id = :project_id
            ORDER BY created_at DESC, id DESC
        """),
        {"project_id": project_id},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/reports/schedule")
def create_report_schedule(
    project_id: int = Body(...),
    report_type: str = Body(default="donor"),
    report_format: str = Body(default="pdf"),
    recipients: str = Body(default=""),
    cron_expr: str | None = Body(default=None),
    timezone: str = Body(default="Africa/Lagos"),
    webhook_url: str | None = Body(default=None),
    created_by: str | None = Body(default=None),
    db: Session = Depends(get_db),
):
    row = db.execute(
        text("""
            INSERT INTO green_scheduled_reports (
                project_id, report_type, report_format, recipients, cron_expr, timezone, webhook_url, created_by
            )
            VALUES (
                :project_id, :report_type, :report_format, :recipients, :cron_expr, :timezone, :webhook_url, :created_by
            )
            RETURNING id, project_id, report_type, report_format, recipients, cron_expr, timezone, webhook_url,
                      is_enabled, created_by, last_run_at, next_run_at, created_at, updated_at
        """),
        {
            "project_id": project_id,
            "report_type": (_normalize_name(report_type) or "donor"),
            "report_format": (_normalize_name(report_format) or "pdf"),
            "recipients": recipients or "",
            "cron_expr": cron_expr,
            "timezone": timezone or "Africa/Lagos",
            "webhook_url": webhook_url,
            "created_by": created_by,
        },
    ).mappings().first()
    db.commit()
    return dict(row)


@router.patch("/reports/schedule/{schedule_id}")
def update_report_schedule(
    schedule_id: int,
    report_type: str | None = Body(default=None),
    report_format: str | None = Body(default=None),
    recipients: str | None = Body(default=None),
    cron_expr: str | None = Body(default=None),
    timezone: str | None = Body(default=None),
    webhook_url: str | None = Body(default=None),
    is_enabled: bool | None = Body(default=None),
    db: Session = Depends(get_db),
):
    row = db.execute(
        text("""
            UPDATE green_scheduled_reports
            SET report_type = COALESCE(:report_type, report_type),
                report_format = COALESCE(:report_format, report_format),
                recipients = COALESCE(:recipients, recipients),
                cron_expr = COALESCE(:cron_expr, cron_expr),
                timezone = COALESCE(:timezone, timezone),
                webhook_url = COALESCE(:webhook_url, webhook_url),
                is_enabled = COALESCE(:is_enabled, is_enabled),
                updated_at = NOW()
            WHERE id = :schedule_id
            RETURNING id, project_id, report_type, report_format, recipients, cron_expr, timezone, webhook_url,
                      is_enabled, created_by, last_run_at, next_run_at, created_at, updated_at
        """),
        {
            "schedule_id": schedule_id,
            "report_type": _normalize_name(report_type) if report_type is not None else None,
            "report_format": _normalize_name(report_format) if report_format is not None else None,
            "recipients": recipients,
            "cron_expr": cron_expr,
            "timezone": timezone,
            "webhook_url": webhook_url,
            "is_enabled": is_enabled,
        },
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Schedule not found")
    db.commit()
    return dict(row)


@router.delete("/reports/schedule/{schedule_id}")
def delete_report_schedule(schedule_id: int, db: Session = Depends(get_db)):
    deleted = db.execute(
        text("DELETE FROM green_scheduled_reports WHERE id = :schedule_id RETURNING id"),
        {"schedule_id": schedule_id},
    ).scalar()
    if not deleted:
        raise HTTPException(status_code=404, detail="Schedule not found")
    db.commit()
    return {"status": "ok", "id": int(deleted)}


@router.get("/alerts/rules")
def list_alert_rules(
    project_id: int = Query(...),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text("""
            SELECT id, project_id, rule_name, metric_key, comparator, threshold, severity,
                   message_template, is_enabled, created_by, created_at, updated_at
            FROM green_alert_rules
            WHERE project_id = :project_id
            ORDER BY created_at DESC, id DESC
        """),
        {"project_id": project_id},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.post("/alerts/rules")
def create_alert_rule(
    project_id: int = Body(...),
    rule_name: str = Body(...),
    metric_key: str = Body(...),
    comparator: str = Body(default="gte"),
    threshold: float = Body(...),
    severity: str = Body(default="warning"),
    message_template: str | None = Body(default=None),
    created_by: str | None = Body(default=None),
    db: Session = Depends(get_db),
):
    cmp_key = _normalize_name(comparator)
    if cmp_key not in {"gt", "gte", "lt", "lte", "eq"}:
        raise HTTPException(status_code=400, detail="Invalid comparator")
    sev = _normalize_name(severity) or "warning"
    if sev not in {"info", "warning", "danger"}:
        sev = "warning"
    row = db.execute(
        text("""
            INSERT INTO green_alert_rules (
                project_id, rule_name, metric_key, comparator, threshold, severity, message_template, created_by
            )
            VALUES (
                :project_id, :rule_name, :metric_key, :comparator, :threshold, :severity, :message_template, :created_by
            )
            RETURNING id, project_id, rule_name, metric_key, comparator, threshold, severity,
                      message_template, is_enabled, created_by, created_at, updated_at
        """),
        {
            "project_id": project_id,
            "rule_name": rule_name.strip(),
            "metric_key": _normalize_name(metric_key),
            "comparator": cmp_key,
            "threshold": threshold,
            "severity": sev,
            "message_template": message_template,
            "created_by": created_by,
        },
    ).mappings().first()
    db.commit()
    return dict(row)


@router.patch("/alerts/rules/{rule_id}")
def update_alert_rule(
    rule_id: int,
    rule_name: str | None = Body(default=None),
    metric_key: str | None = Body(default=None),
    comparator: str | None = Body(default=None),
    threshold: float | None = Body(default=None),
    severity: str | None = Body(default=None),
    message_template: str | None = Body(default=None),
    is_enabled: bool | None = Body(default=None),
    db: Session = Depends(get_db),
):
    cmp_key = _normalize_name(comparator) if comparator is not None else None
    if cmp_key is not None and cmp_key not in {"gt", "gte", "lt", "lte", "eq"}:
        raise HTTPException(status_code=400, detail="Invalid comparator")
    sev = _normalize_name(severity) if severity is not None else None
    if sev is not None and sev not in {"info", "warning", "danger"}:
        raise HTTPException(status_code=400, detail="Invalid severity")
    row = db.execute(
        text("""
            UPDATE green_alert_rules
            SET rule_name = COALESCE(:rule_name, rule_name),
                metric_key = COALESCE(:metric_key, metric_key),
                comparator = COALESCE(:comparator, comparator),
                threshold = COALESCE(:threshold, threshold),
                severity = COALESCE(:severity, severity),
                message_template = COALESCE(:message_template, message_template),
                is_enabled = COALESCE(:is_enabled, is_enabled),
                updated_at = NOW()
            WHERE id = :rule_id
            RETURNING id, project_id, rule_name, metric_key, comparator, threshold, severity,
                      message_template, is_enabled, created_by, created_at, updated_at
        """),
        {
            "rule_id": rule_id,
            "rule_name": rule_name,
            "metric_key": _normalize_name(metric_key) if metric_key is not None else None,
            "comparator": cmp_key,
            "threshold": threshold,
            "severity": sev,
            "message_template": message_template,
            "is_enabled": is_enabled,
        },
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Rule not found")
    db.commit()
    return dict(row)


@router.post("/alerts/evaluate")
def evaluate_alert_rules(
    project_id: int = Body(...),
    db: Session = Depends(get_db),
):
    metrics = _compute_kpi_snapshot(project_id, db)
    rules = db.execute(
        text("""
            SELECT id, rule_name, metric_key, comparator, threshold, severity, message_template
            FROM green_alert_rules
            WHERE project_id = :project_id
              AND is_enabled = TRUE
            ORDER BY id ASC
        """),
        {"project_id": project_id},
    ).mappings().all()
    webhook_targets = [
        str(row.get("webhook_url") or "").strip()
        for row in db.execute(
            text(
                """
                SELECT DISTINCT webhook_url
                FROM green_scheduled_reports
                WHERE project_id = :project_id
                  AND is_enabled = TRUE
                  AND webhook_url IS NOT NULL
                  AND TRIM(webhook_url) <> ''
                """
            ),
            {"project_id": project_id},
        ).mappings().all()
    ]
    created_events: list[dict] = []
    created_deliveries = 0
    for rule in rules:
        metric_key = _normalize_name(rule.get("metric_key"))
        raw_value = metrics.get(metric_key)
        if raw_value is None:
            continue
        metric_value = float(raw_value)
        threshold = float(rule.get("threshold") or 0)
        if not _compare_metric(metric_value, rule.get("comparator"), threshold):
            continue
        msg = (rule.get("message_template") or "").strip()
        if not msg:
            msg = f"{rule.get('rule_name')}: {metric_key}={metric_value} vs threshold {threshold} ({rule.get('comparator')})."
        event = db.execute(
            text("""
                INSERT INTO green_alert_events (
                    project_id, rule_id, severity, metric_key, metric_value, threshold, message, payload
                )
                VALUES (
                    :project_id, :rule_id, :severity, :metric_key, :metric_value, :threshold, :message, CAST(:payload AS JSONB)
                )
                RETURNING id, project_id, rule_id, severity, status, metric_key, metric_value, threshold,
                          message, payload, triggered_at, resolved_at
            """),
            {
                "project_id": project_id,
                "rule_id": int(rule["id"]),
                "severity": _normalize_name(rule.get("severity") or "warning"),
                "metric_key": metric_key,
                "metric_value": metric_value,
                "threshold": threshold,
                "message": msg,
                "payload": _safe_json({"metrics": metrics}),
            },
        ).mappings().first()
        event_payload = dict(event)
        created_events.append(event_payload)
        event_id = int(event_payload.get("id") or 0)
        if event_id > 0 and webhook_targets:
            for target_url in webhook_targets:
                db.execute(
                    text(
                        """
                        INSERT INTO green_webhook_deliveries (event_id, target_url, status, attempt_count)
                        VALUES (:event_id, :target_url, 'pending', 0)
                        """
                    ),
                    {"event_id": event_id, "target_url": target_url},
                )
                created_deliveries += 1
    db.commit()
    return {
        "project_id": project_id,
        "created": len(created_events),
        "events": created_events,
        "webhook_deliveries_created": created_deliveries,
    }


@router.get("/alerts/events")
def list_alert_events(
    project_id: int = Query(...),
    status: str = Query(default="all"),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text("""
            SELECT id, project_id, rule_id, severity, status, metric_key, metric_value, threshold,
                   message, payload, triggered_at, resolved_at
            FROM green_alert_events
            WHERE project_id = :project_id
              AND (:status = 'all' OR status = :status)
            ORDER BY triggered_at DESC, id DESC
            LIMIT 500
        """),
        {"project_id": project_id, "status": status},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.get("/alerts/webhook-deliveries")
def list_webhook_deliveries(
    project_id: int = Query(...),
    status: str = Query(default="all"),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text(
            """
            SELECT d.id, d.event_id, d.target_url, d.status, d.response_code, d.response_body,
                   d.attempt_count, d.delivered_at, d.created_at,
                   e.rule_id, e.severity, e.metric_key, e.metric_value, e.threshold, e.message, e.triggered_at
            FROM green_webhook_deliveries d
            JOIN green_alert_events e ON e.id = d.event_id
            WHERE e.project_id = :project_id
              AND (:status = 'all' OR d.status = :status)
            ORDER BY d.created_at DESC, d.id DESC
            LIMIT 500
            """
        ),
        {"project_id": project_id, "status": status},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.patch("/alerts/webhook-deliveries/{delivery_id}")
def update_webhook_delivery(
    delivery_id: int,
    status: str | None = Body(default=None),
    response_code: int | None = Body(default=None),
    response_body: str | None = Body(default=None),
    increment_attempt: bool = Body(default=False),
    db: Session = Depends(get_db),
):
    status_key = _normalize_name(status) if status is not None else None
    if status_key is not None and status_key not in {"pending", "failed", "delivered"}:
        raise HTTPException(status_code=400, detail="Invalid delivery status")
    row = db.execute(
        text(
            """
            UPDATE green_webhook_deliveries
            SET status = COALESCE(:status, status),
                response_code = COALESCE(:response_code, response_code),
                response_body = COALESCE(:response_body, response_body),
                attempt_count = CASE WHEN :increment_attempt THEN attempt_count + 1 ELSE attempt_count END,
                delivered_at = CASE
                    WHEN COALESCE(:status, status) = 'delivered' THEN NOW()
                    ELSE delivered_at
                END
            WHERE id = :delivery_id
            RETURNING id, event_id, target_url, status, response_code, response_body,
                      attempt_count, delivered_at, created_at
            """
        ),
        {
            "delivery_id": delivery_id,
            "status": status_key,
            "response_code": response_code,
            "response_body": response_body,
            "increment_attempt": bool(increment_attempt),
        },
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Webhook delivery not found")
    db.commit()
    return dict(row)


@router.get("/projects/{project_id}/audit-events")
def project_audit_events(
    project_id: int,
    limit: int = Query(default=200, ge=1, le=1000),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text("""
            SELECT id, project_id, entity_type, entity_id, action, actor, details, created_at
            FROM green_audit_events
            WHERE project_id = :project_id
            ORDER BY created_at DESC, id DESC
            LIMIT :limit
        """),
        {"project_id": project_id, "limit": limit},
    ).mappings().all()
    return [dict(row) for row in rows]


@router.get("/projects/{project_id}/live-maintenance")
def live_maintenance_rows(
    project_id: int,
    season_mode: str = Query(default="rainy"),
    assignee_name: str | None = Query(default=None),
    tree_scope: str = Query(default="new_planting"),
    db: Session = Depends(get_db),
):
    payload = _compute_live_maintenance_rows(
        db=db,
        project_id=project_id,
        season_mode=season_mode,
        assignee_name=assignee_name,
        tree_scope=tree_scope,
    )
    return {
        "project_id": project_id,
        "season_mode": "dry" if _normalize_name(season_mode) == "dry" else "rainy",
        "tree_scope": _normalize_name(tree_scope or "new_planting"),
        "computed_at": datetime.utcnow().isoformat(),
        "summary": payload["summary"],
        "rows": payload["rows"],
        "sources": LIVE_SOURCE_REFERENCES,
    }


def _build_donor_report_rows(project_id: int, db: Session) -> list[dict]:
    rows = db.execute(
        text("""
            SELECT
                   t.id AS task_id,
                   t.tree_id,
                   tr.species,
                   tr.tree_origin,
                   tr.attribution_scope,
                   tr.tree_height_m,
                   tr.count_in_planting_kpis,
                   tr.count_in_carbon_scope,
                   tr.custodian_id,
                   c.name AS custodian_name,
                   t.assignee_name,
                   t.task_type,
                   t.priority,
                   t.status, t.review_state, t.due_date, t.created_at, t.submitted_at, t.reviewed_at,
                   t.reviewed_by, t.review_notes, t.completed_at, t.notes, t.photo_url,
                   t.reported_tree_status, tr.status AS tree_status
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            LEFT JOIN green_custodians c ON c.id = tr.custodian_id
            WHERE tr.project_id = :project_id
            ORDER BY COALESCE(t.reviewed_at, t.submitted_at, t.created_at) DESC, t.id DESC
        """),
        {"project_id": project_id},
    ).mappings().all()
    report_rows: list[dict] = []
    today = date.today()
    for row in rows:
        due_date = _parse_date_value(row.get("due_date"))
        completed_date = _parse_date_value(row.get("completed_at"))
        delay_context = None
        if due_date and completed_date:
            delay_days = _day_diff(completed_date, due_date)
            delay_context = "completion"
        elif due_date and not completed_date:
            delay_days = _day_diff(today, due_date)
            delay_context = "schedule"
        else:
            delay_days = None
        evidence_ok, _ = _has_required_evidence(row.get("task_type"), row.get("notes"), row.get("photo_url"))
        item = dict(row)
        item["evidence_status"] = "complete" if evidence_ok else "missing"
        item["delay_days"] = delay_days
        item["delay_context"] = delay_context
        report_rows.append(item)
    return report_rows


def _review_summary_by_tree(project_id: int, db: Session, assignee_name: str | None = None) -> dict[int, dict]:
    rows = db.execute(
        text("""
            SELECT tr.id AS tree_id,
                   SUM(CASE WHEN LOWER(COALESCE(t.review_state, 'none')) = 'submitted' THEN 1 ELSE 0 END) AS submitted_count,
                   SUM(CASE WHEN LOWER(COALESCE(t.review_state, 'none')) = 'approved' THEN 1 ELSE 0 END) AS approved_count,
                   SUM(CASE WHEN LOWER(COALESCE(t.review_state, 'none')) = 'rejected' THEN 1 ELSE 0 END) AS rejected_count,
                   MAX(t.submitted_at) AS last_submitted_at,
                   MAX(t.reviewed_at) AS last_reviewed_at,
                   (ARRAY_AGG(t.review_notes ORDER BY COALESCE(t.reviewed_at, t.submitted_at, t.created_at) DESC NULLS LAST, t.id DESC))[1]
                       AS last_review_note,
                   (ARRAY_AGG(t.review_state ORDER BY COALESCE(t.reviewed_at, t.submitted_at, t.created_at) DESC NULLS LAST, t.id DESC))[1]
                       AS last_review_state
            FROM trees tr
            LEFT JOIN tree_tasks t ON t.tree_id = tr.id
            WHERE tr.project_id = :project_id
              AND (:assignee_name IS NULL OR tr.created_by = :assignee_name)
            GROUP BY tr.id
        """),
        {"project_id": project_id, "assignee_name": assignee_name},
    ).mappings().all()
    result: dict[int, dict] = {}
    for row in rows:
        tree_id = int(row.get("tree_id"))
        result[tree_id] = {
            "review_submitted": int(row.get("submitted_count") or 0),
            "review_approved": int(row.get("approved_count") or 0),
            "review_rejected": int(row.get("rejected_count") or 0),
            "last_submitted_at": row.get("last_submitted_at"),
            "last_reviewed_at": row.get("last_reviewed_at"),
            "last_review_note": row.get("last_review_note") or "",
            "last_review_state": row.get("last_review_state") or "",
        }
    return result


def _compute_age_based_survival(
    project_id: int,
    db: Session,
    checkpoints_days: tuple[int, ...] = AGE_SURVIVAL_CHECKPOINTS_DAYS,
    as_of_date: date | None = None,
) -> dict:
    as_of = as_of_date or date.today()

    tree_rows = db.execute(
        text(
            """
            SELECT id, planting_date, status, species
            FROM trees
            WHERE project_id = :project_id
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    history_rows = db.execute(
        text(
            """
            SELECT tree_id, status, status_date, created_at, id
            FROM green_tree_status_history
            WHERE project_id = :project_id
            ORDER BY tree_id ASC, status_date ASC, created_at ASC, id ASC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    history_by_tree: dict[int, list[tuple[date, str]]] = {}
    for row in history_rows:
        tree_id = int(row.get("tree_id") or 0)
        status_date = _parse_date_value(row.get("status_date"))
        if tree_id <= 0 or status_date is None:
            continue
        status_value = _normalize_tree_status(row.get("status"))
        if status_value not in TREE_STATUS_VALUES:
            continue
        history_by_tree.setdefault(tree_id, []).append((status_date, status_value))

    checkpoint_metrics: dict[int, dict] = {
        int(day): {
            "eligible_trees": 0,
            "survived_trees": 0,
            "missing_status_trees": 0,
        }
        for day in checkpoints_days
    }
    species_metrics: dict[str, dict] = {}
    missing_planting_date_trees = 0

    for row in tree_rows:
        tree_id = int(row.get("id") or 0)
        if tree_id <= 0:
            continue
        species_label_raw = str(row.get("species") or "").strip()
        species_key = _normalize_name(species_label_raw) or "__unknown__"
        species_label = species_label_raw or "Unknown Species"
        if species_key not in species_metrics:
            species_metrics[species_key] = {
                "species_key": species_key,
                "species_label": species_label,
                "trees_with_planting_date": 0,
                "current_total_trees": 0,
                "current_healthy_trees": 0,
                "max_tree_age_days": 0,
                "checkpoints": {
                    int(day): {
                        "eligible_trees": 0,
                        "survived_trees": 0,
                        "missing_status_trees": 0,
                    }
                    for day in checkpoints_days
                },
            }
        planting_ref = _parse_date_value(row.get("planting_date"))
        if planting_ref is None:
            missing_planting_date_trees += 1
            continue
        species_metrics[species_key]["trees_with_planting_date"] += 1
        history = history_by_tree.get(tree_id) or []
        fallback_status = _normalize_tree_status(row.get("status"))
        species_metrics[species_key]["current_total_trees"] += 1
        if fallback_status in HEALTHY_TREE_STATUSES:
            species_metrics[species_key]["current_healthy_trees"] += 1
        tree_age_days = max((as_of - planting_ref).days, 0)
        species_metrics[species_key]["max_tree_age_days"] = max(
            int(species_metrics[species_key].get("max_tree_age_days") or 0),
            int(tree_age_days),
        )

        for day in checkpoints_days:
            checkpoint = int(day)
            target_date = planting_ref + timedelta(days=checkpoint)
            if target_date > as_of:
                continue

            metric = checkpoint_metrics[checkpoint]
            species_metric = species_metrics[species_key]["checkpoints"][checkpoint]
            metric["eligible_trees"] += 1
            species_metric["eligible_trees"] += 1

            status_at_target = None
            for status_date, status_value in history:
                if status_date <= target_date:
                    status_at_target = status_value
                else:
                    break
            if status_at_target is None:
                metric["missing_status_trees"] += 1
                species_metric["missing_status_trees"] += 1
                status_at_target = fallback_status
            if status_at_target in HEALTHY_TREE_STATUSES:
                metric["survived_trees"] += 1
                species_metric["survived_trees"] += 1

    result = {
        "as_of_date": as_of.isoformat(),
        "checkpoints_days": [int(day) for day in checkpoints_days],
        "trees_missing_planting_date": int(missing_planting_date_trees),
    }
    for day in checkpoints_days:
        metric = checkpoint_metrics[int(day)]
        eligible = int(metric.get("eligible_trees") or 0)
        survived = int(metric.get("survived_trees") or 0)
        missing = int(metric.get("missing_status_trees") or 0)
        rate = round((survived / eligible) * 100, 1) if eligible > 0 else 0.0
        result[f"day_{int(day)}"] = {
            "eligible_trees": eligible,
            "survived_trees": survived,
            "survival_rate": rate,
            "missing_status_trees": missing,
        }
    species_rows: list[dict] = []
    for _, item in species_metrics.items():
        checkpoints = item.get("checkpoints") or {}
        row_payload = {
            "species_key": item.get("species_key"),
            "species_label": item.get("species_label"),
            "trees_with_planting_date": int(item.get("trees_with_planting_date") or 0),
            "current_total_trees": int(item.get("current_total_trees") or 0),
            "current_healthy_trees": int(item.get("current_healthy_trees") or 0),
            "max_tree_age_days": int(item.get("max_tree_age_days") or 0),
        }
        current_total = int(item.get("current_total_trees") or 0)
        current_healthy = int(item.get("current_healthy_trees") or 0)
        row_payload["current_survival_rate"] = round((current_healthy / current_total) * 100, 1) if current_total > 0 else 0.0
        for day in checkpoints_days:
            bucket = checkpoints.get(int(day)) or {}
            eligible = int(bucket.get("eligible_trees") or 0)
            survived = int(bucket.get("survived_trees") or 0)
            missing = int(bucket.get("missing_status_trees") or 0)
            rate = round((survived / eligible) * 100, 1) if eligible > 0 else 0.0
            row_payload[f"day_{int(day)}"] = {
                "eligible_trees": eligible,
                "survived_trees": survived,
                "survival_rate": rate,
                "missing_status_trees": missing,
            }
        species_rows.append(row_payload)
    species_rows.sort(
        key=lambda row: (
            -int(row.get("trees_with_planting_date") or 0),
            str(row.get("species_label") or "").lower(),
        )
    )
    result["species_breakdown"] = species_rows
    return result


def _compute_kpi_snapshot(project_id: int, db: Session) -> dict:
    tree_rows = db.execute(
        text(
            """
            SELECT
                status,
                tree_origin,
                COALESCE(count_in_planting_kpis, TRUE) AS in_planting_scope,
                COALESCE(count_in_carbon_scope, TRUE) AS in_carbon_scope
            FROM trees
            WHERE project_id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().all()
    task_rows = db.execute(
        text("""
            SELECT t.status, t.review_state, t.due_date, t.notes, t.photo_url, t.task_type
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE tr.project_id = :project_id
        """),
        {"project_id": project_id},
    ).mappings().all()

    total_trees_all = len(tree_rows)
    in_scope_tree_rows = [row for row in tree_rows if bool(row.get("in_planting_scope"))]
    total_trees = len(in_scope_tree_rows)
    healthy_trees = sum(1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in HEALTHY_TREE_STATUSES)
    dead_trees = sum(1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in DEAD_TREE_STATUSES)
    attention_trees = sum(1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in ATTENTION_TREE_STATUSES)
    pending_planting = sum(1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) == "pending_planting")
    survival_rate = round((healthy_trees / total_trees) * 100, 1) if total_trees else 0.0

    today = date.today()
    total_tasks = len(task_rows)
    submitted_tasks = 0
    approved_tasks = 0
    rejected_tasks = 0
    open_tasks = 0
    overdue_tasks = 0
    evidence_required = 0
    evidence_complete = 0
    for task in task_rows:
        state = _normalize_name(task.get("review_state") or "none")
        status = _normalize_name(task.get("status") or "pending")
        due = _parse_date_value(task.get("due_date"))
        if state == "submitted":
            submitted_tasks += 1
        if state == "approved":
            approved_tasks += 1
        if state == "rejected":
            rejected_tasks += 1
        if not (_is_done_status(status) and state in {"approved", "none"}):
            open_tasks += 1
        if due and due < today and not (_is_done_status(status) and state in {"approved", "none"}):
            overdue_tasks += 1
        policy = _task_needs_evidence(task.get("task_type"))
        evidence_ok, _ = _has_required_evidence(task.get("task_type"), task.get("notes"), task.get("photo_url"))
        evidence_in_scope = _is_done_status(status) or state in {"submitted", "approved", "rejected"}
        if (policy.get("require_notes") or policy.get("require_photo")) and evidence_in_scope:
            evidence_required += 1
            if evidence_ok:
                evidence_complete += 1

    if evidence_required > 0:
        evidence_rate = round((evidence_complete / evidence_required) * 100, 1)
    elif total_tasks > 0:
        evidence_rate = 100.0
    else:
        evidence_rate = 0.0

    # Carbon data for KPI
    carbon_tree_rows = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :project_id
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"project_id": project_id}).mappings().all()
    carbon = compute_project_carbon([dict(r) for r in carbon_tree_rows])
    age_survival = _compute_age_based_survival(project_id, db)

    return {
        "project_id": project_id,
        "snapshot_date": datetime.utcnow().isoformat(),
        "trees_total": total_trees,
        "trees_total_all": total_trees_all,
        "trees_healthy": healthy_trees,
        "trees_dead_or_removed": dead_trees,
        "trees_attention": attention_trees,
        "trees_pending_planting": pending_planting,
        "survival_rate": survival_rate,
        "tasks_total": total_tasks,
        "tasks_open": open_tasks,
        "tasks_submitted": submitted_tasks,
        "tasks_approved": approved_tasks,
        "tasks_rejected": rejected_tasks,
        "tasks_overdue": overdue_tasks,
        "evidence_complete_rate": evidence_rate,
        "evidence_required_tasks": evidence_required,
        "evidence_complete_tasks": evidence_complete,
        "co2_current_tonnes": carbon["current_co2_tonnes"],
        "co2_annual_tonnes": carbon["annual_co2_tonnes"],
        "co2_projected_lifetime_tonnes": carbon["projected_lifetime_co2_tonnes"],
        "age_survival": age_survival,
        "age_survival_30d": age_survival.get("day_30", {}),
        "age_survival_90d": age_survival.get("day_90", {}),
        "age_survival_180d": age_survival.get("day_180", {}),
    }


def _month_start(value: date) -> date:
    return date(value.year, value.month, 1)


def _next_month_start(value: date) -> date:
    if value.month == 12:
        return date(value.year + 1, 1, 1)
    return date(value.year, value.month + 1, 1)


def _survival_phase_label(age_days: int) -> str:
    age = max(int(age_days or 0), 0)
    if age >= 180:
        return "past 180 days"
    if age >= 90:
        return "past 90 days"
    if age >= 30:
        return "past 30 days"
    return "0-29 days"


def _build_species_daily_survival_series(project_id: int, db: Session) -> dict:
    """
    Build per-species daily survival lines from planting date to today.
    Survival uses status history timeline (maintenance/task-review/manual updates)
    with current tree status as fallback baseline.
    """
    today = date.today()

    tree_rows = db.execute(
        text(
            """
            SELECT id, species, planting_date, status
            FROM trees
            WHERE project_id = :project_id
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    history_rows = db.execute(
        text(
            """
            SELECT tree_id, status, status_date, created_at, id
            FROM green_tree_status_history
            WHERE project_id = :project_id
            ORDER BY tree_id ASC, status_date ASC, created_at ASC, id ASC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    history_by_tree: dict[int, list[tuple[date, str]]] = {}
    for row in history_rows:
        tree_id = int(row.get("tree_id") or 0)
        status_date = _parse_date_value(row.get("status_date"))
        if tree_id <= 0 or status_date is None:
            continue
        status_value = _normalize_tree_status(row.get("status"))
        if status_value not in TREE_STATUS_VALUES:
            continue
        history_by_tree.setdefault(tree_id, []).append((status_date, status_value))

    total_deltas_by_day: dict[date, dict[str, int]] = {}
    healthy_deltas_by_day: dict[date, dict[str, int]] = {}

    def _apply_delta(store: dict[date, dict[str, int]], when: date, species_key: str, delta: int) -> None:
        day_bucket = store.setdefault(when, {})
        day_bucket[species_key] = int(day_bucket.get(species_key) or 0) + int(delta)

    species_labels: dict[str, str] = {}
    species_tree_counts: dict[str, int] = {}
    species_first_planting: dict[str, date] = {}
    trees_missing_planting_date = 0

    for row in tree_rows:
        tree_id = int(row.get("id") or 0)
        if tree_id <= 0:
            continue

        planting_ref = _parse_date_value(row.get("planting_date"))
        if planting_ref is None:
            trees_missing_planting_date += 1
            continue

        species_label_raw = str(row.get("species") or "").strip()
        species_key = _normalize_name(species_label_raw) or "__unknown__"
        species_label = species_label_raw or "Unknown Species"
        species_labels[species_key] = species_labels.get(species_key) or species_label
        species_tree_counts[species_key] = int(species_tree_counts.get(species_key) or 0) + 1

        first_for_species = species_first_planting.get(species_key)
        if first_for_species is None or planting_ref < first_for_species:
            species_first_planting[species_key] = planting_ref

        fallback_status = _normalize_tree_status(row.get("status"))
        if fallback_status not in TREE_STATUS_VALUES:
            fallback_status = "alive"

        timeline_raw = history_by_tree.get(tree_id) or []
        timeline_by_day: dict[date, str] = {}
        for event_date, event_status in timeline_raw:
            # Keep last status event of the day.
            timeline_by_day[event_date] = event_status
        timeline = sorted(timeline_by_day.items(), key=lambda item: item[0])

        baseline_status = fallback_status
        for event_date, event_status in timeline:
            if event_date <= planting_ref:
                baseline_status = event_status
            else:
                break

        _apply_delta(total_deltas_by_day, planting_ref, species_key, 1)
        if baseline_status in HEALTHY_TREE_STATUSES:
            _apply_delta(healthy_deltas_by_day, planting_ref, species_key, 1)

        prev_status = baseline_status
        for event_date, event_status in timeline:
            if event_date <= planting_ref:
                continue
            if event_status == prev_status:
                continue
            was_healthy = prev_status in HEALTHY_TREE_STATUSES
            is_healthy = event_status in HEALTHY_TREE_STATUSES
            if was_healthy != is_healthy:
                _apply_delta(healthy_deltas_by_day, event_date, species_key, 1 if is_healthy else -1)
            prev_status = event_status

    if not species_first_planting:
        return {
            "as_of_date": today.isoformat(),
            "start_date": None,
            "species_count": 0,
            "trees_missing_planting_date": int(trees_missing_planting_date),
            "day_markers": {"day_30": 30, "day_90": 90, "day_180": 180},
            "species": [],
        }

    project_start_date = min(species_first_planting.values())
    species_keys = sorted(
        species_first_planting.keys(),
        key=lambda key: (
            -int(species_tree_counts.get(key) or 0),
            str(species_labels.get(key) or key).lower(),
        ),
    )

    running_total: dict[str, int] = {key: 0 for key in species_keys}
    running_healthy: dict[str, int] = {key: 0 for key in species_keys}
    points_by_species: dict[str, list[dict]] = {key: [] for key in species_keys}

    cursor = project_start_date
    while cursor <= today:
        total_bucket = total_deltas_by_day.get(cursor) or {}
        for species_key, delta in total_bucket.items():
            running_total[species_key] = max(int(running_total.get(species_key) or 0) + int(delta), 0)

        healthy_bucket = healthy_deltas_by_day.get(cursor) or {}
        for species_key, delta in healthy_bucket.items():
            running_healthy[species_key] = int(running_healthy.get(species_key) or 0) + int(delta)

        for species_key in species_keys:
            species_start = species_first_planting.get(species_key)
            if species_start is None or cursor < species_start:
                continue

            eligible = int(running_total.get(species_key) or 0)
            if eligible <= 0:
                continue

            survived = int(running_healthy.get(species_key) or 0)
            survived = min(max(survived, 0), eligible)
            day_since_species_start = (cursor - species_start).days
            day_since_project_start = (cursor - project_start_date).days
            survival_rate = round((survived / eligible) * 100, 1) if eligible > 0 else 0.0

            points_by_species[species_key].append(
                {
                    "date": cursor.isoformat(),
                    "day_since_species_start": int(day_since_species_start),
                    "day_since_project_start": int(day_since_project_start),
                    "survival_rate": survival_rate,
                    "eligible_trees": eligible,
                    "survived_trees": survived,
                    "phase": _survival_phase_label(day_since_species_start),
                }
            )
        cursor += timedelta(days=1)

    species_rows: list[dict] = []
    for species_key in species_keys:
        points = points_by_species.get(species_key) or []
        if not points:
            continue
        species_rows.append(
            {
                "species_key": species_key,
                "species_label": species_labels.get(species_key) or "Unknown Species",
                "trees_with_planting_date": int(species_tree_counts.get(species_key) or 0),
                "start_date": species_first_planting.get(species_key).isoformat(),
                "max_age_days": int(points[-1].get("day_since_species_start") or 0),
                "points": points,
            }
        )

    return {
        "as_of_date": today.isoformat(),
        "start_date": project_start_date.isoformat(),
        "species_count": len(species_rows),
        "trees_missing_planting_date": int(trees_missing_planting_date),
        "day_markers": {"day_30": 30, "day_90": 90, "day_180": 180},
        "species": species_rows,
    }


def _build_kpi_trend_series(project_id: int, db: Session, days: int = 180) -> list[dict]:
    """
    Build meaningful KPI trend points by month:
    - Survival: cumulative healthy share across planting cohorts over time.
    - Evidence: cumulative proof-complete share across in-scope task activity over time.
    """
    window_days = max(int(days), 1)
    today = date.today()
    window_start = today - timedelta(days=window_days - 1)

    earliest_planting = db.execute(
        text(
            """
            SELECT MIN(planting_date) AS first_planting_date
            FROM trees
            WHERE project_id = :project_id
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
              AND planting_date IS NOT NULL
            """
        ),
        {"project_id": project_id},
    ).scalar()
    earliest_planting_date = _parse_date_value(earliest_planting)
    trend_start_date = earliest_planting_date or window_start
    if earliest_planting_date and earliest_planting_date > today:
        trend_start_date = today

    start_month = _month_start(trend_start_date)
    end_month = _month_start(today)

    months: list[date] = []
    cursor = start_month
    while cursor <= end_month:
        months.append(cursor)
        cursor = _next_month_start(cursor)

    tree_rows = db.execute(
        text(
            """
            SELECT planting_date, status
            FROM trees
            WHERE project_id = :project_id
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    tree_month_totals: dict[date, int] = {}
    tree_month_healthy: dict[date, int] = {}
    baseline_tree_total = 0
    baseline_tree_healthy = 0

    for row in tree_rows:
        event_date = _parse_date_value(row.get("planting_date"))
        if event_date is None:
            continue
        bucket = _month_start(event_date)
        is_healthy = _normalize_tree_status(row.get("status")) in HEALTHY_TREE_STATUSES
        if bucket < start_month:
            baseline_tree_total += 1
            if is_healthy:
                baseline_tree_healthy += 1
            continue
        if bucket > end_month:
            continue
        tree_month_totals[bucket] = tree_month_totals.get(bucket, 0) + 1
        if is_healthy:
            tree_month_healthy[bucket] = tree_month_healthy.get(bucket, 0) + 1

    task_rows = db.execute(
        text(
            """
            SELECT t.task_type, t.status, t.review_state, t.notes, t.photo_url,
                   COALESCE(t.completed_at::date, t.submitted_at::date, t.reviewed_at::date, t.created_at::date)
                     AS activity_date
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE tr.project_id = :project_id
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    task_month_required: dict[date, int] = {}
    task_month_complete: dict[date, int] = {}
    baseline_task_required = 0
    baseline_task_complete = 0

    for task in task_rows:
        policy = _task_needs_evidence(task.get("task_type"))
        if not (policy.get("require_notes") or policy.get("require_photo")):
            continue
        state = _normalize_name(task.get("review_state") or "none")
        status = _normalize_name(task.get("status") or "pending")
        evidence_in_scope = _is_done_status(status) or state in {"submitted", "approved", "rejected"}
        if not evidence_in_scope:
            continue
        activity_date = _parse_date_value(task.get("activity_date"))
        if activity_date is None:
            continue
        evidence_ok, _ = _has_required_evidence(task.get("task_type"), task.get("notes"), task.get("photo_url"))
        bucket = _month_start(activity_date)
        if bucket < start_month:
            baseline_task_required += 1
            if evidence_ok:
                baseline_task_complete += 1
            continue
        if bucket > end_month:
            continue
        task_month_required[bucket] = task_month_required.get(bucket, 0) + 1
        if evidence_ok:
            task_month_complete[bucket] = task_month_complete.get(bucket, 0) + 1

    trend: list[dict] = []
    cumulative_tree_total = baseline_tree_total
    cumulative_tree_healthy = baseline_tree_healthy
    cumulative_task_required = baseline_task_required
    cumulative_task_complete = baseline_task_complete

    for month in months:
        cumulative_tree_total += tree_month_totals.get(month, 0)
        cumulative_tree_healthy += tree_month_healthy.get(month, 0)
        cumulative_task_required += task_month_required.get(month, 0)
        cumulative_task_complete += task_month_complete.get(month, 0)

        survival_rate = (
            round((cumulative_tree_healthy / cumulative_tree_total) * 100, 1)
            if cumulative_tree_total > 0
            else 0.0
        )
        evidence_rate = (
            round((cumulative_task_complete / cumulative_task_required) * 100, 1)
            if cumulative_task_required > 0
            else 0.0
        )

        trend.append(
            {
                "snapshot_at": month.isoformat(),
                "metrics": {
                    "survival_rate": survival_rate,
                    "evidence_complete_rate": evidence_rate,
                    "cohort_trees_total": cumulative_tree_total,
                    "evidence_required_tasks": cumulative_task_required,
                },
            }
        )

    return trend


def _store_kpi_snapshot(project_id: int, metrics: dict, db: Session):
    latest = db.execute(
        text(
            """
            SELECT snapshot_at, metrics
            FROM green_kpi_snapshots
            WHERE project_id = :project_id
            ORDER BY snapshot_at DESC, id DESC
            LIMIT 1
            """
        ),
        {"project_id": project_id},
    ).mappings().first()
    if latest:
        previous_metrics = dict(latest.get("metrics") or {})
        is_same = _safe_json(previous_metrics) == _safe_json(metrics)
        latest_at = latest.get("snapshot_at")
        if is_same and isinstance(latest_at, datetime):
            if latest_at >= datetime.utcnow() - timedelta(minutes=30):
                return
    db.execute(
        text("""
            INSERT INTO green_kpi_snapshots (project_id, metrics)
            VALUES (:project_id, CAST(:metrics AS JSONB))
        """),
        {"project_id": project_id, "metrics": _safe_json(metrics)},
    )


def _to_iso_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    return str(value)


def _excel_csv_writer(target) -> csv.writer:
    # Ensures Excel opens comma-separated exports in proper columns across locale settings.
    target.write("sep=,\n")
    return csv.writer(target)


def _is_within_monitoring_period(
    candidate: date | None,
    monitoring_start: date | None,
    monitoring_end: date | None,
) -> bool:
    if candidate is None:
        return monitoring_start is None and monitoring_end is None
    if monitoring_start and candidate < monitoring_start:
        return False
    if monitoring_end and candidate > monitoring_end:
        return False
    return True


def _build_verra_vcs_payload(
    project_id: int,
    db: Session,
    season_mode: str = "rainy",
    assignee_name: str | None = None,
    monitoring_start: date | None = None,
    monitoring_end: date | None = None,
    methodology_id: str | None = None,
    verifier_notes: str | None = None,
) -> dict:
    project = get_project(project_id, db)
    season = "dry" if _normalize_name(season_mode) == "dry" else "rainy"
    assignee_clean = (assignee_name or "").strip() or None

    tree_rows_raw = db.execute(
        text(
            """
            SELECT
                   t.id,
                   t.project_id,
                   t.species,
                   t.planting_date,
              t.status,
              t.notes,
              t.photo_url,
              t.photo_urls,
              t.created_by,
              t.created_at,
                   t.tree_origin,
                   t.tree_height_m,
                   t.attribution_scope,
                   t.count_in_planting_kpis,
                   t.count_in_carbon_scope,
                   t.custodian_id,
                   c.name AS custodian_name,
                   ST_X(t.geom) AS lng,
                   ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id
              AND (:assignee_name IS NULL OR t.created_by = :assignee_name)
            ORDER BY t.created_at ASC, t.id ASC
            """
        ),
        {"project_id": project_id, "assignee_name": assignee_clean},
    ).mappings().all()
    tree_rows = [dict(row) for row in tree_rows_raw]
    if monitoring_start or monitoring_end:
        tree_rows = [
            row
            for row in tree_rows
            if _is_within_monitoring_period(
                _parse_date_value(row.get("planting_date")) or _parse_date_value(row.get("created_at")),
                monitoring_start,
                monitoring_end,
            )
        ]
    maintenance_rows = _maintenance_summary_by_tree(project_id, db, assignee_clean)
    tree_rows = _attach_maintenance_to_tree_rows(tree_rows, maintenance_rows)
    review_summary = _review_summary_by_tree(project_id, db, assignee_clean)
    for tree in tree_rows:
        review = review_summary.get(int(tree.get("id") or 0), {})
        tree["review_submitted"] = int(review.get("review_submitted", 0))
        tree["review_approved"] = int(review.get("review_approved", 0))
        tree["review_rejected"] = int(review.get("review_rejected", 0))
        tree["last_review_state"] = review.get("last_review_state", "")
        tree["last_review_note"] = review.get("last_review_note", "")
        tree["last_submitted_at"] = review.get("last_submitted_at")
        tree["last_reviewed_at"] = review.get("last_reviewed_at")

    task_rows_raw = db.execute(
        text(
            """
            SELECT t.id, t.tree_id, t.task_type, t.assignee_name, t.due_date, t.priority,
                   t.status, t.notes, t.photo_url, t.created_at, t.completed_at, t.review_state,
                   t.submitted_at, t.reviewed_at, t.reviewed_by, t.review_notes, t.auto_generated,
                   t.model_season, t.source_task_id, t.reported_tree_status,
                   tr.status AS tree_status, tr.species AS tree_species
            FROM tree_tasks t
            JOIN trees tr ON tr.id = t.tree_id
            WHERE tr.project_id = :project_id
              AND COALESCE(t.auto_generated, FALSE) = FALSE
              AND (:assignee_name IS NULL OR t.assignee_name = :assignee_name)
            ORDER BY t.created_at ASC, t.id ASC
            """
        ),
        {"project_id": project_id, "assignee_name": assignee_clean},
    ).mappings().all()
    task_rows = [dict(row) for row in task_rows_raw]
    filtered_tree_ids = {int(row.get("id")) for row in tree_rows}
    if filtered_tree_ids:
        task_rows = [row for row in task_rows if int(row.get("tree_id") or 0) in filtered_tree_ids]
    elif monitoring_start or monitoring_end:
        task_rows = []
    if monitoring_start or monitoring_end:
        task_rows = [
            row
            for row in task_rows
            if _is_within_monitoring_period(
                _parse_date_value(
                    row.get("reviewed_at")
                    or row.get("submitted_at")
                    or row.get("completed_at")
                    or row.get("due_date")
                    or row.get("created_at")
                ),
                monitoring_start,
                monitoring_end,
            )
        ]

    donor_rows = _build_donor_report_rows(project_id, db)
    if assignee_clean:
        assignee_key = _normalize_name(assignee_clean)
        donor_rows = [row for row in donor_rows if _normalize_name(row.get("assignee_name")) == assignee_key]
    if filtered_tree_ids:
        donor_rows = [row for row in donor_rows if int(row.get("tree_id") or 0) in filtered_tree_ids]
    elif monitoring_start or monitoring_end:
        donor_rows = []
    if monitoring_start or monitoring_end:
        donor_rows = [
            row
            for row in donor_rows
            if _is_within_monitoring_period(
                _parse_date_value(
                    row.get("reviewed_at")
                    or row.get("submitted_at")
                    or row.get("completed_at")
                    or row.get("due_date")
                ),
                monitoring_start,
                monitoring_end,
            )
        ]

    live_payload = _compute_live_maintenance_rows(
        db=db,
        project_id=project_id,
        season_mode=season,
        assignee_name=assignee_clean,
    )
    if filtered_tree_ids:
        live_rows_filtered = [
            row for row in (live_payload.get("rows") or []) if int(row.get("treeId") or 0) in filtered_tree_ids
        ]
        live_payload = {
            "rows": live_rows_filtered,
            "summary": {
                "total": len(live_rows_filtered),
                "danger": sum(1 for item in live_rows_filtered if item.get("tone") == "danger"),
                "warning": sum(1 for item in live_rows_filtered if item.get("tone") == "warning"),
                "ok": sum(1 for item in live_rows_filtered if item.get("tone") == "ok"),
                "info": sum(1 for item in live_rows_filtered if item.get("tone") == "info"),
                "dueSoon": sum(
                    1
                    for item in live_rows_filtered
                    if isinstance(item.get("countdownDays"), int) and 0 <= int(item.get("countdownDays")) <= 7
                ),
            },
        }

    species_maturity_rows = db.execute(
        text(
            """
            SELECT species_key, species_label, maturity_years, updated_at
            FROM green_species_maturity
            WHERE project_id = :project_id
            ORDER BY COALESCE(species_label, species_key) ASC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()
    species_maturity = [dict(row) for row in species_maturity_rows]

    scope_tree_rows_for_carbon = [
        {
            "id": row.get("id"),
            "species": row.get("species"),
            "planting_date": row.get("planting_date"),
            "status": row.get("status"),
            "created_at": row.get("created_at"),
        }
        for row in tree_rows
        if bool(row.get("count_in_carbon_scope", True))
    ]
    carbon = compute_project_carbon(scope_tree_rows_for_carbon, projection_years=40)
    carbon_projection = generate_co2_projection_table(scope_tree_rows_for_carbon, years=40)

    in_scope_tree_rows = [row for row in tree_rows if bool(row.get("count_in_planting_kpis", True))]
    total_trees = len(in_scope_tree_rows)
    trees_healthy = sum(
        1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in HEALTHY_TREE_STATUSES
    )
    trees_dead_or_removed = sum(
        1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in DEAD_TREE_STATUSES
    )
    trees_attention = sum(
        1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) in ATTENTION_TREE_STATUSES
    )
    trees_pending_planting = sum(
        1 for row in in_scope_tree_rows if _normalize_tree_status(row.get("status")) == "pending_planting"
    )
    survival_rate = round((trees_healthy / total_trees) * 100, 1) if total_trees else 0.0

    today = date.today()
    tasks_total = len(task_rows)
    tasks_submitted = 0
    tasks_approved = 0
    tasks_rejected = 0
    tasks_open = 0
    tasks_overdue = 0
    evidence_required = 0
    evidence_complete = 0
    for task in task_rows:
        review_state = _normalize_name(task.get("review_state") or "none")
        status = _normalize_name(task.get("status") or "pending")
        due_date = _parse_date_value(task.get("due_date"))
        if review_state == "submitted":
            tasks_submitted += 1
        if review_state == "approved":
            tasks_approved += 1
        if review_state == "rejected":
            tasks_rejected += 1
        if not (_is_done_status(status) and review_state in {"approved", "none"}):
            tasks_open += 1
        if due_date and due_date < today and not (_is_done_status(status) and review_state in {"approved", "none"}):
            tasks_overdue += 1
        policy = _task_needs_evidence(task.get("task_type"))
        evidence_ok, _ = _has_required_evidence(task.get("task_type"), task.get("notes"), task.get("photo_url"))
        evidence_in_scope = _is_done_status(status) or review_state in {"submitted", "approved", "rejected"}
        if (policy.get("require_notes") or policy.get("require_photo")) and evidence_in_scope:
            evidence_required += 1
            if evidence_ok:
                evidence_complete += 1
    if evidence_required > 0:
        evidence_complete_rate = round((evidence_complete / evidence_required) * 100, 1)
    elif tasks_total > 0:
        evidence_complete_rate = 100.0
    else:
        evidence_complete_rate = 0.0

    monitoring_start_candidates: list[date] = []
    for row in in_scope_tree_rows:
        planting_date = _parse_date_value(row.get("planting_date"))
        created_stamp = _parse_date_value(row.get("created_at"))
        if planting_date:
            monitoring_start_candidates.append(planting_date)
        elif created_stamp:
            monitoring_start_candidates.append(created_stamp)
    monitoring_start = min(monitoring_start_candidates) if monitoring_start_candidates else None

    species_map: dict[str, dict] = {}
    for row in in_scope_tree_rows:
        raw_species = (row.get("species") or "").strip()
        species_label = raw_species if raw_species else "Unspecified"
        species_key = f"{species_label.lower()}::{_normalize_species_key(raw_species)}"
        model_species = _get_species_params(raw_species).get("label", "Medium-growth tropical (default)")
        status_key = _normalize_tree_status(row.get("status"))
        entry = species_map.setdefault(
            species_key,
            {
                "species_input": species_label,
                "model_species": model_species,
                "tree_count": 0,
                "healthy": 0,
                "attention": 0,
                "dead_or_removed": 0,
                "pending_planting": 0,
                "last_recorded_date": "",
            },
        )
        entry["tree_count"] += 1
        if status_key in HEALTHY_TREE_STATUSES:
            entry["healthy"] += 1
        if status_key in ATTENTION_TREE_STATUSES:
            entry["attention"] += 1
        if status_key in DEAD_TREE_STATUSES:
            entry["dead_or_removed"] += 1
        if status_key == "pending_planting":
            entry["pending_planting"] += 1
        last_date = _to_iso_text(row.get("created_at"))
        if last_date and last_date > entry["last_recorded_date"]:
            entry["last_recorded_date"] = last_date
    species_summary = sorted(
        species_map.values(),
        key=lambda item: (-(item.get("tree_count") or 0), item.get("species_input") or ""),
    )

    task_type_map: dict[str, dict] = {}
    for row in donor_rows:
        activity = _normalize_name(row.get("task_type")) or "unknown"
        entry = task_type_map.setdefault(
            activity,
            {
                "task_type": activity,
                "total": 0,
                "open": 0,
                "submitted": 0,
                "approved": 0,
                "rejected": 0,
                "overdue": 0,
                "with_photo": 0,
                "with_notes": 0,
                "avg_delay_days": 0.0,
                "_delay_sum": 0.0,
                "_delay_count": 0,
            },
        )
        entry["total"] += 1
        review_state = _normalize_name(row.get("review_state") or "none")
        status = _normalize_name(row.get("status") or "pending")
        due_date = _parse_date_value(row.get("due_date"))
        if review_state == "submitted":
            entry["submitted"] += 1
        if review_state == "approved":
            entry["approved"] += 1
        if review_state == "rejected":
            entry["rejected"] += 1
        if not (_is_done_status(status) and review_state in {"approved", "none"}):
            entry["open"] += 1
        if due_date and due_date < today and not (_is_done_status(status) and review_state in {"approved", "none"}):
            entry["overdue"] += 1
        if (row.get("photo_url") or "").strip():
            entry["with_photo"] += 1
        if (row.get("notes") or "").strip():
            entry["with_notes"] += 1
        delay_days = row.get("delay_days")
        if isinstance(delay_days, int):
            entry["_delay_sum"] += float(delay_days)
            entry["_delay_count"] += 1
    for entry in task_type_map.values():
        if entry["_delay_count"] > 0:
            entry["avg_delay_days"] = round(entry["_delay_sum"] / entry["_delay_count"], 1)
        entry.pop("_delay_sum", None)
        entry.pop("_delay_count", None)
    task_type_summary = sorted(task_type_map.values(), key=lambda item: item.get("task_type") or "")

    order_rows = db.execute(
        text(
            """
            SELECT assignee_name, work_type, target_trees, maintenance_schedule, due_date, status, created_at
            FROM green_work_orders
            WHERE project_id = :project_id
              AND (:assignee_name IS NULL OR LOWER(TRIM(assignee_name)) = LOWER(TRIM(:assignee_name)))
            ORDER BY created_at DESC, id DESC
            """
        ),
        {"project_id": project_id, "assignee_name": assignee_clean},
    ).mappings().all()
    staff_map: dict[str, dict] = {}

    def _staff_entry(name: str) -> dict:
        key = _normalize_name(name)
        label = name.strip() if name.strip() else "Unassigned"
        if key not in staff_map:
            staff_map[key] = {
                "staff_name": label,
                "trees_recorded": 0,
                "trees_approved": 0,
                "tasks_total": 0,
                "tasks_open": 0,
                "tasks_submitted": 0,
                "tasks_approved": 0,
                "tasks_rejected": 0,
                "orders_total": 0,
                "planting_target_trees": 0,
                "maintenance_orders": 0,
                "last_activity_at": "",
            }
        return staff_map[key]

    for row in in_scope_tree_rows:
        name = str(row.get("created_by") or "Unassigned")
        entry = _staff_entry(name)
        entry["trees_recorded"] += 1
        if _normalize_tree_status(row.get("status")) != "pending_planting":
            entry["trees_approved"] += 1
        created_at = _to_iso_text(row.get("created_at"))
        if created_at and created_at > entry["last_activity_at"]:
            entry["last_activity_at"] = created_at

    for row in task_rows:
        name = str(row.get("assignee_name") or "Unassigned")
        entry = _staff_entry(name)
        entry["tasks_total"] += 1
        review_state = _normalize_name(row.get("review_state") or "none")
        status = _normalize_name(row.get("status") or "pending")
        if review_state == "submitted":
            entry["tasks_submitted"] += 1
        if review_state == "approved":
            entry["tasks_approved"] += 1
        if review_state == "rejected":
            entry["tasks_rejected"] += 1
        if not (_is_done_status(status) and review_state in {"approved", "none"}):
            entry["tasks_open"] += 1
        timestamps = [
            _to_iso_text(row.get("created_at")),
            _to_iso_text(row.get("completed_at")),
            _to_iso_text(row.get("submitted_at")),
            _to_iso_text(row.get("reviewed_at")),
        ]
        for stamp in timestamps:
            if stamp and stamp > entry["last_activity_at"]:
                entry["last_activity_at"] = stamp

    for row in order_rows:
        name = str(row.get("assignee_name") or "Unassigned")
        entry = _staff_entry(name)
        entry["orders_total"] += 1
        if _normalize_name(row.get("work_type")) == "planting":
            entry["planting_target_trees"] += int(row.get("target_trees") or 0)
        elif _normalize_name(row.get("work_type")) == "maintenance":
            entry["maintenance_orders"] += 1
        created_at = _to_iso_text(row.get("created_at"))
        if created_at and created_at > entry["last_activity_at"]:
            entry["last_activity_at"] = created_at
    staff_summary = sorted(staff_map.values(), key=lambda item: item.get("staff_name") or "")

    risk_items: list[dict] = []
    for item in live_payload.get("rows", []):
        tone = item.get("tone") or "ok"
        if tone not in {"danger", "warning"}:
            continue
        risk_items.append(
            {
                "tree_id": item.get("treeId"),
                "activity": item.get("activity"),
                "status_text": item.get("statusText"),
                "indicator": item.get("indicator"),
                "effective_due_date": item.get("effectiveDueDate"),
                "countdown_days": item.get("countdownDays"),
                "open_task_id": item.get("openTaskId"),
                "severity": "high" if tone == "danger" else "medium",
            }
        )

    manual_fields_required = [
        {
            "field": "VCS methodology reference (e.g., VMxxxx)",
            "status": "manual_input_required",
            "note": "Attach the approved methodology and version used for this project.",
        },
        {
            "field": "Project boundary and strata definitions",
            "status": "manual_input_required",
            "note": "Provide shapefiles/boundary narrative aligned with Verra requirements.",
        },
        {
            "field": "Leakage, non-permanence, and uncertainty treatment",
            "status": "manual_input_required",
            "note": "Document assumptions and verifier-ready calculations.",
        },
        {
            "field": "Validation/verification body statements",
            "status": "manual_input_required",
            "note": "To be filled during third-party assurance workflow.",
        },
    ]

    monitoring_period_end = monitoring_end or today
    payload = {
        "template": {
            "name": "LandCheck Verra VCS Structured Monitoring Template",
            "version": "1.0",
            "aligned_standard": "Verra Verified Carbon Standard (VCS)",
            "generated_at_utc": datetime.utcnow().isoformat() + "Z",
            "data_mode": "live_project_snapshot",
            "refresh_behavior": "Automatically recomputed from project records on each export.",
            "scope_note": "This package pre-fills structured monitoring data and annex tables for drafting. Final VCS submission text and verifier evidence remain required.",
        },
        "project": {
            "id": project.get("id"),
            "name": project.get("name") or "",
            "location_text": project.get("location_text") or "",
            "sponsor": project.get("sponsor") or "",
            "created_at": _to_iso_text(project.get("created_at")),
        },
        "monitoring_period": {
            "start_date": _to_date_input(monitoring_start),
            "end_date": _to_date_input(monitoring_period_end),
            "duration_days": _day_diff(monitoring_period_end, monitoring_start) if monitoring_start else 0,
            "season_model": season,
            "assignee_filter": assignee_clean or "all",
            "is_custom_period_filter": bool(monitoring_start or monitoring_end),
        },
        "verifier_metadata": {
            "methodology_id": (methodology_id or "").strip(),
            "verifier_notes": (verifier_notes or "").strip(),
        },
        "section_1_project_identification": {
            "project_summary": {
                "total_trees_all": len(tree_rows),
                "total_trees": total_trees,
                "trees_healthy": trees_healthy,
                "trees_dead_or_removed": trees_dead_or_removed,
                "trees_attention": trees_attention,
                "trees_pending_planting": trees_pending_planting,
                "survival_rate_percent": survival_rate,
                "attribution_scope_note": "Trees counted in planting KPIs use count_in_planting_kpis=true scope.",
            },
            "species_count": len(species_summary),
            "staff_count": len(staff_summary),
        },
        "section_2_activity_monitoring": {
            "task_snapshot": {
                "tasks_total": tasks_total,
                "tasks_open": tasks_open,
                "tasks_submitted": tasks_submitted,
                "tasks_approved": tasks_approved,
                "tasks_rejected": tasks_rejected,
                "tasks_overdue": tasks_overdue,
            },
            "task_type_summary": task_type_summary,
            "live_maintenance_summary": live_payload.get("summary", {}),
            "high_risk_items": risk_items[:200],
        },
        "section_3_ghg_quantification": {
            "co2_current_tonnes": carbon.get("current_co2_tonnes", 0),
            "co2_annual_tonnes": carbon.get("annual_co2_tonnes", 0),
            "co2_projected_lifetime_tonnes": carbon.get("projected_lifetime_co2_tonnes", 0),
            "co2_average_per_tree_kg": carbon.get("co2_per_tree_avg_kg", 0),
            "methodology": carbon.get("methodology"),
            "top_species_by_co2": carbon.get("top_species", []),
            "projection_table": carbon_projection,
            "carbon_data_quality": {
                "trees_missing_age_data": carbon.get("trees_missing_age_data", 0),
                "trees_with_fallback_age": carbon.get("trees_with_fallback_age", 0),
                "trees_pending_review": carbon.get("trees_pending_review", 0),
            },
        },
        "section_4_qa_qc_and_evidence": {
            "evidence_required_tasks": evidence_required,
            "evidence_complete_tasks": evidence_complete,
            "evidence_complete_rate_percent": evidence_complete_rate,
            "recent_review_timeline": donor_rows[:500],
        },
        "section_5_reversal_and_risk_tracking": {
            "tree_status_distribution": {
                "healthy": trees_healthy,
                "attention": trees_attention,
                "dead_or_removed": trees_dead_or_removed,
                "pending_planting": trees_pending_planting,
            },
            "risk_indicators": {
                "live_danger_rows": int(live_payload.get("summary", {}).get("danger", 0)),
                "live_warning_rows": int(live_payload.get("summary", {}).get("warning", 0)),
                "overdue_tasks": tasks_overdue,
                "rejected_tasks": tasks_rejected,
            },
        },
        "section_6_annex_data_tables": {
            "tree_inventory_count": len(tree_rows),
            "task_timeline_count": len(donor_rows),
            "live_maintenance_count": len(live_payload.get("rows", [])),
            "species_summary_count": len(species_summary),
            "staff_summary_count": len(staff_summary),
            "species_maturity_rules_count": len(species_maturity),
        },
        "species_summary": species_summary,
        "staff_summary": staff_summary,
        "species_maturity_rules": species_maturity,
        "manual_fields_required_for_submission": manual_fields_required,
        "source_references": [*VERRA_VCS_REFERENCES, *LIVE_SOURCE_REFERENCES],
    }

    return {
        "payload": payload,
        "trees": tree_rows,
        "tasks": task_rows,
        "donor_rows": donor_rows,
        "live_rows": live_payload.get("rows", []),
        "species_summary": species_summary,
        "staff_summary": staff_summary,
        "species_maturity": species_maturity,
    }


def _write_csv_to_zip(zf: zipfile.ZipFile, filename: str, headers: list[str], rows: list[list[object]]):
    sio = io.StringIO()
    writer = _excel_csv_writer(sio)
    writer.writerow(headers)
    for row in rows:
        writer.writerow([_to_iso_text(value) for value in row])
    zf.writestr(filename, "\ufeff" + sio.getvalue())


def _render_verra_vcs_zip(package: dict) -> io.BytesIO:
    payload = package.get("payload") or {}
    trees = package.get("trees") or []
    tasks = package.get("tasks") or []
    donor_rows = package.get("donor_rows") or []
    live_rows = package.get("live_rows") or []
    species_summary = package.get("species_summary") or []
    staff_summary = package.get("staff_summary") or []
    species_maturity = package.get("species_maturity") or []

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("01_verra_vcs_template.json", json.dumps(payload, indent=2, default=str))
        readme_lines = [
            "LandCheck Verra VCS Structured Export Package",
            "",
            "Purpose:",
            "- Pre-fills monitoring data in a Verra-aligned structure using live project records.",
            "- Regenerates on every export request to reflect latest project growth and reviews.",
            "",
            "Included files:",
            "- 01_verra_vcs_template.json",
            "- 02_tree_inventory.csv",
            "- 03_task_activity.csv",
            "- 04_review_timeline.csv",
            "- 05_live_maintenance_table.csv",
            "- 06_species_summary.csv",
            "- 07_staff_summary.csv",
            "- 08_species_maturity_rules.csv",
            "- 09_source_references.csv",
            "",
            "Note:",
            "This is a structured drafting package and does not replace final verifier-reviewed VCS submission requirements.",
        ]
        zf.writestr("00_README.txt", "\n".join(readme_lines))

        _write_csv_to_zip(
            zf,
            "02_tree_inventory.csv",
            [
                "tree_id",
                "project_id",
                "created_by",
                "tree_origin",
                "attribution_scope",
                "count_in_planting_kpis",
                "count_in_carbon_scope",
                "custodian_id",
                "custodian_name",
                "tree_height_m",
                "species",
                "planting_date",
                "status",
                "longitude",
                "latitude",
                "maintenance_count",
                "maintenance_done",
                "maintenance_pending",
                "maintenance_overdue",
                "maintenance_types",
                "last_maintenance_type",
                "last_maintenance_date",
                "review_submitted",
                "review_approved",
                "review_rejected",
                "last_review_state",
                "last_review_note",
                "last_submitted_at",
                "last_reviewed_at",
                "photo_url",
                "notes",
                "created_at",
            ],
            [
                [
                    row.get("id"),
                    row.get("project_id"),
                    row.get("created_by"),
                    row.get("tree_origin"),
                    row.get("attribution_scope"),
                    row.get("count_in_planting_kpis"),
                    row.get("count_in_carbon_scope"),
                    row.get("custodian_id"),
                    row.get("custodian_name"),
                    row.get("tree_height_m"),
                    row.get("species"),
                    row.get("planting_date"),
                    row.get("status"),
                    row.get("lng"),
                    row.get("lat"),
                    row.get("maintenance_count"),
                    row.get("maintenance_done"),
                    row.get("maintenance_pending"),
                    row.get("maintenance_overdue"),
                    row.get("maintenance_types"),
                    row.get("last_maintenance_type"),
                    row.get("last_maintenance_date"),
                    row.get("review_submitted"),
                    row.get("review_approved"),
                    row.get("review_rejected"),
                    row.get("last_review_state"),
                    row.get("last_review_note"),
                    row.get("last_submitted_at"),
                    row.get("last_reviewed_at"),
                    row.get("photo_url"),
                    row.get("notes"),
                    row.get("created_at"),
                ]
                for row in trees
            ],
        )

        _write_csv_to_zip(
            zf,
            "03_task_activity.csv",
            [
                "task_id",
                "tree_id",
                "task_type",
                "assignee_name",
                "status",
                "review_state",
                "priority",
                "due_date",
                "completed_at",
                "submitted_at",
                "reviewed_at",
                "reviewed_by",
                "reported_tree_status",
                "tree_status",
                "photo_url",
                "notes",
                "created_at",
            ],
            [
                [
                    row.get("id"),
                    row.get("tree_id"),
                    row.get("task_type"),
                    row.get("assignee_name"),
                    row.get("status"),
                    row.get("review_state"),
                    row.get("priority"),
                    row.get("due_date"),
                    row.get("completed_at"),
                    row.get("submitted_at"),
                    row.get("reviewed_at"),
                    row.get("reviewed_by"),
                    row.get("reported_tree_status"),
                    row.get("tree_status"),
                    row.get("photo_url"),
                    row.get("notes"),
                    row.get("created_at"),
                ]
                for row in tasks
            ],
        )

        _write_csv_to_zip(
            zf,
            "04_review_timeline.csv",
            [
                "task_id",
                "tree_id",
                "species",
                "tree_origin",
                "attribution_scope",
                "count_in_planting_kpis",
                "count_in_carbon_scope",
                "custodian_id",
                "custodian_name",
                "tree_height_m",
                "assignee_name",
                "task_type",
                "priority",
                "status",
                "review_state",
                "due_date",
                "completed_at",
                "submitted_at",
                "reviewed_at",
                "reviewed_by",
                "review_notes",
                "delay_days",
                "delay_context",
                "evidence_status",
                "reported_tree_status",
                "tree_status",
                "photo_url",
                "notes",
            ],
            [
                [
                    row.get("task_id"),
                    row.get("tree_id"),
                    row.get("species"),
                    row.get("tree_origin"),
                    row.get("attribution_scope"),
                    row.get("count_in_planting_kpis"),
                    row.get("count_in_carbon_scope"),
                    row.get("custodian_id"),
                    row.get("custodian_name"),
                    row.get("tree_height_m"),
                    row.get("assignee_name"),
                    row.get("task_type"),
                    row.get("priority"),
                    row.get("status"),
                    row.get("review_state"),
                    row.get("due_date"),
                    row.get("completed_at"),
                    row.get("submitted_at"),
                    row.get("reviewed_at"),
                    row.get("reviewed_by"),
                    row.get("review_notes"),
                    row.get("delay_days"),
                    row.get("delay_context"),
                    row.get("evidence_status"),
                    row.get("reported_tree_status"),
                    row.get("tree_status"),
                    row.get("photo_url"),
                    row.get("notes"),
                ]
                for row in donor_rows
            ],
        )

        _write_csv_to_zip(
            zf,
            "05_live_maintenance_table.csv",
            [
                "tree_id",
                "assignee",
                "activity",
                "activity_label",
                "planting_date",
                "tree_age_days",
                "last_done_at",
                "model_due_date",
                "assigned_due_date",
                "effective_due_date",
                "countdown_days",
                "tone",
                "status_text",
                "indicator",
                "done_count",
                "pending_count",
                "overdue_count",
                "open_task_id",
                "model_rationale",
            ],
            [
                [
                    row.get("treeId"),
                    row.get("assignee"),
                    row.get("activity"),
                    row.get("activityLabel"),
                    row.get("plantingDate"),
                    row.get("treeAgeDays"),
                    row.get("lastDoneAt"),
                    row.get("modelDueDate"),
                    row.get("assignedDueDate"),
                    row.get("effectiveDueDate"),
                    row.get("countdownDays"),
                    row.get("tone"),
                    row.get("statusText"),
                    row.get("indicator"),
                    row.get("doneCount"),
                    row.get("pendingCount"),
                    row.get("overdueCount"),
                    row.get("openTaskId"),
                    row.get("modelRationale"),
                ]
                for row in live_rows
            ],
        )

        _write_csv_to_zip(
            zf,
            "06_species_summary.csv",
            [
                "species_input",
                "model_species",
                "tree_count",
                "healthy",
                "attention",
                "dead_or_removed",
                "pending_planting",
                "last_recorded_date",
            ],
            [
                [
                    row.get("species_input"),
                    row.get("model_species"),
                    row.get("tree_count"),
                    row.get("healthy"),
                    row.get("attention"),
                    row.get("dead_or_removed"),
                    row.get("pending_planting"),
                    row.get("last_recorded_date"),
                ]
                for row in species_summary
            ],
        )

        _write_csv_to_zip(
            zf,
            "07_staff_summary.csv",
            [
                "staff_name",
                "trees_recorded",
                "trees_approved",
                "tasks_total",
                "tasks_open",
                "tasks_submitted",
                "tasks_approved",
                "tasks_rejected",
                "orders_total",
                "planting_target_trees",
                "maintenance_orders",
                "last_activity_at",
            ],
            [
                [
                    row.get("staff_name"),
                    row.get("trees_recorded"),
                    row.get("trees_approved"),
                    row.get("tasks_total"),
                    row.get("tasks_open"),
                    row.get("tasks_submitted"),
                    row.get("tasks_approved"),
                    row.get("tasks_rejected"),
                    row.get("orders_total"),
                    row.get("planting_target_trees"),
                    row.get("maintenance_orders"),
                    row.get("last_activity_at"),
                ]
                for row in staff_summary
            ],
        )

        _write_csv_to_zip(
            zf,
            "08_species_maturity_rules.csv",
            ["species_key", "species_label", "maturity_years", "updated_at"],
            [
                [
                    row.get("species_key"),
                    row.get("species_label"),
                    row.get("maturity_years"),
                    row.get("updated_at"),
                ]
                for row in species_maturity
            ],
        )

        _write_csv_to_zip(
            zf,
            "09_source_references.csv",
            ["label", "url"],
            [
                [ref.get("label"), ref.get("url")]
                for ref in (payload.get("source_references") or [])
            ],
        )
    buffer.seek(0)
    return buffer


def _render_verra_vcs_docx(package: dict) -> io.BytesIO:
    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=501, detail="DOCX export requires python-docx.") from exc

    payload = package.get("payload") or {}
    project_section = payload.get("section_1_project_identification") or {}
    monitoring_section = payload.get("section_2_monitoring_summary") or {}
    ghg_section = payload.get("section_3_ghg_quantification") or {}
    annex_section = payload.get("section_6_annex_data_tables") or {}

    trees = package.get("trees") or []
    tasks = package.get("tasks") or []
    donor_rows = package.get("donor_rows") or []
    top_species = ghg_section.get("top_species_by_co2") or []
    sources = payload.get("source_references") or []

    document = Document()
    document.add_heading("LandCheck Verra VCS Structured Report", level=0)
    document.add_paragraph(
        "\n".join(
            [
                f"Generated at: {_to_iso_text(payload.get('generated_at'))}",
                f"Project ID: {_to_iso_text(project_section.get('project_id'))}",
                f"Project name: {_to_iso_text(project_section.get('project_name'))}",
                f"Location: {_to_iso_text(project_section.get('project_location'))}",
                f"Monitoring period: {_to_iso_text(project_section.get('monitoring_period_start'))} to {_to_iso_text(project_section.get('monitoring_period_end'))}",
                f"Methodology: {_to_iso_text(project_section.get('methodology_reference'))}",
            ]
        )
    )

    def add_kv_table(title: str, rows: list[tuple[str, object]]):
        document.add_heading(title, level=2)
        table = document.add_table(rows=len(rows) + 1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for idx, (label, value) in enumerate(rows, start=1):
            table.rows[idx].cells[0].text = str(label)
            table.rows[idx].cells[1].text = _to_iso_text(value)

    add_kv_table(
        "Monitoring Summary",
        [
            ("Trees total", monitoring_section.get("trees_total", 0)),
            ("Trees healthy", monitoring_section.get("trees_healthy", 0)),
            ("Trees attention", monitoring_section.get("trees_attention", 0)),
            ("Trees dead/removed", monitoring_section.get("trees_dead_or_removed", 0)),
            ("Open tasks", monitoring_section.get("tasks_open", 0)),
            ("Submitted tasks", monitoring_section.get("tasks_submitted", 0)),
            ("Approved tasks", monitoring_section.get("tasks_approved", 0)),
            ("Rejected tasks", monitoring_section.get("tasks_rejected", 0)),
            ("Overdue tasks", monitoring_section.get("tasks_overdue", 0)),
            ("Evidence complete rate (%)", monitoring_section.get("evidence_complete_rate_percent", 0)),
        ],
    )

    add_kv_table(
        "GHG Quantification",
        [
            ("CO2 current (tonnes)", ghg_section.get("co2_current_tonnes", 0)),
            ("CO2 annual (tonnes)", ghg_section.get("co2_annual_tonnes", 0)),
            ("CO2 projected lifetime (tonnes)", ghg_section.get("co2_projected_lifetime_tonnes", 0)),
            ("CO2 average per tree (kg)", ghg_section.get("co2_average_per_tree_kg", 0)),
            ("Methodology", ghg_section.get("methodology", "")),
        ],
    )

    add_kv_table(
        "Annex Counts",
        [
            ("Tree inventory rows", annex_section.get("tree_inventory_count", len(trees))),
            ("Task timeline rows", annex_section.get("task_timeline_count", len(tasks))),
            ("Review timeline rows", len(donor_rows)),
            ("Species summary rows", annex_section.get("species_summary_count", 0)),
            ("Staff summary rows", annex_section.get("staff_summary_count", 0)),
            ("Species maturity rows", annex_section.get("species_maturity_rules_count", 0)),
        ],
    )

    if top_species:
        document.add_heading("Top Species by CO2", level=2)
        species_table = document.add_table(rows=min(len(top_species), 20) + 1, cols=4)
        species_table.style = "Table Grid"
        species_table.rows[0].cells[0].text = "Species input"
        species_table.rows[0].cells[1].text = "Model species"
        species_table.rows[0].cells[2].text = "Tree count"
        species_table.rows[0].cells[3].text = "Current CO2 (kg)"
        for idx, species_row in enumerate(top_species[:20], start=1):
            species_table.rows[idx].cells[0].text = _to_iso_text(species_row.get("species"))
            species_table.rows[idx].cells[1].text = _to_iso_text(species_row.get("model_species"))
            species_table.rows[idx].cells[2].text = _to_iso_text(species_row.get("count"))
            species_table.rows[idx].cells[3].text = _to_iso_text(species_row.get("co2_kg"))

    if donor_rows:
        document.add_heading("Recent Review Timeline (sample)", level=2)
        review_table = document.add_table(rows=min(len(donor_rows), 25) + 1, cols=6)
        review_table.style = "Table Grid"
        review_table.rows[0].cells[0].text = "Task ID"
        review_table.rows[0].cells[1].text = "Tree ID"
        review_table.rows[0].cells[2].text = "Task type"
        review_table.rows[0].cells[3].text = "Status"
        review_table.rows[0].cells[4].text = "Review state"
        review_table.rows[0].cells[5].text = "Reviewed at"
        for idx, timeline_row in enumerate(donor_rows[:25], start=1):
            review_table.rows[idx].cells[0].text = _to_iso_text(timeline_row.get("task_id"))
            review_table.rows[idx].cells[1].text = _to_iso_text(timeline_row.get("tree_id"))
            review_table.rows[idx].cells[2].text = _to_iso_text(timeline_row.get("task_type"))
            review_table.rows[idx].cells[3].text = _to_iso_text(timeline_row.get("status"))
            review_table.rows[idx].cells[4].text = _to_iso_text(timeline_row.get("review_state"))
            review_table.rows[idx].cells[5].text = _to_iso_text(timeline_row.get("reviewed_at"))

    if sources:
        document.add_heading("Source References", level=2)
        for source in sources:
            label = _to_iso_text(source.get("label"))
            url = _to_iso_text(source.get("url"))
            document.add_paragraph(f"- {label}: {url}")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("/projects/{project_id}/export/verra-vcs")
def export_project_verra_vcs(
    project_id: int,
    season_mode: str = Query(default="rainy"),
    assignee_name: str | None = Query(default=None),
    monitoring_start: str | None = Query(default=None),
    monitoring_end: str | None = Query(default=None),
    methodology_id: str | None = Query(default=None),
    verifier_notes: str | None = Query(default=None),
    generated_by: str | None = Query(default=None),
    output_format: str = Query(default="zip", alias="format"),
    db: Session = Depends(get_db),
):
    monitoring_start_date = _parse_date_value(monitoring_start)
    monitoring_end_date = _parse_date_value(monitoring_end)
    if monitoring_start and monitoring_start_date is None:
        raise HTTPException(status_code=400, detail="Invalid monitoring_start date.")
    if monitoring_end and monitoring_end_date is None:
        raise HTTPException(status_code=400, detail="Invalid monitoring_end date.")
    if monitoring_start_date and monitoring_end_date and monitoring_end_date < monitoring_start_date:
        raise HTTPException(status_code=400, detail="monitoring_end cannot be before monitoring_start.")

    package = _build_verra_vcs_payload(
        project_id=project_id,
        db=db,
        season_mode=season_mode,
        assignee_name=assignee_name,
        monitoring_start=monitoring_start_date,
        monitoring_end=monitoring_end_date,
        methodology_id=methodology_id,
        verifier_notes=verifier_notes,
    )
    format_key = _normalize_name(output_format)
    if format_key not in {"zip", "json", "docx"}:
        format_key = "zip"
    project_token = f"project_{project_id}"
    if format_key == "json":
        file_name = f"{project_token}_verra_vcs_template.json"
    elif format_key == "docx":
        file_name = f"{project_token}_verra_vcs_report.docx"
    else:
        file_name = f"{project_token}_verra_vcs_package.zip"

    payload = package.get("payload") or {}
    payload_summary = {
        "tree_inventory_count": int(payload.get("section_6_annex_data_tables", {}).get("tree_inventory_count", 0)),
        "task_timeline_count": int(payload.get("section_6_annex_data_tables", {}).get("task_timeline_count", 0)),
        "live_maintenance_count": int(payload.get("section_6_annex_data_tables", {}).get("live_maintenance_count", 0)),
        "co2_current_tonnes": float(payload.get("section_3_ghg_quantification", {}).get("co2_current_tonnes", 0) or 0),
        "co2_projected_lifetime_tonnes": float(
            payload.get("section_3_ghg_quantification", {}).get("co2_projected_lifetime_tonnes", 0) or 0
        ),
    }
    db.execute(
        text(
            """
            INSERT INTO green_verra_exports (
                project_id, season_mode, assignee_name, output_format, monitoring_start, monitoring_end,
                methodology_id, verifier_notes, generated_by, file_name, payload_summary
            )
            VALUES (
                :project_id, :season_mode, :assignee_name, :output_format, :monitoring_start, :monitoring_end,
                :methodology_id, :verifier_notes, :generated_by, :file_name, CAST(:payload_summary AS JSONB)
            )
            """
        ),
        {
            "project_id": project_id,
            "season_mode": "dry" if _normalize_name(season_mode) == "dry" else "rainy",
            "assignee_name": (assignee_name or "").strip() or None,
            "output_format": format_key,
            "monitoring_start": monitoring_start_date,
            "monitoring_end": monitoring_end_date,
            "methodology_id": (methodology_id or "").strip() or None,
            "verifier_notes": (verifier_notes or "").strip() or None,
            "generated_by": (generated_by or "").strip() or None,
            "file_name": file_name,
            "payload_summary": _safe_json(payload_summary),
        },
    )
    db.commit()

    if format_key == "json":
        content = json.dumps(payload, indent=2, default=str).encode("utf-8")
        headers = {"Content-Disposition": f'attachment; filename="{file_name}"'}
        return StreamingResponse(io.BytesIO(content), media_type="application/json", headers=headers)

    if format_key == "docx":
        docx_buffer = _render_verra_vcs_docx(package)
        headers = {"Content-Disposition": f'attachment; filename="{file_name}"'}
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    zip_buffer = _render_verra_vcs_zip(package)
    headers = {"Content-Disposition": f'attachment; filename="{file_name}"'}
    return StreamingResponse(zip_buffer, media_type="application/zip", headers=headers)


@router.get("/donor/export/verra-vcs")
def export_project_verra_vcs_alias(
    project_id: int = Query(...),
    season_mode: str = Query(default="rainy"),
    assignee_name: str | None = Query(default=None),
    monitoring_start: str | None = Query(default=None),
    monitoring_end: str | None = Query(default=None),
    methodology_id: str | None = Query(default=None),
    verifier_notes: str | None = Query(default=None),
    generated_by: str | None = Query(default=None),
    output_format: str = Query(default="zip", alias="format"),
    db: Session = Depends(get_db),
):
    return export_project_verra_vcs(
        project_id=project_id,
        season_mode=season_mode,
        assignee_name=assignee_name,
        monitoring_start=monitoring_start,
        monitoring_end=monitoring_end,
        methodology_id=methodology_id,
        verifier_notes=verifier_notes,
        generated_by=generated_by,
        output_format=output_format,
        db=db,
    )


@router.get("/projects/{project_id}/verra/exports")
def list_verra_export_history(
    project_id: int,
    limit: int = Query(default=100, ge=1, le=500),
    db: Session = Depends(get_db),
):
    rows = db.execute(
        text(
            """
            SELECT id, project_id, season_mode, assignee_name, output_format,
                   monitoring_start, monitoring_end, methodology_id, verifier_notes,
                   generated_by, file_name, payload_summary, created_at
            FROM green_verra_exports
            WHERE project_id = :project_id
            ORDER BY created_at DESC, id DESC
            LIMIT :limit
            """
        ),
        {"project_id": project_id, "limit": int(limit)},
    ).mappings().all()
    return [dict(row) for row in rows]

@router.get("/projects/{project_id}/donor-report/csv")
def export_donor_report_csv(project_id: int, db: Session = Depends(get_db)):
    project = get_project(project_id, db)
    rows = db.execute(text("""
        SELECT
            t.id,
            t.project_id,
            t.species,
            t.planting_date,
            t.status,
            t.notes,
            t.photo_url,
            t.created_by,
            t.created_at,
            t.tree_origin,
            t.attribution_scope,
            t.count_in_planting_kpis,
            t.count_in_carbon_scope,
            t.custodian_id,
            c.name AS custodian_name,
            t.tree_height_m,
            ST_X(t.geom) AS lng,
            ST_Y(t.geom) AS lat
        FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.project_id = :project_id
        ORDER BY t.created_at DESC
    """), {"project_id": project_id}).mappings().all()
    maintenance_rows = _maintenance_summary_by_tree(project_id, db)
    rows = _attach_maintenance_to_tree_rows(rows, maintenance_rows)
    review_summary = _review_summary_by_tree(project_id, db)
    kpi = _compute_kpi_snapshot(project_id, db)
    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_csv = tempfile.NamedTemporaryFile(suffix="_donor_report.csv", delete=False)
    csv_path = tmp_csv.name
    tmp_csv.close()

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = _excel_csv_writer(f)
        # Carbon summary for CSV header
        carbon_trees_csv = db.execute(text("""
            SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
            FROM trees
            WHERE project_id = :pid
              AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
        """), {"pid": project_id}).mappings().all()
        carbon_csv = compute_project_carbon([dict(r) for r in carbon_trees_csv])

        writer.writerow(["LandCheck Donor + Operations Report"])
        writer.writerow(["Project", project.get("name") or "", "Location", project.get("location_text") or "", "Sponsor", project.get("sponsor") or ""])
        writer.writerow([
            "KPI",
            f"Trees {kpi.get('trees_total', 0)}",
            f"Healthy {kpi.get('trees_healthy', 0)}",
            f"Attention {kpi.get('trees_attention', 0)}",
            f"Open Tasks {kpi.get('tasks_open', 0)}",
            f"Submitted {kpi.get('tasks_submitted', 0)}",
            f"Rejected {kpi.get('tasks_rejected', 0)}",
            f"Overdue {kpi.get('tasks_overdue', 0)}",
        ])
        writer.writerow([
            "Carbon Impact",
            f"CO2 Sequestered: {carbon_csv.get('current_co2_tonnes', 0)} tonnes",
            f"Annual Rate: {carbon_csv.get('annual_co2_tonnes', 0)} t/yr",
            f"40-Year Projection: {carbon_csv.get('projected_lifetime_co2_tonnes', 0)} tonnes",
            f"Avg per Tree: {carbon_csv.get('co2_per_tree_avg_kg', 0)} kg",
            "Methodology: IPCC Tier 1 + Chave et al. (2014)",
        ])
        writer.writerow([
            "Carbon Data Quality",
            f"Missing age data: {carbon_csv.get('trees_missing_age_data', 0)}",
            f"Fallback age used: {carbon_csv.get('trees_with_fallback_age', 0)}",
            f"Pending review trees: {carbon_csv.get('trees_pending_review', 0)}",
        ])
        top_species_csv = carbon_csv.get("top_species", []) or []
        if top_species_csv:
            writer.writerow(["Top Species CO2 Table"])
            writer.writerow(["species_input", "model_species", "tree_count", "current_co2_kg"])
            for sp in top_species_csv[:10]:
                writer.writerow([
                    sp.get("species", ""),
                    sp.get("model_species", ""),
                    sp.get("count", 0),
                    sp.get("co2_kg", 0),
                ])
        writer.writerow([])
        writer.writerow([
            "tree_id",
            "project_id",
            "lng",
            "lat",
            "tree_origin",
            "attribution_scope",
            "count_in_planting_kpis",
            "count_in_carbon_scope",
            "custodian_id",
            "custodian_name",
            "tree_height_m",
            "species",
            "planting_date",
            "tree_status",
            "created_by",
            "maintenance_count",
            "maintenance_done",
            "maintenance_pending",
            "maintenance_overdue",
            "maintenance_types",
            "last_maintenance_type",
            "last_maintenance_date",
            "review_submitted",
            "review_approved",
            "review_rejected",
            "last_review_state",
            "last_review_note",
            "last_submitted_at",
            "last_reviewed_at",
            "tree_notes",
            "tree_photo_url",
            "tree_created_at",
        ])
        for row in rows:
            tree_id = int(row.get("id"))
            review = review_summary.get(tree_id, {})
            writer.writerow([
                row.get("id"),
                row.get("project_id"),
                row.get("lng"),
                row.get("lat"),
                row.get("tree_origin"),
                row.get("attribution_scope"),
                row.get("count_in_planting_kpis"),
                row.get("count_in_carbon_scope"),
                row.get("custodian_id"),
                row.get("custodian_name"),
                row.get("tree_height_m"),
                row.get("species"),
                row.get("planting_date"),
                row.get("status"),
                row.get("created_by"),
                row.get("maintenance_count"),
                row.get("maintenance_done"),
                row.get("maintenance_pending"),
                row.get("maintenance_overdue"),
                row.get("maintenance_types"),
                row.get("last_maintenance_type"),
                row.get("last_maintenance_date"),
                review.get("review_submitted", 0),
                review.get("review_approved", 0),
                review.get("review_rejected", 0),
                review.get("last_review_state", ""),
                review.get("last_review_note", ""),
                review.get("last_submitted_at", ""),
                review.get("last_reviewed_at", ""),
                row.get("notes"),
                row.get("photo_url"),
                row.get("created_at"),
            ])

        writer.writerow([])
        writer.writerow(["Recent Task Review Timeline"])
        writer.writerow([
            "task_id",
            "tree_id",
            "species",
            "tree_origin",
            "attribution_scope",
            "count_in_planting_kpis",
            "count_in_carbon_scope",
            "custodian_id",
            "custodian_name",
            "tree_height_m",
            "assignee_name",
            "task_type",
            "priority",
            "status",
            "review_state",
            "due_date",
            "completed_at",
            "submitted_at",
            "reviewed_at",
            "reviewed_by",
            "review_notes",
            "delay_days",
            "delay_context",
            "evidence_status",
            "reported_tree_status",
            "tree_status",
            "photo_url",
            "notes",
        ])
        donor_rows = _build_donor_report_rows(project_id, db)
        for row in donor_rows:
            writer.writerow([
                row.get("task_id"),
                row.get("tree_id"),
                row.get("species"),
                row.get("tree_origin"),
                row.get("attribution_scope"),
                row.get("count_in_planting_kpis"),
                row.get("count_in_carbon_scope"),
                row.get("custodian_id"),
                row.get("custodian_name"),
                row.get("tree_height_m"),
                row.get("assignee_name"),
                row.get("task_type"),
                row.get("priority"),
                row.get("status"),
                row.get("review_state"),
                row.get("due_date"),
                row.get("completed_at"),
                row.get("submitted_at"),
                row.get("reviewed_at"),
                row.get("reviewed_by"),
                row.get("review_notes"),
                row.get("delay_days"),
                row.get("delay_context"),
                row.get("evidence_status"),
                row.get("reported_tree_status"),
                row.get("tree_status"),
                row.get("photo_url"),
                row.get("notes"),
            ])

    filename = f"project_{project_id}_donor_report.csv"
    return FileResponse(csv_path, media_type="text/csv", filename=filename)


@router.get("/projects/{project_id}/donor-report/pdf")
def export_donor_report_pdf(
    project_id: int,
    assignee_name: str | None = Query(default=None),
    include_photos: bool = Query(default=False),
    lng: float | None = Query(default=None),
    lat: float | None = Query(default=None),
    zoom: float | None = Query(default=None),
    bearing: float | None = Query(default=0.0),
    pitch: float | None = Query(default=0.0),
    db: Session = Depends(get_db),
):
    # Use the comprehensive map report and include donor/review details in additional pages.
    return export_work_report_pdf(
        project_id=project_id,
        assignee_name=assignee_name,
        include_photos=include_photos,
        lng=lng,
        lat=lat,
        zoom=zoom,
        bearing=bearing,
        pitch=pitch,
        db=db,
    )


@router.get("/donor/export/csv")
def export_donor_report_csv_alias(
    project_id: int = Query(...),
    db: Session = Depends(get_db),
):
    return export_donor_report_csv(project_id=project_id, db=db)


@router.get("/donor/export/pdf")
def export_donor_report_pdf_alias(
    project_id: int = Query(...),
    assignee_name: str | None = Query(default=None),
    include_photos: bool = Query(default=False),
    lng: float | None = Query(default=None),
    lat: float | None = Query(default=None),
    zoom: float | None = Query(default=None),
    bearing: float | None = Query(default=0.0),
    pitch: float | None = Query(default=0.0),
    db: Session = Depends(get_db),
):
    return export_donor_report_pdf(
        project_id=project_id,
        assignee_name=assignee_name,
        include_photos=include_photos,
        lng=lng,
        lat=lat,
        zoom=zoom,
        bearing=bearing,
        pitch=pitch,
        db=db,
    )


@router.get("/trees/{tree_id}/timeline")
def tree_timeline(tree_id: int, db: Session = Depends(get_db)):
    tree = db.execute(text("""
        SELECT
            t.id,
            t.species,
            t.planting_date,
            t.status,
            t.notes,
            t.photo_url,
            t.created_by,
            t.created_at,
            t.tree_origin,
            t.custodian_id,
            c.name AS custodian_name,
            t.custody_started_at,
            t.attribution_scope,
              t.count_in_planting_kpis,
              t.count_in_carbon_scope,
              t.source_project_id,
              t.tree_height_m,
              t.tree_age_months,
              COALESCE(t.inventory_tree_count, 1) AS inventory_tree_count,
              t.existing_area_geojson,
              t.existing_area_sqm
          FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.id = :tree_id
    """), {"tree_id": tree_id}).mappings().first()
    tasks = list_tree_tasks(tree_id, db)
    visits = db.execute(text("""
        SELECT visit_date, status, notes, photo_url, created_by, created_at
        FROM tree_visits
        WHERE tree_id = :tree_id
        ORDER BY visit_date DESC
    """), {"tree_id": tree_id}).mappings().all()
    return {
        "tree": dict(tree) if tree else None,
        "tasks": tasks,
        "visits": [dict(v) for v in visits],
    }


@router.get("/projects/{project_id}/task-stats")
def task_stats(project_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT
            COUNT(*) AS total,
            SUM(
                CASE
                    WHEN LOWER(COALESCE(status, '')) IN ('done', 'completed', 'closed')
                         AND LOWER(COALESCE(review_state, 'none')) IN ('approved', 'none')
                    THEN 1 ELSE 0
                END
            ) AS done,
            SUM(
                CASE
                    WHEN NOT (
                        LOWER(COALESCE(status, '')) IN ('done', 'completed', 'closed')
                        AND LOWER(COALESCE(review_state, 'none')) IN ('approved', 'none')
                    ) AND LOWER(COALESCE(status, 'pending')) <> 'overdue'
                    THEN 1 ELSE 0
                END
            ) AS pending,
            SUM(
                CASE
                    WHEN LOWER(COALESCE(status, 'pending')) = 'overdue'
                    THEN 1 ELSE 0
                END
            ) AS overdue
        FROM tree_tasks
        WHERE tree_id IN (SELECT id FROM trees WHERE project_id = :project_id)
    """), {"project_id": project_id}).mappings().first()
    return dict(rows)


@router.post("/work-orders")
def create_work_order(
    db: Session = Depends(get_db),
    project_id: int = Body(...),
    assignee_name: str = Body(...),
    work_type: str = Body(...),
    target_trees: int = Body(default=0),
    species_allocations: list[dict] | None = Body(default=None),
    maintenance_schedule: str = Body(default=""),
    auto_assign_first_cycle_maintenance: bool = Body(default=False),
    due_date: str | None = Body(default=None),
    area_enabled: bool = Body(default=False),
    area_label: str | None = Body(default=None),
    area_geojson: dict | str | None = Body(default=None),
    allow_existing_tree_area_reuse: bool = Body(default=False),
):
    if work_type not in {"planting", "maintenance"}:
        raise HTTPException(status_code=400, detail="Invalid work_type")
    assignee_clean = (assignee_name or "").strip()
    if not assignee_clean:
        raise HTTPException(status_code=400, detail="Assignee name required")
    if area_enabled and work_type != "planting":
        raise HTTPException(status_code=400, detail="Area assignment is only available for planting orders")
    normalized_species_allocations = _normalize_species_allocations(species_allocations)
    if work_type != "planting" and normalized_species_allocations:
        raise HTTPException(status_code=400, detail="Species-based allocation is only available for planting orders")
    if work_type == "planting" and normalized_species_allocations:
        target_trees = int(sum(int(item.get("count") or 0) for item in normalized_species_allocations))
    if work_type == "planting" and int(target_trees or 0) <= 0:
        raise HTTPException(status_code=400, detail="Target trees must be greater than 0 for planting orders")
    auto_assign_first_cycle_maintenance = bool(auto_assign_first_cycle_maintenance) and work_type == "planting"

    normalized_area_geojson = _normalize_work_area_geojson(area_geojson)
    area_label_clean = (area_label or "").strip() or None
    if area_enabled and normalized_area_geojson is None:
        raise HTTPException(
            status_code=400,
            detail="Invalid planting area polygon. Draw a valid Polygon/MultiPolygon before assigning.",
        )
    if not area_enabled:
        area_label_clean = None
        normalized_area_geojson = None
        allow_existing_tree_area_reuse = False
    else:
        allow_existing_tree_area_reuse = bool(allow_existing_tree_area_reuse) and work_type == "planting"

    can_accumulate = (
        work_type == "planting"
        and not area_enabled
        and normalized_area_geojson is None
        and not normalized_species_allocations
    )
    existing_id = None
    if can_accumulate:
        existing_id = db.execute(
            text(
                """
                SELECT id
                FROM green_work_orders
                WHERE project_id = :project_id
                  AND LOWER(TRIM(assignee_name)) = LOWER(TRIM(:assignee_name))
                  AND work_type = :work_type
                  AND COALESCE(area_enabled, FALSE) = FALSE
                  AND area_geojson IS NULL
                  AND COALESCE(species_allocations, '[]'::jsonb) = '[]'::jsonb
                  AND COALESCE(auto_assign_first_cycle_maintenance, FALSE) = :auto_assign_first_cycle_maintenance
                  AND LOWER(COALESCE(status, 'assigned')) NOT IN ('done', 'completed', 'closed', 'cancelled')
                ORDER BY created_at DESC, id DESC
                LIMIT 1
                """
            ),
            {
                "project_id": project_id,
                "assignee_name": assignee_clean,
                "work_type": work_type,
                "auto_assign_first_cycle_maintenance": auto_assign_first_cycle_maintenance,
            },
        ).scalar()

    if existing_id and work_type == "planting":
        row = db.execute(
            text(
                """
                UPDATE green_work_orders
                SET target_trees = COALESCE(target_trees, 0) + :target_trees,
                    due_date = COALESCE(:due_date, due_date),
                    last_update = NOW()
                WHERE id = :id
                RETURNING id
                """
            ),
            {"id": int(existing_id), "target_trees": int(target_trees or 0), "due_date": due_date},
        ).scalar()
        action_name = "work_order_accumulated"
    else:
        row = db.execute(text("""
            INSERT INTO green_work_orders (
                project_id, assignee_name, work_type, target_trees, species_allocations, maintenance_schedule,
                auto_assign_first_cycle_maintenance, due_date,
                area_enabled, area_label, area_geojson, allow_existing_tree_area_reuse
            )
            VALUES (
                :project_id, :assignee_name, :work_type, :target_trees, CAST(:species_allocations AS JSONB), :maintenance_schedule,
                :auto_assign_first_cycle_maintenance, :due_date,
                :area_enabled, :area_label, CAST(:area_geojson AS JSONB), :allow_existing_tree_area_reuse
            )
            RETURNING id
        """), {
            "project_id": project_id,
            "assignee_name": assignee_clean,
            "work_type": work_type,
            "target_trees": target_trees,
            "species_allocations": _safe_json(normalized_species_allocations) if normalized_species_allocations else None,
            "maintenance_schedule": maintenance_schedule or None,
            "auto_assign_first_cycle_maintenance": auto_assign_first_cycle_maintenance,
            "due_date": due_date,
            "area_enabled": bool(area_enabled),
            "area_label": area_label_clean,
            "area_geojson": _safe_json(normalized_area_geojson) if normalized_area_geojson else None,
            "allow_existing_tree_area_reuse": bool(allow_existing_tree_area_reuse),
        }).scalar()
        action_name = "work_order_created"
    _log_audit_event(
        db,
        project_id=project_id,
        entity_type="work_order",
        entity_id=int(row),
        action=action_name,
        actor=assignee_clean,
        details={
            "work_type": work_type,
            "target_trees": target_trees,
            "species_allocations": normalized_species_allocations,
            "maintenance_schedule": maintenance_schedule or None,
            "auto_assign_first_cycle_maintenance": auto_assign_first_cycle_maintenance,
            "due_date": due_date,
            "area_enabled": bool(area_enabled),
            "area_label": area_label_clean,
            "area_geojson": normalized_area_geojson,
            "allow_existing_tree_area_reuse": bool(allow_existing_tree_area_reuse),
        },
    )
    db.commit()
    return {"id": row}


@router.get("/work-orders")
def list_work_orders(
    project_id: int,
    assignee_name: str | None = None,
    db: Session = Depends(get_db),
):
    rows = db.execute(text("""
        WITH tree_counts AS (
            SELECT LOWER(TRIM(created_by)) AS assignee_key, COUNT(*) AS planted
            FROM trees
            WHERE project_id = :project_id
              AND LOWER(COALESCE(tree_origin, 'new_planting')) = 'new_planting'
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
              AND LOWER(REPLACE(REPLACE(COALESCE(status, ''), '-', '_'), ' ', '_')) <> 'pending_planting'
            GROUP BY created_by
        )
        SELECT o.id, o.project_id, o.assignee_name, o.work_type, o.target_trees,
               o.species_allocations,
               o.maintenance_schedule,
               COALESCE(o.auto_assign_first_cycle_maintenance, FALSE) AS auto_assign_first_cycle_maintenance,
               o.due_date, o.status,
               COALESCE(o.area_enabled, FALSE) AS area_enabled,
               o.area_label, o.area_geojson,
               COALESCE(o.allow_existing_tree_area_reuse, FALSE) AS allow_existing_tree_area_reuse,
               CASE
                   WHEN o.work_type = 'planting' THEN COALESCE(t.planted, 0)
                   ELSE o.planted_count
               END AS planted_count,
               o.last_update, o.created_at
        FROM green_work_orders o
        LEFT JOIN tree_counts t ON t.assignee_key = LOWER(TRIM(o.assignee_name))
        WHERE o.project_id = :project_id
          AND (:assignee_name IS NULL OR LOWER(TRIM(o.assignee_name)) = LOWER(TRIM(:assignee_name)))
        ORDER BY o.created_at DESC
    """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()
    payload: list[dict] = []
    for row in rows:
        item = dict(row)
        raw_geojson = item.get("area_geojson")
        if isinstance(raw_geojson, str):
            try:
                item["area_geojson"] = json.loads(raw_geojson)
            except Exception:
                item["area_geojson"] = None
        raw_species_allocations = item.get("species_allocations")
        if isinstance(raw_species_allocations, str):
            try:
                raw_species_allocations = json.loads(raw_species_allocations)
            except Exception:
                raw_species_allocations = []
        item["species_allocations"] = _normalize_species_allocations(raw_species_allocations)
        item["auto_assign_first_cycle_maintenance"] = bool(item.get("auto_assign_first_cycle_maintenance"))
        item["area_enabled"] = bool(item.get("area_enabled")) and item.get("work_type") == "planting"
        item["allow_existing_tree_area_reuse"] = bool(item.get("allow_existing_tree_area_reuse")) and item["area_enabled"]
        if not item["area_enabled"]:
            item["area_geojson"] = None
            item["area_label"] = None
            item["allow_existing_tree_area_reuse"] = False
        payload.append(item)
    return payload


@router.patch("/work-orders/{work_id}")
def update_work_order(
    work_id: int,
    db: Session = Depends(get_db),
    status: str | None = Body(default=None),
    planted_count: int | None = Body(default=None),
):
    # Auto-calc planted_count from trees created by assignee for planting orders.
    row = db.execute(text("""
        SELECT id, project_id, assignee_name, work_type, status
        FROM green_work_orders
        WHERE id = :work_id
    """), {"work_id": work_id}).mappings().first()

    planted_value = planted_count
    if row and row["work_type"] == "planting":
        planted_value = db.execute(text("""
            SELECT COUNT(*) FROM trees
            WHERE project_id = :project_id AND created_by = :assignee_name
              AND LOWER(COALESCE(tree_origin, 'new_planting')) = 'new_planting'
              AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
              AND LOWER(REPLACE(REPLACE(COALESCE(status, ''), '-', '_'), ' ', '_')) <> 'pending_planting'
        """), {"project_id": row["project_id"], "assignee_name": row["assignee_name"]}).scalar()

    existing_status = row.get("status") if row else None
    db.execute(text("""
        UPDATE green_work_orders
        SET status = COALESCE(:status, status),
            planted_count = COALESCE(:planted_count, planted_count),
            last_update = NOW()
        WHERE id = :work_id
    """), {
        "status": status,
        "planted_count": planted_value,
        "work_id": work_id,
    })
    if row:
        _log_audit_event(
            db,
            project_id=int(row["project_id"]),
            entity_type="work_order",
            entity_id=work_id,
            action="work_order_updated",
            actor=row.get("assignee_name"),
            details={
                "status": status,
                "planted_count": planted_value,
                "work_type": row.get("work_type"),
                "existing_status": existing_status,
            },
        )
    db.commit()
    return {"status": "ok"}


@router.get("/work-stats")
def work_stats(project_id: int, db: Session = Depends(get_db)):
    orders = db.execute(text("""
        SELECT assignee_name,
               COUNT(*) AS orders,
               SUM(target_trees) AS target_trees
        FROM green_work_orders
        WHERE project_id = :project_id
        GROUP BY assignee_name
    """), {"project_id": project_id}).mappings().all()

    tree_counts = db.execute(text("""
        SELECT created_by AS assignee_name, COUNT(*) AS trees_logged
        FROM trees
        WHERE project_id = :project_id
          AND LOWER(COALESCE(tree_origin, 'new_planting')) = 'new_planting'
          AND COALESCE(count_in_planting_kpis, TRUE) = TRUE
        GROUP BY created_by
    """), {"project_id": project_id}).mappings().all()
    tree_map = {r["assignee_name"]: r["trees_logged"] for r in tree_counts}
    merged = []
    for r in orders:
        row = dict(r)
        row["planted_count"] = tree_map.get(row.get("assignee_name"), 0)
        merged.append(row)

    maintenance_by_assignee = db.execute(text("""
        SELECT t.assignee_name,
               COUNT(*) AS maintenance_total,
               SUM(
                   CASE
                       WHEN LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                            AND LOWER(COALESCE(t.review_state, 'none')) IN ('approved', 'none')
                       THEN 1 ELSE 0
                   END
               ) AS maintenance_done,
               SUM(
                   CASE
                       WHEN NOT (
                           LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                           AND LOWER(COALESCE(t.review_state, 'none')) IN ('approved', 'none')
                       ) AND LOWER(COALESCE(t.status, 'pending')) <> 'overdue'
                       THEN 1 ELSE 0
                   END
               ) AS maintenance_pending,
               SUM(CASE WHEN t.status = 'overdue' THEN 1 ELSE 0 END) AS maintenance_overdue,
               COALESCE(STRING_AGG(DISTINCT t.task_type, ', ' ORDER BY t.task_type), '') AS maintenance_types,
               MAX(COALESCE(t.completed_at::date, t.due_date, t.created_at::date)) AS last_maintenance_date
        FROM tree_tasks t
        JOIN trees tr ON tr.id = t.tree_id
        WHERE tr.project_id = :project_id
        GROUP BY t.assignee_name
        ORDER BY t.assignee_name
    """), {"project_id": project_id}).mappings().all()

    maintenance_by_type = db.execute(text("""
        SELECT t.assignee_name,
               t.task_type,
               COUNT(*) AS maintenance_times,
               MAX(COALESCE(t.completed_at::date, t.due_date, t.created_at::date)) AS last_maintenance_date
        FROM tree_tasks t
        JOIN trees tr ON tr.id = t.tree_id
        WHERE tr.project_id = :project_id
        GROUP BY t.assignee_name, t.task_type
        ORDER BY t.assignee_name, maintenance_times DESC, t.task_type
    """), {"project_id": project_id}).mappings().all()

    return {
        "orders": merged,
        "trees_by_user": [dict(r) for r in tree_counts],
        "maintenance_by_assignee": [dict(r) for r in maintenance_by_assignee],
        "maintenance_by_type": [dict(r) for r in maintenance_by_type],
    }


@router.get("/work-stats/export/csv")
def export_work_stats_csv(project_id: int, db: Session = Depends(get_db)):
    stats = work_stats(project_id, db)
    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_csv = tempfile.NamedTemporaryFile(suffix="_work_stats.csv", delete=False)
    csv_path = tmp_csv.name
    tmp_csv.close()

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = _excel_csv_writer(f)
        writer.writerow(["assignee", "orders", "target_trees", "planted_count"])
        for r in stats["orders"]:
            writer.writerow([
                r.get("assignee_name", ""),
                r.get("orders", 0),
                r.get("target_trees", 0),
                r.get("planted_count", 0),
            ])

    filename = f"project_{project_id}_work_stats.csv"
    return FileResponse(csv_path, media_type="text/csv", filename=filename)


@router.get("/work-stats/export/pdf")
def export_work_stats_pdf(project_id: int, db: Session = Depends(get_db)):
    project = get_project(project_id, db)
    stats = work_stats(project_id, db)
    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_work_report.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()
    render_green_work_report_pdf(pdf_path, project, stats)
    filename = f"project_{project_id}_work_report.pdf"
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


@router.get("/projects/{project_id}/custodians/export/pdf")
def export_custodian_report_pdf(
    project_id: int,
    include_photos: bool = Query(default=True),
    db: Session = Depends(get_db),
):
    project = get_project(project_id, db)
    summary = db.execute(
        text(
            """
            SELECT
                (SELECT COUNT(*) FROM green_custodians WHERE project_id = :project_id) AS total_custodians,
                (SELECT COUNT(*) FROM green_custodians WHERE project_id = :project_id
                    AND LOWER(COALESCE(verification_status, 'pending')) = 'verified') AS verified_custodians,
                (SELECT COUNT(*) FROM green_distribution_events WHERE project_id = :project_id) AS distribution_events,
                (SELECT COALESCE(SUM(quantity), 0) FROM green_distribution_events WHERE project_id = :project_id) AS seedlings_distributed,
                (SELECT COUNT(*) FROM green_distribution_allocations WHERE project_id = :project_id) AS distribution_allocations,
                (SELECT COALESCE(SUM(supervision_target), 0) FROM green_distribution_allocations WHERE project_id = :project_id) AS supervision_target_total,
                (
                    SELECT COUNT(*)
                    FROM tree_tasks t
                    JOIN green_distribution_allocations a ON a.id = t.distribution_allocation_id
                    WHERE a.project_id = :project_id
                      AND LOWER(COALESCE(t.task_type, '')) = :supervision_type
                ) AS supervision_assigned,
                (
                    SELECT COALESCE(SUM(
                        CASE
                            WHEN LOWER(COALESCE(t.review_state, 'none')) = 'approved'
                                 OR (
                                     LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                                     AND LOWER(COALESCE(t.review_state, 'none')) = 'none'
                                 )
                            THEN 1 ELSE 0
                        END
                    ), 0)
                    FROM tree_tasks t
                    JOIN green_distribution_allocations a ON a.id = t.distribution_allocation_id
                    WHERE a.project_id = :project_id
                      AND LOWER(COALESCE(t.task_type, '')) = :supervision_type
                ) AS supervision_done,
                (SELECT COUNT(*) FROM trees
                    WHERE project_id = :project_id
                      AND LOWER(COALESCE(tree_origin, 'new_planting')) <> 'new_planting') AS existing_trees
            """
        ),
        {"project_id": project_id, "supervision_type": SUPERVISION_TASK_TYPE},
    ).mappings().first()

    custodians = db.execute(
        text(
            """
            SELECT
                c.id,
                c.project_id,
                c.custodian_type,
                c.name,
                c.contact_person,
                c.phone,
                c.alt_phone,
                c.email,
                c.address_text,
                c.local_government,
                c.community_name,
                c.verification_status,
                c.notes,
                c.created_at,
                COALESCE(SUM(a.quantity_allocated), 0) AS allocated_seedlings,
                COALESCE(SUM(a.supervision_target), 0) AS supervision_target,
                COUNT(DISTINCT t.id) AS linked_trees,
                COALESCE(SUM(CASE
                      WHEN LOWER(COALESCE(t.status, '')) IN ('alive', 'healthy') THEN 1
                      ELSE 0
                    END), 0) AS healthy_trees,
                COALESCE(MAX(s.supervision_assigned), 0) AS supervision_assigned,
                COALESCE(MAX(s.supervision_done), 0) AS supervision_done
            FROM green_custodians c
            LEFT JOIN green_distribution_allocations a ON a.custodian_id = c.id
            LEFT JOIN trees t ON t.custodian_id = c.id AND t.project_id = c.project_id
            LEFT JOIN (
                SELECT
                    t.custodian_id,
                    COUNT(*) AS supervision_assigned,
                    SUM(
                        CASE
                            WHEN LOWER(COALESCE(t.review_state, 'none')) = 'approved'
                                 OR (
                                     LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                                     AND LOWER(COALESCE(t.review_state, 'none')) = 'none'
                                 )
                            THEN 1 ELSE 0
                        END
                    ) AS supervision_done
                FROM tree_tasks t
                JOIN trees tr ON tr.id = t.tree_id
                WHERE tr.project_id = :project_id
                  AND LOWER(COALESCE(t.task_type, '')) = :supervision_type
                  AND t.custodian_id IS NOT NULL
                GROUP BY t.custodian_id
            ) s ON s.custodian_id = c.id
            WHERE c.project_id = :project_id
            GROUP BY
                c.id,
                c.project_id,
                c.custodian_type,
                c.name,
                c.contact_person,
                c.phone,
                c.alt_phone,
                c.email,
                c.address_text,
                c.local_government,
                c.community_name,
                c.verification_status,
                c.notes,
                c.created_at
            ORDER BY c.created_at DESC, c.id DESC
            """
        ),
        {"project_id": project_id, "supervision_type": SUPERVISION_TASK_TYPE},
    ).mappings().all()

    distribution_events = db.execute(
        text(
            """
            SELECT
                id,
                event_date,
                species,
                quantity,
                source_batch_ref,
                distributed_by,
                notes,
                created_at
            FROM green_distribution_events
            WHERE project_id = :project_id
            ORDER BY event_date DESC, id DESC
            LIMIT 500
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    existing_trees = db.execute(
        text(
            """
            SELECT
                t.id,
                t.species,
                t.tree_height_m,
                t.status,
                t.created_by,
                t.created_at,
                c.name AS custodian_name
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id
              AND LOWER(COALESCE(t.tree_origin, 'new_planting')) <> 'new_planting'
            ORDER BY t.created_at DESC, t.id DESC
            LIMIT 800
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    supervision_photo_rows: list[dict] = []
    if include_photos:
        supervision_photo_task_rows = db.execute(
            text(
                """
                SELECT
                    t.id AS task_id,
                    t.tree_id,
                    t.task_type,
                    t.assignee_name,
                    t.notes,
                    t.photo_url,
                    t.photo_urls,
                    t.supervision_visit_no,
                    t.supervision_total_visits,
                    t.created_at,
                    t.submitted_at,
                    t.reviewed_at,
                    tr.species,
                    tr.status,
                    tr.planting_date,
                    c.name AS custodian_name
                FROM tree_tasks t
                JOIN trees tr ON tr.id = t.tree_id
                LEFT JOIN green_custodians c ON c.id = COALESCE(t.custodian_id, tr.custodian_id)
                WHERE tr.project_id = :project_id
                  AND LOWER(COALESCE(t.task_type, '')) = :supervision_type
                  AND (
                        COALESCE(TRIM(t.photo_url), '') <> ''
                        OR COALESCE(
                            CASE
                                WHEN jsonb_typeof(t.photo_urls) = 'array' THEN jsonb_array_length(t.photo_urls)
                                ELSE 0
                            END,
                            0
                        ) > 0
                      )
                ORDER BY COALESCE(t.reviewed_at, t.submitted_at, t.created_at) DESC, t.id DESC
                LIMIT 1200
                """
            ),
            {"project_id": project_id, "supervision_type": SUPERVISION_TASK_TYPE},
        ).mappings().all()
        for row in supervision_photo_task_rows:
            photo_urls = _normalize_photo_urls(row.get("photo_urls"))
            photo_url = str(row.get("photo_url") or "").strip()
            if photo_url and photo_url not in photo_urls:
                photo_urls.append(photo_url)
            for idx, url in enumerate(photo_urls):
                supervision_photo_rows.append(
                    {
                        "id": row.get("tree_id"),
                        "species": row.get("species"),
                        "status": row.get("status"),
                        "planting_date": row.get("planting_date"),
                        "created_by": row.get("assignee_name"),
                        "custodian_name": row.get("custodian_name"),
                        "photo_url": url,
                        "visit_label": f"Visit {int(row.get('supervision_visit_no') or 0)}/{int(row.get('supervision_total_visits') or 0)}",
                        "task_id": row.get("task_id"),
                        "photo_index": idx + 1,
                    }
                )

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_custodian_report.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()

    render_green_custodian_report_pdf(
        output_path=pdf_path,
        project=project,
        summary=dict(summary or {}),
        custodians=[dict(row) for row in custodians],
        distribution_events=[dict(row) for row in distribution_events],
        existing_trees=[dict(row) for row in existing_trees],
        supervision_photo_rows=supervision_photo_rows,
    )
    filename = f"project_{project_id}_custodian_report.pdf"
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


@router.get("/projects/{project_id}/tasks/export/csv")
def export_tasks_csv(project_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT t.id, t.task_type, t.assignee_name, t.due_date, t.priority, t.status,
               t.review_state, t.submitted_at, t.reviewed_at, t.reviewed_by, t.review_notes,
               t.model_season, t.source_task_id, t.auto_generated,
               t.notes, t.photo_url, t.created_at, t.completed_at
        FROM tree_tasks t
        JOIN trees tr ON tr.id = t.tree_id
        WHERE tr.project_id = :project_id
        ORDER BY t.created_at DESC
    """), {"project_id": project_id}).mappings().all()

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_csv = tempfile.NamedTemporaryFile(suffix="_tasks.csv", delete=False)
    csv_path = tmp_csv.name
    tmp_csv.close()

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = _excel_csv_writer(f)
        writer.writerow([
            "task_id", "task_type", "assignee_name", "due_date", "priority", "status",
            "review_state", "submitted_at", "reviewed_at", "reviewed_by", "review_notes",
            "model_season", "source_task_id", "auto_generated",
            "notes", "photo_url", "created_at", "completed_at"
        ])
        for r in rows:
            writer.writerow([
                r["id"], r["task_type"], r["assignee_name"], r["due_date"], r["priority"], r["status"],
                r["review_state"], r["submitted_at"], r["reviewed_at"], r["reviewed_by"], r["review_notes"],
                r["model_season"], r["source_task_id"], r["auto_generated"],
                r["notes"], r["photo_url"], r["created_at"], r["completed_at"],
            ])

    filename = f"project_{project_id}_tasks.csv"
    return FileResponse(csv_path, media_type="text/csv", filename=filename)


@router.get("/projects/{project_id}/tasks/export/pdf")
def export_tasks_pdf(project_id: int, db: Session = Depends(get_db)):
    project = get_project(project_id, db)
    stats = task_stats(project_id, db)
    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_tasks_report.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()
    render_green_work_report_pdf(pdf_path, project, {"orders": [], **stats})
    filename = f"project_{project_id}_tasks_report.pdf"
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


def _estimate_tree_co2_height_aware_kg(
    species: str | None,
    age_years: float,
    tree_height_m: float | None,
) -> tuple[float, bool]:
    """Estimate current tree CO2, preferring measured height where available."""
    if age_years <= 0:
        return 0.0, False

    try:
        measured_height = float(tree_height_m) if tree_height_m is not None else None
    except Exception:
        measured_height = None
    if measured_height is None or measured_height <= 0:
        return estimate_tree_co2_kg(species, age_years), False

    try:
        params = _get_species_params(species)
        dbh_cm = project_dbh(params, age_years)
        agb = calculate_agb_chave(params["wood_density"], dbh_cm, measured_height)
        if agb <= 0:
            return estimate_tree_co2_kg(species, age_years), False
        bgb = agb * float(params.get("root_shoot_ratio", 0.24))
        total_biomass = agb + bgb
        carbon_kg = total_biomass * float(params.get("carbon_fraction", 0.47))
        co2_kg = carbon_kg * (44.0 / 12.0)
        return round(max(co2_kg, 0.0), 2), True
    except Exception:
        return estimate_tree_co2_kg(species, age_years), False


def _coerce_tree_age_months(value: object) -> float | None:
    try:
        if value is None or value == "":
            return None
        months = float(value)
    except Exception:
        return None
    if months < 0:
        return None
    return months


def _infer_tree_age_days_for_maintenance(tree_row: dict, as_of_date: date | None = None) -> int | None:
    today = as_of_date or date.today()
    origin_key = _normalize_tree_origin(tree_row.get("tree_origin"))
    age_months = _coerce_tree_age_months(tree_row.get("tree_age_months"))
    if origin_key == "existing_inventory" and age_months is not None:
        capture_ref = (
            _parse_date_value(tree_row.get("created_at"))
            or _parse_date_value(tree_row.get("submitted_at"))
            or _parse_date_value(tree_row.get("reviewed_at"))
        )
        elapsed_days = max(_day_diff(today, capture_ref), 0) if capture_ref is not None else 0
        base_days = int(round(age_months * 30.4375))
        return max(base_days + elapsed_days, 0)

    planting_ref = _parse_date_value(tree_row.get("planting_date"))
    if planting_ref is not None:
        return max(_day_diff(today, planting_ref), 0)

    if age_months is None:
        return None

    capture_ref = (
        _parse_date_value(tree_row.get("created_at"))
        or _parse_date_value(tree_row.get("submitted_at"))
        or _parse_date_value(tree_row.get("reviewed_at"))
    )
    elapsed_days = max(_day_diff(today, capture_ref), 0) if capture_ref is not None else 0
    base_days = int(round(age_months * 30.4375))
    return max(base_days + elapsed_days, 0)


def _infer_tree_age_for_carbon(tree_row: dict) -> tuple[float, str]:
    """Infer tree age in years with existing-tree age-months priority."""
    tree = dict(tree_row)
    today = date.today()
    origin_key = _normalize_tree_origin(tree.get("tree_origin"))
    age_months = _coerce_tree_age_months(tree.get("tree_age_months"))

    # Existing Tree records often store a survey/reference date in planting_date.
    # When explicit age months is provided, prefer it to avoid zero-age CO2 from a recent reference date.
    if origin_key == "existing_inventory" and age_months is not None:
        elapsed_years = 0.0
        capture_ref = _parse_date_value(tree.get("created_at")) or _parse_date_value(tree.get("submitted_at")) or _parse_date_value(tree.get("reviewed_at"))
        if capture_ref is not None:
            elapsed_years = tree_age_years(capture_ref, today)
        age_years = max((age_months / 12.0) + elapsed_years, 0.0)
        return age_years, "tree_age_months"

    planting_ref = _parse_date_value(tree.get("planting_date"))
    if planting_ref is not None:
        return tree_age_years(planting_ref, today), "planting_date"

    if age_months is None:
        ref_date, ref_source = _infer_tree_reference_date(tree)
        if ref_date is None:
            return 0.0, "none"
        return tree_age_years(ref_date, today), ref_source

    elapsed_years = 0.0
    capture_ref = _parse_date_value(tree.get("created_at")) or _parse_date_value(tree.get("submitted_at")) or _parse_date_value(tree.get("reviewed_at"))
    if capture_ref is not None:
        elapsed_years = tree_age_years(capture_ref, today)
    age_years = max((age_months / 12.0) + elapsed_years, 0.0)
    return age_years, "tree_age_months"


def _build_tree_carbon_metrics(
    tree_row: dict,
    projection_years: int = 40,
    enforce_carbon_scope: bool = False,
) -> dict:
    """Per-tree CO2 metrics with measured-height preference for current stock."""
    tree = dict(tree_row)
    age_years, ref_source = _infer_tree_age_for_carbon(tree)
    species = tree.get("species")
    params = _get_species_params(species)
    status_key = _normalize_tree_status(tree.get("status"))
    is_alive = status_key not in DEAD_TREE_STATUSES
    raw_carbon_scope = tree.get("count_in_carbon_scope")
    in_carbon_scope = True if raw_carbon_scope is None else bool(raw_carbon_scope)
    scope_applies = in_carbon_scope or not enforce_carbon_scope
    try:
        inventory_tree_count = int(tree.get("inventory_tree_count") or 1)
    except Exception:
        inventory_tree_count = 1
    inventory_tree_count = max(inventory_tree_count, 1)

    current_co2_per_tree_kg = 0.0
    annual_co2_per_tree_kg = 0.0
    lifetime_co2_per_tree_kg = 0.0
    height_used = False

    if scope_applies and age_years > 0:
        current_co2_per_tree_kg, height_used = _estimate_tree_co2_height_aware_kg(
            species,
            age_years,
            tree.get("tree_height_m"),
        )
        if is_alive:
            if height_used and age_years > 1:
                previous_modeled = estimate_tree_co2_kg(species, max(age_years - 1.0, 0.0))
                annual_co2_per_tree_kg = round(max(current_co2_per_tree_kg - previous_modeled, 0.0), 2)
            else:
                annual_co2_per_tree_kg = estimate_annual_co2_kg(species, age_years)
            lifetime_co2_per_tree_kg = estimate_lifetime_co2_kg(species, projection_years)

    current_co2_kg = max(current_co2_per_tree_kg, 0.0) * inventory_tree_count
    annual_co2_kg = max(annual_co2_per_tree_kg, 0.0) * inventory_tree_count
    lifetime_co2_kg = max(lifetime_co2_per_tree_kg, 0.0) * inventory_tree_count

    return {
        "species_matched": params.get("label", "Unknown"),
        "growth_class": params.get("growth_class", "medium"),
        "age_years": round(age_years, 2),
        "age_source": ref_source,
        "inventory_tree_count": inventory_tree_count,
        "current_co2_per_tree_kg": round(max(current_co2_per_tree_kg, 0.0), 2),
        "annual_co2_per_tree_kg": round(max(annual_co2_per_tree_kg, 0.0), 2),
        "lifetime_co2_per_tree_kg": round(max(lifetime_co2_per_tree_kg, 0.0), 2),
        "current_co2_kg": round(current_co2_kg, 2),
        "annual_co2_kg": round(max(annual_co2_kg, 0.0), 2),
        "lifetime_co2_kg": round(max(lifetime_co2_kg, 0.0), 2),
        "lifetime_co2_tonnes": round(max(lifetime_co2_kg, 0.0) / 1000.0, 4),
        "height_used_for_co2": height_used,
        "co2_height_source": "measured_tree_height_m" if height_used else "modeled_height",
        "co2_in_scope": in_carbon_scope,
        "is_alive_for_co2": is_alive,
        "co2_methodology": (
            "IPCC Tier 1 + Chave et al. (2014); measured height used for current stock where available, "
            "otherwise modeled height from species curve"
        ),
    }


def _fetch_existing_tree_export_rows(project_id: int, db: Session) -> list[dict]:
    rows = db.execute(
        text(
            """
            SELECT
                t.id,
                t.project_id,
                t.project_tree_no,
                t.species,
                t.planting_date,
                t.status,
                t.notes,
                t.photo_url,
                t.photo_urls,
                t.created_by,
                t.created_at,
                t.tree_origin,
                t.attribution_scope,
                t.count_in_planting_kpis,
                t.count_in_carbon_scope,
                t.source_project_id,
                t.tree_age_months,
                t.custodian_id,
                c.name AS custodian_name,
                t.tree_height_m,
                COALESCE(t.inventory_tree_count, 1) AS inventory_tree_count,
                t.existing_area_geojson,
                t.existing_area_sqm,
                ST_X(t.geom) AS lng,
                ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id
              AND (
                    (
                        LOWER(REPLACE(COALESCE(t.tree_origin, ''), ' ', '_')) <> ''
                        AND LOWER(REPLACE(COALESCE(t.tree_origin, ''), ' ', '_')) <> 'new_planting'
                    )
                    OR LOWER(REPLACE(COALESCE(t.attribution_scope, ''), ' ', '_')) = 'monitor_only'
                    OR COALESCE(t.count_in_planting_kpis, TRUE) = FALSE
                    OR COALESCE(t.source_project_id, 0) > 0
                  )
            ORDER BY t.created_at DESC, t.id DESC
            """
        ),
        {"project_id": project_id},
    ).mappings().all()

    merged_rows = _attach_maintenance_to_tree_rows(rows, _maintenance_summary_by_tree(project_id, db))
    review_summary = _review_summary_by_tree(project_id, db)

    enriched: list[dict] = []
    for row in merged_rows:
        item = dict(row)
        item["photo_urls"] = _normalize_photo_urls(item.get("photo_urls"))
        try:
            item["inventory_tree_count"] = max(int(item.get("inventory_tree_count") or 1), 1)
        except Exception:
            item["inventory_tree_count"] = 1
        try:
            item["existing_area_sqm"] = (
                round(float(item.get("existing_area_sqm")), 2) if item.get("existing_area_sqm") is not None else None
            )
        except Exception:
            item["existing_area_sqm"] = None
        item["existing_area_ha"] = (
            round(float(item.get("existing_area_sqm") or 0.0) / 10000.0, 4)
            if item.get("existing_area_sqm") is not None
            else None
        )
        item.update(review_summary.get(int(item.get("id") or 0), {}))
        item.update(_build_tree_carbon_metrics(item, projection_years=40, enforce_carbon_scope=True))
        enriched.append(item)
    return enriched


def _summarize_existing_tree_export_rows(rows: list[dict]) -> dict:
    total_rows = len(rows)
    total_existing_trees = 0
    alive_rows = 0
    dead_rows = 0
    attention_rows = 0
    pending_rows = 0
    rows_with_height = 0
    rows_missing_height = 0
    trees_missing_age_data = 0
    trees_with_fallback_age = 0
    carbon_scope_rows = 0
    carbon_excluded_rows = 0
    current_co2_kg = 0.0
    annual_co2_kg = 0.0
    lifetime_co2_kg = 0.0
    total_existing_area_sqm = 0.0
    rows_with_existing_area = 0
    species_agg: dict[str, dict] = {}

    for row in rows:
        try:
            tree_units = max(int(row.get("inventory_tree_count") or 1), 1)
        except Exception:
            tree_units = 1
        total_existing_trees += tree_units

        status_key = _normalize_tree_status(row.get("status"))
        if status_key in HEALTHY_TREE_STATUSES:
            alive_rows += tree_units
        elif status_key in DEAD_TREE_STATUSES:
            dead_rows += tree_units
        elif status_key in ATTENTION_TREE_STATUSES:
            attention_rows += tree_units
        elif status_key == "pending_planting":
            pending_rows += tree_units

        try:
            height_val = float(row.get("tree_height_m")) if row.get("tree_height_m") is not None else None
        except Exception:
            height_val = None
        if height_val is not None and height_val > 0:
            rows_with_height += 1
        else:
            rows_missing_height += 1

        age_source = str(row.get("age_source") or "")
        if age_source == "none":
            trees_missing_age_data += 1
        elif age_source and age_source != "planting_date":
            trees_with_fallback_age += 1

        in_scope = bool(row.get("co2_in_scope", True))
        if in_scope:
            carbon_scope_rows += tree_units
            current_co2_kg += float(row.get("current_co2_kg") or 0.0)
            annual_co2_kg += float(row.get("annual_co2_kg") or 0.0)
            lifetime_co2_kg += float(row.get("lifetime_co2_kg") or 0.0)
        else:
            carbon_excluded_rows += tree_units

        try:
            area_sqm = float(row.get("existing_area_sqm")) if row.get("existing_area_sqm") is not None else None
        except Exception:
            area_sqm = None
        if area_sqm is not None and area_sqm > 0:
            total_existing_area_sqm += area_sqm
            rows_with_existing_area += 1

        species_label = str(row.get("species") or "").strip() or str(row.get("species_matched") or "Unknown")
        species_key = species_label.lower()
        if species_key not in species_agg:
            species_agg[species_key] = {"species": species_label, "count": 0, "co2_kg": 0.0}
        species_agg[species_key]["count"] += tree_units
        if in_scope:
            species_agg[species_key]["co2_kg"] += float(row.get("current_co2_kg") or 0.0)

    top_species = sorted(
        species_agg.values(),
        key=lambda item: (float(item.get("co2_kg") or 0.0), int(item.get("count") or 0)),
        reverse=True,
    )[:12]
    for item in top_species:
        item["co2_kg"] = round(float(item.get("co2_kg") or 0.0), 2)

    return {
        "total_existing_trees": total_existing_trees,
        "total_existing_rows": total_rows,
        "alive_trees": alive_rows,
        "dead_trees": dead_rows,
        "attention_trees": attention_rows,
        "pending_trees": pending_rows,
        "rows_with_height": rows_with_height,
        "rows_missing_height": rows_missing_height,
        "rows_with_existing_area": rows_with_existing_area,
        "total_existing_area_sqm": round(total_existing_area_sqm, 2),
        "total_existing_area_ha": round(total_existing_area_sqm / 10000.0, 4),
        "trees_missing_age_data": trees_missing_age_data,
        "trees_with_fallback_age": trees_with_fallback_age,
        "carbon_scope_rows": carbon_scope_rows,
        "carbon_excluded_rows": carbon_excluded_rows,
        "current_co2_kg": round(current_co2_kg, 2),
        "current_co2_tonnes": round(current_co2_kg / 1000.0, 3),
        "annual_co2_kg": round(annual_co2_kg, 2),
        "annual_co2_tonnes": round(annual_co2_kg / 1000.0, 3),
        "projected_lifetime_co2_kg": round(lifetime_co2_kg, 2),
        "projected_lifetime_co2_tonnes": round(lifetime_co2_kg / 1000.0, 3),
        "projection_years": 40,
        "methodology": (
            "IPCC Tier 1 + Chave et al. (2014); per-tree current CO2 uses measured height where available "
            "(tree_height_m), otherwise modeled height. Annual and 40-year values use species-age growth model."
        ),
        "top_species": top_species,
    }


@router.get("/projects/{project_id}/existing-trees/metrics")
def get_existing_tree_metrics(project_id: int, db: Session = Depends(get_db)):
    _get_project_settings(db, project_id)
    rows = _fetch_existing_tree_export_rows(project_id, db)
    items = []
    for row in rows:
        items.append(
            {
                "tree_id": int(row.get("id") or 0),
                "project_tree_no": row.get("project_tree_no"),
                "inventory_tree_count": row.get("inventory_tree_count"),
                "existing_area_sqm": row.get("existing_area_sqm"),
                "existing_area_ha": row.get("existing_area_ha"),
                "photo_count": len(_normalize_photo_urls(row.get("photo_urls"))),
                "tree_age_months": row.get("tree_age_months"),
                "age_years": row.get("age_years"),
                "age_source": row.get("age_source"),
                "current_co2_kg": row.get("current_co2_kg"),
                "annual_co2_kg": row.get("annual_co2_kg"),
                "lifetime_co2_kg": row.get("lifetime_co2_kg"),
                "co2_in_scope": bool(row.get("co2_in_scope", True)),
                "co2_height_source": row.get("co2_height_source"),
                "height_used_for_co2": bool(row.get("height_used_for_co2")),
            }
        )
    return {
        "project_id": project_id,
        "count": len(items),
        "items": items,
        "generated_at": datetime.utcnow().isoformat(),
    }


@router.get("/projects/{project_id}/existing-trees/export/csv")
def export_existing_trees_csv(project_id: int, db: Session = Depends(get_db)):
    project = get_project(project_id, db)
    rows = _fetch_existing_tree_export_rows(project_id, db)
    summary = _summarize_existing_tree_export_rows(rows)

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_csv = tempfile.NamedTemporaryFile(suffix="_existing_trees_detailed.csv", delete=False)
    csv_path = tmp_csv.name
    tmp_csv.close()

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = _excel_csv_writer(f)
        writer.writerow(["LandCheck Existing Trees Detailed Export"])
        writer.writerow(["Project", project.get("name", "")])
        writer.writerow(["Location", project.get("location_text", "")])
        writer.writerow(["Generated At (UTC)", datetime.utcnow().isoformat()])
        writer.writerow([])
        writer.writerow([
            "Summary",
            f"Existing Trees {summary.get('total_existing_trees', 0)}",
            f"Rows {summary.get('total_existing_rows', len(rows))}",
            f"Area {summary.get('total_existing_area_ha', 0)} ha",
            f"Carbon Scope {summary.get('carbon_scope_rows', 0)}",
            f"Excluded {summary.get('carbon_excluded_rows', 0)}",
            f"Current CO2 {summary.get('current_co2_tonnes', 0)} t",
            f"Annual CO2 {summary.get('annual_co2_tonnes', 0)} t/yr",
            f"40Y Projection {summary.get('projected_lifetime_co2_tonnes', 0)} t",
        ])
        writer.writerow(["Methodology", summary.get("methodology", "")])
        writer.writerow([])
        writer.writerow([
            "tree_id",
            "project_tree_no",
            "project_id",
            "lng",
            "lat",
            "inventory_tree_count",
            "existing_area_sqm",
            "existing_area_ha",
            "species",
            "species_matched",
            "growth_class",
            "planting_date",
            "age_years",
            "age_source",
            "status",
            "tree_height_m",
            "tree_age_months",
            "height_used_for_co2",
            "co2_height_source",
            "count_in_carbon_scope",
            "current_co2_kg",
            "annual_co2_kg",
            "lifetime_co2_kg_40y",
            "tree_origin",
            "attribution_scope",
            "count_in_planting_kpis",
            "source_project_id",
            "custodian_id",
            "custodian_name",
            "created_by",
            "created_at",
            "notes",
            "photo_url",
            "photo_urls",
            "existing_area_geojson",
            "maintenance_count",
            "maintenance_done",
            "maintenance_pending",
            "maintenance_overdue",
            "maintenance_types",
            "last_maintenance_type",
            "last_maintenance_date",
            "review_submitted",
            "review_approved",
            "review_rejected",
            "last_review_state",
            "last_review_note",
            "last_submitted_at",
            "last_reviewed_at",
        ])
        for row in rows:
            writer.writerow([
                row.get("id"),
                row.get("project_tree_no"),
                row.get("project_id"),
                row.get("lng"),
                row.get("lat"),
                row.get("inventory_tree_count"),
                row.get("existing_area_sqm"),
                row.get("existing_area_ha"),
                row.get("species"),
                row.get("species_matched"),
                row.get("growth_class"),
                _to_iso_text(row.get("planting_date")),
                row.get("age_years"),
                row.get("age_source"),
                row.get("status"),
                row.get("tree_height_m"),
                row.get("tree_age_months"),
                row.get("height_used_for_co2"),
                row.get("co2_height_source"),
                row.get("count_in_carbon_scope"),
                row.get("current_co2_kg"),
                row.get("annual_co2_kg"),
                row.get("lifetime_co2_kg"),
                row.get("tree_origin"),
                row.get("attribution_scope"),
                row.get("count_in_planting_kpis"),
                row.get("source_project_id"),
                row.get("custodian_id"),
                row.get("custodian_name"),
                row.get("created_by"),
                _to_iso_text(row.get("created_at")),
                row.get("notes"),
                row.get("photo_url"),
                json.dumps(_normalize_photo_urls(row.get("photo_urls"))),
                _safe_json(row.get("existing_area_geojson")) if row.get("existing_area_geojson") is not None else "",
                row.get("maintenance_count", 0),
                row.get("maintenance_done", 0),
                row.get("maintenance_pending", 0),
                row.get("maintenance_overdue", 0),
                row.get("maintenance_types", ""),
                row.get("last_maintenance_type", ""),
                _to_iso_text(row.get("last_maintenance_date")),
                row.get("review_submitted", 0),
                row.get("review_approved", 0),
                row.get("review_rejected", 0),
                row.get("last_review_state", ""),
                row.get("last_review_note", ""),
                _to_iso_text(row.get("last_submitted_at")),
                _to_iso_text(row.get("last_reviewed_at")),
            ])

    filename = f"project_{project_id}_existing_trees_detailed.csv"
    return FileResponse(csv_path, media_type="text/csv", filename=filename)


@router.get("/projects/{project_id}/existing-trees/export/pdf")
def export_existing_trees_pdf(
    project_id: int,
    include_photos: bool = Query(default=False),
    db: Session = Depends(get_db),
):
    project = get_project(project_id, db)
    rows = _fetch_existing_tree_export_rows(project_id, db)
    summary = _summarize_existing_tree_export_rows(rows)

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_existing_trees_detailed.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()

    photo_rows: list[dict] = []
    if include_photos:
        for row in rows:
            merged_urls = _normalize_photo_urls(row.get("photo_urls"))
            if not merged_urls and str(row.get("photo_url") or "").strip():
                merged_urls = [str(row.get("photo_url") or "").strip()]
            if not merged_urls:
                continue
            for index, photo_url in enumerate(merged_urls, start=1):
                item = dict(row)
                item["photo_url"] = photo_url
                item["photo_index"] = index
                item["photo_total"] = len(merged_urls)
                photo_rows.append(item)
    render_green_existing_trees_report_pdf(
        pdf_path,
        project=project,
        rows=[dict(r) for r in rows],
        summary=summary,
        include_photos=include_photos,
        photo_rows=photo_rows,
    )
    filename = f"project_{project_id}_existing_trees_detailed.pdf"
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


@router.get("/projects/{project_id}/export/csv")
def export_project_csv(project_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT
            t.id,
            t.project_id,
            t.species,
            t.planting_date,
            t.status,
            t.notes,
            t.photo_url,
            t.created_by,
            t.created_at,
            t.tree_origin,
            t.attribution_scope,
            t.count_in_planting_kpis,
            t.count_in_carbon_scope,
            t.custodian_id,
            c.name AS custodian_name,
            t.tree_height_m,
            ST_X(t.geom) AS lng,
            ST_Y(t.geom) AS lat
        FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.project_id = :project_id
        ORDER BY t.created_at DESC
    """), {"project_id": project_id}).mappings().all()
    maintenance_rows = _maintenance_summary_by_tree(project_id, db)
    rows = _attach_maintenance_to_tree_rows(rows, maintenance_rows)
    review_summary = _review_summary_by_tree(project_id, db)
    kpi = _compute_kpi_snapshot(project_id, db)

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_csv = tempfile.NamedTemporaryFile(suffix="_trees.csv", delete=False)
    csv_path = tmp_csv.name
    tmp_csv.close()

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = _excel_csv_writer(f)
        writer.writerow(["LandCheck Project Export"])
        writer.writerow([
            "KPI",
            f"Trees {kpi.get('trees_total', 0)}",
            f"Healthy {kpi.get('trees_healthy', 0)}",
            f"Attention {kpi.get('trees_attention', 0)}",
            f"Open Tasks {kpi.get('tasks_open', 0)}",
            f"Submitted {kpi.get('tasks_submitted', 0)}",
            f"Rejected {kpi.get('tasks_rejected', 0)}",
        ])
        writer.writerow([])
        writer.writerow([
            "tree_id", "project_id", "lng", "lat", "tree_origin", "attribution_scope",
            "count_in_planting_kpis", "count_in_carbon_scope", "custodian_id", "custodian_name", "tree_height_m",
            "species", "planting_date",
            "status", "notes", "photo_url", "created_by", "created_at",
            "maintenance_count", "maintenance_done", "maintenance_pending", "maintenance_overdue",
            "maintenance_types", "last_maintenance_type", "last_maintenance_date",
            "review_submitted", "review_approved", "review_rejected", "last_review_state", "last_review_note",
            "last_submitted_at", "last_reviewed_at"
        ])
        for r in rows:
            review = review_summary.get(int(r["id"]), {})
            writer.writerow([
                r["id"], r["project_id"], r["lng"], r["lat"], r.get("tree_origin"), r.get("attribution_scope"),
                r.get("count_in_planting_kpis"), r.get("count_in_carbon_scope"), r.get("custodian_id"), r.get("custodian_name"), r.get("tree_height_m"), r["species"],
                r["planting_date"], r["status"], r["notes"], r["photo_url"],
                r["created_by"], r["created_at"],
                r.get("maintenance_count", 0), r.get("maintenance_done", 0), r.get("maintenance_pending", 0), r.get("maintenance_overdue", 0),
                r.get("maintenance_types", ""), r.get("last_maintenance_type", ""), r.get("last_maintenance_date", ""),
                review.get("review_submitted", 0), review.get("review_approved", 0), review.get("review_rejected", 0),
                review.get("last_review_state", ""), review.get("last_review_note", ""),
                review.get("last_submitted_at", ""), review.get("last_reviewed_at", ""),
            ])

    filename = f"project_{project_id}_trees.csv"
    return FileResponse(csv_path, media_type="text/csv", filename=filename)


def _maintenance_summary_by_tree(project_id: int, db: Session, assignee_name: str | None = None) -> list[dict]:
    rows = db.execute(text("""
        SELECT tr.id AS tree_id,
               COUNT(t.id) AS maintenance_count,
               SUM(
                   CASE
                       WHEN LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                            AND LOWER(COALESCE(t.review_state, 'none')) IN ('approved', 'none')
                       THEN 1 ELSE 0
                   END
               ) AS maintenance_done,
               SUM(
                   CASE
                       WHEN NOT (
                           LOWER(COALESCE(t.status, '')) IN ('done', 'completed', 'closed')
                           AND LOWER(COALESCE(t.review_state, 'none')) IN ('approved', 'none')
                       ) AND LOWER(COALESCE(t.status, 'pending')) <> 'overdue'
                       THEN 1 ELSE 0
                   END
               ) AS maintenance_pending,
               SUM(CASE WHEN t.status = 'overdue' THEN 1 ELSE 0 END) AS maintenance_overdue,
               COALESCE(STRING_AGG(DISTINCT t.task_type, ', ' ORDER BY t.task_type), '') AS maintenance_types,
               MAX(COALESCE(t.completed_at::date, t.due_date, t.created_at::date)) AS last_maintenance_date,
               (ARRAY_AGG(t.task_type ORDER BY COALESCE(t.completed_at, t.created_at) DESC NULLS LAST, t.id DESC))[1]
                 AS last_maintenance_type
        FROM trees tr
        LEFT JOIN tree_tasks t ON t.tree_id = tr.id
        WHERE tr.project_id = :project_id
          AND (:assignee_name IS NULL OR tr.created_by = :assignee_name)
        GROUP BY tr.id
        ORDER BY tr.id
    """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()

    cleaned = []
    for row in rows:
        item = dict(row)
        item["maintenance_count"] = int(item.get("maintenance_count") or 0)
        item["maintenance_done"] = int(item.get("maintenance_done") or 0)
        item["maintenance_pending"] = int(item.get("maintenance_pending") or 0)
        item["maintenance_overdue"] = int(item.get("maintenance_overdue") or 0)
        item["maintenance_types"] = item.get("maintenance_types") or ""
        item["last_maintenance_type"] = item.get("last_maintenance_type") or ""
        cleaned.append(item)
    return cleaned


def _attach_maintenance_to_tree_rows(rows: list[dict], summary_rows: list[dict]) -> list[dict]:
    summary_by_tree = {r["tree_id"]: r for r in summary_rows}
    merged = []
    for row in rows:
        item = dict(row)
        summary = summary_by_tree.get(item.get("id")) or {}
        item["maintenance_count"] = summary.get("maintenance_count", 0)
        item["maintenance_done"] = summary.get("maintenance_done", 0)
        item["maintenance_pending"] = summary.get("maintenance_pending", 0)
        item["maintenance_overdue"] = summary.get("maintenance_overdue", 0)
        item["maintenance_types"] = summary.get("maintenance_types", "")
        item["last_maintenance_type"] = summary.get("last_maintenance_type", "")
        item["last_maintenance_date"] = summary.get("last_maintenance_date")
        merged.append(item)
    return merged


@router.get("/projects/{project_id}/export/pdf")
def export_project_pdf(
    project_id: int,
    lng: float | None = Query(default=None),
    lat: float | None = Query(default=None),
    zoom: float | None = Query(default=None),
    bearing: float | None = Query(default=0.0),
    pitch: float | None = Query(default=0.0),
    db: Session = Depends(get_db),
):
    def _coerce_optional_float(value: object) -> float | None:
        if value is None:
            return None
        try:
            if isinstance(value, bool):
                return None
            return float(value)
        except Exception:
            return None

    lng_value = _coerce_optional_float(lng)
    lat_value = _coerce_optional_float(lat)
    zoom_value = _coerce_optional_float(zoom)
    bearing_value = _coerce_optional_float(bearing)
    pitch_value = _coerce_optional_float(pitch)

    project = get_project(project_id, db)
    rows = db.execute(text("""
        SELECT
               t.id,
               t.species,
               t.planting_date,
               t.status,
               t.notes,
               t.tree_origin,
               t.tree_height_m,
               t.attribution_scope,
               t.count_in_planting_kpis,
               t.count_in_carbon_scope,
               t.custodian_id,
               c.name AS custodian_name,
               ST_X(t.geom) AS lng,
               ST_Y(t.geom) AS lat
        FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.project_id = :project_id
        ORDER BY t.created_at DESC
        LIMIT 200
    """), {"project_id": project_id}).mappings().all()
    map_rows = db.execute(text("""
        SELECT
               t.id,
               t.species,
               t.planting_date,
               t.status,
               t.notes,
               t.tree_origin,
               t.tree_height_m,
               t.attribution_scope,
               t.count_in_planting_kpis,
               t.count_in_carbon_scope,
               t.custodian_id,
               c.name AS custodian_name,
               ST_X(t.geom) AS lng,
               ST_Y(t.geom) AS lat
        FROM trees t
        LEFT JOIN green_custodians c ON c.id = t.custodian_id
        WHERE t.project_id = :project_id
        ORDER BY t.created_at DESC
        LIMIT 1000
    """), {"project_id": project_id}).mappings().all()
    maintenance_rows = _maintenance_summary_by_tree(project_id, db)
    rows = _attach_maintenance_to_tree_rows(rows, maintenance_rows)
    map_rows = _attach_maintenance_to_tree_rows(map_rows, maintenance_rows)

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_project_report.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()

    map_png = _build_report_map_png(
        map_rows=map_rows,
        lng=lng_value,
        lat=lat_value,
        zoom=zoom_value,
        bearing=bearing_value,
        pitch=pitch_value,
    )

    map_view = None
    if lng_value is not None and lat_value is not None and zoom_value is not None:
        map_view = {"lng": lng_value, "lat": lat_value, "zoom": zoom_value}
    donor_rows = _build_donor_report_rows(project_id, db)
    kpi_snapshot = _compute_kpi_snapshot(project_id, db)
    try:
        _store_kpi_snapshot(project_id, kpi_snapshot, db)
        db.commit()
    except Exception:
        db.rollback()

    # Carbon data for executive summary
    carbon_trees = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :pid
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"pid": project_id}).mappings().all()
    carbon_data = compute_project_carbon([dict(r) for r in carbon_trees])
    carbon_data["projection"] = generate_co2_projection_table([dict(r) for r in carbon_trees], 30)

    # KPI trend for survival chart
    kpi_trend = _fetch_kpi_trend(project_id, db, days=90)
    species_daily_survival = _build_species_daily_survival_series(project_id, db)

    try:
        render_green_report_pdf(
            pdf_path,
            project,
            rows,
            map_png=map_png,
            map_rows=map_rows,
            map_view=map_view,
            maintenance_rows=maintenance_rows,
            donor_rows=donor_rows,
            kpi_snapshot=kpi_snapshot,
            carbon_data=carbon_data,
            kpi_trend=kpi_trend,
            species_daily_survival=species_daily_survival,
        )
    except Exception:
        render_green_report_pdf(
            pdf_path,
            project,
            rows,
            map_png=None,
            map_rows=map_rows,
            map_view=map_view,
            maintenance_rows=maintenance_rows,
            donor_rows=donor_rows,
            kpi_snapshot=kpi_snapshot,
            carbon_data=carbon_data,
            kpi_trend=kpi_trend,
            species_daily_survival=species_daily_survival,
        )
    filename = f"project_{project_id}_report.pdf"
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


def _fetch_kpi_trend(project_id: int, db: Session, days: int = 90) -> list[dict]:
    """Fetch KPI trend series for charts using cohort/activity monthly basis."""
    return _build_kpi_trend_series(project_id, db, days=days)


def _build_tree_stats(rows: list[dict]) -> dict:
    scoped_rows = [r for r in rows if bool(r.get("count_in_planting_kpis", True))]
    total = len(scoped_rows)
    alive = sum(1 for r in scoped_rows if _normalize_tree_status(r.get("status")) in HEALTHY_TREE_STATUSES)
    dead = sum(1 for r in scoped_rows if _normalize_tree_status(r.get("status")) in DEAD_TREE_STATUSES)
    needs_attention = sum(1 for r in scoped_rows if _normalize_tree_status(r.get("status")) in ATTENTION_TREE_STATUSES)
    pending = sum(1 for r in scoped_rows if _normalize_tree_status(r.get("status")) == "pending_planting")
    survival_rate = round((alive / total) * 100, 1) if total else 0.0
    return {
        "total": total,
        "total_all": len(rows),
        "alive": alive,
        "dead": dead,
        "needs_attention": needs_attention,
        "pending_planting": pending,
        "survival_rate": survival_rate,
    }


@router.get("/work-report/pdf")
def export_work_report_pdf(
    project_id: int,
    assignee_name: str | None = None,
    include_photos: bool = Query(default=False),
    lng: float | None = Query(default=None),
    lat: float | None = Query(default=None),
    zoom: float | None = Query(default=None),
    bearing: float | None = Query(default=0.0),
    pitch: float | None = Query(default=0.0),
    db: Session = Depends(get_db),
):
    def _coerce_optional_float(value: object) -> float | None:
        if value is None:
            return None
        try:
            if isinstance(value, bool):
                return None
            return float(value)
        except Exception:
            return None

    lng_value = _coerce_optional_float(lng)
    lat_value = _coerce_optional_float(lat)
    zoom_value = _coerce_optional_float(zoom)
    bearing_value = _coerce_optional_float(bearing)
    pitch_value = _coerce_optional_float(pitch)

    project = get_project(project_id, db)
    if assignee_name:
        rows = db.execute(text("""
            SELECT
                   t.id,
                   t.species,
                   t.planting_date,
                   t.status,
                   t.notes,
                   t.photo_url,
                   t.created_by,
                   t.tree_origin,
                   t.tree_height_m,
                   t.attribution_scope,
                   t.count_in_planting_kpis,
                   t.count_in_carbon_scope,
                   t.custodian_id,
                   c.name AS custodian_name,
                   ST_X(t.geom) AS lng,
                   ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id AND t.created_by = :assignee_name
            ORDER BY t.created_at DESC
            LIMIT 200
        """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()
        map_rows = db.execute(text("""
            SELECT
                   t.id,
                   t.species,
                   t.planting_date,
                   t.status,
                   t.notes,
                   t.photo_url,
                   t.created_by,
                   t.tree_origin,
                   t.tree_height_m,
                   t.attribution_scope,
                   t.count_in_planting_kpis,
                   t.count_in_carbon_scope,
                   t.custodian_id,
                   c.name AS custodian_name,
                   ST_X(t.geom) AS lng,
                   ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id AND t.created_by = :assignee_name
            ORDER BY t.created_at DESC
            LIMIT 1000
        """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()
    else:
        rows = db.execute(text("""
            SELECT
                   t.id,
                   t.species,
                   t.planting_date,
                   t.status,
                   t.notes,
                   t.photo_url,
                   t.created_by,
                   t.tree_origin,
                   t.tree_height_m,
                   t.attribution_scope,
                   t.count_in_planting_kpis,
                   t.count_in_carbon_scope,
                   t.custodian_id,
                   c.name AS custodian_name,
                   ST_X(t.geom) AS lng,
                   ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id
            ORDER BY t.created_at DESC
            LIMIT 200
        """), {"project_id": project_id}).mappings().all()
        map_rows = db.execute(text("""
            SELECT
                   t.id,
                   t.species,
                   t.planting_date,
                   t.status,
                   t.notes,
                   t.photo_url,
                   t.created_by,
                   t.tree_origin,
                   t.tree_height_m,
                   t.attribution_scope,
                   t.count_in_planting_kpis,
                   t.count_in_carbon_scope,
                   t.custodian_id,
                   c.name AS custodian_name,
                   ST_X(t.geom) AS lng,
                   ST_Y(t.geom) AS lat
            FROM trees t
            LEFT JOIN green_custodians c ON c.id = t.custodian_id
            WHERE t.project_id = :project_id
            ORDER BY t.created_at DESC
            LIMIT 1000
        """), {"project_id": project_id}).mappings().all()

    photo_rows: list[dict] = []
    if include_photos:
        if assignee_name:
            photo_rows = db.execute(text("""
                SELECT
                       t.id,
                       t.species,
                       t.planting_date,
                       t.status,
                       t.tree_origin,
                       t.tree_height_m,
                       t.attribution_scope,
                       t.created_by,
                       t.custodian_id,
                       c.name AS custodian_name,
                       t.photo_url
                FROM trees t
                LEFT JOIN green_custodians c ON c.id = t.custodian_id
                WHERE t.project_id = :project_id
                  AND t.created_by = :assignee_name
                  AND COALESCE(TRIM(t.photo_url), '') <> ''
                ORDER BY t.created_at DESC, t.id DESC
            """), {"project_id": project_id, "assignee_name": assignee_name}).mappings().all()
        else:
            photo_rows = db.execute(text("""
                SELECT
                       t.id,
                       t.species,
                       t.planting_date,
                       t.status,
                       t.tree_origin,
                       t.tree_height_m,
                       t.attribution_scope,
                       t.created_by,
                       t.custodian_id,
                       c.name AS custodian_name,
                       t.photo_url
                FROM trees t
                LEFT JOIN green_custodians c ON c.id = t.custodian_id
                WHERE t.project_id = :project_id
                  AND COALESCE(TRIM(t.photo_url), '') <> ''
                ORDER BY t.created_at DESC, t.id DESC
            """), {"project_id": project_id}).mappings().all()

    maintenance_rows = _maintenance_summary_by_tree(project_id, db, assignee_name)
    rows = _attach_maintenance_to_tree_rows(rows, maintenance_rows)
    map_rows = _attach_maintenance_to_tree_rows(map_rows, maintenance_rows)

    project_copy = dict(project)
    project_copy["stats"] = _build_tree_stats(map_rows)
    project_copy["report_assignee"] = assignee_name

    os.makedirs(REPORTS_DIR, exist_ok=True)
    tmp_pdf = tempfile.NamedTemporaryFile(suffix="_work_map_report.pdf", delete=False)
    pdf_path = tmp_pdf.name
    tmp_pdf.close()

    map_png = _build_report_map_png(
        map_rows=map_rows,
        lng=lng_value,
        lat=lat_value,
        zoom=zoom_value,
        bearing=bearing_value,
        pitch=pitch_value,
    )

    map_view = None
    if lng_value is not None and lat_value is not None and zoom_value is not None:
        map_view = {"lng": lng_value, "lat": lat_value, "zoom": zoom_value}
    donor_rows = _build_donor_report_rows(project_id, db)
    kpi_snapshot = _compute_kpi_snapshot(project_id, db)
    try:
        _store_kpi_snapshot(project_id, kpi_snapshot, db)
        db.commit()
    except Exception:
        db.rollback()

    # Carbon data for executive summary
    carbon_trees = db.execute(text("""
        SELECT id, species, planting_date, status, created_at, tree_age_months, COALESCE(inventory_tree_count, 1) AS inventory_tree_count
        FROM trees
        WHERE project_id = :pid
          AND COALESCE(count_in_carbon_scope, TRUE) = TRUE
    """), {"pid": project_id}).mappings().all()
    carbon_data = compute_project_carbon([dict(r) for r in carbon_trees])
    carbon_data["projection"] = generate_co2_projection_table([dict(r) for r in carbon_trees], 30)

    # KPI trend for survival chart
    kpi_trend = _fetch_kpi_trend(project_id, db, days=90)
    species_daily_survival = _build_species_daily_survival_series(project_id, db)

    try:
        render_green_report_pdf(
            pdf_path,
            project_copy,
            rows,
            map_png=map_png,
            map_rows=map_rows,
            map_view=map_view,
            maintenance_rows=maintenance_rows,
            donor_rows=donor_rows,
            kpi_snapshot=kpi_snapshot,
            carbon_data=carbon_data,
            kpi_trend=kpi_trend,
            species_daily_survival=species_daily_survival,
            photo_rows=photo_rows,
            include_photos=include_photos,
        )
    except Exception:
        render_green_report_pdf(
            pdf_path,
            project_copy,
            rows,
            map_png=None,
            map_rows=map_rows,
            map_view=map_view,
            maintenance_rows=maintenance_rows,
            donor_rows=donor_rows,
            kpi_snapshot=kpi_snapshot,
            carbon_data=carbon_data,
            kpi_trend=kpi_trend,
            species_daily_survival=species_daily_survival,
            photo_rows=photo_rows,
            include_photos=include_photos,
        )
    filename = (
        f"project_{project_id}_work_report_{assignee_name}.pdf"
        if assignee_name
        else f"project_{project_id}_work_report_all.pdf"
    )
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)
